# -*- coding: utf-8 -*-
"""
Created on Fri Feb  1 18:17:55 2019

@author: Marion
"""

import pandas as pd
import docx
from docx.api import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
import qn
#from docx.shared import RGBColor

#docmuents du cpp pour les dispositifs médicaux

def main_cpp_hps():
    document = docx.Document()
    cpp_hps(document)
    page2_cpp_hps(document)
    cpp_hps_annexe2(document)
    document.save("soumission-cpp-hps.docx")

def cpp_hps(document):
    

    
    '''Marge de la page'''
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    '''Titre CPP'''
  #  paragraph = document.add_paragraph()
    styles= document.styles
    style1 = styles.add_style('Debut', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style1.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style1.font
    fontdebut.name = 'Book Antiqua'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(20) 
    
    
    
    paragraph1 = document.add_paragraph('Comité de Protection des Personnes', style='Debut')
    paragraph1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    #Sous titre
     
    
    paragraph = document.add_paragraph('OUEST III', style='Debut')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    #ajouter le trait et les ombres
    
    #Infos promoteur
    
    style2=styles.add_style('Promoteur', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style2.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style2.font
    fontdebut.name = 'Book Antiqua'
    fontdebut.italic = True
    fontdebut.size = docx.shared.Pt(10) 
      
    paragraph = document.add_paragraph("Agréé par arrêté ministériel en date du 31 mai 2012, \nConstitué selon l'arrêté du Directeur Général de l'ARS Poitou Charentes en date du 25 juin 2012.\n\n"
                                       "C.H.U La Milétrie\nPavillon Administratif - Porte 213\n "
                                       "2 rue de le milétrie - CS 90 577 - 86021 POITIERS CEDEX\n"
                                       "Tel : 05.49.45.21.57\nFax : 05.49.46.12.62 \nE-mail : cpp-ouest3@chu-poitiers.fr \n", style='Promoteur')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    
    #titre du milieu qui dit pour quel proto
    
    paragraph = document.add_paragraph()
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = paragraph.add_run("Demande d'avis au CPP")
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence.font.name = 'Arial Narrow'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(12)
    sentence2 = paragraph.add_run(" (arrêté du 2 décembre 2016)\n")
    sentence2.font.name = 'Arial Narrow'
    sentence2.font.size = docx.shared.Pt(12)
    sentence3 = paragraph.add_run("sur un projet de recherche mentionnée au 1o ou au 2o de l’article L. 1121-1 du code de la santé publique ne portant pas sur un produit mentionné à l’article L. 5311-1 du même code.")
    sentence3.font.name = 'Arial Narrow'
    sentence3.bold = True
    sentence3.font.size = docx.shared.Pt(12)
    sentence4 = paragraph.add_run("(les médicaments, les produits contraceptifs, les biomatériaux et les dispositifs médicaux …)\n")
    sentence4.font.name = 'Arial Narrow'
    sentence4.font.size = docx.shared.Pt(12)
    
 ###########################################   
 
    #ENtre le titre et le tableau
 
    style3=styles.add_style('Avant_tableau', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style3.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
#    paragraph_format.left_indent = Inches(10)
    fontdebut = style3.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.italic = True
    fontdebut.size = docx.shared.Pt(10) 
    
    paragraph = document.add_paragraph("Préalablement au dépôt du dossier le promoteur obtient un numéro d’enregistrement sur le site internet de l’ANSM. Ce numéro identifie chaque recherche réalisée en France.", style='Avant_tableau')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
  
 ##############################

    #Tableau central
    table=document.add_table(rows=15, cols=1, style='Table Grid')
    

    table.cell(0,0).text=("DOSSIER ADMINISTRATIF")
    table.cell(1,0).text=("1 Courrier de demande d’avis daté et signé")
    table.cell(2,0).text=("2 Formulaire de demande d’avis (annexe 1)")
    table.cell(3,0).text=("3 Document additionnel (annexe 2) + supports pour recrutement des personnes")
    table.cell(4,0).text=("8.2 Pour les recherches mentionnées au 1o de l’article L. 1121-1, si nécessaire, la copie de la ou des autorisations de lieux de recherches mentionnées à l’article L. 1121-13 du CSP")
    table.cell(5,0).text=("DOSSIER SUR LA RECHERCHE")
    table.cell(6,0).text=("4 Protocole de recherche (daté + numéro de version)")
    table.cell(7,0).text=("5 Résumé du protocole (daté + numéro de version)")
    table.cell(8,0).text=("Le cas échéant, la brochure pour l’investigateur mentionnée à l’article R. 1123-20 du code de la santé publique, datée et comportant un numéro de version, lorsque la recherche porte sur un produit autre que ceux mentionnés à l’article L. 5311-1 du CSP")
    table.cell(9,0).text=("6.1 Document d’information sauf situation art. L. 1122-1-4")
    table.cell(10,0).text=("6.2 Formulaire de consentement sauf situation art. L. 1122-1-4")
    table.cell(11,0).text=("7 Attestation d’assurance (Décret n°2016-1537 du 16 novembre 2016 - art. 3)")
    table.cell(12,0).text=("8 Une justification de l’adéquation des moyens humains, matériels et techniques au projet de recherche et de leur compatibilité avec les impératifs de sécurité des personnes qui s’y prêtent, sauf si le lieu bénéficie de l’autorisation mentionnée à l’article L. 1121-13 du CSP")
    table.cell(13,0).text=("8.1 Curriculum vitae signé du ou des investigateurs datant d’un an maximum")
    table.cell(14,0).text=("Le cas échéant, la nature de la décision finale de l’ANSM, si disponible.")
    n=1
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    if n==1 or n==6:
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        #run.bold = True
                        font = run.font
                        font.size= docx.shared.Pt(11)
                        font.name = 'Arial'
                    else:
                        font = run.font
                        font.size= docx.shared.Pt(10)
                        font.name = 'Arial Narrow'
                        #paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
                    n=n+1
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="AFAFAF"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="AFAFAF"/>'.format(nsdecls('w')))
    table.rows[5].cells[0]._tc.get_or_add_tcPr().append(shading_elm_2)
    
    style5=styles.add_style('fin_tableau', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style5.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style5.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.italic = True
    fontdebut.size = docx.shared.Pt(10) 
    document.add_paragraph('Forme : 4 dossiers complets + 1 version électronique\n\n', style='fin_tableau')
    
    
def page2_cpp_hps(document):
    document.add_page_break()
    styles= document.styles
    style=styles.add_style('debut_page', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(10) 
    
    paragraph = document.add_paragraph("Annexe 1\nFormulaire de damande d'avis au comité de protection des personnes pour une recherche\nmentionnée au 1° ou au 2° de l'article L.1121-1 du code de la santé publique et ne portant pas\nsur un produit de santé\n", style="debut_page")
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sentence=("Demande d'avis au comite de protection des personnes pour une recherche\nmentionnee au 1° ou 2° de l'article L.1121-1 du code de la sante publique et ne\nportant pas sur un produit mentionne a\nl'article L. 5311-1 du code de la sante publique\n")
    sentence.upper()
    
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=sentence
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.bold = True
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
                    fontdebut.color.rgb = RGBColor(0x0,0x70,0xC0)
    
    style=styles.add_style('gras_tableau', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style.font
    fontdebut.bold = True
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(10) 
    
    paragraph = document.add_paragraph("Partie réservée au Comité de protection des personnes (CPP)", style="gras_tableau")
               
    table = document.add_table(rows=2, cols=3, style='Table Grid')
    table.cell(0,0).text=("Date d'enregistrement de la\ndemande considérée complète :")
    table.cell(0,1).text=("Date de réception des informations\ncomplémentaires / amendées :")
    table.cell(0,2).text=("Avis du CPP :")
    a=table.cell(1,0)
    b=table.cell(1,2)
    a.merge(b)
    table.cell(1,0).text=("Date du début de procédure :")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
    
    paragraph = document.add_paragraph("Partie à compléter par le demandeur :", style="gras_tableau")
    paragraph=document.add_paragraph()
    sentence = paragraph.add_run("RECHERCHE MENTIONNEE AU 1° de l'article L.1121-1 □             RECHERCHE MENTIONNEE AU 2° DE L'ARTICLE L.1121-1□\nDEMANDE D'AUTORISATION A L'ANSM :")
    fontdebut = sentence.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(10) 
    sentence = paragraph.add_run("     oui        non\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.bold=True
    fontdebut.size = docx.shared.Pt(10)
    sentence = paragraph.add_run("DEMANDE D'AVIS AU CPP :     ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(10) 
    sentence = paragraph.add_run("     oui        non\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.bold=True
    fontdebut.size = docx.shared.Pt(10)
    
    paragraph = document.add_paragraph("A. IDENTIFICATION DE LA RECHERCHE", style="gras_tableau")
    paragraph=document.add_paragraph()
    sentence = paragraph.add_run("Titre complet de la recherche :\n \nNuméro d'enregistrement de la recherche (délivré par l'ANSM) : \n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(10) 
    
    table = document.add_table(rows=1, cols=3, style='Table Grid')
    table.cell(0,0).text=("Numéro de code du promoteur de\nla recherche donné par le\npromoteur")
    table.cell(0,1).text=("Version")
    table.cell(0,2).text=("Date :")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
    
    style=styles.add_style('normal', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(10) 
    
    
    paragraph = document.add_paragraph("Nom ou titre abrégé de la recherche,\nle cas échéant:\nJustifier la catégorie de votre recherche\n",style="normal")
    paragraph=document.add_paragraph()
    sentence = paragraph.add_run("Inscription au fichier VRB")
    fontdebut = sentence.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(10) 
    sentence = paragraph.add_run("     oui        non\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.bold=True
    fontdebut.size = docx.shared.Pt(10)
    
    
    paragraph = document.add_paragraph("B. IDENTIFICATION DU PROMOTEUR RESPONSABLE DE LA DEMANDE\n    B1. Promoteur",style="gras_tableau")
    table = document.add_table(rows=5, cols=2, style='Table Grid')
    a=table.cell(0,0)
    b=table.cell(0,1)
    a.merge(b)
    c=table.cell(1,0)
    d=table.cell(1,1)
    c.merge(d)
    table.cell(0,0).text=("Nom de l'organisme :")
    table.cell(1,0).text=("Nom de la personne à contacter :")
    table.cell(0,2).text=("Avis du CPP :")
    e=table.cell(2,0)
    f=table.cell(4,0)
    e.merge(f)
    table.cell(2,0).text=("Adresse :")
    table.cell(2,1).text=("Numéro de téléphone :")
    table.cell(3,1).text=("Numéro de télécopie :")
    table.cell(4,1).text=("Courriel :")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
    paragraph = document.add_paragraph("B2. Représentant légal du promoteur dans l'Union européenne pour la recherche (si différent du promoteur)",style="gras_tableau")
    table = document.add_table(rows=6, cols=2, style='Table Grid')
    a=table.cell(0,0)
    b=table.cell(0,1)
    a.merge(b)
    c=table.cell(1,0)
    d=table.cell(1,1)
    c.merge(d)
    table.cell(0,0).text=("Nom de l'organisme :")
    table.cell(1,0).text=("Nom de la personne à contacter :")
    table.cell(0,2).text=("Avis du CPP :")
    e=table.cell(2,0)
    f=table.cell(4,0)
    e.merge(f)
    table.cell(2,0).text=("Adresse :")
    table.cell(2,1).text=("Numéro de téléphone :")
    table.cell(3,1).text=("Numéro de télécopie :")
    table.cell(4,1).text=("Courriel :")
    g=table.cell(5,0)
    h=table.cell(5,1)
    g.merge(h)
    table.cell(5,0).text=("Statut du promoteur :         commercial       non commercial")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
    
    paragraph = document.add_paragraph("\nC. IDENTIFICATION DU DEMANDEUR",style="gras_tableau")
    paragraph = document.add_paragraph("Nom de l'organisme : \n   Adresse :\n\n   Numéro de téléphone :\n   Numéro de télécopie :\n   Courriel :\n\nNom de la personne à contacter :\n\n   Adresse :\n   Numéro de téléphone :\n   Numéro de télécopie :\n   Courriel :\n",style="normal")
    paragraph = document.add_paragraph("\nD. DONNEES SUR LE(S) PRODUIT(S) EXPERIMENTAL(AUX) UTILISE(S) DANS LA RECHERCHE:\nPRODUIT(S) ETUDIE(S) OU UTILISE(S) COMME COMPARATEUR(S)",style="gras_tableau")
    
    table = document.add_table(rows=4, cols=1, style='Table Grid')
    table.cell(0,0).text=("Indiquer ici quel PE est concerné par cette section D ; si nécessaire, utiliser d'autres fiches pour chaque PE utilisé dans l'essai (à numéroter de 1 à n) :")
    table.cell(1,0).text=("Cette section concerne le PE numéro :")
    table.cell(2,0).text=("PE étudié        oui         non")
    table.cell(3,0).text=("PE utilisé comme comparateur         oui        non")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
    
    paragraph = document.add_paragraph("DESCRIPTION DU PRODUIT EXPERIMENTAL",style="gras_tableau")
    table = document.add_table(rows=5, cols=1, style='Table Grid')
    table.cell(0,0).text=("Nom du produit, le cas échéant :")
    table.cell(1,0).text=("Nom de code, le cas échéant :")
    table.cell(2,0).text=("Voie d'adiministration (utiliser les termes standard):")
    table.cell(3,0).text=("Dosage (préciser tous les dosages utilisés) : \n-Concentration (nombre) :\n-Unité de concentration :")
    table.cell(4,0).text=("Le produit expérimental contient-il une substance active :\n"
                          "\n-d'origine chimique ?       oui       non"
                          "\n-d'otirigine biologique ?   oui       non"
                          "\n\nEst-ce :\n"
                          "-un produit à base de plantes ?    oui      non\n"
                          "\n-un médicament contenant des organismes génétiquement modifiés ?       oui      non"
                          "\n\n"
                          "         Si oui,\n"
                          "         l'autorisation relative au confinement et à la dissémination volontaire de l'OGM a-t-elle été accordée ?\n"
                          "         ou est-elle en attente ?       oui        non\n\n"
                          "- un autre type de produit ?  oui       non\n\n"
                          "         Si oui, préciser :")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
                    
    paragraph = document.add_paragraph("E. INFORMATIONS SUR LE PLACEBO (le cas échéant) (répéter la section nécessaire)",style="gras_tableau")
    table = document.add_table(rows=6, cols=1, style='Table Grid')
    table.cell(0,0).text=("Cette section se rapport au placebo n° :")
    table.cell(1,0).text=("Un placebo est_il utilisé ?       oui       non")
    table.cell(2,0).text=("De quel produit expérimental est-ce un placebo ?")
    table.cell(3,0).text=("Préciser le(s) numéro(s) de PE selon la section D.")
    table.cell(4,0).text=("Voie d'admission :")
    table.cell(5,0).text=("Composition, hormis la (les) substance(s) active(s) :"
                          "-est-elle identique à celle du produit expérimental étudié?      oui       non\n\n"
                          "        Si non, préciser les principaux composants :")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
      
    
    paragraph = document.add_paragraph("FABRICANT DU PLACEBO",style="gras_tableau")
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("Fabricant\n-Nom de l'établissement :\n-Adresse :")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
                    
    paragraph = document.add_paragraph("G. INFORMATIONS GENERALES SUR LA RECHERCHE",style="gras_tableau")
    table = document.add_table(rows=5, cols=1, style='Table Grid')
    table.cell(0,0).text=("Condition médicale ou pathologique étudiée\n"
                          "Préciser la condition médicale :\n"
                          "Classification CIM :\n"
                          "Classification MedDRA :\n\n"
                          "Est-ce une maladie rare ?      oui      non\n"
                          "\nObjectif(s) de l'essai\n"
                          "Objectif principal :\n"
                          "Objectifs secondaires :\n")
    table.cell(1,0).text=("Principaux critères d'inclusion (énumérer les plus importants)")
    table.cell(2,0).text=("Principaux critères de non inclusion (énumérer les plus importants)")
    table.cell(3,0).text=("Critère(s) d'évaluation principal(aux)")
    table.cell(4,0).text=("Domaine(s) d'étude :\n"
                          "-Physiologie\n-Physiopathologie\n-Epidémiologie\n-Génétique\n"
                          "-Science du comportement\n-Produits à visée nutritionnelle\n"
                          "-Stratégies diagnostiques\n-Stratégies thérapeutiques et préventives\n\n"
                          "          Si autres préciser :")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
                    
    paragraph = document.add_paragraph("I. INVESTIGATEURS ET LIEUX DE RECHERCHE\n       I.1. Investigateur coordonnateur",style="gras_tableau")
    table = document.add_table(rows=5, cols=2, style='Table Grid')
    a=table.cell(0,1)
    b=table.cell(1,1)
    a.merge(b)
    c=table.cell(2,0)
    d=table.cell(2,1)
    c.merge(d)
    e=table.cell(4,0)
    f=table.cell(4,1)
    e.merge(f)
    table.cell(0,0).text=("Nom :")
    table.cell(0,1).text=("Adresse :")
    table.cell(1,0).text=("Prénoms :")
    table.cell(2,0).text=("Qualification, spécialité :")
    table.cell(3,0).text=("Courriel :")
    table.cell(3,1).text=("Téléphone :")
    table.cell(4,0).text=("N°ADELI :")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
    
    paragraph = document.add_paragraph("\n       I.2. Autres investigateurs",style="gras_tableau")
    table = document.add_table(rows=5, cols=2, style='Table Grid')
    a=table.cell(0,1)
    b=table.cell(1,1)
    a.merge(b)
    c=table.cell(2,0)
    d=table.cell(2,1)
    c.merge(d)
    e=table.cell(4,0)
    f=table.cell(4,1)
    e.merge(f)
    table.cell(0,0).text=("Nom :")
    table.cell(0,1).text=("Adresse :")
    table.cell(1,0).text=("Prénoms :")
    table.cell(2,0).text=("Qualification, spécialité :")
    table.cell(3,0).text=("Courriel :")
    table.cell(3,1).text=("Téléphone :")
    table.cell(4,0).text=("N°ADELI :")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
    
    paragraph = document.add_paragraph("\n       I.3. Lieu de recherche (le cas échéant, si la recherche doit se dérouler dans un lieu nécessitant une autorisation de l'ARS) :",style="gras_tableau")
    table = document.add_table(rows=5, cols=1, style='Table Grid')
    table.cell(0,0).text=("Intitulé du lieu :")
    table.cell(1,0).text=("N° d'autorisation :")
    table.cell(2,0).text=("délivré le :")
    table.cell(3,0).text=("date de limite de validité :")
    table.cell(4,0).text=("Nom et adresse :")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(10)
                    
    paragraph = document.add_paragraph("K. SIGNATURE DU DEMANDEUR EN FRANCE\n",style="gras_tableau")
    paragraph= document.add_paragraph()
    sentence = paragraph.add_run("Par la présent, j'atteste/j'atteste au nom du promoteur (rayer la mention inutile) ce qui suit\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.bold=True
    fontdebut.size = docx.shared.Pt(10)
    sentence = paragraph.add_run("-les informations fournies ci-dessus à l'appui de la demande sont exactes ;\n"
                                 "-la recherche sera réalisée conformément au protocole, à la réglementation nationale et aux principes de bonnes pratiques cliniques ;\n"
                                 "-il est raisonnable d'entreprendre la recherche proposée ;\n"
                                 "-je soumettrai un résumé du rapport final de la recherche à l'Ansm et au Comité de protection des personnes concerné au plus tard 1 an après la fin de la recherche dans tous les pays ;"
                                 "-je déclarerai la date effective du commencement de la recherche à l'Ansm et au Comité de protection des personnes concerné dès qu'elle sera connue.\n"
                                 "\n\n"
                                 "                                                        DEMANDEUR auprès du CPP\n"
                                 "                                                        (comme indiqué à la section C) :\n"
                                 "                                                        Date :\n"
                                 "                                                        Signature :\n"
                                 "                                                        Nom:")
    fontdebut = sentence.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(10) 
    
 #   modifyBorder(table)


#def modifyBorder(table):
#    tbl = table._tbl # get xml element in table
#    for cell in tbl.iter_tcs():
#        tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
#        tcBorders = OxmlElement('w:tcBorders')
#        top = OxmlElement('w:top')
#        top.set(qn('w:val'), 'nil')
#        
#        left = OxmlElement('w:left')
#        left.set(qn('w:val'), 'nil')
#        
#        bottom = OxmlElement('w:bottom')
#        bottom.set(qn('w:color'), 'blue')
#
#        right = OxmlElement('w:right')
#        right.set(qn('w:color'), 'blue')
#        
#        left = OxmlElement('w:right')
#        left.set(qn('w:color'), 'blue')
#        
#        top = OxmlElement('w:right')
#        top.set(qn('w:color'), 'blue')
#
#        tcBorders.append(top)
#        tcBorders.append(left)
#        tcBorders.append(bottom)
#        tcBorders.append(right)
#        tcPr.append(tcBorders)
    
 
def cpp_hps_annexe2(document):
    document.add_page_break()
    '''Marge de la page'''
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        
    styles= document.styles
    style1 = styles.add_style('Debut_page2', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style1.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style1.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(12) 
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run('Annexe 2\n')
    sentence.font.name = 'Arial Narrow'
    sentence.font.size = docx.shared.Pt(10.5)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph=document.add_paragraph()
    sentence3=paragraph.add_run("DOCUMENT ADDITIONNEL À LA DEMANDE D’AVIS AU COMITÉ DE PROTECTION DES PERSONNES SUR UN PROJET DE RECHERCHE MENTIONNÉE AU 1° OU AU 2° DE L’ARTICLE L. 1121-1 DU CODE DE LA SANTE PUBLIQUE NE PORTANT PAS SUR UN PRODUIT MENTIONNE A L'ARTICLE L.5311-1 DU MEME CODE\n")
    sentence3.font.name = 'Arial Narrow'
    sentence3.font.size = docx.shared.Pt(12)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW 
    
    style3 = styles.add_style('page2_normal', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style3.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style3.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(11)
    
    paragraph=document.add_paragraph('\nCe document doit être complété de façon claire, compréhensible et en français.\n', style='page2_normal')
  
    
    paragraph=document.add_paragraph('1. Numéro d’enregistrement de la recherche :', style='page2_normal')
    paragraph=document.add_paragraph('2. Titre complet de la recherche :', style='page2_normal')
    paragraph=document.add_paragraph('3. Justification de la recherche :', style='page2_normal')
    paragraph=document.add_paragraph('4. Hypothèse principale de la recherche et objectifs :', style='page2_normal')
    paragraph=document.add_paragraph("5. Evaluation des bénéfices et des risques que présente la recherche, notamment les bénéfices escomptés pour les personnes qui se prêtent à la recherche et les risques prévisibles liés au traitement et aux procédures d'investigation de la recherche (incluant notamment la douleur, l'inconfort, l'atteinte à l'intégrité physique des personnes se prêtant à la recherche, les mesures visant à éviter et/ou prendre en charge les évènements inattendus):", style='page2_normal')
    paragraph=document.add_paragraph("6. Justifications de l’inclusion de personnes visées aux articles L. 1121-5 à L. 1121-8 et L. 1122-1-2 du code de la santé publique (ex. : mineurs, majeurs protégés, recherches mises en oeuvre dans des situations d'urgence, etc.) et procédure mise en oeuvre afin d’informer et de recueillir le consentement de ces personnes ou de leurs représentants légaux :", style='page2_normal')
    paragraph=document.add_paragraph('7. Description des modalités de recrutement des personnes (joindre notamment tous les supports publicitaires utilisés pour la recherche en vue du recrutement des personnes) :', style='page2_normal')
    paragraph=document.add_paragraph('8. Procédures d’investigation menées et différences par rapport à la prise en charge habituelle, le cas échéant :', style='page2_normal')
    paragraph=document.add_paragraph('9. Justification de l’existence ou non :\n-d’une interdiction de participer simultanément à une autre recherche\n-d’une période d’exclusion pendant laquelle la participation à une autre recherche est interdite.', style='page2_normal')   
    paragraph=document.add_paragraph('10. Modalités et montant de l’indemnisation des personnes se prêtant à la recherche, le cas échéant :', style='page2_normal')
    paragraph=document.add_paragraph('11. Motifs de constitution ou non d’un comité de surveillance indépendant :', style='page2_normal')
    paragraph=document.add_paragraph('12. Nombre prévu de personnes à inclure dans la recherche :\n', style='page2_normal')
    paragraph=document.add_paragraph('Par la présente, j’atteste/j’atteste au nom du promoteur (rayer la mention inutile) que les informations fournies ci-dessus à l’appui de la demande d’avis sont exactes.\n', style='page2_normal')
    paragraph=document.add_paragraph('Nom :\nPrénom :\nAdresse :\nFonction :\nDate :\nSignature :', style='page2_normal')
    