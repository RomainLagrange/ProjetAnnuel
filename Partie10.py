# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 15:05:06 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document) --> écrire en minuscule !!!
#    TexteGrisJustif(texte,document)

def Partie10(document,extract):
    'Creation de la partie 10 du protcole de catégorie 1'
 #   document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

#    Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
     #ecriture du premier titre 
    Titre1('10	SURVEILLANCE DE LA RECHERCHE',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le Comité de Surveillance Indépendant (CSI) est un comité consultatif chargé de donner au promoteur d’un essai clinique et au conseil scientifique le cas échéant, un avis sur la conduite de l’essai. ')
    run1.style='Paragraphe'

    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Le CSI est composé de ')
    run1.style='Paragraphe'
    run2=p.add_run('[donner le nombre] ')
    run2.style='Paragraphe'
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('membres indépendants de la recherche, les noms et les fonctions sont décrits dans la charte du CSI, jointe au protocole. Les membres se verront remettre une charte dans laquelle sont décrites les modalités de fonctionnement du comité.')
    run3.style='Paragraphe'
#    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Le comité se réunira au moins ')
    run1.style='Paragraphe'
    run2=p.add_run('[donner le nombre] fois ')
    run2.style='Paragraphe'
    run2.font.italic = True
    run2.font.bold= True
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('par an. Des réunions extraordinaires peuvent être demandées par le promoteur ou le conseil scientifique de l’essai, notamment en cas :')
    run3.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('de fréquence accrue des effets indésirables attendus')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('d’événements indésirables graves liés aux procédures de l’essai clinique dont la fréquence et / ou la gravité serait susceptible de modifier le rapport bénéfice-risque de l’essai.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’avis du CSI est transmis dans les meilleurs délais au promoteur et à l’investigateur coordonnateur de l’essai. Il sera transmis par le promoteur au Comité de Protection des Personnes et aux autorités compétentes dans le cadre du rapport annuel de sécurité.')
    run1.style='Paragraphe'
    
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['comite_surveillance_independant'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

#   document.save("Partie10.docx")