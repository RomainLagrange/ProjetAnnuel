# -*- coding: utf-8 -*-
"""
Created on Sat Feb 16 14:07:53 2019

@author: Julie
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre3, TexteGris
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#revoir titre1 et texte encardé gris

def Page15():
    'Creation de la page 15 du protcole de catégorie 1'
    document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
 
    StyleProt1.Style(document)

    
    #ecriture du premier titre 
    document.add_paragraph('5	CRITERES D’ELIGIBILITE\n', style='Titre1') #titre                

    

    
   # Ecriture du 5.1  
    document.add_paragraph('5.1	Critères d’inclusion\n', style='Titre2') 
    
    # Ecriture du 5.2  
    document.add_paragraph('5.2	Critères de non inclusion\n', style='Titre2') 
    
    # Ecriture du 5.3  
    document.add_paragraph('5.3	Faisabilité et modalités de recrutement\n', style='Titre2') 
   
#    
#    tableTEST = document.add_table(rows = 1, cols = 1)
#    tableTEST.style = "Table Grid"
#    row = tableTEST.rows[0]
#    cell = row.cells[0]
#    cell.text = "text"

    
    
    document.save("page15.docx")   


