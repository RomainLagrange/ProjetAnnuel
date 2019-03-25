# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 15:17:45 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document) --> écrire en minuscule !!!
#    TexteGrisJustif(texte,document)

def Partie14(document,extract):
    'Creation de la partie 14 du protcole de catégorie 1'
   # document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

  #  Style(document)
    
#---------------------------------------------------------------ECRITURE
    
    
     #ecriture du premier titre 
    Titre1('14	CONSIDERATIONS ETHIQUES ET REGLEMENTAIRES',document)
    
     #Texte gris centré
    TexteGris('prendre contact avec la promotion interne \n pour aide a la redaction de ces chapitres',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le promoteur et l’(es) investigateur(s) s’engagent à ce que cette recherche soit réalisée en conformité avec la loi n°2004-806 du 9 août 2004, ainsi qu’en accord avec les Bonnes Pratiques Cliniques (I.C.H. version 4 du 1er mai 1996 et décision du 24 novembre 2006) et la déclaration d’Helsinki (qui peut être retrouvée dans sa version intégrale sur le site ')
    run1.style='Paragraphe'
    run2=p.add_run('http://www.wma.net).')
    run2.style='Paragraphe'
    run2.font.underline=True

    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('La recherche est conduite conformément au présent protocole. Hormis dans les situations d’urgence nécessitant la mise en place d’actes thérapeutiques précis, l’(es) investigateur(s) s’engage(nt) à respecter le protocole en tous points en particulier en ce qui concerne le recueil du consentement et la notification et le suivi des événements indésirables graves.')
    run1.style='Paragraphe'
    
    #Ecriture du 14.1  
    Titre2('14.1	Approbation de la recherche',document)
    #nom du CPP : extract['CPP']
    #nom du promoteur : extract['promoteur_nom_organisme']
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Cette recherche a reçu l’avis favorable du Comité de Protection des Personnes (CPP) de ')
    run1.style='Paragraphe'
    run2=p.add_run('nom du CPP ')
    run2.style='Paragraphe'
    run2.font.italic=True
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('et l’autorisation de l’ANSM.')
    run3.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le nom du promoteur ')
    run1.style='Paragraphe'
    run1.font.italic=True
    run1.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run2=p.add_run(', promoteur de cette recherche, a souscrit un contrat d’assurance en responsabilité civile auprès de ')
    run2.style='Paragraphe'
    run3=p.add_run('nom de la société d’assurance ')
    run3.style='Paragraphe'
    run3.font.italic=True
    run3.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run4=p.add_run('conformément aux dispositions de l’article L1121-10 du code de la santé publique.')
    run4.style='Paragraphe'


    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les données enregistrées à l’occasion de cette recherche font l’objet d’un traitement informatisé à ')
    run1.style='Paragraphe'
    run2=p.add_run('nom de la structure responsable du traitement des données ')
    run2.style='Paragraphe'
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run2.font.italic=True
    run3=p.add_run('dans le respect de la loi n°78-17 du 6 janvier 1978 relative à l’informatique, aux fichiers et aux libertés  modifiée par la loi 2004-801 du 6 août 2004. ')
    run3.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Cette recherche entre dans le cadre de la « Méthodologie de référence » (MR-001) en application des dispositions de l’article 54 alinéa 5 de la loi du 6 janvier 1978 modifiée relative à l’information, aux fichiers et aux libertés. Ce changement a été homologué par décision du 5 janvier 2006, mise à jour le 21 juillet 2016. ')
    run1.style='Paragraphe'
    run2=p.add_run('Le nom de la structure responsable du traitement des données ')
    run2.style='Paragraphe'
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run2.font.italic=True
    run3=p.add_run('a signé un engagement de conformité à cette « Méthodologie de référence ». ')
    run3.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Si la recherche ne rentre pas dans le champ d’application de la MR-001 : Le nom de la structure responsable du traitement des données ')
    run1.style='Paragraphe'
    run1.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run1.font.italic=True
    run2=p.add_run('a déclaré la recherche à la Commission Nationale de l’Informatique et des Libertés (CNIL).')
    run2.style='Paragraphe'


    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Si applicable :')
    run1.style='Paragraphe'
    run1.font.italic=True

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Cette recherche est enregistrée dans la base européenne ')
    run1.style='Paragraphe'
    run2=p.add_run('EudraCT ')
    run2.style='Paragraphe'
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('sous le ')
    run3.style='Paragraphe'
    run4=p.add_run('n° numéro enregistrement XXX')
    run4.style='Paragraphe'
    run4.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run4.font.italic=True

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Cette recherche est enregistrée sur le site http://clinicaltrials.gov/ ')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Après la recherche, la conservation de la collection d’échantillons biologiques sera déclarée au ministre chargé de la recherche et au directeur de l’Agence Régionale de Santé (et soumise au CPP pour avis si changement de finalité de recherche).')
    run1.style='Paragraphe'

    #Ecriture du 14.2  
    Titre2('14.2	Modifications au protocole',document)

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toute modification substantielle, c’est à dire toute modification de nature à avoir un impact significatif sur la protection des personnes, sur les conditions de validité et sur les résultats de la recherche, sur la qualité et la sécurité des produits expérimentés, sur l’interprétation des documents scientifiques qui viennent appuyer le déroulement de la recherche ou sur les modalités de conduite de celle-ci, fait l’objet d’un amendement écrit qui est soumis au promoteur ; celui-ci doit obtenir, préalablement à sa mise en œuvre, un avis favorable du CPP et, le cas échéant, une autorisation de l’ANSM.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les modifications non substantielles, c\'est à dire celles n’ayant pas d’impact significatif sur quelque aspect de la recherche que ce soit, sont communiquées au CPP à titre d’information.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toutes les modifications sont validées par le promoteur, et par tous les intervenants de la recherche concernés par la modification, avant soumission au CPP et, le cas échéant, à l’ANSM. Cette validation peut nécessiter la réunion de tout comité constitué pour la recherche.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toutes les modifications au protocole doivent être portées à la connaissance de tous les investigateurs qui participent à la recherche. Les investigateurs s’engagent à en respecter le contenu.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toute modification qui modifie la prise en charge des participants ou les bénéfices, risques et contraintes de la recherche fait l’objet d’une nouvelle note d’information et d’un nouveau formulaire de consentement dont le recueil suit la même procédure que celle précitée.')
    run1.style='Paragraphe'
    
    #Ecriture du 14.3  
    Titre2('14.3	Information du patient et formulaire de consentement éclairé écrit',document)

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les patients seront informés de façon complète et loyale, en des termes compréhensibles, des objectifs et des contraintes de l\'étude, des risques éventuels encourus, des mesures de surveillance et de sécurité nécessaires, de leurs droits de refuser de participer à l\'étude ou de la possibilité de se rétracter à tout moment.\nToutes ces informations figurent sur un formulaire d’information et de consentement remis au patient. \nLe consentement libre, éclairé et écrit du patient sera recueilli par l’investigateur, ou un médecin qui le représente avant l’inclusion définitive dans l’étude. \nLe formulaire est signé par les deux parties :')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Un exemplaire original est conservé par l’investigateur,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Un exemplaire (une copie ou un deuxième original) sera remis au patient.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’investigateur devra s’assurer que la personne qui se prête à la recherche aura eu le temps de prendre sa décision librement et aura pu lire et comprendre la notice d’information et le formulaire de consentement.')
    run1.style='Paragraphe'

     #Ecriture du 14.4 
    Titre2('14.4	Inscription au fichier national des personnes se prêtant à une recherche   ',document)

    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    #document.save("Partie7.docx") 
    
   # document.save("Partie14.docx")
    
    