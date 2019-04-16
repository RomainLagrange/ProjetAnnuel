# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:20:01 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:13:09 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Mon Feb 18 11:51:47 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie9(document):
    'Creation de la partie 9 du protcole de catégorie 3'
   # document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

   # Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('9	DROITS D’ACCES AUX DONNEES ET AUX DOCUMENTS SOURCE',document)
    
    
   # Ecriture du 9.1  
    Titre2('9.1	Accès aux données',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’acceptation de la participation au protocole implique que les personnes qui réalisent la recherche mettront à disposition les documents et données individuelles strictement nécessaires au suivi, au contrôle de qualité et à l’audit de la recherche, à la disposition des personnes ayant un accès à ces documents conformément aux dispositions législatives et réglementaires en vigueur.')
    run1.style='Paragraphe'
    
    # Ecriture du 9.2  
    Titre2('9.2	Données sources',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Ensemble des informations figurant dans des documents originaux, ou dans des copies authentifiées de ces documents, relatif aux examens cliniques, aux observations ou à d’autres activités menées dans le cadre d’une recherche et nécessaires à la reconstitution et à l’évaluation de la recherche. Les documents dans lesquels les données sources sont enregistrées sont appelés les documents sources.')
    run1.style='Paragraphe'
    
    # Ecriture du 9.3  
    Titre2('9.3	Confidentialité des données',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Conformément aux dispositions législatives en vigueur, les personnes ayant un accès direct aux données sources prendront toutes les précautions nécessaires en vue d\'assurer la confidentialité des informations relatives aux recherches, aux personnes qui s\'y prêtent et notamment en ce qui concerne leur identité ainsi qu’aux résultats obtenus.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Ces personnes, au même titre que les personnes qui dirigent et surveillent la recherche, sont soumises au secret professionnel.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Pendant la recherche ou à son issue, les données recueillies sur les personnes qui s’y prêtent et transmises au promoteur par les personnes qui dirigent et surveillent la recherche (ou tous autres intervenants spécialisés) seront codifiées. Elles ne doivent en aucun cas faire apparaître en clair les noms des personnes concernées ni leur adresse.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le promoteur s’assurera que chaque personne qui se prête à la recherche a été informée de l’accès aux données individuelles la concernant et strictement nécessaires au contrôle de qualité de la recherche.')
    run1.style='Paragraphe'
    
    #Ecriture du titre 9.4
    Titre2('9.4	Origine et nature des données recueillies :',document)

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les données nécessaires à la réalisation de la recherche sont saisies à partir du dossier médical du patient qui constitue le dossier source.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les données sources (dossier médical, résultats d’examens, questionnaires, correspondance médicale ')
    run1.style='Paragraphe'
    run2=p.add_run('xxxxxxxxx')
    run2.style='Paragraphe'
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run(') sont conservées par l’investigateur pour une durée de 25 ans après la fin de l’étude. ')
    run3.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’investigateur s’engage à autoriser un accès direct aux données sources de l’étude lors des visites de contrôle, d’audit ou d’inspection.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Pour chaque patient éligible, sera colligée dans un cahier d’observation standard papier l’observation des données biologiques, cliniques, et d’imagerie ……… rétrospectives/prospectives contenues dans son dossier médical. ')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les données suivantes seront relevées :')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('démographiques : âge, sexe, poids et taille (ou IMC)')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('xxxxxxx')
    run1.style='Paragraphe'
    run1.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('La base de données sera installée dans le service de ')
    run1.style='Paragraphe'
    run2=p.add_run('xxxxxx ')
    run2.style='Paragraphe'
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('du CHU de Poitiers.\nLe Dr ')
    run3.style='Paragraphe'
    run4=p.add_run('xxxxxx ')
    run4.style='Paragraphe'
    run4.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run5=p.add_run('sera responsable de la collecte et de la saisie des données, le data manager de la gestion de la base de données et le biostatisticien des analyses statistiques.\nLes données seront alors reportées et enregistrées sous format électronique dans une base de données sécurisée dont l’accès sera limité aux responsables de l’étude, au data manager, au biostatisticien ainsi qu’aux autorités de santé si nécessaire.')
    run5.style='Paragraphe'
    

    #Ecriture du titre 9.5
    Titre2('9.5	Mode de circulation des données',document)

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’ensemble de ces données recueillies pour la recherche seront retranscrites dans base de données électronique par les investigateurs ou leur technicien d’étude clinique. \nUne copie du cahier d’observation est conservée par l’investigateur pour une durée de 25 ans après la fin de l’étude sous format papier ou sous format numérique.')
    run1.style='Paragraphe'
    
    
    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
  
  #  document.save("Cat3Part9.docx")   