# -*- coding: utf-8 -*-
"""
Created on Sat Feb 16 14:07:53 2019

@author: Julie
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)


def Partie5(document,extract):
#def Partie5():
    'Creation de la partie 5 du protcole de catégorie 1'
 #   document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
 
  #  Style(document)

    
    #ecriture du premier titre 
    Titre1('5	CRITERES D’ELIGIBILITE',document)
    
   # Ecriture du 5.1  
    Titre2('5.1	Critères d’inclusion',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Tous les patients inclus dans cette recherche devront vérifier tous les critères d’inclusion listés ci-dessous :')
    run1.style='Paragraphe'
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['critere_inclusion_longue'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['justification_inclusion'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    
    # Ecriture du 5.2  
    Titre2('5.2	Critères de non inclusion',document)

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Tous les patients inclus dans cette recherche ne devront avoir aucun des critères de non inclusion listés ci-dessous :')
    run1.style='Paragraphe'
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['critere_non_inclusion_longue'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    
    # Ecriture du 5.3  
    Titre2('5.3	Faisabilité et modalités de recrutement',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['modalite_recrutement'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['duree_inclusion'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
   
                    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

    
    
  #  document.save("Partie5.docx")   


