# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 14:01:09 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie7(document,extract):
#def Partie7():
    'Creation de la partie 7 du protcole de catégorie 1'
 #   document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

   # Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('7	TRAITEMENT(S) / STRATEGIE(S) / PROCEDURES DE LA RECHERCHE',document)
    
    
   # Ecriture du 7.1  
    Titre2('7.1	Traitement / stratégie / procédure expérimental(e)',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['traitement_strategie_longue'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Médicament expérimental : médicament expérimenté ou utilisé comme référence, y compris en tant que placebo, lors d’un essai clinique (article 2 du règlement européen).')
    run1.style='Paragraphe'
    
    #Ecriture des trois textes gris justifiés
    TexteGrisJustif('Pour un traitement de type médicament',document)
    TexteGrisJustif('Pour un placebo',document)
    TexteGrisJustif('Pour un traitement de type dispositif médical (DM)',document)
    TexteGrisJustif('Pour une stratégie/procédure',document)
    
    # Ecriture du 7.2
    Titre2('7.2	Traitement / Stratégie / Procédure de comparaison',document)
    
    #Ecriture des deux textes gris justifiés
    TexteGrisJustif('Pour un traitement de type dispositif médical (DM)',document)
    TexteGrisJustif('Pour une stratégie/procédure',document)
    
    # Ecriture du 7.3
    Titre2('7.3	Circuit des produits',document)
    
    #Texte gris
    TexteGris('prendre contact avec la pharmacie du chu de poitiers \n pour aide a la redaction de ces chapitres',document)
    #Ecriture du 7.3.1
    Titre3('7.3.1','Libération et distribution des produits',document)
    #Ecriture du 7.3.2
    Titre3('7.3.2','Fourniture des produits',document)
    #Ecriture du 7.3.3
    Titre3('7.3.3','Conditionnement des produits',document)
    #Ecriture du 7.3.4
    Titre3('7.3.4','Etiquetage des produits',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Chaque patient se verra remettre ')
    run1.style='Paragraphe'
    run2=p.add_run('XXXX boîtes, flacons ')
    run2.style='Paragraphe'
    run2.font.italic= True
    run3=p.add_run('pour la totalité de la durée du traitement.')
    run3.style='Paragraphe'
    
    #Ecriture du 7.3.5
    Titre3('7.3.5','Expédition et gestion des produits',document)
    #Ecriture du 7.3.6
    StyleProt1.Titre3('7.3.6','Dispensation des produits et observance',document)
    #Ecriture du 7.3.7
    Titre3('7.3.7','Stockage ',document)
    #Ecriture du 7.3.8
    Titre3('7.3.8','Retour et destruction des produits non utilisés',document)
    
    # Ecriture du 7.4
    Titre2('7.4	Insu',document)
    
    #Texte gris centré
    TexteGris('prendre contact avec la plateforme de methodologie \n pour aide a la redaction de ce chapitre', document)

    #Ecriture du 7.4.1
    Titre3('7.4.1','Organisation de l’insu',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('La pharmacie est destinataire de la liste de randomisation.')
    run1.style='Paragraphe'
    
    #Ecriture du 7.4.2
    Titre3('7.4.2','Levée de l’insu',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('En situation d’urgence médicale nécessitant une levée d’aveugle, la procédure DRC-VIGI-003 du promoteur sera suivie. ')
    run1.style='Paragraphe'
   
    # Ecriture du 7.5
    Titre2('7.5	Réductions et ajustements de dose',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les retards et modifications de dose seront effectués selon les recommandations suivantes. L’évaluation des toxicités se fera selon la classification CTCAE (Common Terminology Criteria for Adverse Events) du NCI (National Cancer Institute).')
    run1.style='Paragraphe'
    
    #Ecriture du 7.5.1
    Titre3('7.5.1','Réductions/ajustements de doses',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les tableaux suivants résument les modifications de dose du médicament 1, médicament 2,… pour gérer d’éventuelles toxicités.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Tableau 1 : diminutions de dose pour ')
    run1.style='Paragraphe'
    run2=p.add_run('médicament 1')
    run2.style='Paragraphe'
    run2.font.italic= True
    
    table = document.add_table(3, 6)
    table.style = 'Table Grid'
    a = table.cell(0, 0)
    b = table.cell(1, 0)
    A = a.merge(b)
    B=table.cell(0, 1).merge(table.cell(0, 5))
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)
    shading_elm_3 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[1].cells[1]._tc.get_or_add_tcPr().append(shading_elm_3)
    shading_elm_4 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[1].cells[2]._tc.get_or_add_tcPr().append(shading_elm_4)
    shading_elm_5 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[1].cells[3]._tc.get_or_add_tcPr().append(shading_elm_5)
    shading_elm_6 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[1].cells[4]._tc.get_or_add_tcPr().append(shading_elm_6)
    shading_elm_7 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[1].cells[5]._tc.get_or_add_tcPr().append(shading_elm_7)
    
    
    row = table.rows[0].cells
    para_text  = 'Dose initiale'
    cell = row[0]
    pt = cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'
    p.bold = True

    row = table.rows[0]
    texte='Réductions de dose du Médicament 1'
    text_formatted = row.cells[1].paragraphs[0].add_run(texte)
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True

    row = table.rows[1].cells
    para_text  = 'Dose -1'
    cell = row[1]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'
    
    para_text  = 'Dose -2'
    cell = row[2]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'

    para_text  = 'Dose -3'
    cell = row[3]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'
    
    

    para_text  = 'Dose -4'
    cell = row[4]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'
    

    para_text  = 'Dose -5'
    cell = row[5]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'

    row = table.rows[2].cells
    para_text  = 'discontinue'
    cell = row[5]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'
    
    p=document.add_paragraph('')

    table = document.add_table(3, 6)
    table.style = 'Table Grid'
    a = table.cell(0, 0)
    b = table.cell(1, 0)
    A = a.merge(b)
    B=table.cell(0, 1).merge(table.cell(0, 5))
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)
    for y in range(1,6):
        shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        table.rows[1].cells[y]._tc.get_or_add_tcPr().append(shading_elm)

    row = table.rows[0].cells
    para_text  = 'Dose initiale'
    cell = row[0]
    pt = cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'
    p.bold = True

    row = table.rows[0]
    texte='Réductions de dose du Médicament 1'
    text_formatted = row.cells[1].paragraphs[0].add_run(texte)
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True

    row = table.rows[1].cells
    para_text  = 'Dose -1'
    cell = row[1]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'
    
    para_text  = 'Dose -2'
    cell = row[2]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'

    para_text  = 'Dose -3'
    cell = row[3]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'  

    para_text  = 'Dose -4'
    cell = row[4]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'    

    para_text  = 'Dose -5'
    cell = row[5]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'

    row = table.rows[2].cells
    para_text  = 'discontinue'
    cell = row[5]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    p.font.name = 'Times New Roman'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Tableau 2 : diminutions de dose pour ')
    run1.style='Paragraphe'
    run2=p.add_run('médicament 2')
    run2.style='Paragraphe'
    run2.font.italic= True
    
    p=document.add_paragraph('')
    
    
    #Ecriture du 7.5.2
    Titre3('7.5.2','Réductions de dose pour les toxicités hématologiques',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les tableaux suivants décrivent les recommandations de réduction de dose pour ')
    run1.style='Paragraphe'
    run2=p.add_run('médicament 1 / médicament… ')
    run2.style='Paragraphe'
    run2.font.italic= True
    run3=p.add_run('en cas de thrombopénie, neutropénie et anémie.')
    run3.style='Paragraphe'
    
    table = document.add_table(7, 3)
    table.style = 'Table Grid'
    a = table.cell(0, 1)
    b = table.cell(0, 2)
    A = a.merge(b)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)
    for i in range(1,7):
        for y in range(0,3):
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            table.rows[i].cells[y]._tc.get_or_add_tcPr().append(shading_elm)
    
    row = table.rows[0]
    text_formatted = row.cells[1].paragraphs[0].add_run('Actions recommandées')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    
    row = table.rows[1]
    text_formatted = row.cells[0].paragraphs[0].add_run('Taux plaquettes')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    text_formatted = row.cells[1].paragraphs[0].add_run('Médicament 1')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    text_formatted = row.cells[2].paragraphs[0].add_run('Médicament 2')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    
        
    row = table.rows[3]
    text_formatted = row.cells[0].paragraphs[0].add_run('Taux PNN')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    text_formatted = row.cells[1].paragraphs[0].add_run('Médicament 1')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    text_formatted = row.cells[2].paragraphs[0].add_run('Médicament 2')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    
    row = table.rows[5]
    text_formatted = row.cells[0].paragraphs[0].add_run('Taux hémoglobine')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    text_formatted = row.cells[1].paragraphs[0].add_run('Médicament 1')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    text_formatted = row.cells[2].paragraphs[0].add_run('Médicament 2')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True


    #Ecriture du 7.5.3
    Titre3('7.5.3','Réductions de dose pour les toxicités non hématologiques',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les lignes directrices d’ajustement de dose pour ')
    run1.style='Paragraphe'
    run2=p.add_run('médicament 1 / médicament 2… ')
    run2.style='Paragraphe'
    run2.font.italic= True
    run3=p.add_run('en cas de toxicités non hématologiques sont résumées comme suit :')
    run3.style='Paragraphe'
    
    table = document.add_table(8, 3)
    table.style = 'Table Grid'
    a = table.cell(0, 1)
    b = table.cell(0, 2)
    A = a.merge(b)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)
    for i in range(1,8):
        for y in range(0,3):
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            table.rows[i].cells[y]._tc.get_or_add_tcPr().append(shading_elm)
    
    row = table.rows[0]
    text_formatted = row.cells[0].paragraphs[0].add_run('Symptômes')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    text_formatted = row.cells[1].paragraphs[0].add_run('Actions recommandées')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    
    row = table.rows[1]
    text_formatted = row.cells[1].paragraphs[0].add_run('Médicament 1')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    text_formatted = row.cells[2].paragraphs[0].add_run('Médicament 2')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    
        
    row = table.rows[2]
    text_formatted = row.cells[0].paragraphs[0].add_run('Symptôme 1')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True
    
    row = table.rows[3]
    text_formatted = row.cells[0].paragraphs[0].add_run('Grade 2-3')
    text_formatted.font.name = 'Times New Roman'
    
    row = table.rows[4]
    text_formatted = row.cells[0].paragraphs[0].add_run('Grade 4')
    text_formatted.font.name = 'Times New Roman'
    
    row = table.rows[5]
    text_formatted = row.cells[0].paragraphs[0].add_run('Symptôme 2')
    text_formatted.font.name = 'Times New Roman'
    text_formatted.bold = True

    row = table.rows[6]
    text_formatted = row.cells[0].paragraphs[0].add_run('Grade 3')
    text_formatted.font.name = 'Times New Roman'
    
    row = table.rows[7]
    text_formatted = row.cells[0].paragraphs[0].add_run('Grade 4')
    text_formatted.font.name = 'Times New Roman'


        #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    
 #   document.save("Partie7.docx")   
    
    