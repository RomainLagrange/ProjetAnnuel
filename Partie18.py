# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 15:24:24 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document) --> écrire en minuscule !!!
#    TexteGrisJustif(texte,document)

def Partie18(document,extract):
    'Creation de la partie 18 du protcole de catégorie 1'
  #  document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

  #  Style(document)
    
#---------------------------------------------------------------ECRITURE
    
    
     #ecriture du premier titre 
    Titre1('18	FAISABILITE DE L’ETUDE',document)
    
    #Ecriture du 17.1  
    Titre2('18.1	Expertise scientifique',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['investigateur_coordinateur_nom_etablissement'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    
    
    
    #Ecriture du 17.2  
    Titre2('18.2	Collaborations ',document)
    
#    paragraph2 = document.add_paragraph()
#    sentence2 = paragraph2.add_run(extract['titre_complet'])
#    sentence2.font.name = 'Times New Roman'
#    sentence2.font.size = docx.shared.Pt(10)
#    paragraph2 = document.add_paragraph()
#    sentence2 = paragraph2.add_run(extract['taille_etude_courte'])
#    sentence2.font.name = 'Times New Roman'
#    sentence2.font.size = docx.shared.Pt(10)
#    paragraph2 = document.add_paragraph()
#    sentence2 = paragraph2.add_run(extract['duree_inclusion'])
#    sentence2.font.name = 'Times New Roman'
#    sentence2.font.size = docx.shared.Pt(10)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le  projet ')
    run1.style='Paragraphe'
    run2=p.add_run(extract['titre_abrege'])
    run2.style='Paragraphe'
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('bénéfice…. ')
    run3.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('En effet, le nombre de sujets nécessaires a été évalué à '+extract['taille_etude_courte'])
    run1.style='Paragraphe'
    run2=p.add_run('XXX ')
    run2.style='Paragraphe'
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('patients sur une période d’inclusion de '+extract['duree_inclusion'])
    run3.style='Paragraphe'
    run4=p.add_run('XX ')
    run4.style='Paragraphe'
    run4.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run5=p.add_run('mois ce qui correspond aux capacités de recrutement des centres investigateurs pressentis')
    run5.style='Paragraphe'
    
    #Ecriture du 18.3  
    Titre2('18.3	Financement du projet (si ce point ne fait pas partie d’un document distinct)',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['titre_complet'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le projet ')
    run1.style='Paragraphe'
    run2=p.add_run(extract['titre_abrege'])
    run2.style='Paragraphe'
    run2.font.italic=True
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('bénéficie d’un financement ')
    run3.style='Paragraphe'    

    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
 
#    
#    document.save("Partie18.docx")