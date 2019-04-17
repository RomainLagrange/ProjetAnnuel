# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:11:16 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:09:16 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:07:33 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:05:49 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 16:19:49 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style, Titre1,Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK

from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)
    

def Partie5(document,extract):
    'Creation de la partie 5 du protcole de catégorie 3'
   # document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)
        
   # Style(document)

 
    Titre1('5	CRITERES D’ELIGIBILITE',document)
    #ecriture du premier titre 

     #ecriture du titre5.1
    Titre2('5.1	Critères d’inclusion',document) 

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Tous les patients inclus dans cette recherche devront vérifier tous les critères d’inclusion listés ci-dessous :')
    run1.style='Paragraphe'
    
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['criteres_inclusion'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run('- Recueil de la non-opposition du participant')
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)

    # Ecriture du 5.2  
    Titre2('5.2	Critères de non inclusion',document)
  
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Tous les patients inclus dans cette recherche ne devront avoir aucun des critères de non inclusion listés ci-dessous :')
    run1.style='Paragraphe'
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['criteres_non_inclusion'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
                                        
    #ecriture du titre5.3
    Titre2('5.3	Faisabilité et modalités de recrutement ',document) 
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['taille_etude_courte'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['duree_inclusion'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)


    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
     

   # document.save("Cat3Partie5.docx")   