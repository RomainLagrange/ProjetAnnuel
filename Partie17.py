# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 15:22:09 2019

@author: Asuspc
"""


import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document) --> écrire en minuscule !!!
#    TexteGrisJustif(texte,document)

def Partie17(document):
    'Creation de la partie 17 du protcole de catégorie 1'
  #  document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

 #   Style(document)
    
#---------------------------------------------------------------ECRITURE
    
    
     #ecriture du premier titre 
    Titre1('17	REGLES RELATIVES A LA PUBLICATION',document)
    
     #Texte gris centré
    TexteGris('prendre contact avec la plateforme de methodologie \n pour aide a la redaction de ces chapitres',document)
    
    #Ecriture du 17.1  
    Titre2('17.1	Communications scientifiques',document)
#    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’analyse des données fournies par les centres investigateurs est réalisée par ')
    run1.style='Paragraphe'
    run2=p.add_run('nom de la structure')
    run2.style='Paragraphe'
    run2.font.italic=True
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('. Cette analyse donne lieu à un rapport écrit qui est soumis au promoteur, qui transmettra au Comité de Protection des Personnes et à l’autorité compétente.')
    run3.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toute communication écrite ou orale des résultats de la recherche doit recevoir l’accord préalable de l’investigateur coordonnateur et, le cas échéant, de tout comité constitué pour la recherche.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’investigateur coordonnateur/principal s’engage à mettre à disposition du public les résultats de la recherche aussi bien négatifs et non concluants que positifs.')
    run1.style='Paragraphe'

    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('La publication des résultats principaux mentionne le nom du promoteur, de tous les investigateurs ayant inclus ou suivi des participants dans la recherche, des méthodologistes, biostatisticiens et data managers ayant participé à la recherche, des vigilants ayant participé à l’analyse de la sécurité des participants, des membres du(des) comité(s) constitué(s) pour la recherche et la participation éventuelle du laboratoire ')
    run1.style='Paragraphe'
    run2=p.add_run('nom du laboratoire pharmaceutique ')
    run2.style='Paragraphe'
    run2.font.italic=True
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('/ la source de financement. Il sera tenu compte des règles internationales d’écriture et de publication (')
    run3.style='Paragraphe'
    run4=p.add_run('The Uniform Requirements for Manuscripts ')
    run4.style='Paragraphe'
    run4.font.italic=True
    run5=p.add_run('de l’ICMJE, avril 2010).')
    run5.style='Paragraphe'
    
    
    #Ecriture du 17.2  
    Titre2('17.2	Communication des résultats aux participants',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Conformément à la loi n°2002-303 du 4 mars 2002, les participants sont informés, à leur demande, des résultats globaux de la recherche.')
    run1.style='Paragraphe'
    
    #Ecriture du 17.3  
    Titre2('17.3	Cession des données',document)

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('La gestion des données est assurée par ')
    run1.style='Paragraphe'
    run2=p.add_run('nom de la structure')
    run2.style='Paragraphe'
    run2.font.italic=True
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('. Les conditions de cession de tout ou partie de la base de données de la recherche sont décidées par le promoteur de la recherche et font l’objet d’un contrat écrit.')
    run3.style='Paragraphe'

    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    #document.save("Partie7.docx") 
    
 #   document.save("Partie17.docx")
    