# -*- coding: utf-8 -*-
"""
Created on Fri Feb 15 17:50:09 2019

@author: Julie
"""
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.section import WD_ORIENT, WD_SECTION




def Titre1(texte, document):
    paragraph=document.add_paragraph(texte+'\n', style='Titre1') #titre
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER #centrer
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE #espacement simple entre les lignes de texte
    
def Titre2(texte,document):
    p=document.add_paragraph(texte+'\n', style='Titre2')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def Titre3(num, texte, document):
    styles=document.styles
    p=document.add_heading()
    p.paragraph_format.left_indent = Inches(0.98) #indentation en pouce, ici 1,5cm
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run()
    run1.text=num+'	    '
    run1.style='ListeTitre3'
    run2=p.add_run()
    run2.text=texte+'\n'
    run2.style='Titre3'
    p.next_paragraph_style = styles['Normal']

def TexteGris(texte,document):
    styles=document.styles
    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table = document.add_table(rows = 1, cols = 1)
    row = table.rows[0].cells
    para_text =texte
    cell = row[0]
    pt = cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    cell._tc.get_or_add_tcPr().append(shading_elm)
    p.style='BackgroundGrey'
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.next_paragraph_style = styles['Normal']
    
def TexteGrisJustif(texte,document):
    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table = document.add_table(rows = 1, cols = 1)
    row = table.rows[0].cells
    para_text =texte
    cell = row[0]
    pt = cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
  #  pt.add_run('\n')
    cell._tc.get_or_add_tcPr().append(shading_elm)
    p.style='BackgroundGreyJustif'
    paragraph=document.add_paragraph(' ')

   # pt.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
   
def Titre2Paysage(texte,document):
    document.add_paragraph(texte+'\n', style='paysage')

def Style(document):
    'Défintion des styles du protocole de catégorie 1'
    
# TROUVER COMMENT PARTAGER LES STYLES A TOUTES LES PAGES SONT COPIER TOUT
# LE CODE    
    
    styles=document.styles
    


    #   definition du style Paragraphe, modification du style Normal 
    stylePara = styles.add_style('Paragraphe', WD_STYLE_TYPE.CHARACTER)
    stylePara.base_style = styles['Normal']
    fontPara = stylePara.font
    fontPara.name = 'Times New Roman' #police
    fontPara.size = docx.shared.Pt(11) #taille

#   definition du style Titre1
    styleTitre1 = styles.add_style('Titre1', WD_STYLE_TYPE.PARAGRAPH, WD_ALIGN_PARAGRAPH.CENTER) #le style créé s'appelle Titre1 et est un paragraphe centré
    styleTitre1.base_style = styles['Heading1']  #se base sur le style heading 1 du package
    fontTitre1 = styleTitre1.font
    fontTitre1.name = 'Times New Roman' #police
    fontTitre1.size = docx.shared.Pt(12) #taille
    fontTitre1.all_caps = True #toujours en majuscule
    fontTitre1.bold= True #en gras
    fontTitre1.underline= True #souligné
    fontTitre1.color.rgb = RGBColor(0x0,0x70,0xC0) #couleur bleu, en base 16
    styles['Titre1'].next_paragraph_style = styles['Normal']

    #Definition du Titre2, correspond par exemple au 1.1 ou 1.2
    styleTitre2 = styles.add_style('Titre2', WD_STYLE_TYPE.PARAGRAPH) #le style va s'appeler Titre2
    styleTitre2.base_style = styles['Heading2']  #se base sur le style heading 2 du package
    fontTitre2 = styleTitre2.font
    fontTitre2.name = 'Times New Roman' #police d'écriture
    fontTitre2.size = docx.shared.Pt(14)  #taille de la police
    fontTitre2.bold= True  #en gras
    fontTitre2.color.rgb = RGBColor(0x0,0x0,0x0)   #couleur noire
    styleTitre2.paragraph_format.left_indent = Inches(0.59)  #indentation
    styles['Titre2'].next_paragraph_style = styles['Normal']  #le texte écrit à la suite sera en style Normal
    
    
     #Definition du style Titre3; correspond aux 1.1.1 ou 1.1.2...
    styleTitre3 = styles.add_style('Titre3', WD_STYLE_TYPE.CHARACTER)  #le style va s'appeler Titre3
    styleTitre3.base_style = styles['Heading3']  #se base sur le style heading 3 du package python_docx
    fontTitre3 = styleTitre3.font
    fontTitre3.name = 'Times New Roman'  #police d'écriture
    fontTitre3.size = docx.shared.Pt(12)  #taille de la police
    fontTitre3.bold= False  #ne sera pas en gras
    fontTitre3.underline= True  #souligné 
    fontTitre3.color.rgb = RGBColor(0x0,0x0,0x0)  #couleur noire
    
    
   #Définition du ListeTitre3, correspond au nom du titre après le 1.1.1 ou 1.2.1    
    styleTitreListe3 = styles.add_style('ListeTitre3', WD_STYLE_TYPE.CHARACTER)
    styleTitreListe3.base_style = styles['Heading3']
    fontTitreListe3 = styleTitreListe3.font
    fontTitreListe3.name = 'Times New Roman'
    fontTitreListe3.size = docx.shared.Pt(12)
    fontTitreListe3.bold= True
    fontTitreListe3.underline= False
    fontTitreListe3.color.rgb = RGBColor(0x0,0x0,0x0)  
    
    
    #Definition style texte surligné en gris centré   --> SUPPRIMER ESPACE EN BAS
#    styles = document.styles
    styleBackgroundGrey = styles.add_style('BackgroundGrey', WD_STYLE_TYPE.CHARACTER)
    styleBackgroundGrey.base_style = styles['No Spacing']
    fontBackgroundGrey = styleBackgroundGrey.font
    fontBackgroundGrey.name = 'Times New Roman'
    fontBackgroundGrey.size = docx.shared.Pt(11)
    fontBackgroundGrey.bold = True
    fontBackgroundGrey.small_caps = True
    
    
    #Definition style texte surligné en gris justifié   --> SUPPRIMER ESPACE EN BAS
  #  styles = document.styles
    styleBackgroundGrey = styles.add_style('BackgroundGreyJustif', WD_STYLE_TYPE.CHARACTER)
    styleBackgroundGrey.base_style = styles['No Spacing']
    fontBackgroundGrey = styleBackgroundGrey.font
    fontBackgroundGrey.name = 'Times New Roman'
    fontBackgroundGrey.size = docx.shared.Pt(11)
    fontBackgroundGrey.italic = True
    fontBackgroundGrey.bold = True
    
    
    
    #definition du style pour le texte indicatif -->  ESPACEMENT LIGNES 
    styleIndic = styles.add_style('TexteItalic', WD_STYLE_TYPE.PARAGRAPH)
    styleIndic.base_style = styles['Normal']
    fontIndic = styleIndic.font
    fontIndic.name = 'Times New Roman'   
    fontIndic.size = docx.shared.Pt(11)
    fontIndic.italic= True
    
    
    #style specifique protocle 1 partie 6.2 pour indentation
    stylepaysage = styles.add_style('paysage', WD_STYLE_TYPE.PARAGRAPH)
    stylepaysage.base_style = styles['Heading2']
    fontpaysage = stylepaysage.font
    fontpaysage.name = 'Times New Roman'
    fontpaysage.size = docx.shared.Pt(14)
    fontpaysage.bold= True
    fontpaysage.color.rgb = RGBColor(0x0,0x0,0x0)
    stylepaysage.paragraph_format.left_indent = Cm(1.25)
    styles['paysage'].next_paragraph_style = styles['Normal']
    
    

def change_orientation(document):
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    return new_section
    

