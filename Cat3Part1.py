# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:05:49 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 16:19:49 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style, Titre1,Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK

from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)
    

def Partie1(document,extract):
    'Creation de la partie 1 du protcole de catégorie 3'
   # document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)
        
  #  Style(document) 

 
    Titre1('1	JUSTICATION SCIENTIFIQUE ET DESCRIPTION GENERALE',document)
    #ecriture du premier titre 
#    paragraph=document.add_paragraph('1	JUSTICATION SCIENTIFIQUE ET DESCRIPTION GENERALE\n', style='Titre1') #titre
#    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER #centrer

    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['justification_etude_longue'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    
    # Ecriture du 1.1  
    Titre2('1.1	Etat actuel des connaissances',document)
  
              
                                        
    #ecriture du titre1.2
    Titre2('1.2	Hypothèse de la recherche',document) 
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['critere_jugement_principal_courte'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['traitement_strategie_courte'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    
     #Ecriture du titre1.3
    Titre2('1.3 Retombées attendues',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['retombee_attenduees_longue'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    

    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
     

   # document.save("Cat3Partie1.docx")   