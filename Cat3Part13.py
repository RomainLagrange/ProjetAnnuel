# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:33:39 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:20:01 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:13:09 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Mon Feb 18 11:51:47 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie13(document):
    'Creation de la partie 13 du protcole de catégorie 3'
   # document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

   # Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('13	REGLES RELATIVES A LA PUBLICATION',document)
    
    
   # Ecriture du 13.1  
    Titre2('13.1	Communications scientifiques',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’analyse des données fournies par les centres est réalisée par ')
    run1.style='Paragraphe'
    run2=p.add_run('nom de la structure')
    run2.style='Paragraphe'
    run2.font.italic=True
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('. Cette analyse donne lieu à un rapport écrit qui est soumis au promoteur. Ce rapport permet la préparation d’une ou plusieurs publication(s).')
    run3.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toute communication écrite ou orale des résultats de la recherche doit recevoir l’accord préalable de la personne qui dirige et surveille la recherche et, le cas échéant, de tout comité constitué pour la recherche.')
    run1.style='Paragraphe'
    
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('La publication des résultats principaux mentionne le nom du promoteur, de toutes les personnes  ayant inclus ou suivi des patients dans la recherche, des méthodologistes, biostatisticiens et data managers ayant participé à la recherche, des membres du(des) comité(s) constitué(s) pour la recherche et la participation éventuelle du laboratoire ')
    run1.style='Paragraphe'
    run2=p.add_run('nom du laboratoire pharmaceutique ')
    run2.style='Paragraphe'
    run2.font.italic=True
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('// la source de financement. Il sera tenu compte des règles internationales d’écriture et de publication (')
    run3.style='Paragraphe'
    run4=p.add_run('The Uniform Requirements for Manuscripts ')
    run4.style='Paragraphe'
    run4.font.italic=True
    run5=p.add_run('de l’ICMJE, avril 2010).')
    run5.style='Paragraphe'
    
    # Ecriture du 13.2  
    Titre2('13.2	Communication des résultats aux patients',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('A leur demande, les participants à la recherche sont informés des résultats globaux de celle-ci.')
    run1.style='Paragraphe'

    
    # Ecriture du 13.3  
    Titre2('13.3	Cession des données',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le recueil et la gestion des données sont assurés par ')
    run1.style='Paragraphe'
    run2=p.add_run('nom de la structure')
    run2.style='Paragraphe'
    run2.font.italic=True
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('. Les conditions de cession de tout ou partie de la base de données de la recherche sont décidées par le promoteur de la recherche et font l’objet d’un contrat écrit.')
    run3.style='Paragraphe'
    
    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
  
  #  document.save("Cat3Part13.docx")   