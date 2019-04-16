# -*- coding: utf-8 -*-
"""
Created on Wed Feb 20 22:04:40 2019

@author: Marion
"""

import pandas as pd
import docx
from docx.api import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
import time
from time import gmtime, strftime


def main_ansm_dm(extract):
    document = docx.Document()
    partie_une_ansm_dm(document)
    partie_B_C(document)
    partie_D(document)
    partie_E_F(document)
    partieF5_suite(document)
    a_partir_F8(document)
    partie_F10(document)
    parties_H_I(document)
    date = (strftime('%d-%m-%Y',time.localtime()))
    document.save("soumission_ansm_dm"+extract['titre_abrege']+"_"+date+".docx")

def partie_une_ansm_dm(document):
    
    '''Marge de la page'''
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        
    '''Titre du document'''
    styles= document.styles
    style=styles.add_style('debut_page', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(11) 
    
    paragraph=document.add_paragraph("Formulaire de demande d’autorisation auprès de l'ANSM et demande d’avis du comite de protection des personnes (CPP) pour une recherche mentionnée au 1° ou au 2° de l’article L. 1121-1 du code de la santé publique portant sur un dispositif médical (DM) ou un dispositif médical de diagnostic in vitro (DMDIV)\n", style='debut_page')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("Toutes les rubriques du formulaire doivent être complétées.\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.italic=True
    fontdebut.size = docx.shared.Pt(10.5)
    sentence=paragraph.add_run("Partie réservée à l’ANSM / au CPP :\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.italic=True
    fontdebut.size = docx.shared.Pt(10)
    
    table = document.add_table(rows=3, cols=3, style='Table Grid')
    a=table.cell(1,0)
    b=table.cell(2,0)
    a.merge(b)
    c=table.cell(1,1)
    d=table.cell(2,1)
    c.merge(d)
    table.cell(0,0).text=("Date de réception de la demande :\n\n\n     /     /     ")
    table.cell(0,1).text=("Date de demande d’informations complémentaires :\n\n     /     /     ")
    table.cell(0,2).text=("Refus d’autorisation / avis défavorable :	□\n\n	Préciser la  date :\n     /     /     ")
    table.cell(1,0).text=("\nDate du début de la procédure :\n\n     /     /     ")
    table.cell(1,1).text=("Date de réception des informations complémentaires :\n\n     /     /     ")
    table.cell(1,2).text=("Autorisation / avis favorable :	□\n\nPréciser la date : \n     /     /     ")
    table.cell(2,2).text=("Retrait de la demande :	□\nDate :      /     /     ")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nPartie à compléter par le demandeur :\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.italic=True
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("Ce formulaire est commun pour la demande d’autorisation auprès de l’ANSM et pour la demande d’avis auprès du CPP. Veuillez cocher ci-après la case correspondant à l’objet de la demande.\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("□ Recherche interventionnelle mentionnée au 1° de l’article L. 1121-1 du code de la santé publique\n\n"
                               "□ Recherche interventionnelle ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(12)
    sentence=paragraph.add_run("comportant des risques et contraintes minimes ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold=True
    fontdebut.size = docx.shared.Pt(12)
    sentence=paragraph.add_run("mentionnée au 2° de l’article L. 1121-1 du code de la santé publique\n\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(12)
    sentence=paragraph.add_run("Demande d’autorisation à l’ANSM :  □				Demande d’avis au CPP :  □\n\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold=True
    fontdebut.size = docx.shared.Pt(11)
    
    
    '''A identification de la recherche'''
    paragraph=document.add_paragraph("A. Identification de la recherche \n", style='debut_page')
    table = document.add_table(rows=9, cols=2, style='Table Grid')
    table.cell(0,0).text=("Numéro d’enregistrement de la recherche auprès de l'ANSM (n°IDRCB)")
    table.cell(1,0).text=("Numéro EUDAMED  (le cas échéant)")
    table.cell(2,0).text=("Titre complet de la recherche ")
    table.cell(3,0).text=("Numéro de code du protocole attribué par le promoteur, version et date")
    table.cell(4,0).text=("Nom ou titre abrégé (le cas échéant) ")
    table.cell(5,0).text=("S’agit-il d’une resoumission ?")
    table.cell(5,1).text=("□ oui       □ non")
    table.cell(6,0).text=("Si oui, indiquer la lettre de resoumission  ")
    table.cell(7,0).text=("Préciser par ailleurs si les documents précédemment versés ont été modifiés ?")
    table.cell(7,1).text=("□ oui       □ non")
    a=table.cell(8,0)
    b=table.cell(8,1)
    a.merge(b)
    table.cell(8,0).text=("(Si oui, joindre au dossier de demande d’AEC un tableau comparatif)")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==2:
                        fontdebut.bold=True
                    n=n+1
 
    
# Parties B et C du document
def partie_B_C(document):
    
    paragraph=document.add_paragraph("\n\nB. Identification du promoteur responsable de la recherche \n", style='debut_page')
    table = document.add_table(rows=7, cols=3, style='Table Grid')
    a=table.cell(0,0)
    b=table.cell(0,2)
    a.merge(b)
    table.cell(0,0).text=("B1. Promoteur")
    table.cell(1,0).text=("Nom de l'organisme")
    table.cell(2,0).text=("Nom de la personne à contacter ")
    table.cell(3,0).text=("Adresse")
    table.cell(4,0).text=("Numéro de téléphone")
    table.cell(5,0).text=("Numéro de télécopie")
    table.cell(6,0).text=("Courriel")
    n=0
    for i in range (1,7):
        c=table.cell(i, 1)
        d=table.cell(i,2)
        c.merge(d)
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    if n==0:
                        fontdebut.bold=True
                        fontdebut.size = docx.shared.Pt(11)
                    else:
                        fontdebut.size = docx.shared.Pt(10)
                    n=n+1
    #separe les deux tableaux
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=7, cols=3, style='Table Grid')
    a=table.cell(0,0)
    b=table.cell(0,2)
    a.merge(b)
    table.cell(0,0).text=("B2. Représentant légal  du promoteur dans la Communauté européenne pour la recherche (si différent du promoteur)")
    table.cell(1,0).text=("Nom de l'organisme")
    table.cell(2,0).text=("Nom de la personne à contacter ")
    table.cell(3,0).text=("Adresse")
    table.cell(4,0).text=("Numéro de téléphone")
    table.cell(5,0).text=("Numéro de télécopie")
    table.cell(6,0).text=("Courriel")
    n=0
    for i in range (1,7):
        c=table.cell(i, 1)
        d=table.cell(i,2)
        c.merge(d)
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    if n==0:
                        fontdebut.bold=True
                        fontdebut.size = docx.shared.Pt(11)
                    else:
                        fontdebut.size = docx.shared.Pt(10)
                    n=n+1
    #separe les deux tableaux
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("B3. Statut du promoteur \nCommercial	□            	Non commercial	□")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.bold=True
                    fontdebut.size = docx.shared.Pt(11)
    
    '''Partie C'''
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\n\nC. Identification du demandeur ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    fontdebut.bold=True
    sentence=paragraph.add_run("(cocher les cases appropriées)\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    
    table = document.add_table(rows=2, cols=2, style='Table Grid')
    a=table.cell(1,0)
    b=table.cell(1,1)
    a.merge(b)
    table.cell(0,0).text=("C1. Demandeur pour l'ANSM 		□")
    table.cell(0,1).text=("C2. Demandeur pour le CPP  		□")
    table.cell(1,0).text=("\nPromoteur ……………………………………………………………………………………………□\n"
                          "\nReprésentant légal du promoteur …………………………………………………………………□\n"
                          "\nPersonne ou organisme délégué par le promoteur pour soumettre la demande…………….□\n"
                          "\nDans ce cas, compléter ci-après :")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    if n==0:
                        fontdebut.bold=True
                        fontdebut.size = docx.shared.Pt(11)
                    else:
                        fontdebut.size = docx.shared.Pt(10)
                    n=n+1
                    
    table = document.add_table(rows=6, cols=3, style='Table Grid')
    table.cell(0,0).text=("Nom de l'organisme")
    table.cell(1,0).text=("Nom de la personne à contacter")
    table.cell(2,0).text=("Adresse")
    table.cell(3,0).text=("Numéro de téléphone")
    table.cell(4,0).text=("Numéro de télécopie")
    table.cell(5,0).text=("Courriel")
    for i in range (0,6):
        c=table.cell(i, 1)
        d=table.cell(i,2)
        c.merge(d)   
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
       
        
#partie D
def partie_D(document):
    
    paragraph=document.add_paragraph("\n\nD. Fiche de données sur le(s) DM (s)/ DM-DIV (s) faisant l'objet de la recherche, y compris les comparateurs \n", style='debut_page')
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("Indiquer ici quel DM / DM-DIV est concerné par cette section D (utiliser une fiche pour chaque DM / DM-DIV)\n")
    table.cell(1,0).text=("Dispositif sur lequel porte la recherche ……………………………………………….□\nDispositif utilisé comme comparateur………………………………….………………□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    if n==0:
                        fontdebut.italic=True
                        fontdebut.size = docx.shared.Pt(10)
                    else:
                        fontdebut.bold=True
                        fontdebut.size = docx.shared.Pt(10)
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,0)
    a.merge(b)
    
    paragraph=document.add_paragraph("\nD1. Statut du DM / DM-DIV\n", style='debut_page')
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("Le dispositif est-il marqué CE ?	                                       		□ oui   □ non\n\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.bold=True
    sentence=paragraph.add_run("Si le dispositif est marqué CE, compléter la rubrique ci-dessous ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    
    table = document.add_table(rows=10, cols=2, style='Table Grid')
    table.cell(0,0).text=("Nom de l'Organisme Notifié")
    table.cell(1,0).text=("Numéro de l'Organisme Notifié")
    table.cell(2,0).text=("Date du marquage CE")
    table.cell(2,1).text=("     /     /     ")
    table.cell(3,0).text=("Destination(s) du marquage CE (telles que mentionnées dans la notice) ")
    table.cell(4,0).text=("Destination(s) du dispositif dans l’essai  ")
    table.cell(5,0).text=("Est-ce que l’utilisation du dispositif, dans le cadre de la recherche, se fait dans l’indication de son marquage CE ?")
    table.cell(5,1).text=("□ oui   □ non")
    table.cell(6,0).text=("La destination du dispositif dans l’essai est-elle conforme à des recommandations publiées (HAS, ANSM, Sociétés savantes, etc..) ?")
    table.cell(6,1).text=("□ oui   □ non   □ NA")
    table.cell(7,0).text=("Si oui, citer les références : ")
    table.cell(8,0).text=("Le dispositif est-il commercialisé dans un Etat membre de la Communauté européenne ou dans un pays tiers ? ")
    table.cell(8,1).text=("□ oui   □ non")
    table.cell(9,0).text=("Si oui, citer les pays concernés  ")
    n=0
    for col in table.columns:
        for cell in col.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n>9:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    n=n+1
    
    '''Partie D2'''
    
    paragraph=document.add_paragraph("\nD2. Identification du dispositif à étudier \n", style='debut_page')
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("Rubriques à compléter dans tous les cas")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    
    table = document.add_table(rows=10, cols=2, style='Table Grid')
    table.cell(0,0).text=("Dénomination commune \n(exemple : stent artériel…)")
    table.cell(1,0).text=("Dénomination commerciale")
    table.cell(2,0).text=("Modèle")
    table.cell(3,0).text=("Version (y compris version du logiciel) ")
    table.cell(4,0).text=("Classification européenne")
    table.cell(5,0).text=("Classe du DM :\n")
    table.cell(5,1).text=("Classe du DMDIV :\n")
    table.cell(6,0).text=("Classe I                                                    □\nClasse IIa invasif à long terme                 □\n"
                          "Autres IIa                                                  □\nClasse IIb                                                 □\n"
                          "Classe III                                                  □\nDMIA                                                        □\n")
    table.cell(6,1).text=("Hors annexe II      		□\nAnnexe II liste A    		□\n"
                          "Annexe II liste B    		□\nAutotest                              	□\n")
    a=table.cell(7,0)
    b=table.cell(7,1)
    a.merge(b)
    table.cell(7,0).text=("En cas de dispositif non pourvu du marquage CE, renseigner la classe du dispositif et joindre une justification de la classification ")
    table.cell(8,0).text=("S’agit-il d’un dispositif implantable ?")
    table.cell(8,1).text=("□ oui   □ non")
    table.cell(9,0).text=("S’agit-il d’un dispositif « sur mesure » ?")
    table.cell(9,1).text=("□ oui   □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==12 or n==14:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==5 or n==6:
                        fontdebut.bold=True
                    elif n==9:
                        fontdebut.italic=True
                    n=n+1
    c=table.cell(5,0)
    d=table.cell(6,0)
    c.merge(d)
    e=table.cell(5,1)
    f=table.cell(6,1)
    e.merge(f)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=6, cols=2, style='Table Grid')
    table.cell(0,0).text=("Fabricant du dispositif à étudier ")
    table.cell(0,1).text=("(à compléter quel que soit le statut du promoteur) ")
    table.cell(1,0).text=("Nom")
    table.cell(2,0).text=("Adresse")
    table.cell(3,0).text=("Numéro de téléphone")
    table.cell(4,0).text=("Numéro de télécopie")
    table.cell(5,0).text=("Courriel")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold=True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,1)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=5, cols=8, style='Table Grid')
    for i in range (0,5):
        a=table.cell(i,0)
        b=table.cell(i,5)
        a.merge(b)
    table.cell(0,0).text=("Le dispositif sur lequel porte la recherche contient-il une des substances suivantes :")
    table.cell(0,6).text=("oui")
    table.cell(0,7).text=("non")
    table.cell(1,0).text=("-	Substance qui, si elle est utilisée séparément, est susceptible d'être considérée comme un médicament ? ")    
    table.cell(2,0).text=("-	Produits d’origine biologique (DMOA) ou dans la fabrication duquel interviennent de tels produits ?")
    table.cell(3,0).text=("-	OGM ?")
    table.cell(4,0).text=("-	Radioélément ?")
    for i in range (1,5):
        table.cell(i,6).text=("□")
        table.cell(i,7).text=("□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==5 or n==13 or n==21 or n==29 or n==37:
                   paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                   paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    if n==0:
                        fontdebut.bold=True
                        fontdebut.size = docx.shared.Pt(11)
                    elif n==6 or n==7:
                        fontdebut.size = docx.shared.Pt(11)
                    else:
                        fontdebut.size = docx.shared.Pt(10)
                    n=n+1

    '''Partie D3'''
    paragraph=document.add_paragraph("\nD3. Cas particulier : utilisation de dispositifs à étudier commercialisés et ayant la même dénomination commune, dans le cadre d’un essai dont le protocole n’impose pas l’utilisation d’un dispositif en particulier", style='debut_page')
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nEst-ce que ce cas particulier est applicable à l’essai concerné ?			□ oui   □ non\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("Si oui,  compléter la rubrique ci-dessous")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    
    table = document.add_table(rows=3, cols=8, style='Table Grid')
    for i in range (0,3):
        a=table.cell(i,0)
        b=table.cell(i,5)
        a.merge(b)
    table.cell(0,6).text=("oui")
    table.cell(0,7).text=("non")
    table.cell(1,0).text=("DM")
    table.cell(1,6).text=("□")
    table.cell(1,7).text=("□")
    table.cell(2,0).text=("DMDIV")
    table.cell(2,6).text=("□")
    table.cell(2,7).text=("□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==0 or n==1 or n==8 or n==9 or n==16 or n==17:
                   paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(11)
                    n=n+1
    
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("Dans l’affirmative, préciser les informations mentionnées ci-après")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    
    table = document.add_table(rows=5, cols=4, style='Table Grid')
    c=table.cell(0,0)
    d=table.cell(1,0)
    c.merge(d)
    table.cell(0,0).text=("Nom du dispositif")
    e=table.cell(0,1)
    f=table.cell(1,1)
    e.merge(f)
    table.cell(0,1).text=("sans marquage CE")
    a=table.cell(0,2)
    b=table.cell(0,3)
    a.merge(b)
    table.cell(0,2).text=("disposant d'un marquage CE")
    table.cell(1,2).text=("Utilisation conforme au marquage CE")
    table.cell(1,3).text=("Utilisation dans une autre destination que celle du marquage CE")
    for i in range (2,5):
        table.cell(i,1).text=("□")
        table.cell(i,2).text=("□")
        table.cell(i,3).text=("□")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(11)
    
    '''Partie D4'''
    paragraph=document.add_paragraph("\nD4. Dossier technique du dispositif faisant l'objet de la recherche\n", style='debut_page')
    table = document.add_table(rows=10, cols=8, style='Table Grid')
    for i in range (0,10):
        a=table.cell(i,0)
        b=table.cell(i,5)
        a.merge(b)
    table.cell(0,6).text=("oui")
    table.cell(0,7).text=("non")
    table.cell(1,0).text=("Dossier technique complet ")
    table.cell(2,0).text=("Dossier technique simplifié")
    a=table.cell(3,0)
    b=table.cell(3,7)
    a.merge(b)
    table.cell(4,0).text=("1. Dispositif marqué CE utilisé dans la destination du marquage")
    table.cell(5,0).text=("2. DM de classe I ou IIa (à l'exception des classes IIa invasifs à long terme) marqué CE hors indication") 
    f=table.cell(7,0)
    g=table.cell(7,7)
    g.merge(f)
    e=table.cell(8,0)
    d=table.cell(9,0)
    e.merge(d)
    table.cell(6,0).text=("3. Dispositif ayant fait l'objet d'une précédente demande d’autorisation de recherche auprès de l'ANSM\n")
    table.cell(8,0).text=("- dans la même destination et dans les mêmes conditions "
                          "(Mentionner le N° IDRCB) :\n"
                          "- dans une autre destination "
                          "(Mentionner le N° IDRCB) :")
    for i in range (1,10):
        if not(i==3) or not(i==7):
            table.cell(i,6).text=("□")
            table.cell(i,7).text=("□")
    table.cell(3,0).text=("En cas de dossier technique simplifié, cocher la ou les cases ci-dessous :")
    table.cell(7,0).text=("Dans l’affirmative, préciser si le dispositif était utilisé dans la précédente demande :")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==7 or n==15 or n==25 or n==31 or n==39 or n==47 or n==57 or n==63 or n==71:
                   paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==2 or n==5:
                        fontdebut.bold=True
                    n=n+1
    
    paragraph=document.add_paragraph("\nD5. Informations supplémentaires sur le dispositif à étudier ou comparateur\n", style='debut_page')
    table = document.add_table(rows=6, cols=8, style='Table Grid')
    for i in range(0,5):
        a=table.cell(i,0)
        b=table.cell(i,5)
        a.merge(b)
    table.cell(0,6).text=("oui")
    table.cell(0,7).text=("non")
    table.cell(1,0).text=("Est-ce que le dispositif sur lequel porte la recherche appartient à un tiers ?")
    table.cell(1,6).text=("□")
    table.cell(1,7).text=("□")
    a=table.cell(2,0)
    b=table.cell(2,7)
    a.merge(b)
    table.cell(2,0).text=("si oui, joindre l'autorisation délivrée par ce dernier au promoteur pour communiquer les données relatives au dispositif concerné")
    table.cell(3,0).text=("Est-ce que la brochure pour l'investigateur appartient à un tiers ?")
    table.cell(3,6).text=("□")
    table.cell(3,7).text=("□")
    table.cell(4,0).text=("Est-ce que le dossier technique appartient à un tiers ?")
    table.cell(4,6).text=("□")
    table.cell(4,7).text=("□")
    d=table.cell(5,0)
    e=table.cell(5,7)
    d.merge(e)
    table.cell(5,0).text=("si oui dans l’un ou les deux cas précédents, joindre l'autorisation du tiers délivrée au promoteur pour l'utiliser cette brochure pour l’investigateur et/ou le dossier technique")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==7 or n==17 or n==23 or n==31 or n==41:
                   paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                   paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    if n==0 or n==1:
                        fontdebut.size = docx.shared.Pt(11)
                    else:
                        fontdebut.size = docx.shared.Pt(10)
                    n=n+1
    
    #parties E et F
def partie_E_F(document):
    
    '''Partie E'''
    document.add_page_break()
    
    paragraph=document.add_paragraph("E. Informations sur le dispositif utilisé comme placebo\n", style='debut_page')
    table = document.add_table(rows=2, cols=2, style='Table Grid')
    table.cell(0,0).text=("Description / Composition ")
    table.cell(1,0).text=("Mode d'utilisation / Indication ")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
    
    paragraph=document.add_paragraph()
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=6, cols=2, style='Table Grid')
    table.cell(0,0).text=("Fabricant  du placebo ")
    table.cell(0,1).text=("(à compléter quel que soit le statut du promoteur)")
    table.cell(1,0).text=("Nom")
    table.cell(2,0).text=("Adresse")
    table.cell(3,0).text=("Numéro de téléphone")
    table.cell(4,0).text=("Numéro de télécopie")
    table.cell(5,0).text=("Courriel")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    if n==0 or n==1:
                        fontdebut.size = docx.shared.Pt(11)
                        if n==0:
                            fontdebut.bold=True
                    else:
                        fontdebut.size = docx.shared.Pt(10)
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,1)
    a.merge(b)
    
    '''Partie F'''
    paragraph=document.add_paragraph("\n\nF. Données générales sur la recherche\nF1. Domaine concerné par la recherche\n", style='debut_page')
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("1)	Domaine médical ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.bold=True
    sentence=paragraph.add_run("(cocher 1 seule case)\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    sentence=paragraph.add_run("Médecine    □  				Chirurgie   □  				Imagerie / diagnostic   □\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("\n2)	Domaine thérapeutique principal  ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.bold=True
    sentence=paragraph.add_run("(cocher 1 seule case)\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    
    table = document.add_table(rows=1, cols=3, style='Table Grid')
    table.cell(0,0).text=("Anesthésie/ Réanimation	□\n"
                          "Cancérologie                  	□\n"
                          "Cardiologie/vasculaire    	□\n"
                          "Dermatologie                  	□\n"
                          "Endocrinologie/Diabétologie	□\n")
    table.cell(0,1).text=("Gastro-entérologie          	□\n"
                          "Gynécologie                   	□\n"
                          "Neurologie                      	□\n"
                          "Ophtalmologie                	□\n"
                          "Orthopédie                      	□\n")
    table.cell(0,2).text=("ORL                                	□\n"
                          "Pneumologie                  	□\n"
                          "Urologie/Néphrologie      	□\n"
                          "Autre (à préciser) :")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)

    '''Partie F2'''
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nF2. S’agit-il d’une recherche de première utilisation chez l’homme dans la ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    fontdebut.bold=True
    sentence=paragraph.add_run("                 □ oui    □ non\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    sentence=paragraph.add_run("destination de l’essai ?\n\nF3. Procédures prévues pour les seuls besoins de la recherche\n\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    fontdebut.bold=True
    sentence=paragraph.add_run("1)	Prélèvements biologiques pour les seuls besoins de la recherche (c’est à dire prélèvements qui n’auraient pas été réalisés si le sujet ne se prêtait pas à cette recherche)\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.bold=True
    sentence=paragraph.add_run("exemple : sang, urine, salive, tissus, liquide céphalorachidien ...\n\n"
                               "Est-ce que de tels prélèvements sont prévus dans le cadre de la recherche ? 				□ oui   □ non\n\n"
                               "si oui, compléter le tableau ci-dessous")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("(rajouter autant de ligne que nécessaire)")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    table = document.add_table(rows=5, cols=4, style='Table Grid')
    table.cell(0,0).text=("Type de prélèvements")
    table.cell(0,1).text=("Fréquence de réalisation")
    table.cell(0,2).text=("Volume / Diamètre unitaire")
    table.cell(0,3).text=("Volume / Nombre cumulé")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)

    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("2)	Examens spécifiques pour les seuls besoins de la recherche (c’est à dire examens qui n’auraient pas été réalisés si le sujet ne se prêtait pas à cette recherche)\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.bold=True
    sentence=paragraph.add_run("Est-ce que de tels examens sont prévus dans le cadre de la recherche ? 					□ oui   □ non\n"
              "•	si oui, s’agit-il d’examens irradiants et/ou invasifs ?	                            □ oui   □ non\n\n\n\n"
              "Dans l’affirmative, compléter le tableau ci-dessous ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("(rajouter autant de ligne que nécessaire)")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    table = document.add_table(rows=4, cols=6, style='Table Grid')
    table.cell(0,0).text=("Examens")
    table.cell(0,1).text=("Fréquence de réalisation de ces examens")
    table.cell(0,2).text=("Fréquence usuelle (oui/non)")
    table.cell(0,3).text=("Délai entre les examens")
    table.cell(0,4).text=("Dose administrée par examen (si applicable)")
    table.cell(0,5).text=("Dose cumulée sur la durée de l’étude (si applicable)")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)

    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nDans la négative, lister les autres examens :\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("\n\n\n")
    
    '''Partie F4'''
    paragraph=document.add_paragraph("\nF4. Informations sur les produits de santé non expérimentaux\n", style='debut_page')
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("1) L’utilisation d’un médicament non expérimental ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.bold=True
    sentence=paragraph.add_run("(MNE) est-elle prévue dans cet essai ?		□ oui   □ non\n\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("Dans l’affirmative, préciser les informations suivantes (rajouter autant de ligne que nécessaire)")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    
    table = document.add_table(rows=4, cols=6, style='Table Grid')
    a=table.cell(0,0)
    b=table.cell(1,1)
    a.merge(b)
    c=table.cell(2,0)
    d=table.cell(2,1)
    c.merge(d)
    e=table.cell(3,0)
    f=table.cell(3,1)
    e.merge(f)
    table.cell(0,0).text=("MNE concerné")
    g=table.cell(0,2)
    h=table.cell(0,3)
    g.merge(h)
    table.cell(0,2).text=("MNE disposant d’une AMM \n(en France, en Europe, aux Etats-Unis ou au Japon)")
    i=table.cell(0,4)
    j=table.cell(0,5)
    i.merge(j)
    table.cell(0,4).text=("Si le MNE dispose d’une AMM, son utilisation dans l’essai est-elle divergente par rapport à l’AMM? ")
    table.cell(1,2).text=("oui")
    table.cell(1,3).text=("non")
    table.cell(1,4).text=("oui")
    table.cell(1,5).text=("non")
    for i in range(2,4):
        for j in range(2,6):
            table.cell(i,j).text=("□")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)

    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("En cas d’utilisation d’un MNE ne disposant pas d’une AMM (en France, UE, Etats-Unis ou Japon), veuillez en indiquer ci-dessous les justifications ou préciser où se trouve cette information dans le dossier soumis")
    table.cell(1,0).text=("\n\n\n")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)

    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("Est-il prévu d’importer des médicaments non expérimentaux pour les besoins de la recherche ? ")
    table.cell(1,0).text=("Si oui, joindre le Formulaire « Attestation en vue de l’importation de médicaments nécessaires à la réalisation d’une recherche » ")
    c=table.cell(0,5)
    d=table.cell(1,5)
    c.merge(d)
    table.cell(0,5).text=("□ oui □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==3:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold=True
                    elif n==3:
                        fontdebut.size = docx.shared.Pt(9)
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,4)
    a.merge(b)
    
    
    document.add_page_break()
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("2)	L’utilisation d’un dispositif non expérimental ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.bold=True
    sentence=paragraph.add_run("est-elle prévue dans cet essai ?			□ oui   □ non\n\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("Dans l’affirmative, préciser les informations suivantes (rajouter autant de ligne que nécessaire)\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    
    table = document.add_table(rows=3, cols=8, style='Table Grid')
    for i in range(0,3):
        a=table.cell(i,0)
        b=table.cell(i,5)
        a.merge(b)
    table.cell(0,6).text=("oui")
    table.cell(0,7).text=("non")
    table.cell(1,0).text=("DM")
    table.cell(1,6).text=("□")
    table.cell(1,7).text=("□")
    table.cell(2,0).text=("DMDIV")
    table.cell(2,6).text=("□")
    table.cell(2,7).text=("□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==0 or n==1 or n==8 or n==9 or n==16 or n==17:
                   paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(11)
                    n=n+1
    
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("Dans l’affirmative, préciser les informations mentionnées ci-après")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    
    table = document.add_table(rows=5, cols=4, style='Table Grid')
    c=table.cell(0,0)
    d=table.cell(1,0)
    c.merge(d)
    table.cell(0,0).text=("Nom du dispositif")
    e=table.cell(0,1)
    f=table.cell(1,1)
    e.merge(f)
    table.cell(0,1).text=("sans marquage CE")
    a=table.cell(0,2)
    b=table.cell(0,3)
    a.merge(b)
    table.cell(0,2).text=("disposant d'un marquage CE")
    table.cell(1,2).text=("Utilisation conforme au marquage CE")
    table.cell(1,3).text=("Utilisation dans une autre destination que celle du marquage CE")
    for i in range(2,5):
        for j in range(1,4):
            table.cell(i,j).text=("□")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(11)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("En cas d’utilisation d’un dispositif ne disposant pas d’un marquage CE, veuillez en indiquer ci-dessous les justifications ou préciser où se trouve cette information dans le dossier soumis (un dossier technique est requis)")
    table.cell(0,1).text=("\n\n\n")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\n3)	L’utilisation d’un produit cosmétique ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.bold=True
    sentence=paragraph.add_run("est-elle prévue dans cet essai ?			□ oui   □ non\n\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("Dans l’affirmative, Dans l’affirmative, préciser pour chacun d’eux s’ils sont commercialisés en France, UE, ou autre")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("(rajouter autant de ligne que nécessaire)\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("\n\n\n")
    
    
    document.add_page_break()
    
#a partir de F5
def partieF5_suite(document):
    ''' Partie F5'''
    paragraph=document.add_paragraph("F5 Méthodologie de la recherche ", style='debut_page')
    table = document.add_table(rows=36, cols=8, style='Table Grid')
    #on merge les cellules
    for i in range(0,36):
        if i==26 or i==28 or i==35:
            a=table.cell(i,0)
            b=table.cell(i,7)
            a.merge(b)
        else:
            a=table.cell(i,0)
            b=table.cell(i,5)
            a.merge(b)
    table.cell(0,6).text=("oui")
    table.cell(0,7).text=("non")
    table.cell(1,0).text=("Type d’investigation clinique ?")
    a=table.cell(2,0)
    b=table.cell(3,0)
    a.merge(b)
    table.cell(2,0).text=("Exploratoire (preuve de concept / faisabilité / pilote / FIM ) ……………………………...\n"
                          "Confirmatoire (démonstrative / pivot)………………………………………………………..")
    table.cell(4,0).text=("Effectif de l’étude justifié par le calcul ?")
    i=table.cell(5,0)
    j=table.cell(7,0)
    i.merge(j)
    table.cell(5,0).text=("Si oui, préciser les risques  de 1ère et  de 2ème espèce associés,\n"
                          "     risque  ≤ 0.05……………………………………………………………………….………\n"
                          "     puissance (1-) ≥ 80%…………………………………………………….………………..")
    table.cell(8,0).text=("Randomisation ?")
    table.cell(9,0).text=("Stratification ?")
    table.cell(10,0).text=("L’essai est-il comparatif ?")
    q=table.cell(11,0)
    r=table.cell(25,0)
    q.merge(r)
    table.cell(11,0).text=("Si oui préciser,\n"
                           "L’objectif de la comparaison :\n"
                           "\nSupériorité………………………………………………………………………………………\n"
                           "Non-infériorité…………………………………………………………………………………..\n"
                           "Equivalence…………………………………………………………………………………….\n"
                           "\nLe plan expérimental :\n"
                           "Ouvert…………………………………………………………………………………………...\n"
                           "Simple insu…………………………………………………………………………………….\n"
                           "Double insu…………………………………………………………………………………….\n"
                           "Autre insu : procédures d’évaluation centralisée………………………………………….\n"
                           "Groupes parallèles……………………………………………………………………………\n"
                           "Simple bras / apparié / propre témoin.………………………………………………………\n"
                           "Plan croisé……………………………………………………………………………………...\n"
                           "Autre ……………………………………………………………………………………………\n"
                           "Si autre préciser :")
    table.cell(26,0).text=("L’essai n’est pas comparatif, justifier brièvement pourquoi :\n\n\n\n")
    table.cell(27,0).text=("Comparateur(s) utilisé(s) ?")
    table.cell(28,0).text=("Si oui préciser :")
    table.cell(29,0).text=("Autre DM / DMDIV.………………………………………………………………………….")
    table.cell(30,0).text=("Médicament………………………………………………………………………………….")
    table.cell(31,0).text=("Placebo……………………………………………………………………………………….")
    table.cell(32,0).text=("Procédure « sham » / fantôme…………………………………………………………….")
    table.cell(33,0).text=("Procédure de confirmation diagnostique (« Gold Standard »)…………………………...")
    table.cell(34,0).text=("Autre……………………………………………………………………………………………")
    table.cell(35,0).text=("Si autre préciser :")
    #on met les carrés
    for i in range(2,35):
        for j in range(6,8):
            if not(i==1 or i==5 or i==11 or i==12 or i==16 or i==25 or i==26 or i==28 or i==35):
                table.cell(i,j).text=("□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==197 or n==203 or n==213 or n==219 or n==227 or n==235 or n==243 or n==251 or n==259 or n==269 or n==189 or n==7 or n==21 or n==29 or n==51 or n==59 or n==67 or n==75:
                   paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                   paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==197 or n==203 or n==7 or n==29 or n==59 or n==67 or n==75:
                        fontdebut.bold=True
                    n=n+1
    a=table.cell(1,0)
    b=table.cell(2,0)
    a.merge(b)
    a=table.cell(4,0)
    b=table.cell(5,0)
    a.merge(b)
    a=table.cell(11,0)
    b=table.cell(12,0)
    a.merge(b)


    '''Partie F6'''
    paragraph=document.add_paragraph("\nF6. Design de l’essai\n", style='debut_page')
    table = document.add_table(rows=9, cols=3, style='Table Grid')
    for i in range(0, 9):
        a=table.cell(i,1)
        b=table.cell(i,2)
        a.merge(b)
    table.cell(0,0).text=("Objectif principal")
    table.cell(1,0).text=("Objectifs secondaires")
    table.cell(2,0).text=("Critère principal de jugement")
    table.cell(3,0).text=("Critères secondaires de jugement")
    table.cell(4,0).text=("Principaux critères d’inclusion")
    table.cell(5,0).text=("Principaux critères de non inclusion")
    table.cell(6,0).text=("Durée de participation pour une personne se prêtant à la recherche")
    table.cell(7,0).text=("Durée de suivi pour une personne se prêtant à la recherche")
    table.cell(8,0).text=("Durée de l’essai")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
    
    '''Partie F7'''
    paragraph=document.add_paragraph("\nF7. Autres informations sur l’essai\n", style='debut_page')
    table = document.add_table(rows=4, cols=4, style='Table Grid')
    for i in range(0, 2):
        a=table.cell(i,0)
        b=table.cell(i,2)
        a.merge(b)
    for i in range(2, 4):
        a=table.cell(i,0)
        b=table.cell(i,3)
        a.merge(b)
    table.cell(0,0).text=("La constitution d’un comité indépendant d’évaluation de données de performance est-elle prévue ?")
    table.cell(1,0).text=("La constitution d’un comité de surveillance indépendant est-elle prévue ?")
    table.cell(0,3).text=("□ oui □ non")
    table.cell(1,3).text=("□ oui □ non")
    table.cell(2,0).text=("Dans la négative, veuillez indiquer ci-dessous les justifications de non constitution d’un tel comité (justifications requises conformément aux dispositions de l’article L. 1123-7 du code de la santé publique (CSP)).\n\n\n")
    table.cell(3,0).text=("Si cette justification n’est pas apportée ici, préciser où se trouve cette information dans le dossier soumis.\n\n\n")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==3 or n==7:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==2 or n==6:
                        fontdebut.bold=True
                    n=n+1
    
    document.add_page_break()
    
def a_partir_F8(document):
    
    '''Partie F8'''
    paragraph=document.add_paragraph("F8. Nombre de lieux de recherche et de pays concernés par la recherche\n", style='debut_page')
    table = document.add_table(rows=7, cols=8, style='Table Grid')
    for i in range(0, 7):
        a=table.cell(i,0)
        b=table.cell(i,5)
        a.merge(b)
    c=table.cell(4,0)
    d=table.cell(4,7)
    c.merge(d)
    e=table.cell(6,0)
    f=table.cell(6,7)
    e.merge(f)
    table.cell(0,6).text=("oui")
    table.cell(0,7).text=("non")
    table.cell(1,0).text=("L’essai est-il monocentrique (voir aussi section G) ?")
    table.cell(2,0).text=("L’essai est-il multicentrique (voir aussi section G) ?")
    table.cell(3,0).text=("L’essai est-il prévu pour être mené dans plusieurs Etats membres de l’Union européenne?")
    table.cell(4,0).text=("Si oui, lesquels ?")
    table.cell(5,0).text=("Cet essai implique-t-il des pays tiers à la Communauté européenne ?")
    table.cell(6,0).text=("Si oui, lesquels ?")
    for i in range(1,4):
        table.cell(i,6).text=("□")
        table.cell(i,7).text=("□")
    table.cell(5,6).text=("□")
    table.cell(5,7).text=("□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==7 or n==15 or n==23 or n==33 or n==39 or n==49:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    n=n+1
    
    '''Partie F9'''
    paragraph=document.add_paragraph("\nF9. Décision rendue par d’autres autorités compétentes dans l’UE\n", style='debut_page')
    table = document.add_table(rows=9, cols=5, style='Table Grid')
    a=table.cell(0,0)
    b=table.cell(0,3)
    a.merge(b)
    table.cell(0,0).text=("S’agit-il d’un essai mené dans d’autres Etats membres de l’Union européenne (EM) ?")
    table.cell(0,4).text=("□ oui □ non")
    c=table.cell(1,0)
    d=table.cell(1,4)
    c.merge(d)
    table.cell(1,0).text=("Si oui, lister les pays concernés :")
    a=table.cell(2,0)
    b=table.cell(2,3)
    a.merge(b)
    table.cell(2,0).text=("Dans ce cas, est ce qu’une autorité compétente (AC) a déjà rendu une décision finale sur l’essai (au moment du dépôt de la demande d’AEC à l’ANSM) ?")
    table.cell(2,4).text=("□ oui □ non")
    c=table.cell(3,0)
    d=table.cell(3,4)
    c.merge(d)
    table.cell(3,0).text=("Dans l’affirmative, préciser uniquement pour chaque EM où l’AC a déjà rendu une décision finale, si les documents suivants soumis dans le dossier de demande d’AEC sont identiques à ceux sur lesquels a reposé la décision rendue par cette AC")
    table.cell(4,0).text=("Etat Membre")
    table.cell(4,1).text=("Décision finale de l’AC")
    table.cell(4,2).text=("Même Protocole")
    table.cell(4,3).text=("Même Brochure pour l’investigateur (BI)")
    table.cell(4,4).text=("Même Dossier Technique (DT)")
    for i in range (5,8):
        for j in range(1,5):
            if j==1:
                table.cell(i,j).text=("□ autorisation □ refus")
            else:
                table.cell(i,j).text=("□ oui □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==4 or n==14:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    n=n+1
    document.add_page_break()
    
    #partie F10
    
def partie_F10(document):
    
    paragraph=document.add_paragraph("F10. Personnes se prêtant à la recherche\n", style='debut_page')
    table = document.add_table(rows=27, cols=8, style='Table Grid')
    for i in range (0,27):
        if i==1 or i==12 or i==15:
            a=table.cell(i,0)
            b=table.cell(i,7)
            a.merge(b)
        else:
            a=table.cell(i,0)
            b=table.cell(i,5)
            a.merge(b)
    table.cell(0,6).text=("oui")
    table.cell(0,7).text=("non")
    table.cell(1,0).text=("Tranche d'âge étudiée")
    table.cell(2,0).text=("Moins de 18 ans")
    table.cell(10,0).text=("De 18 à 65 ans……………………………………………………………………………………")
    table.cell(11,0).text=("Plus de 65 ans……………………………………………………………………………………")
    table.cell(12,0).text=("Sexe")
    table.cell(15,0).text=("Population")
    table.cell(18,0).text=("En particulier : ")
    for i in range (2,27):
        if i<3 or 3<i<12 or 12<i<15 or 15<i<18 or 18<i<24 or 24<i<26:
            table.cell(i,6).text=("□")
            table.cell(i,7).text=("□")
    a=table.cell(3,0)
    b=table.cell(9,5)
    a.merge(b)
    c=table.cell(13,0)
    d=table.cell(14,5)
    c.merge(d)
    e=table.cell(16,0)
    f=table.cell(17,5)
    e.merge(f)
    g=table.cell(19,0)
    h=table.cell(26,5)
    g.merge(h)
    table.cell(3,0).text=("Si oui, préciser :\n"
                          "In Utero	...............................................................................................................................\n"
                          "Nouveau-nés prématurés (jusqu’à l’âge gestationnel ≤ 37 semaines)…………………\n"
                          "Nouveau-nés (0-27 jours)……………………………………………………………………..\n"
                          "Nourrissons (28 jours - 23 mois)………………………………………………………………\n"
                          "Enfants (2-11 ans)………………………………………………………………………………\n"
                          "Adolescents (12-17 ans)…………………………………………………………………………")
    table.cell(13,0).text=("Femmes……………………………………………………………………………………………\n"
                           "Hommes…………………………………………………………………………………………..")
    table.cell(16,0).text=("Sujets sains………………………………………………………………………………………..\n"
                           "Sujets malades...………………………………………………………………………………….")
    table.cell(19,0).text=("- femmes en âge de procréer…………………………………………………………………..\n"
                           "- femmes enceintes……………………………………………………………………………...\n"
                           "- femmes allaitantes……………………………………………………………………………..\n"
                           "- personnes en situation d’urgence…………………………………………………………….\n"
                           "- personnes incapables de donner personnellement leur consentement………………..\n"
                           "\nSi oui, préciser : \n"
                           "- autres……………………………………………………………………………………………..\n"
                           "Si oui, préciser :\n")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==9 or n==15 or n==69 or n==77 or n==85 or n==95 or n==109 or n==119 or n==133 or n==141 or n==201:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==9 or n==15 or n==77 or n==85 or n==95 or n==119 or n==141:
                        fontdebut.bold=True
                    n=n+1
    k=table.cell(2,0)
    a.merge(k)
    l=table.cell(10,0)
    m=table.cell(11,0)
    a.merge(l)
    a.merge(m)
    n=table.cell(18,0)
    n.merge(g)
    
    paragraph=document.add_paragraph()
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=4, cols=4, style='Table Grid')
    a=table.cell(0,0)
    b=table.cell(0,3)
    a.merge(b)
    c=table.cell(1,0)
    d=table.cell(3,2)
    c.merge(d)
    table.cell(0,0).text=("Nombre prévu de personnes à inclure :")
    table.cell(1,0).text=("-	en France……………………………………………………………………………………\n"
                          "-	dans la Communauté européenne………………………………………………………\n"
                          "-	pour l’ensemble de la recherche…………………………………………………………")
    table.cell(1,3).text=("     ")
    table.cell(2,3).text=("     ")
    table.cell(3,3).text=("     ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==7 or n==11 or n==15:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold=True
                    n=n+1
    document.add_page_break()
    
    paragraph=document.add_paragraph()
    sentence= paragraph.add_run("F10. Traitements ou soins ou examens, procédures, prévus pour les personnes se prêtant à la recherche à la fin de leur participation à l’essai\n\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    fontdebut.bold=True
    sentence=paragraph.add_run("Est-ce qu’il diffère du traitement habituel de la pathologie étudiée ?				□ oui   □ non\n\nSi oui, à préciser et justifier : \n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("\n\n\n")
    
    paragraph=document.add_paragraph()
    paragraph=document.add_paragraph()
    
    '''Partie G'''
    paragraph=document.add_paragraph("G. Lieux de recherche envisagés en France\n\nG1. Investigateur coordonnateur \n", style='debut_page')
    
    table = document.add_table(rows=9, cols=2, style='Table Grid')
    table.cell(0,0).text=("Nom / Prénoms")
    table.cell(1,0).text=("Qualification, Spécialité ")
    table.cell(2,0).text=("Adresse professionnelle")
    table.cell(3,0).text=("Nom de l’établissement")
    table.cell(4,0).text=("Service ")
    table.cell(5,0).text=("Adresse")
    table.cell(6,0).text=("Numéro de téléphone ")
    table.cell(7,0).text=("Numéro de télécopie ")
    table.cell(8,0).text=("Courriel")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold=True
                        fontdebut.size = docx.shared.Pt(11)
                    n=n+1
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\n\nG2. Autres investigateurs : ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    fontdebut.bold=True
    sentence=paragraph.add_run("(à compléter pour chaque investigateur)"
                               "\n(En cas d’essai multicentrique avec un grand nombre de lieux de recherche, il est possible de, fournir une liste contenant les informations de la rubrique pour chaque investigateur en annexe de ce document)  \n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    
    table = document.add_table(rows=9, cols=2, style='Table Grid')
    table.cell(0,0).text=("Nom / Prénoms")
    table.cell(1,0).text=("Qualification, Spécialité ")
    table.cell(2,0).text=("Adresse professionnelle")
    table.cell(3,0).text=("Nom de l’établissement")
    table.cell(4,0).text=("Service ")
    table.cell(5,0).text=("Adresse")
    table.cell(6,0).text=("Numéro de téléphone ")
    table.cell(7,0).text=("Numéro de télécopie ")
    table.cell(8,0).text=("Courriel")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold=True
                        fontdebut.size = docx.shared.Pt(11)
                    n=n+1
    
    document.add_page_break()
    
    '''Partie G3'''
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("G3. Plateau technique utilisé au cours de l'essai\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    fontdebut.bold=True
    sentence=paragraph.add_run("Y-a-t-il un laboratoire ou autre plateau technique où sont effectuées de façon centralisée les                 □ oui   □ non\n"
                               "mesures ou évaluations des paramètres ou critères principaux étudiés dans l’essai ?\n\n"
                               "Si oui, compléter pour chaque organisme ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("(répéter la section si nécessaire)")
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    
    table = document.add_table(rows=7, cols=2, style='Table Grid')
    table.cell(0,0).text=("Organisme")
    table.cell(1,0).text=("Nom de la personne à contacter")
    table.cell(2,0).text=("Adresse")
    table.cell(3,0).text=("Numéro de téléphone")
    table.cell(4,0).text=("Numéro de télécopie")
    table.cell(5,0).text=("Courriel")
    table.cell(6,0).text=("Tâches confiées")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold=True
                    n=n+1
    
    '''Partie G4'''
    paragraph=document.add_paragraph()
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("G4. Personne physique ou morale à qui le promoteur a confié certaines tâches et fonctions afférentes à l’essai\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    fontdebut.bold=True
    sentence=paragraph.add_run("Le promoteur a-t-il confié en partie ou en totalité des  tâches et des fonctions lui incombant au                □ oui   □ non\n"
                               "titre de l’essai à un tiers \n\n"
                               "Si oui, compléter pour chaque organisme ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("(répéter la section si nécessaire)")
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    
    table = document.add_table(rows=6, cols=2, style='Table Grid')
    table.cell(0,0).text=("Organisme")
    table.cell(1,0).text=("Nom de la personne à contacter")
    table.cell(2,0).text=("Adresse")
    table.cell(3,0).text=("Numéro de téléphone")
    table.cell(4,0).text=("Numéro de télécopie")
    table.cell(5,0).text=("Courriel")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold=True
                    n=n+1
    table = document.add_table(rows=8, cols=8, style='Table Grid')
    for i in range (0,8):
        a=table.cell(i,0)
        b=table.cell(i,5)
        a.merge(b)
    table.cell(0,0).text=("Tâches / fonctions confiées")
    table.cell(0,6).text=("oui")
    table.cell(0,7).text=("non")
    table.cell(1,0).text=("Ensemble des tâches du promoteur……………………………………………………")
    table.cell(2,0).text=("Monitoring…………………………………………………………………………………")
    table.cell(3,0).text=("Réglementaire (ex : préparation des dossiers soumis à l'ANSM et/ou au CPP)…..")
    table.cell(4,0).text=("Gestion/collecte des données…………………………………………………………...")
    table.cell(5,0).text=("Déclaration de vigilance (EIG, faits nouveaux, mesures urgentes de sécurité)……")
    table.cell(6,0).text=("Autres fonctions confiées……………………………………………………………….")
    table.cell(7,0).text=("Si oui, veuillez préciser : ")
    for i in range (1,7):
        table.cell(i,6).text=("□")
        table.cell(i,7).text=("□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==5 or n==13 or n==21 or n==29 or n==37 or n==45 or n==53 or n==61:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold=True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(7,5)
    a.merge(b)
    
    
    document.add_page_break()
    
    '''Partie G5'''
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("G5. Informations relatives à la vigilance\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    fontdebut.bold=True
    sentence=paragraph.add_run("1)	Localisation des informations de référence sur la sécurité (IRS) pour la qualification du caractère attendu/inattendu des effets indésirables graves")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.bold=True
    table = document.add_table(rows=4, cols=4, style='Table Grid')
    table.cell(0,0).text=("Dispositif concerné")
    a=table.cell(0,1)
    b=table.cell(0,3)
    a.merge(b)
    table.cell(0,1).text=("Localisation des IRS dans le protocole, la brochure pour l’investigateur (BI) ou la notice ")
    for i in range (1,4):
        table.cell(i,1).text=("□ Protocole")
    for i in range (1,4):
        table.cell(i,2).text=("□ BI")
    for i in range (1,4):
        table.cell(i,3).text=("□ Notice")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==0 or n==3:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    n=n+1
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("2)	Modalités de déclaration des données de vigilance (entre promoteur et ANSM\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.bold=True
    table = document.add_table(rows=3, cols=4, style='Table Grid')
    a=table.cell(0,0)
    b=table.cell(0,3)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(1,2)
    a.merge(b)
    a=table.cell(2,0)
    b=table.cell(2,2)
    a.merge(b)
    table.cell(0,0).text=("Les modalités de déclaration des données de vigilance sont-elles conformes :")
    table.cell(1,0).text=("- à la réglementation nationale ?")
    table.cell(2,0).text=("- à la phase pilote européenne ? (cf guide MEDDEV 2.7.3)")
    table.cell(1,3).text=("□ oui   □ non")
    table.cell(2,3).text=("□ oui   □ non")
    
#parties H et I
def parties_H_I(document):
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\n\nH. Information sur le Comité de Protection des Personnes (CPP) / l’Autorité compétente\n\nH.1. Informations sur le CPP concerné ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    fontdebut.bold=True
    sentence=paragraph.add_run("(A compléter si la demande est adressée à l’ANSM) :\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    
    table = document.add_table(rows=5, cols=4, style='Table Grid')
    for i in range (0,3):
        a=table.cell(i,1)
        b=table.cell(i,3)
        a.merge(b)
    table.cell(0,0).text=("Nom du CPP")
    table.cell(1,0).text=("Adresse du CPP")
    table.cell(2,0).text=("Date de soumission")
    table.cell(3,0).text=("Avis du CPP")
    table.cell(4,0).text=("Si avis donné, préciser")
    table.cell(2,1).text=("     /     /     ")
    table.cell(3,1).text=("A demander	    □")
    table.cell(3,2).text=("En cours         □")
    table.cell(3,3).text=("Donné            □")
    table.cell(4,1).text=("Date de l’avis   	     /     /     ")
    table.cell(4,2).text=("Avis favorable   □")
    table.cell(4,3).text=("Avis défavorable □")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold=True
                    n=n+1
    
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nSi un avis défavorable a été rendu, indiquer :")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    table = document.add_table(rows=4, cols=2, style='Table Grid')
    table.cell(0,0).text=("Les raisons")
    table.cell(1,0).text=("(une copie du courrier doit être jointe au dossier)")
    table.cell(2,0).text=("un second examen a-t-il été demandé à un autre CPP ?")
    table.cell(2,1).text=("□ oui   □ non")
    table.cell(3,0).text=("si oui, date prévue de dépôt")
    table.cell(3,1).text=("     /     /     ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==1:
                        fontdebut.italic=True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,0)
    a.merge(b)
    a=table.cell(0,1)
    b=table.cell(1,1)
    a.merge(b)
    
    '''Partie H2'''
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\n\nH.2. Informations sur l’Autorité compétente concernée ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    fontdebut.bold=True
    sentence=paragraph.add_run("(A compléter si la demande est adressée à un CPP :\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(11)
    
    table = document.add_table(rows=3, cols=4, style='Table Grid')
    table.cell(0,0).text=("Date de soumission")
    table.cell(1,0).text=("Décision de l'ANSM")
    table.cell(2,0).text=("Si avis donné, préciser")
    table.cell(0,1).text=("     /     /     ")
    table.cell(1,1).text=("A demander	    □")
    table.cell(1,2).text=("En cours         □")
    table.cell(1,3).text=("Donné            □")
    table.cell(2,1).text=("Date de l’avis   	     /     /     ")
    table.cell(2,2).text=("Avis favorable   □")
    table.cell(2,3).text=("Avis défavorable □")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)

    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nSi une décision de refus a été rendue, indiquer les raisons : \n ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("(une copie du courrier doit être jointe au dossier) ")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.italic=True
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=(" \n\n\n")
    
    
    '''Partie I'''
    paragraph=document.add_paragraph("\n\n I. Engagement du demandeur\n", style='debut_page')
    table = document.add_table(rows=4, cols=2, style='Table Grid')
    a=table.cell(0,0)
    b=table.cell(0,1)
    a.merge(b)
    table.cell(0,0).text=("Par la présente, j’atteste / j’atteste au nom du promoteur (rayer la mention inutile) ce qui suit :\n"
                          "-	les informations fournies ci-dessus à l’appui de la demande sont exactes ;\n"
                          "-	la recherche sera réalisée conformément au protocole et à la réglementation nationale ;\n"
                          "-	il est raisonnable d’entreprendre la recherche proposée ;\n"
                          "-	je déclarerai la date effective du commencement de la recherche à l'ANSM et au CPP concerné dès qu’elle sera connue.\n")
    table.cell(1,0).text=("Demandeur auprès de l'ANSM")
    table.cell(1,1).text=("Demandeur auprès du CPP")
    table.cell(2,0).text=("(comme indiqué à la section C.1) :	□")
    table.cell(2,1).text=("(comme indiqué à la section C.2) :	□")
    table.cell(3,0).text=("\nDate :      /     /     \n\nNom :      \n")
    table.cell(3,1).text=("Signature :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==3 or n==4 or n==5 or n==2:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==2 or n==3:
                        fontdebut.bold=True
                    n=n+1
    a=table.cell(1,0)
    b=table.cell(2,0)
    a.merge(b)
    a=table.cell(1,1)
    b=table.cell(2,1)
    a.merge(b)
    