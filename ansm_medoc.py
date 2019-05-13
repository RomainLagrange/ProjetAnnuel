# -*- coding: utf-8 -*-
"""
Created on Mon Mar  4 22:56:32 2019

@author: Marion
"""


#import pandas as pd
import docx
from docx.api import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
import time
from time import gmtime, strftime
import os
from os import sys

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)


def main_ansm_medoc(extract):
     document = docx.Document()
     partie_A_B(document, extract)
     partie_C(document, extract)
     Partie_D(document, extract)
     Partie_E(document, extract)
     Partie_F_G(document, extract)
     Partie_H_I(document, extract)
     date = (strftime('%d-%m-%Y',time.localtime()))
     document.save("soumission_ansm_medicament_"+extract['titre_abrege']+"_"+date+".docx")

def partie_A_B(document, extract):
    
    '''Marge de la page'''
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        
    '''Introduction'''
    table = document.add_table(rows=1, cols=6, style='Table Grid')
    cell=table.cell(0,0)
    paragraph = cell.paragraphs[0]
    ca = paragraph.add_run()
    ca.add_picture(resource_path('ansm.jpg'))
    table.cell(0,1).text=("Formulaire de demande d’autorisation auprès de l’ANSM et de demande d’avis à un Comité de protection des personnes d'une recherche mentionnée au 1° de l’article L. 1121-1 du code de la santé publique portant sur un médicament à usage humain ")
    table.cell(0,5).text=("FAEC")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.bold = True
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(11)
    a=table.cell(0,1)
    b=table.cell(0,4)
    a.merge(b)
    
    styles= document.styles
    style=styles.add_style('debut_page', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(10) 
    
    paragraph=document.add_paragraph("\nPARTIE A COMPLETER PAR L’ANSM / LE COMITE DE PROTECTION DES PERSONNES (CPP)\n", style='debut_page')
    
    table = document.add_table(rows=7, cols=3, style='Table Grid')
    table.cell(0,0).text=("Date de réception de la demande :     \n"
                          "Date de demande d’information pour validation :      ")
    table.cell(0,1).text=("Date de demande d’informations complémentaires :      ")
    table.cell(0,2).text=("Refus d’autorisation / avis défavorable :    □\nDate :      ")
    table.cell(5,0).text=("Date d’enregistrement du dossier complet :      \nDate du début d'évaluation :      ")
    table.cell(5,1).text=("Date de réception des informations complémentaires / amendées :      ")
    table.cell(5,2).text=("Autorisation / avis favorable :      □\nDate :      ")
    table.cell(6,0).text=("Référence attribuée par l'ANSM :      \nRéférence attribuée par le CPP :      ")
    table.cell(6,1).text=("Retrait de la demande    □\nDate :      ")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
    a=table.cell(6,1)
    b=table.cell(6,2)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nPARTIE A COMPLETER PAR LE DEMANDEUR\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("Ce formulaire est destiné à la fois à la demande d’autorisation auprès de l’ANSM et à la demande d’avis au comité de protection des personnes (CPP). Veuillez cocher ci-après la case correspondant à l’objet de la demande.\n\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    sentence=paragraph.add_run("\nDEMANDE D’AUTORISATION À L’ANSM : 	                                                                         □\n"
                               "DEMANDE D’AVIS AU CPP	                                                                                                   □\n"
                               "\nA. IDENTIFICATION DE L’ESSAI CLINIQUE")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("A.1 Etat membre dans lequel la demande est soumise : FRANCE\n"
                          "A.2 Numéro EudraCT  :  "+extract['num_eudract']+"    \n"
                          "A.3 Titre complet de l’essai clinique :  "+extract['titre_complet']+"    \n"
                          "A.4 Numéro de code du protocole de l’essai attribué par le promoteur, version et date  :   "+extract['code_protocole']+"   \n"
                          "A.5 Nom ou titre abrégé de l’essai, le cas échéant :  "+extract['titre_abrege']+"    \n"
                          "A.6 Numérotation ISRCTN , le cas échéant :      \n"
                          "A.7 S'agit-il d'une resoumission de la demande ?  □ oui  □ non  Si oui, indiquer la lettre de resoumission  :      ")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
    
    '''Partie B'''
    paragraph=document.add_paragraph("\nB. IDENTIFICATION DU PROMOTEUR RESPONSABLE DE LA DEMANDE", style='debut_page')
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("B.1	Promoteur ")
    table.cell(1,0).text=("B.1.1	Organisme :      \n"
                          "B.1.2	Nom de la personne à contacter :     \n" 
                          "B.1.3	Adresse :     \n" 
                          "B.1.4	Numéro de téléphone :      \n"
                          "B.1.5	Numéro de télécopie :    \n"  
                          "B.1.6	Mél :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("B.2 Représentant légal  du promoteur dans l’Union européenne pour l’essai concerné")
    table.cell(1,0).text=("B.2.1	Organisme :      \n"
                          "B.2.2	Nom de la personne à contacter :    \n"  
                          "B.2.3	Adresse :      \n"
                          "B.2.4	Numéro de téléphone :  \n"    
                          "B.2.5	Numéro de télécopie :   \n"   
                          "B.2.6	Mail :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("B.3 Statut du promoteur")
    table.cell(1,0).text=("B.3.1    Privé (commercial)                                                                      	□\n"
                          "B.3.2    Institutionnel  (non commercial)                                              	□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
    
def partie_C(document, extract):
    
    '''Partie C'''
    paragraph=document.add_paragraph("\nC. IDENTIFICATION DU DEMANDEUR (cocher les cases appropriées)\n", style='debut_page')
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("C.1	Demande auprès de l’ANSM	                                                                      □")
    table.cell(1,0).text=("C.1.1       Promoteur                                                                                                                   □\n"
                          "C.1.2       Représentant légal du promoteur                                                                               □\n"
                          "C.1.3	Personne ou organisme délégué par le promoteur pour soumettre la demande	     □\n"
                          "C.1.4 	Préciser ci-après les informations relatives au demandeur, même si elles figurent ailleurs dans le formulaire : Si promoteur, partie B1, si représentant légal du promoteur, partie B2\n"
                          "C.1.4.1 	Organisme :  "+extract['demandeur_nom_organisme']+"    \n"
                          "C.1.4.2 	Nom de la personne à contacter :  "+extract['demandeur_nom_personne_contact']+"    \n"
                          "C.1.4.3 	Adresse :  "+extract['demandeur_UE_adresse']+"    \n"
                          "C.1.4.4 	Numéro de téléphone :  "+extract['demandeur_UE_num_telephone']+"    \n"
                          "C.1.4.5 	Numéro de télécopie :  "+extract['demandeur_UE_num_telecopie']+"    \n"
                          "C.1.4.6	Mail :  "+extract['demandeur_UE_courriel']+"    \n"
                          "C.1.5 Demande d'envoi d'une copie des données du formulaire sous format xml :\n"
                          "C.1.5.1 Souhaitez-vous recevoir une copie du fichier xml des données du formulaire sauvegardées sur la base EudraCT ?	□ oui    □ non \n"
                          "C.1.5.1.1 Si oui, indiquer les adresses mél auxquelles cette copie doit être adressée (5 adresses maximum) :      \n"
                          "C.1.5.1.2 Souhaitez-vous que cet envoi soit sécurisé  ?	□ oui    □ non\n"
                          "Si non à la question C.1.5.1.2, le fichier xml vous sera transmis par courrier électronique non sécurisé.")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("C.2	DEMANDE AUPRÈS DU CPP                                                                                      □")
    table.cell(1,0).text=("C.2.1       Promoteur                                                                                                                   □\n"
                          "C.2.2       Représentant légal du promoteur                                                                               □\n"
                          "C.2.3	Personne ou organisme délégué par le promoteur pour soumettre la demande	     □\n"
                          "C.2.4 Investigateur chargé de soumettre la demande, si applicable  :\n"
                          "•	Investigateur coordonnateur (en cas d'essai multicentrique)	           □\n"
                          "•	Investigateur principal (en cas d'essai monocentrique)	               □\n"
                          "C.2.5 	Préciser ci-après les informations relatives au demandeur, même si elles figurent ailleurs dans le formulaire : Si promoteur, partie B1, si représentant légal du promoteur, partie B2\n"
                          "C.2.5.1	Organisme :      \n"
                          "C.2.5.2	Nom de la personne à contacter :      \n"
                          "C.2.5.3	Adresse :      \n"
                          "C.2.5.4	Numéro de téléphone :      \n"
                          "C.2.5.5	Numéro de télécopie : \n"     
                          "C.2.5.6	Mél :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
def Partie_D(document, extract):
    
    '''Partie D'''
    paragraph=document.add_paragraph("\nD. DONNEES RELATIVES A CHAQUE MEDICAMENT EXPERIMENTAL", style='debut_page')
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("Les informations concernant chaque 'produit vrac' [c’est-à-dire avant toute opération pharmaceutique spécifique à l’essai (mise en insu, conditionnement et étiquetage)], doivent être indiquées dans cette section, pour chaque médicament expérimental (ME) étudié, y compris pour chaque médicament utilisé comme comparateur et pour chaque placebo, le cas échéant. Si l’essai clinique porte sur plusieurs ME, répéter cette section, en attribuant à chaque ME un numéro d’ordre à l'item D.1.1. Si le médicament est une association, les informations doivent être données pour chaque substance active concernée.\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    
    table = document.add_table(rows=4, cols=1, style='Table Grid')
    table.cell(0,0).text=("D.1	 IDENTIFICATION DU MEDICAMENT EXPERIMENTAL")
    table.cell(1,0).text=("Indiquer ci-dessous quel ME est décrit dans cette section D. Le cas échéant, répéter cette section autant de fois qu'il y a de ME utilisé dans l’essai (numéroter chaque ME de 1 à n)")
    table.cell(2,0).text=("D.1.1         Cette section concerne le ME numéro :	     \n"
                          "D.1.2         ME étudié                                                                       	□\n"
                          "D.1.3         ME utilisé comme comparateur                                   	□")
    table.cell(3,0).text=("Pour le placebo, aller directement en section D.7")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,0)
    a.merge(b)
    a=table.cell(2,0)
    b=table.cell(3,0)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    '''Partie D2'''
    
    table = document.add_table(rows=12, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.2	   STATUT DU PRODUIT SUR LEQUEL PORTE LA RECHERCHE ")
    table.cell(1,0).text=("Si le ME dispose d'une AMM en France, mais que le nom de la spécialité et le titulaire de l’AMM ne sont pas déterminés dans le protocole de l'essai, aller à la section D.2.2.")
    table.cell(2,0).text=("D.2.1	   Le ME utilisé dans l’essai dispose-t-il d'une AMM ?")
    table.cell(3,0).text=("D.2.1.1	   Si oui en D.2.1, préciser pour le médicament utilisé dans l'essai :")
    table.cell(4,0).text=("D.2.1.1.1    Nom de spécialité  :      ")
    table.cell(5,0).text=("D.2.1.1.2    Nom du titulaire de l’AMM :      ")
    table.cell(6,0).text=("D.2.1.1.3    Numéro d’AMM (si AMM délivrée par un Etat membre) :      ")
    table.cell(7,0).text=("D.2.1.1.4    Le ME est-il modifié par rapport à son AMM ?")
    table.cell(8,0).text=("D.2.1.1.4.1 Si oui, veuillez préciser :      ")
    table.cell(9,0).text=("D.2.1.2       Quel pays a délivré l'AMM ?      ")
    table.cell(10,0).text=("D.2.1.2.1    Est-ce la France ?")
    table.cell(11,0).text=("D.2.1.2.2    Est-ce un autre Etat membre ?")
    
    for i in range(1,12):
        if i==2 or i==7 or i==10 or i==11:
            table.cell(i,5).text=("□ oui  □ non")
        else:
            table.cell(i,5).text=(" ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==4 or n==14 or n==20 or n==22:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==3:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,5)
    a.merge(b)
    a=table.cell(2,0)
    b=table.cell(11,4)
    a.merge(b)
    a=table.cell(2,5)
    b=table.cell(11,5)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=5, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.2.2	Cas où le ME utilisé dans l’essai clinique dispose d'une AMM en France, mais le protocole autorise l’utilisation de toute spécialité pour le ME, sous réserve qu’elle dispose d'une AMM en France, et il n’est donc pas possible d’identifier précisément le/les ME avant le début de l’essai")
    table.cell(1,0).text=("D.2.2.1      Dans le protocole, le traitement est-il défini uniquement par la substance active ?\n"
                          "D.2.2.1.1 Si oui, indiquer le nom de la substance active en D.3.8 ou D.3.9")
    table.cell(2,0).text=("D.2.2.2      Dans le protocole, les schémas de traitement permettent-ils différentes combinaisons de médicaments commercialisés, utilisés selon les pratiques cliniques locales dans certains ou dans tous les lieux de recherche en France ?\n"
                          "D.2.2.2.1 Si oui, indiquer le nom de la substance active en D.3.8 ou D.3.9")
    table.cell(3,0).text=("D.2.2.3      Les produits à administrer en tant que ME sont-ils définis comme appartenant à un groupe ATC ?\n"
                          "D.2.2.3.1 Si oui, indiquer ce groupe ATC dans le champ des codes ATC (niveau 3 ou plus jusqu’au niveau pouvant être défini) de la section D.3.3")
    table.cell(4,0).text=("D.2.2.4      Autre :\n"
                          "D.2.2.4.1 Si oui, veuillez préciser :      ")
    table.cell(0,5).text=(" ")
    table.cell(1,5).text=("□ oui  □ non")
    table.cell(2,5).text=("□ oui  □ non")
    table.cell(3,5).text=("□ oui  □ non")
    table.cell(4,5).text=("□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==3 or n==5 or n==7 or n==9:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    for i in range (1,5) :
        a=table.cell(i,0)
        b=table.cell(i,4)
        a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=4, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.2.3	Dossier du médicament expérimental soumis (DME)")
    table.cell(1,0).text=("D.2.3.1	DME complet")
    table.cell(2,0).text=("D.2.3.2	DME simplifié")
    table.cell(3,0).text=("D.2.3.3	Résumé des caractéristiques du produit (RCP) uniquement")
    table.cell(1,5).text=("□ oui  □ non")
    table.cell(2,5).text=("□ oui  □ non")
    table.cell(3,5).text=("□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2 or n==4 or n==6:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(3,4)
    a.merge(b)
    a=table.cell(0,5)
    b=table.cell(3,5)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.2.4	L’utilisation du ME a-t-elle déjà été autorisée dans le cadre d'un essai clinique précédent conduit par le promoteur dans la Communauté européenne ?")
    table.cell(1,0).text=("D.2.4.1	Si oui, préciser dans quel(s) Etat(s) membre(s) :      ")
    table.cell(0,5).text=("\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==1:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,4)
    a.merge(b)
    a=table.cell(0,5)
    b=table.cell(1,5)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.2.5    Le ME est-il désigné, dans l’indication étudiée dans l'essai, comme un médicament orphelin dans la Communauté européenne ?")
    table.cell(1,0).text=("D.2.5.1	Si oui, indiquer le numéro de désignation du médicament orphelin  :      ")
    table.cell(0,5).text=("\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==1:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,4)
    a.merge(b)
    a=table.cell(0,5)
    b=table.cell(1,5)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=3, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.2.6    Un avis scientifique a-t-il été rendu sur le ME dans le cadre de cet essai clinique ?")
    table.cell(1,0).text=("D.2.6.1	Si oui en D.2.6, veuillez préciser qui a rendu l'avis et en joindre une copie à votre dossier :")
    table.cell(2,0).text=("D.2.6.1.1    Avis du CHMP  ?\n"
                          "D.2.6.1.2	Avis d'une autorité compétente d'un Etat membre ?")
    table.cell(0,5).text=("□ oui  □ non")
    table.cell(2,5).text=("\n\n□ oui  □ non\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==1 or n==4:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(2,4)
    a.merge(b)
    a=table.cell(0,5)
    b=table.cell(2,5)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    '''Partie D3'''
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=3, cols=1, style='Table Grid')
    table.cell(0,0).text=("D.3	 DESCRIPTION DU MEDICAMENT EXPERIMENTAL")
    table.cell(1,0).text=("D.3.1         Nom du ME, le cas échéant  :      \n"
                          "D.3.2         Nom de code, le cas échéant  :      \n"
                          "D.3.3         Code ATC, si enregistré officiellement :      \n"
                          "D.3.4         Forme pharmaceutique (utiliser les termes standard) :      \n"
                          "D.3.5         Durée maximale du traitement pour une personne prévue par le protocole :      \n"
                          "D.3.6         Dose maximale permise (préciser : dose journalière ou dose cumulée ; unités et voie d'administration) :      \n"
                          "D.3.7         Voie d’administration (utiliser les termes standard) :      \n"
                          "D.3.8         Nom de chaque substance active (DCI ou DCI proposée, le cas échéant) :      \n"
                          "D.3.9         Autre(s) nom(s) disponible(s) pour chaque substance active (numéro CAS , code précédemment attribué par le promoteur, autre nom descriptif, etc. Indiquer tous les noms disponibles) :      \n"
                          "D.3.10        Dosage (préciser tous les dosages utilisés) : ")
    table.cell(2,0).text=("D.3.10.1      Unité de concentration :      \n"
                          "D.3.10.2      Type de concentration (“nombre exact ”, “intervalle”, “plus que” ou “jusqu’à”) :      \n"
                          "D.3.10.3      Concentration (nombre) :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(1,0)
    b=table.cell(2,0)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.3.11	Catégorie de médicament expérimental		\n"
                          "Le ME contient-il une substance active :")
    table.cell(1,0).text=("D.3.11.1   d’origine chimique ?\n"
                          "D.3.11.2	d’origine biologique / biotechnologique  ?	\n"
                          "Est-ce :\n"
                          "D.3.11.3    un médicament de thérapie cellulaire ?\n"
                          "D.3.11.4    un médicament de thérapie génique ?\n"
                          "D.3.11.5    un médicament radiopharmaceutique ?\n"
                          "D.3.11.6    un médicament immunologique (notamment vaccin, allergène, immun-sérum) ?\n"
                          "D.3.11.7    un médicament dérivé du sang ?\n"
                          "D.3.11.8    un autre médicament d’origine extractive ?\n"
                          "D.3.11.9    un médicament à base de plantes ?\n"
                          "D.3.11.10   un médicament homéopathique ?\n"
                          "D.3.11.11   un médicament contenant des organismes génétiquement modifiés ?	\n"
                          "Si oui en D.3.11.11   \n"
                          "D.3.11.11.1  L’autorisation relative au confinement de l’OGM a-t-elle été accordée ?\n"
                          "D.3.11.11.2  Est-elle en attente ?\n"
                          "D.3.11.12	un autre type de médicament ?	\n"
                          "D.3.11.12.1	Si oui, préciser :      ")
    table.cell(1,5).text=("\n\n□ oui  □ non\n□ oui  □ non\n\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==3:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,4)
    a.merge(b)
    a=table.cell(0,5)
    b=table.cell(1,5)
    a.merge(b)
    document.add_page_break()
    
    '''Partie D4'''
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=3, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.4	   MEDICAMENT EXPERIMENTAL D’ORIGINE BIOLOGIQUE / BIOTECHNOLOGIQUE, Y COMPRIS LES VACCINS")
    table.cell(1,0).text=("D.4.1        Type de médicament")
    table.cell(2,0).text=("D.4.1.1      Produit d'origine extractive\n"
                          "D.4.1.2      Produit recombinant\n"
                          "D.4.1.3      Vaccin\n"
                          "D.4.1.4      Organisme génétiquement modifié\n"
                          "D.4.1.5      Médicament dérivé du sang\n"
                          "D.4.1.6	    Autre	\n"
                          "D.4.1.6.1 Si oui, préciser :      ")
    table.cell(2,5).text=("\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==3:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(2,4)
    a.merge(b)
    a=table.cell(1,5)
    b=table.cell(2,5)
    a.merge(b)
        
    
    '''Partie D5'''
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=5, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.5	   MEDICAMENT EXPERIMENTAL DE THERAPIE CELLULAIRE SOMATIQUE (SANS MODIFICATION GENETIQUE)")
    table.cell(1,0).text=("D.5.1        Origine du tissu, du tissu composite ou de l’organe")
    table.cell(2,0).text=("D.5.1.1     Autologue\n"
                          "D.5.1.2     Allogénique\n"
                          "D.5.1.3	    Xénogénique	\n"
                          "D.5.1.3.1 Si oui, préciser les espèces d’origine :      ")
    table.cell(3,0).text=("D.5.2	  Type de cellules")
    table.cell(4,0).text=("D.5.2.1	  Cellules souches\n"
                          "D.5.2.2	  Cellules différenciées	\n"
                          "D.5.2.2.1 Si oui, préciser le type de cellules (exemple : kératinocytes, fibroblastes, chondrocytes…) :      \n"
                          "D.5.2.3	   Autre	\n"
                          "D.5.2.3.1 Si oui, préciser :      ")
    table.cell(1,5).text=("\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n")
    table.cell(3,5).text=("\n□ oui  □ non\n\n□ oui  □ non\n\n□ oui  □ non\n")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2 or n==5:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1 or n==4:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(2,4)
    a.merge(b)
    a=table.cell(1,5)
    b=table.cell(2,5)
    a.merge(b)
    a=table.cell(3,0)
    b=table.cell(4,4)
    a.merge(b)
    a=table.cell(3,5)
    b=table.cell(4,5)
    a.merge(b)
    
    '''Partie D6'''
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=3, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.6	     MEDICAMENT EXPERIMENTAL DE THERAPIE GENIQUE")
    table.cell(1,0).text=("D.6.1	     Gène(s) d’intérêt :      \n"
                          "D.6.2         Thérapie génique in vivo\n"
                          "D.6.3         Thérapie génique ex vivo\n"
                          "D.6.4         Type de vecteur utilisé")
    table.cell(1,5).text=("\n□ oui  □ non\n□ oui  □ non\n\n□ oui  □ non")
    table.cell(2,0).text=("D.6.4.1	     Acide nucléique (exemple : plasmide)	\n"
                          "Si oui, préciser s’il s’agit :\n"
                          "D.6.4.1.1   d’un acide nucléique nu\n"
                          "D.6.4.1.2   d'un acide nucléique complexe	\n"
                          "D.6.4.2	    Vecteur viral	\n"
                          "D.6.4.2.1   Si oui, préciser le type : adénovirus, rétrovirus, AAV…:      \n"
                          "D.6.4.3	    Autre	\n"
                          "D.6.4.3.1   Si oui, préciser :      ")
    table.cell(2,5).text=("\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2 or n==4:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(2,4)
    a.merge(b)
    a=table.cell(1,5)
    b=table.cell(2,5)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.6.5	    Cellules génétiquement modifiées")
    table.cell(1,0).text=("Si oui, préciser l’origine des cellules\n"
                          "D.6.5.1      Autologue\n"
                          "D.6.5.2      Allogénique\n"
                          "D.6.5.3	    Xénogénique	\n"
                          "D.6.5.3.1   Si oui, préciser les espèces d’origine :      \n"
                          "D.6.5.4	    Autre type de cellules (cellules souches hématopoïétiques, …)	\n"
                          "            Si oui, préciser :      ")
    table.cell(1,5).text=("□ oui  □ non\n\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,4)
    a.merge(b)
    a=table.cell(0,5)
    b=table.cell(1,5)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("D.6.6	Remarques relatives à de nouveaux aspects concernant le ME de thérapie génique (texte libre) :      ")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    fontdebut.bold = True

    '''Partie D7'''
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=3, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.7    DONNEES RELATIVES AU PLACEBO (répéter la section autant de fois que nécessaire, le cas échéant)")
    table.cell(1,0).text=("D.7.1	    Un placebo est-il utilisé ?\n"
                          "D.7.2	    Cette section concerne le placebo numéro : (     )\n"
                          "D.7.3	    Forme pharmaceutique :      \n"
                          "D.7.4	    Voie d’administration :      \n"
                          "D.7.5	    De quel ME est-ce le placebo ? Préciser le numéro du ME, tel qu'indiqué en D.1 : (     )")
    table.cell(2,0).text=("D.7.5.1	    Composition, hormis la ou les substances actives :      \n"
                          "D.7.5.2	     Est-elle identique à celle du ME étudié ?	\n"
                          "D.7.5.2.1  Si non, préciser les principaux composants :      \n")
    table.cell(1,5).text=("□ oui  □ non\n")
    table.cell(2,5).text=("\n\n\n\n\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2 or n==4:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(2,4)
    a.merge(b)
    a=table.cell(1,5)
    b=table.cell(2,5)
    a.merge(b)
    
    '''Partie D8'''
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=8, cols=1, style='Table Grid')
    table.cell(0,0).text=("D.8	ETABLISSEMENT OÙ LA PERSONNE QUALIFIEE LIBERE LES LOTS DE MEDICAMENT EXPERIMENTAL")
    table.cell(1,0).text=("Cette section concerne les ME finis, c’est-à-dire les médicaments (randomisés) conditionnés, étiquetés et libérés spécifiquement pour l’essai clinique. S'il y a plusieurs établissements en charge de la libération ou plusieurs ME à libérer, répéter cette section autant de fois que nécessaire et préciser le numéro du ME concerné, tel qu'indiqué en D.1 ou D.7.2. En cas de pluralité d'établissements libérateurs, préciser le ME libéré par chaque établissement concerné.")
    table.cell(2,0).text=("D.8.1	Ne pas remplir la section D.8.2 si le ME (conditions cumulatives) :\n"
                          "         -	bénéficie d'une AMM dans l'Union européenne et\n"
                          "         -	provient du marché de l'Union européenne et\n"
                          "         -	est utilisé sans modification dans le cadre de l'essai (exemple : non mis en gélule) et\n"
                          "         -	le conditionnement et l'étiquetage sont effectués dans des établissements de santé, pour leur usage exclusif, comme prévu à l'article 9.2 de la directive 2005/28/CE relative aux bonnes pratiques cliniques.\n"
                          "         Si l'ensemble de ces conditions sont réunies, cocher la case ci-contre □ et indiquer le numéro de chaque ME concerné, y compris de chaque placebo, tel qu'indiqué en D.1.1 et D.7.2 : (     )")
    table.cell(3,0).text=("D.8.2	Qui est responsable au sein de l’Union européenne de la libération du ME fini ?")
    table.cell(4,0).text=("         L'établissement est responsable de la libération de (préciser le numéro de chaque ME concerné, y compris de chaque placebo, tel qu'indiqué en D.1.1 et D.7.2) :      \n")
    table.cell(5,0).text=("         Veuillez cocher la case appropriée :\n")
    table.cell(6,0).text=("D.8.2.1  Fabricant		□\n" 
                          "D.8.2.2  Importateur		□\n"
                          "D.8.2.3  Nom de l’établissement :     \n"
                          "D.8.2.3.1 Adresse :      \n"
                          "D.8.2.4   Indiquer le numéro d’autorisation du fabricant :      \n"
                          "D.8.2.4.1	Si pas d’autorisation, préciser les motifs :      \n")
    table.cell(7,0).text=("Si le médicament ne bénéficie pas d'une AMM dans l'Union européenne, mais qu'il est fourni en vrac et que le conditionnement et l'étiquetage sont effectués par un établissement de santé, pour son usage exclusif, conformément aux dispositions de l'article 9.2 de la directive 2005/28/CE relative aux bonnes pratiques cliniques, indiquer l'établissement où le médicament à été certifié en vue de sa libération par la personne qualifiée pour son utilisation dans l'essai clinique en D.8.2.")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==3 or n==5:
                        fontdebut.bold = True
                    elif n==2 or n==7:
                        fontdebut.italic = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,0)
    a.merge(b)
    a=table.cell(3,0)
    b=table.cell(7,0)
    a.merge(b)
    
def Partie_E(document, extract):
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nE. INFORMATIONS GENERALES RELATIVES A L'ESSAI")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(10) 
    sentence=paragraph.add_run("\nCette section est destinée à fournir des informations concernant les objectifs, domaine et méthodologie de l'essai. Si le protocole prévoit la réalisation d'une sous-étude en France, indiquer les informations relatives à cette sous-étude en section E.2.3. Veuillez également cocher la case appropriée en section E.2 relative à l'objectif de l'essai.\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10) 
    
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("E.1	CONDITION MEDICALE OU PATHOLOGIE ETUDIEE")
    table.cell(1,0).text=("E.1.1	Préciser la ou les conditions médicales étudiées  (texte libre) :      \n"
                          "E.1.2	Version MedDRA, niveau, terme et classification  (répéter autant de fois que nécessaire) :      \n"  
                          "E.1.3	L'une des conditions médicales étudiées est-elle une maladie rare  ?	            □ oui   □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie E2'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("E.2	OBJECTIF(S) DE LA RECHERCHE")
    table.cell(1,0).text=("E.2.1	Objectif principal : "+extract['objectif_principal']+"\n"
                          "E.2.2	Objectifs secondaires : "+extract['objectif_secondaire']+"\n"
                          "E.2.3	Une sous-étude est-elle prévue ?	                                                        □ oui   □ non\n"
                          "E.2.3.1	Si oui, préciser le titre complet, la date et la version de chaque sous-étude et leurs objectifs :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
    '''Partie E3'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("E.3	PRINCIPAUX CRITERES D’INCLUSION (énumérer les plus importants)")
    table.cell(1,0).text=(extract['critere_inclusion_courte'])
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
    '''Partie E4'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("E.4	PRINCIPAUX CRITERES DE NON INCLUSION (énumérer les plus importants)")
    table.cell(1,0).text=(extract['critere_non_inclusion_courte'])
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
    '''Partie E5'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("E.5	CRITERE(S) D’EVALUATION PRINCIPAL(AUX)")
    table.cell(1,0).text=(extract['critere_jugement_principal_longue']+"\n"+extract['critere_jugement_secondaire_longue'])
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie E6'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=3, cols=3, style='Table Grid')
    table.cell(0,0).text=("E.6	  DOMAINE(S) D’ETUDE – Cocher la ou les cases appropriées")
    table.cell(1,0).text=("E.6.1     Diagnostic\n"
                          "E.6.2     Prophylaxie\n"
                          "E.6.3     Thérapeutique\n"
                          "E.6.4     Sécurité\n"
                          "E.6.5     Efficacité\n"
                          "E.6.6     Pharmacocinétique\n"
                          "E.6.7     Pharmacodynamie\n"
                          "E.6.8     Bioéquivalence\n"
                          "E.6.9     Dose-effet\n"
                          "E.6.10    Pharmacogénétique\n"
                          "E.6.11    Pharmacogénomie\n"
                          "E.6.12	 Pharmaco-économie\n"
                          "E.6.13    Autre")
    table.cell(2,0).text=("E.6.13.1	Si autre, préciser :      ")
    table.cell(1,2).text=("□ \n□ \n□ \n□ \n□ \n□ \n□ \n□ \n□ \n□ \n□\n□\n□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,2)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(2,1)
    a.merge(b)
                    
    '''Partie E7'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=4, cols=3, style='Table Grid')
    table.cell(0,0).text=("E.7	  TYPE D'ESSAI  ET PHASE")
    table.cell(1,0).text=("E.7.1	  Pharmacologie humaine (Phase I)")
    table.cell(2,0).text=("Il s'agit de :\n"
                          "E.7.1.1    La première administration à l’homme\n"
                          "E.7.1.2    Une étude de bioéquivalence\n"
                          "E.7.1.3	  Autre\n"
                          "E.7.1.3.1 Si autre, préciser :      ")
    table.cell(3,0).text=("E.7.2      Essai thérapeutique exploratoire (Phase II)\n"
                          "E.7.3      Essai thérapeutique de confirmation (Phase III)\n"
                          "E.7.4	  Essai thérapeutique conformément à l’autorisation (Phase IV)")
    table.cell(1,2).text=("□ \n\n□ \n□ \n□ \n\n□ \n□ \n□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1 or n==4:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,2)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(3,1)
    a.merge(b)
    a=table.cell(1,2)
    b=table.cell(3,2)
    a.merge(b)
    
    '''Partie E8'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=12, cols=6, style='Table Grid')
    table.cell(0,0).text=("E.8	METHODOLOGIE DE L'ESSAI")
    table.cell(1,0).text=("E.8.1	  Comparatif	")
    table.cell(2,0).text=("Si oui, préciser :\n"
                          "E.8.1.1    Tirage au sort\n"
                          "E.8.1.2    Ouvert\n"
                          "E.8.1.3    Simple insu\n"
                          "E.8.1.4    Double insu\n"
                          "E.8.1.5    A groupes parallèles\n"
                          "E.8.1.6    Plan croisé\n"
                          "E.8.1.7	 Autre	\n"
                          "E.8.1.7.1 Si autre, préciser :      ")
    table.cell(3,0).text=("E.8.2	  Si comparatif, préciser le comparateur utilisé")
    table.cell(4,0).text=("E.8.2.1    Autre(s) médicament(s)\n"
                          "E.8.2.2    Placebo\n"
                          "E.8.2.4	  Autre\n"
                          "E.8.2.4.1 Si autre, préciser :      ")
    table.cell(1,5).text=("□ oui   □ non\n\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n\n\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non")
    table.cell(5,0).text=("E.8.3      L'essai est-il monocentrique (voir aussi section G) ?\n"
                          "E.8.4	  L'essai est-il multicentrique (voir aussi section G) ?")
    table.cell(6,0).text=("E.8.4.1	  Nombre prévu de lieux de recherche en France :    ")  
    table.cell(7,0).text=("E.8.5	  Est-il prévu de mener l'essai dans plusieurs états membres ?")
    table.cell(8,0).text=("E.8.5.1	  Nombre prévu de lieux de recherche dans la Communauté européenne :  ")
    table.cell(9,0).text=("E.8.6      Est-il prévu de mener la recherche dans des pays tiers ?\n"
                          "E.8.7	  Un comité de surveillance indépendant a-t-il été constitué ?")
    table.cell(5,5).text=("□ oui   □ non\n□ oui   □ non\n\n□ oui   □ non\n\n□ oui   □ non\n□ oui   □ non")
    table.cell(10,0).text=("E.8.8	  Définition de la fin de l'essai, et justification si celle-ci ne correspond pas à la date de la dernière visite de la dernière personne participant à l'essai   :    \n"  
                          "E.8.9      Estimation initiale de la durée de l'essai  (en années, mois et jours) : "+extract['duree_totale_etude']+"\n")
    table.cell(11,0).text=("E.8.9.1   en France : 	      années       mois       jours\n"
                           "E.8.9.2	 dans tous les pays concernés par l’essai : 	      années       mois       jours")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2 or n==7:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1 or n==4 or n==6 or n==9 or n==11 or n==12:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(4,4)
    a.merge(b)
    a=table.cell(1,5)
    b=table.cell(4,5)
    a.merge(b)
    a=table.cell(5,0)
    b=table.cell(9,4)
    a.merge(b)
    a=table.cell(5,5)
    b=table.cell(9,5)
    a.merge(b)
    a=table.cell(10,0)
    b=table.cell(11,5)
    a.merge(b)
    
def Partie_F_G(document, extract):
    
    '''Partie F'''
    paragraph=document.add_paragraph("\nF. PERSONNES PARTICIPANT A L'ESSAI\n", style='debut_page')
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("F.1	Tranche d'âge étudiée")
    table.cell(1,0).text=("F.1.1	Moins de 18 ans\n"
                          "Si oui, préciser :\n"
                          "F.1.1.1  In Utero\n"
                          "F.1.1.2  Nouveaux-nés prématurés (jusqu’à l’âge gestationnel ≤ 37 semaines)\n"
                          "F.1.1.3  Nouveau-nés (0-27 jours)\n"
                          "F.1.1.4  Nourrissons (28 jours - 23 mois)\n"
                          "F.1.1.5  Enfants (2-11 ans)\n"
                          "F.1.1.6  Adolescents (12-17 ans)\n"
                          "F.1.2     De 18 à 65 ans\n"
                          "F.1.3	Plus de 65 ans")
    table.cell(1,5).text=("□ oui   □ non\n\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(1,4)
    a.merge(b)
    
    '''Partie F2'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("F.2	SEXE")
    table.cell(1,0).text=("F.2.1     Femmes    □\n"
                          "F.2.2	Hommes    □")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie F3'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=3, cols=6, style='Table Grid')
    table.cell(0,0).text=("F.3	GROUPE DE PERSONNES PARTICIPANT A L'ESSAI")
    table.cell(1,0).text=("F.3.1         Volontaires sains\n"
                          "F.3.2         Volontaires malades\n"
                          "F.3.3         Populations particulières")
    table.cell(2,0).text=("F.3.3.1      Femmes en âge de procréer\n"
                          "F.3.3.2      Femmes en âge de procréer utilisant un moyen de contraception\n"
                          "F.3.3.3      Femmes enceintes\n"
                          "F.3.3.4      Femmes allaitantes\n"
                          "F.3.3.5      Personnes en situation d’urgence\n"
                          "F.3.3.6	    Personnes incapables de donner personnellement leur consentement	\n"
                          "F.3.3.6.1   Si oui, préciser :   \n"   
                          "F.3.3.7	    Autre\n"
                          "F.3.3.7.1   Si oui, préciser :      ")
    table.cell(1,5).text=("□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n\n□ oui   □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(2,4)
    a.merge(b)
    
    '''Partie F4'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("F.4	NOMBRE PREVU DE PERSONNES A INCLURE")
    table.cell(1,0).text=("F.4.1	En France\n"
                          "F.4.2	En cas d'essai  mené dans plusieurs pays :\n"
                          "F.4.2.1	Dans la Communauté européenne\n"
                          "F.4.2.2	Pour l’ensemble de la recherche")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie F5'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("F.5	TRAITEMENT(S) OU SOIN(S) PREVU(S) POUR LES PERSONNES A LA FIN DE LEUR PARTICIPATION A L'ESSAI  Si cela différe du traitement habituel de la condition médicale étudiée, veuillez préciser (texte libre) :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie G'''
    paragraph=document.add_paragraph("\nG. LIEUX DE RECHERCHES ENVISAGES / INVESTIGATEURS EN FRANCE\n", style='debut_page')
    
    '''Partie G1'''
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("G.1	INVESTIGATEUR COORDONNATEUR (si essai multicentrique) et investigateur principal (si essai monocentrique) ")
    table.cell(1,0).text=("G.1.1	Prénom :  "+extract['investigateur_coordinateur_prenom']+"    \n"
                          "G.1.3	Second prénom, le cas échéant :   "+extract['investigateur_coordinateur_nom']+"   \n"
                          "G.1.4	Qualification, spécialité :   "+extract['investigateur_coordinateur_qualification']+"   \n"
                          "G.1.5	Adresse professionnelle :  "+extract['investigateur_coordinateur_adresse_professionnelle']+"    ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie G2'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("G.2	INVESTIGATEURS PRINCIPAUX (si essai multicentrique ; répéter cette section autant de fois que nécessaire) ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    for i in range(len(extract['autre_investigateur_nom'])):
        table = document.add_table(rows=1, cols=1, style='Table Grid')
        table.cell(0,0).text=("G.2.1	Prénom :  "+extract['autre_investigateur_nom'][i]+"    \n"
                              "G.2.3	Second prénom, le cas échéant :  "+extract['autre_investigateur_prenom'][i]+"   \n" 
                              "G.2.4	Qualification, spécialité :  "+extract['autre_investigateur_qualification'][i]+"    \n"
                              "G.2.5	Adresse professionnelle :  "+extract['autre_investigateur_adresse_professionnelle'][i])
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        fontdebut = run.font
                        fontdebut.name = 'Arial'
                        fontdebut.size = docx.shared.Pt(10)
    '''Partie G3'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("G.3	PLATEAU TECHNIQUE UTILISE AU COURS DE L'ESSAI\nLaboratoire ou autre plateau technique où sont effectuées de façon centralisée les mesures ou évaluations des paramètres ou critères principaux étudiés dans l'essai (à compléter pour chaque organisme, répéter la section si nécessaire)")
    table.cell(1,0).text=("G.3.1	Organisme :  "+extract['plateau_technique_organisme']+"    \n"
                          "G.3.2	Nom de la personne à contacter : "+extract['plateau_technique_personne_contact']+"     \n"
                          "G.3.3	Adresse :  "+extract['plateau_technique_adresse']+"    \n"
                          "G.3.4	Numéro de téléphone : "+extract['plateau_technique_num_telephone']+"     \n"
                          "G.3.5	Tâches confiées : "+extract['plateau_technique_taches_confiees']+"     ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie G4'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("G.4	PRESTATAIRE A QUI LE PROMOTEUR A CONFIE CERTAINES OBLIGATIONS ET FONCTIONS AFFERENTES A L'ESSAI (à compléter pour chaque organisme, répéter la section si nécessaire)")
    table.cell(1,0).text=("G.4.1	    Le promoteur a-t-il confié en partie ou en totalité des obligations et des fonctions majeures lui incombant au titre de l'essai à un autre organisme ou à un tiers ?\n"
                          "Préciser pour chaque organisme :\n"
                          "G.4.1.1	    Organisme :      \n"
                          "G.4.1.2	    Nom de la personne à contacter :    \n"  
                          "G.4.1.3	    Adresse :      \n"
                          "G.4.1.4	    Numéro de téléphone :      \n"
                          "Obligations / fonctions confiées :\n"
                          "G.4.1.5     Ensemble des tâches du promoteur\n"
                          "G.4.1.6     Monitoring\n"
                          "G.4.1.7     Réglementaire (ex : préparation des dossiers soumis à l'Afssaps et au CPP)\n"
                          "G.4.1.8     Recrutement des investigateurs\n"
                          "G.4.1.9     IVRS  - tirage au sort du traitement\n"
                          "G.4.1.10   Gestion/collecte des données\n"
                          "G.4.1.11   Saisie électronique des données\n"
                          "G.4.1.12   Déclaration des effets indésirables graves et/ou incidents graves \n"
                          "G.4.1.13   Audit de l'assurance qualité\n"
                          "G.4.1.14   Analyses statistiques\n"
                          "G.4.1.15   Rédaction médicale\n"
                          "G.4.1.16    Autres devoirs confiés\n"
                          "G.4.1.16.1  Si oui, veuillez préciser :      ")
    table.cell(1,5).text=("□ oui   □ non\n\n\n\n\n\n\n\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(1,4)
    a.merge(b)
    
def Partie_H_I(document, extract):
    
    '''Partie H'''
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nH. ANSM / CPP CONCERNE PAR LA DEMANDE")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(10) 
    
    '''Partie H1'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=3, style='Table Grid')
    table.cell(0,0).text=("H.1	  TYPE DE DEMANDE\n"
                          "Si cette demande est adressée à l’ANSM, cocher la case 'CPP' et indiquer les informations relatives au CPP concerné, et vice-versa.")
    table.cell(1,0).text=("H.1.1	   ANSM\n"
                          "H.1.2	   CPP\n")
    table.cell(1,2).text=("□\n□")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    fontdebut.bold = True
    a=table.cell(0,0)
    b=table.cell(0,2)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(1,1)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("H.2	   INFORMATIONS RELATIVES A L'ANSM / AU CPP")
    table.cell(1,0).text=("H.2.1       Nom et adresse :      \n"
                          "H.2.2	   Date de soumission : \n")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    fontdebut.bold = True
    
    '''Partie H3'''
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nSi cette demande est adressée au CPP remplir les informations ci-dessous :")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10) 
    
    table = document.add_table(rows=3, cols=3, style='Table Grid')
    table.cell(0,0).text=("H.3	AUTORISATION / AVIS")
    table.cell(1,0).text=("H.3.1    A demander\n"
                          "H.3.2	En cours\n"
                          "H.3.3	Obtenu(e)	")
    table.cell(2,0).text=("	Si obtenu(e), préciser :	\n"
                          "H.3.3.1  Date de la décision :   /  /  \n"
                          "H.3.3.2  Autorisation\n"
                          "H.3.3.3	  Refus d'autorisation	\n"
                          "	Si refus d'autorisation, préciser :\n"
                          "H.3.3.3.1	Les motifs :      \n"
                          "H.3.3.3.2	La date éventuelle envisagée de resoumission de le demande :      ")
    table.cell(1,2).text=("□\n□\n□\n\n\n□\n□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,2)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(2,1)
    a.merge(b)
    a=table.cell(1,2)
    b=table.cell(2,2)
    a.merge(b)
    
    document.add_page_break()
    
    '''Partie I'''
    
    paragraph=document.add_paragraph("\nI. SIGNATURE DU DEMANDEUR EN FRANCE", style='debut_page')
    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=1, cols=8, style='Table Grid')
    table.cell(0,0).text=("I.1")
    table.cell(0,1).text=("Par la présente, j’atteste / j’atteste au nom du promoteur (rayer la mention inutile) ce qui suit :\n"
                          "-	les informations fournies ci-dessus à l’appui de la demande sont exactes ;\n"
                          "-	l'essai sera réalisé conformément au protocole, à la réglementation nationale et aux principes de bonnes pratiques ;\n"
                          "-	il est raisonnable de mettre en œuvre la recherche proposée ; \n"
                          "-	je m'engage à déclarer les effets indésirables graves et/ou incidents  et à soumettre les rapports de sécurité, conformément à la réglementation applicable ;\n"
                          "-	je m'engage à soumettre un résumé du rapport final de l'essai à l’ANSM au plus tard 1 an après la fin de l'essai dans tous les pays.")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,1)
    b=table.cell(0,7)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("I.2	DEMANDEUR AUPRES DE L'ANSM (tel qu'indiqué en C.1)")
    table.cell(0,1).text=("I.2.1     Date :   /  /    \n"
                          "I.2.2     Signature  :   \n"
                          "I.2.3	Nom :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("I.3	DEMANDEUR AUPRES DU CPP (tel qu'indiqué en C.2)")
    table.cell(0,1).text=("I.3.1     Date :   /  /    \n"
                          "I.3.2     Signature  :   \n"
                          "I.3.3	Nom :      ")   
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    
    
    
    
    
    
    