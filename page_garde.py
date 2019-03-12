#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Dec  6 16:44:16 2018

@author: romain
"""
import docx
import extraction
from docx.enum.text import WD_ALIGN_PARAGRAPH
#from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
import re
from docx.shared import Cm

#document = docx.Document()
#extract=extraction.extract1()
#'''Marge des page'''
#sections = document.sections
#for section in sections:
#    section.top_margin = Cm(2)
#    section.bottom_margin = Cm(2)
#    section.left_margin = Cm(2)
#    section.right_margin = Cm(2)

def PageGarde(document,extract):
    
 #   document = docx.Document()
    
    sections = document.sections
    page_garde = sections[0]
        
    
    '''Logos de l'en-tete'''
    header = page_garde.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    r = p.add_run() 
    r.add_picture('imageGauche.png')
    r.add_text('                                                                                                                                     ')
    r.add_picture('imageDroite.png')
    
    
    '''Titre de la recherche'''
    paragraph = document.add_paragraph()
    sentence = paragraph.add_run('  \n\n'+extract['titre_complet'])
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence.font.name = 'Times New Roman'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(14) 
    
    '''Acronyme'''
    paragraph2 = document.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run(extract['titre_abrege'])
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(22) 
    
    '''Version protocole'''
    paragraph2 = document.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run(extract['code_protocole'])
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(14) 
    sentence.bold = False
    
    '''Protcole cat 1'''
    paragraph = document.add_paragraph()
    sentence = paragraph.add_run('PROTOCOLE DE RECHERCHE INTERVETIONNELLE IMPLIQUANT LA PERSONNE HUMAINE (catégorie 1)')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence.font.name = 'Times New Roman'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(14) 
    
    '''N° EudraCT : '''
    paragraph2 = document.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run('N° EudraCT : '+extract['num_eudract']+'\n')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(12) 
    sentence.bold = True
    
    '''Promoteur'''
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = paragraph.add_run('PROMOTEUR :\n')
    run1.font.name = 'Times New Roman'
    run1.font.size = docx.shared.Pt(11) 
    run1.bold = True
    run1.underline = True
    run2 = paragraph.add_run(extract['promoteur_nom_organisme']+'\n'+extract['promoteur_adresse']+'\nTél : '+extract['promoteur_num_telephone']+' / Fax : '+extract['promoteur_num_telecopie']+'\n')
    run2.font.name = 'Times New Roman'
    run2.font.size = docx.shared.Pt(11) 
    
    '''investigateur'''
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = paragraph.add_run('INVESTIGATEUR COORDONNATEUR :\n')
    run1.font.name = 'Times New Roman'
    run1.font.size = docx.shared.Pt(11) 
    run1.bold = True
    run1.underline = True
    run2 = paragraph.add_run(extract['investigateur_coordinateur_nom']+'\nService de : '+extract['investigateur_coordinateur_service']+'\n'+extract['investigateur_coordinateur_adresse_professionnelle']+'\nTél : '+extract['investigateur_coordinateur_telephone']+' / Fax : '+extract['investigateur_coordinateur_telecopie']+'\nE-mail : '+extract['investigateur_coordinateur_courriel'])
    run2.font.name = 'Times New Roman'
    run2.font.size = docx.shared.Pt(11) 
    
    '''GIRCI SOHO'''
    paragraph2 = document.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run('Ce protocole a été conçu et rédigé à partir de la version 3.0 du 01/02/2017\ndu protocole-type du GIRCI SOHO\n')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(12) 
    sentence.bold = True
    
    '''Confidentiel'''
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = paragraph.add_run('CE DOCUMENT CONFIDENTIEL')
    run1.font.name = 'Times New Roman'
    run1.font.size = docx.shared.Pt(10) 
    run1.underline = True
    run2 = paragraph.add_run(' EST LA PROPRIETE DU CHU DE POITIERS.\nAUCUNE INFORMATION NON PUBLIEE FIGURANT DANS CE DOCUMENT NE PEUT ETRE DIVULGUEE SANS AUTORISATION ECRITE PREALABLE DU CHU DE POITIERS')
    run2.font.name = 'Times New Roman'
    run2.font.size = docx.shared.Pt(10)
    
    '''Pied de page'''
    footer = page_garde.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    r = p.add_run(' ')
    

    '''Fin de page'''
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

    
   
    #document.save("page_garde.docx")                   #sauvegarde
   
def PageSignature(document):
    
  #  document = docx.Document()
 
    '''Logos de l'en-tete'''
    
    page_sign = document.add_section()
    header2 = page_sign.header
    header2.is_linked_to_previous = False
    p = header2.paragraphs[0]
    r = p.add_run() 
    r.add_text("\t\tACRONYME")
    p2 = header2.add_paragraph()
    r2 = p2.add_run() 
    r2.add_picture('imageGauche.png')
    
#    p.style = document.styles["Header"]
        
    '''Titre'''
    paragraph2 = document.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run('PAGE DE SIGNATURE DU PROTOCOLE')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(16) 
    
    '''Signature investigateur'''
    paragraph2 = document.add_paragraph()
    sentence = paragraph2.add_run('Signature de l’investigateur')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(11) 
    sentence.bold = True
    sentence.underline = True
    
    '''Premiere case'''
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    text1 = ' \nJ\'ai lu ce protocole d’essai clinique dont le CHU de Poitiers est le promoteur. Je confirme qu\'il contient toutes les informations nécessaires à la conduite de l’essai. Je m\'engage à mener cet essai en respectant ses directives et les termes et conditions qui y sont définis.\n'
    text2 = 'Je m\'engage à réaliser l’essai en respectant :\n\n'
    text3 = '    -  les principes de la “Déclaration d’Helsinki”, \n\
    -  les règles et recommandations de bonnes pratiques cliniques internationales (ICH-E6) et française      (règles de bonnes pratiques cliniques pour les recherches portant sur des médicaments à usage humain - décisions du 24 novembre 2006), \n\
    -  la législation nationale et la réglementation relative aux essais cliniques,\n\
    -  la conformité avec la Directive Essais Cliniques de l’UE [2001/20/EC].\n\n\n'
    text4 = "Je m'engage également à ce que les investigateurs et les autres membres qualifiés de mon équipe aient accès au protocole et aux documents relatifs à la conduite de l’essai pour leur permettre de travailler dans le respect des dispositions figurant dans ces documents.\n"
    text5 = "Investigateur : Dr/ Pr XXXXX\n(Prénom NOM)\n\n\n\n"
    text6 = "Signature : ……………………………………………..                          Date : ___________________\n"       

    table.cell(0,0).text = text1 +text2 + text3 + text4 +text5+text6
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= docx.shared.Pt(11)
                    font.name = 'Times New Roman'
                    
    '''Signature investigateur coordonnateur'''
    paragraph2 = document.add_paragraph()
    sentence = paragraph2.add_run(' \nSignature de l’Investigateur Coordonnateur')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(11) 
    sentence.bold = True
    sentence.underline = True
    
    '''Deuxieme case'''
    table2 = document.add_table(rows=1, cols=1, style='Table Grid')
    text5 = "Investigateur Coordonnateur : Dr/ Pr XXXXX\n(Prénom NOM)\n\n\n"
    text6 = "Signature : ……………………………………………..                          Date : ___________________\n" 
    table2.cell(0,0).text = text5 + text6
    for row in table2.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= docx.shared.Pt(11)
                    font.name = 'Times New Roman'
                    
                    
    '''Signature investigateur coordonnateur'''
    paragraph2 = document.add_paragraph()
    sentence = paragraph2.add_run(' \nSignature de l’Investigateur Coordonnateur')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(11) 
    sentence.bold = True
    sentence.underline = True
    
    '''Troisieme case'''
    table3 = document.add_table(rows=1, cols=1, style='Table Grid')
    text5 = "Promoteur : Jean-Pierre DEWITTE\nPour le Directeur Général et par délégation\nle Directeur de la Recherche,\n\n\n"
    text6 = "Signature : ……………………………………………..                          Date : ___________________\n" 
    table3.cell(0,0).text = text5 + text6
    for row in table3.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= docx.shared.Pt(11)
                    font.name = 'Times New Roman'
  
    '''Pied de page'''
    footer = page_sign.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    r = p.add_run('Version n°X du XX/XX/201X\tCONFIDENTIEL\tPage 3 sur 14')
    r.font.name = 'Times New Roman'
    r.font.size = docx.shared.Pt(11)
    
    
#    '''Fin de page'''
#    paragraph = document.add_paragraph()
#    run = paragraph.add_run()
#    run.add_break(WD_BREAK.PAGE)

    
  #  document.save("page_signature.docx")                   #sauvegarde

def liste_abreviation(document,extract):
   
      '''Titre'''
      paragraph2 = document.add_paragraph()
      paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
      sentence = paragraph2.add_run('LISTE DES ABREVIATIONS')
      '''Then format the sentence'''
      sentence.font.name = 'Times New Roman'
      sentence.bold = True
      sentence.font.size = docx.shared.Pt(16) 
      paragraph2 = document.add_paragraph('')
      texte=''
      ansm=amm=arc=bpc=cis=cnil=cpp=crf=e_crf=evi=evig=eig=eigi=ide=mr=rcp=susar=tec=False #pour savoir si on les a deja trouvés une fois
      '''Parcours tout l'extract pour voir si on trouve le mot quelque part dans le texte'''
      for para in extract.values():
          if para !="" and para !=" " and isinstance(para, str):
          #ANSM
              if (re.search(r"ANSM",para) and ansm==False): #si on trouve le mot (re.search renvoi True s'il trouve le mot) pour la première fois
                  texte+='ANSM\t\t\tAgence Nationale de Sécurité du Médicaments et des produits de santé\n' #alors on ajoute la ligne a texte
                  ansm=True #on a trouve ANSM une fois pas besoin de l'ajouter a nouveau les fois suivantes
          #AMM
              if re.search(r"AMM",para) and amm==False:
                 texte+='AMM\t\t\tAutorisation de Mise sur le Marché\n'
                 amm=True
          
          #ARC
              if re.search(r"ARC",para) and arc==False: 
                  texte+='ARC\t\t\tAttaché de Recherche Clinique\n'
                  arc=True
          #BPC
              if re.search(r"BPC",para) and bpc==False:
                  texte+='BPC\t\t\tBonnes Pratiques Cliniques\n'
                  bpc=True
          
          #CIS
              if re.search(r"CIS",para) and cis==False: 
                  texte+='CIS\t\t\tComité Indépendant de Surveillance\n'
                  cis=True
          #CNIL
              if re.search(r"CNIL",para) and cnil==False:
                  texte+='CNIL\t\t\tCommission Nationale de l’Informatique et des Libertés\n'
                  cnil=True
          #CPP
              if re.search(r"CPP",para) and cpp==False: 
                  texte+='CPP\t\t\tComité de Protection des Personnes\n' 
                  cpp=True
          #CRF
              if re.search(r"CRF",para) and crf==False: 
                  texte+='CRF\t\t\tCase Report Form (cahier d’observation)\n' 
                  crf=True
          #e-CRF
              if re.search(r"e-CRF",para) and e_crf==False: 
                  texte+='e-CRF\t\t\tCahier d’observation électronique\n'
                  e_crf=True
          
          #EvI
              if re.search(r"EvI",para) and evi==False: 
                  texte+='EvI\t\t\tEvènement Indésirable\n'
                  evi=True
          #EvIG
              if re.search(r"EvIG",para) and evig==False: 
                  texte+='EvIG\t\t\tEvènement Indésirable Grave\n'
                  evig=True
         #EIG
              if re.search(r"EIG",para) and eig==False: 
                  texte+='EIG\t\t\tEffet Indésirable Grave\n'
                  eig=True
          #EIGI
              if re.search(r"EIGI",para) and eigi==False: 
                  texte+='EIGI\t\t\tEffet Indésirable Grave Inattendu\n'
                  eigi=True
          
          #IDE
              if re.search(r"IDE",para) and ide==False: 
                  texte+='IDE\t\t\tInfirmier (ère) Diplômé(e) d\'Etat\n'
                  ide=True
          
          #MR
              if re.search(r"MR",para) and mr==False: 
                  texte+='MR\t\t\tMéthodologie de Référence\n'
                  mr=True
          
          #RCP
              if re.search(r"RCP",para) and rcp==False: 
                  texte+='RCP\t\t\tRésumé des Caractéristiques d\'un Produit\n'
                  rcp=True
          #SUSAR
              if re.search(r"SUSAR",para) and susar==False: 
                  texte+='SUSAR\t\t\tSssuspected Unexpected Serious Adverse Reaction\n'
                  susar=True
          
          #TEC
              if re.search(r"TEC",para) and tec==False: 
                  texte+='TEC\t\t\tTechnicien d\'Etude Clinique\n'
                  tec=True
                    
      paragraph2 = document.add_paragraph()
      sentence = paragraph2.add_run(texte)
      '''Then format the sentence'''
      sentence.font.name = 'Times New Roman'
      sentence.font.size = docx.shared.Pt(11)
     
     #FIN DU DOC 
      paragraph = document.add_paragraph()
      run = paragraph.add_run()
      run.add_break(WD_BREAK.PAGE)
     
      