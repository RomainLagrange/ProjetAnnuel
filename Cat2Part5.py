# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 16:28:06 2019

@author: Asuspc
"""
import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie5(document):
    'Creation de la partie 5 du protcole de catégorie 2'
   # document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

#    Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('5	DEROULEMENT DE LA RECHERCHE',document)
    
    
   # Ecriture du 5.1  
    Titre2('5.1	Schéma de la recherche',document)
    
    # Ecriture du 5.2  
    Titre2('5.2	Méthode pour la randomisation',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('La liste de randomisation est établie par le statisticien / le méthodologiste de l’étude avant le début de la recherche. Les effectifs des ')
    run1.style='Paragraphe'  
    run2=p.add_run('X groupes ')
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run2.style='Paragraphe'  
    run3 = p.add_run ('de traitement/stratégie/procédure sont équilibrés/déséquilibrés avec un ratio de ')
    run3.style='Paragraphe'  
    run4=p.add_run('x : x')
    run4.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run4.style='Paragraphe' 
    run5=p.add_run('. La randomisation est stratifiée sur - ')
    run5.style='Paragraphe'  
    run6=p.add_run('décrire le facteur de stratification (centre, facteur pronostique)')
    run6.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run6.font.italic=True
    run6.style='Paragraphe' 
    run7=p.add_run('. Un document décrivant la procédure de randomisation est conservé de manière confidentielle par la personne responsable ayant établi la liste')
    run7.style='Paragraphe'  
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Exemple\n')
    run1.style='Paragraphe' 
    run1.font.italic=True
    run2=p.add_run('Vue d’ensemble du déroulement de la recherche :')
    run2.style='Paragraphe' 
    

    # Ecriture du 5.3  
    Titre2('5.3	Calendrier de la recherche',document)
    
    
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Durée de la période d’inclusion : ' + extract['duree_inclusion']) 
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Durée de participation de chaque participant : ' + extract['duree_participation'])
    run1.style='Paragraphe' 

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Durée totale de la recherche ')
    run1.style='Paragraphe' 
    run2=p.add_run('(durée de la période d’inclusion + durée de participation)'+ extract['duree_inclusion'] + extract['duree_participation'])
    run2.font.italic= True
    run2.style='Paragraphe'  
    
    #Ecriture du titre 5.4
    Titre2('5.4	Tableau récapitulatif du suivi d’un participant à la recherche',document)

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le contenu du tableau doit comporter tous les examens du suivi. Distinguer les examens effectués dans le cadre du soin (S) de ceux réalisés spécifiquement pour la recherche (R).')
    run1.style='Paragraphe' 
    run1.font.italic=True

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('(*) ')
    run1.style='Paragraphe'
    run1.font.italic= True
    run1.font.superscript= True
    run2=p.add_run('V-X : unité de temps à adapter en fonction de la recherche : A (année), M (mois), S (semaine), J (jour), H (heure)')
    run2.style='Paragraphe'
    run2.font.italic= True

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('1')
    run1.style='Paragraphe'
    run1.font.italic= True
    run1.font.superscript= True
    run2=p.add_run('Examen clinique : détail de ce que comporte l’examen clinique ')
    run2.style='Paragraphe'
    run2.font.italic= True

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('2')
    run1.style='Paragraphe'
    run1.font.italic= True
    run1.font.superscript= True
    run2=p.add_run('Bilan biologique : liste des examens biologiques')
    run2.style='Paragraphe'
    run2.font.italic= True

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY 
    p.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('3')
    run1.style='Paragraphe'
    run1.font.italic= True
    run1.font.superscript= True
    run2=p.add_run('Examens para-cliniques : liste des examens para-cliniques')
    run2.style='Paragraphe'
    run2.font.italic= True
    #Ecriture du titre 5.5
    Titre2('5.5	Visites de pré-inclusion / inclusion = Visite V0',document)

    #Ecriture du titre5.5.1
    Titre3('5.5.1','Recueil du consentement',document)

    document.add_paragraph(' ')
    
    TexteGris('prendre contact avec la promotion interne \n pour aide a la redaction de ce chapitre', document)

    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run()
    run1.text='Lors de la visite de '
    run1.style='Paragraphe'
    run2=p.add_run()
    run2.text='pré-inclusion (voir selon l’étude si visite d’inclusion),'
    run2.style='Paragraphe'
    run2.font.italic=True
    run3=p.add_run()
    run3.text=' le médecin investigateur informe le patient de la possibilité de participer à cet essai clinique et répond à toutes ses questions concernant l\'objectif, la nature des contraintes, les risques prévisibles et les bénéfices attendus de la recherche. Il précise également les droits du patient dans le cadre d’une recherche et vérifie les critères d’éligibilité. '
    run3.style='Paragraphe'

    
    p=document.add_paragraph()    
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run()
    run1.text=('Un exemplaire de la note d’information et du formulaire de consentement est alors remis au participant par le médecin investigateur.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Après cette séance d’information, le participant dispose d’un délai de réflexion. Le médecin investigateur est responsable de l’obtention du consentement éclairé écrit du participant.\n\nSi le participant donne son accord de participation, ce dernier et l’investigateur inscrivent leurs noms et prénoms en clair, datent et signent le formulaire de consentement. Celui-ci ')
    run1.style='Paragraphe'
    run2=p.add_run('doit être signé avant la réalisation de tout examen ')
    run2.style='Paragraphe'
    run2.font.bold= True
    run2.font.underline= True
    run3=p.add_run('clinique ou biologique ou para-clinique nécessité par la recherche. ')
    run3.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('L’exemplaire ')
    run1.style='Paragraphe'
    run2=p.add_run('original ')
    run2.style='Paragraphe'
    run2.font.underline= True    
    run3=p.add_run('sera conservé dans le classeur de l’investigateur. Un exemplaire (un autre original ou une copie) sera remis au patient. ')
    run3.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('L’investigateur précisera dans le dossier médical du patient sa participation à la recherche, les modalités du recueil du consentement ainsi que celle de l’information. ')
    run1.style='Paragraphe'
        
#    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Le patient se verra attribuer un numéro de patient, selon la règle : ')
    run1.style='Paragraphe'
    
        #IMAGE NUMERO PATIENT
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=p.add_run()
    picture=run.add_picture('num_patient.png')
    

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Procédure d’urgence si applicable')
    run1.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run1.style='Paragraphe'
    run1.font.italic=True
    run1.font.underline=True

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Dans le cas d’une situation d’urgence et conformément à l’article L.1122-1-2, le consentement sera sollicité auprès de « proches » et seulement rétrospectivement auprès du patient dès une récupération suffisante lui permettant de donner son consentement libre et éclairé. ')
    run1.style='Paragraphe'
    
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Dans le cas où les « proches » ne peuvent pas être présents au moment de l’inclusion, une procédure d’urgence sera mise en place dans le cas d’une urgence vitale immédiate. Dans ce cas un médecin indépendant de l’étude, non déclaré comme médecin investigateur, peut donner son consentement d’urgence. L\'intéressé, ou le cas échéant, les membres de la famille ou la personne de confiance sont informés dès que possible et leur consentement leur est demandé pour la poursuite de cette recherche.')
    run1.style='Paragraphe'
    
    
    #Ecriture du titre5.5.2
    Titre3('5.5.2','Déroulement de la visite',document)

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('La visite de pré-inclusion/inclusion est assurée par le médecin investigateur. La visite de pré-inclusion a lieu entre') 
    run1.style='Paragraphe'               
    run2=p.add_run('X jours/semaines/mois')
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0)
    run2.style='Paragraphe'    
    run3=p.add_run('et au plus tard ')
    run3.style='Paragraphe' 
    run4=p.add_run('X jours/semaines/mois')
    run4.style='Paragraphe'
    run4.font.color.rgb = RGBColor(0x0,0xB0,0xF0)
    run5=p.add_run('avant la visite d’inclusion.')
    run5.style='Paragraphe'
    
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Avant tout examen lié à la recherche, l’investigateur recueille le consentement libre, éclairé et écrit du participant (ou de son représentant légal le cas échéant).')
    run1.style='Paragraphe'
    
    
    #Ecriture du titre 5.6
    Titre2('5.6	Visite de randomisation = Visite (Vx, ou Jx, ou Mx…)',document)
    
    #Ecriture du titre5.6.1
    Titre3('5.6.1','Description des examens',document)
    
    #Ecriture du titre5.6.2
    Titre3('5.6.2','Randomisation du patient',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Lorsqu’un investigateur souhaite effectuer la randomisation/l’inclusion après avoir vérifié l’éligibilité du participant/cluster, il se connecte sur le site Internet de l’e-CRF ')
    run1.style='Paragraphe'
    run2=p.add_run('https://www.chu-poitiers.hugo-online.fr/')
    run2.style='Paragraphe'
    run2.font.underline=True
    run3=p.add_run('. L’investigateur complète la page « randomisation » après avoir préalablement confirmé tous les critères d’éligibilité du participant/cluster sur le site. Après validation du contenu, la randomisation/ l’inclusion est effectuée et l’e-CRF :')
    run3.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Si l’étude est en ouvert : ')
    run1.style='Paragraphe'
    run1.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run1.font.italic=True
    run2=p.add_run('communique immédiatement à l’investigateur en clair le résultat de la randomisation l’inclusion, en particulier le groupe de traitement ')
    run2.style='Paragraphe'
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('/stratégie/procédure ')
    run3.style='Paragraphe'
    run3.font.italic=True
    run4=p.add_run('alloué(e) au participant/cluster et le numéro de la boîte de traitement. ')
    run4.style='Paragraphe'    
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Paragraph'
    run1=p.add_run('OU')
    run1.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run1.style='Paragraphe'
    run1.font.italic=True    
    run1.font.bold=True    
    run1.font.underline=True    
   
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Si l’étude est en double aveugle : ')
    run1.style='Paragraphe'
    run1.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run1.font.italic=True
    run2=p.add_run('communique immédiatement à l’investigateur le numéro unique de randomisation/d’inclusion, correspondant à une boîte de traitement ')
    run2.style='Paragraphe'
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('/stratégie/procédure ')
    run3.style='Paragraphe'
    run3.font.italic=True
    run4=p.add_run('alloué(e) au participant/cluster.')
    run4.style='Paragraphe'    
    
 
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Les numéros de randomisation sont établis selon des listes établies à l’avance par le méthodologiste de la recherche (voir paragraphe 5.2).')
    run1.style='Paragraphe'
    
    #Ecriture du titre 5.7
    Titre2('5.7	Visites de suivi = visite (Vx, ou Jx ou Sx ou Mx…)',document)

	
    #Ecriture du titre5.7.1
    Titre3('5.7.1','Visite (Vx, ou Sx, ou Jx, ou Mx…)',document)


    #Ecriture du titre5.7.2
    Titre3('5.7.2','Visite (Vx, ou Sx, ou Jx, ou Mx…)',document)


    #Ecriture du titre 5.8
    Titre2('5.8	Visite de fin de la recherche',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('En cas de sortie prématurée, l’investigateur doit en documenter les raisons de façon aussi complète que possible dans le dossier médical et réaliser une visite de fin d’étude. Il complètera la page de sortie prématurée du CRF. \nEn cas de sujet perdu de vue, l’investigateur mettra tout en œuvre pour reprendre contact avec la personne et connaître les raisons.')
    run1.style='Paragraphe'
    
    
    #Ecriture du titre 5.9
    Titre2('5.9	Règles d’arrêt de la participation d’une personne à la recherche',document)
    
    
    TexteGris('prendre contact avec la plateforme de methodologie \n pour aide a la redaction de ce chapitre', document)

    
    #Ecriture du titre5.9.1
    Titre3('5.9.1','Arrêt de participation définitif ou temporaire d’un patient dans l’étude',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('En fonction de l’état de santé du patient, l’investigateur devra faire son possible pour que les patients continuent à participer à l’étude. Cependant, il pourra interrompre temporairement ou définitivement la participation d’une personne à la recherche ou d’une partie ou de la totalité de la recherche. Les éléments suivants peuvent justifier l’arrêt définitif ou temporaire de la participation d’une personne à la recherche ou d’une partie ou de la totalité de la recherche.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Voir si applicable :')
    run1.style='Paragraphe'
    run1.font.italic=True

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.style='List Bullet 2'
    run1=p.add_run('Evénement(s) indésirable(s) qu’il(s) soi(en)t lié(s) à une procédure du protocole,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Déviation au protocole (apparition d’un critère de non-inclusion du protocole, prise d’un traitement non autorisé),')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.style='List Bullet 2'
    run1=p.add_run('Survenue d’une modification rendant impossible les investigations à effectuer ou la prise du traitement ou modifiant la réponse au traitement à l’étude,')
    run1.style='Paragraphe'


    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Retrait du consentement : les patients peuvent retirer leur consentement et demander à sortir de l’étude à n’importe quel moment et ce, quelle qu’en soit la raison,')
    run1.style='Paragraphe'
    

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Toute raison qui servirait au mieux les intérêts du sujet (par exemple en cas d’événements indésirables graves nécessitant une prise en charge incompatible avec le protocole).')
    run1.style='Paragraphe'


    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('La date et la raison de la sortie d’essai ou de l’arrêt du traitement devront être notées dans le dossier médical du patient et le cahier d’observation. La visite de fin d’étude devra être réalisée dans la mesure du possible.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Dans tous les cas, et dans la mesure du possible, l’investigateur devra compléter la visite de fin d’étude du CRF.')
    run1.style='Paragraphe'
    
    #Ecriture du titre5.9.2
    Titre3('5.9.2','Modalités de remplacement des patients exclus, le cas échéant',document)
    
    #Ecriture du titre5.9.3
    Titre3('5.9.3','Modalités et calendrier de recueil pour ces données',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Les données disponibles des patients sortis d’étude pour raison médicale seront recueillies pour l’analyse. Les données des patients ayant retiré leur consentement seront analysées uniquement si les patients ont donné leur accord. Si le patient n’a pas donné son accord, toutes les données de l’étude le concernant seront supprimées.')
    run1.style='Paragraphe'
    
    #Ecriture du titre 5.9.4
    Titre3('5.9.4','Modalités de suivi de ces personnes',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('La sortie d\'étude d\'un participant ne changera en rien sa prise en charge habituelle par rapport à sa maladie. Il ne bénéficiera cependant pas des évaluations cliniques prévues selon le protocole.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Exemple :\n')
    run1.style='Paragraphe'
    run1.font.italic =True
    run2=p.add_run('Tout participant qui sort de l’étude en raison d’un événement indésirable, grave ou non, ou d’une anomalie biologique devra être suivi jusqu’à ce que cet événement indésirable soit résolu ou que l’anomalie biologique soit normalisée.')
    run2.style='Paragraphe'
        
    
    #Ecriture du titre 5.10
    Titre2('5.10	Contraintes liées à la recherche et indemnisation éventuelle des participants',document)
    
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['indemnisation'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    
     #Ecriture du titre 5.11
    Titre2('5.11	Collection d’échantillons biologiques',document)
    
    TexteGris('prendre contact avec la promotion interne \n pour aide a la redaction de ce chapitre', document)

    p=document.add_paragraph(' ')

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('SI APPLICABLE')
    run1.font.italic=True

    #Ecriture du titre5.11.1
    Titre3('5.11.1','Objectifs',document)
    
    #Ecriture du titre5.11.2
    Titre3('5.11.2','Description de(s) (la) collection(s) ',document)
    
    #Ecriture du titre5.11.3
    Titre3('5.11.3','Conservation',document)
    
    #Ecriture du titre5.11.4
    Titre3('5.11.4','Devenir de la collection',document)    
    
     #Ecriture du titre 5.12
    Titre2('5.12	Arrêt d’une partie ou de la totalité de la recherche',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Le CHU de Poitiers se réserve le droit d\'interrompre l’étude, à tout moment, s\'il s\'avère que les objectifs d’inclusion ne sont pas atteints.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’étude peut être interrompue prématurément en cas de survenue d’événements indésirables inattendus, graves nécessitant une revue du profil d\'innocuité du produit. De même, des événements imprévus ou de nouvelles informations relatives au produit, au vu desquels les objectifs de l\'étude ou du programme clinique ne seront vraisemblablement pas atteints, peuvent amener le promoteur à interrompre prématurément l’étude.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('En cas d’arrêt prématuré de l’étude, l’information sera transmise par le promoteur dans un délai de 15 jours à l’ANSM et au CPP.')
    run1.style='Paragraphe'
    
#        #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

    
  #  document.save("Cat2Partie5.docx")   