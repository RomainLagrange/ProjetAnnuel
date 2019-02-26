# -*- coding: utf-8 -*-
"""
Created on Fri Feb  1 18:17:55 2019

@author: Marion
"""

import pandas as pd
import docx
from docx.api import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
#from docx.shared import RGBColor

#docmuents du cpp pour les dispositifs médicaux

def main_cpp_dm():
    document = docx.Document()
    cpp_dm(document)
    page2_cpp_dm(document)
    document.save("soumission-cpp-dm.docx")

def cpp_dm(document):
    

    
    '''Marge de la page'''
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    '''Titre CPP'''
  #  paragraph = document.add_paragraph()
    styles= document.styles
    style1 = styles.add_style('Debut', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style1.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style1.font
    fontdebut.name = 'Book Antiqua'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(20) 
    
    
    
    paragraph1 = document.add_paragraph('Comité de Protection des Personnes', style='Debut')
    paragraph1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    #Sous titre
     
    
    paragraph = document.add_paragraph('OUEST III', style='Debut')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    #ajouter le trait et les ombres
    
    #Infos promoteur
    
    style2=styles.add_style('Promoteur', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style2.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style2.font
    fontdebut.name = 'Book Antiqua'
    fontdebut.italic = True
    fontdebut.size = docx.shared.Pt(10) 
      
    paragraph = document.add_paragraph("Agréé par arrêté ministériel en date du 31 mai 2012, \nConstitué selon l'arrêté du Directeur Général de l'ARS Poitou Charentes en date du 25 juin 2012.\n\n"
                                       "C.H.U La Milétrie\nPavillon Administratif - Porte 213\n "
                                       "2 rue de le milétrie - CS 90 577 - 86021 POITIERS CEDEX\n"
                                       "Tel : 05.49.45.21.57\nFax : 05.49.46.12.62 \nE-mail : cpp-ouest3@chu-poitiers.fr \n", style='Promoteur')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    
    #titre du milieu qui dit pour quel proto
    
    paragraph = document.add_paragraph()
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = paragraph.add_run("Demande d'avis au CPP")
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence.font.name = 'Arial Narrow'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(12)
    sentence2 = paragraph.add_run(" (arrêté du 2 décembre 2016)\n")
    sentence2.font.name = 'Arial Narrow'
    sentence2.font.size = docx.shared.Pt(12)
    sentence3 = paragraph.add_run("sur un projet de recherche mentionnée au 1° ou 2° de l'article L. 1121-1 du CSP\nportant sur un sur un dispositif médical ou sur un dispositif médical de diagnostic")
    sentence3.font.name = 'Arial Narrow'
    sentence3.bold = True
    sentence3.font.size = docx.shared.Pt(12)
    sentence4 = paragraph.add_run(" in vitro\n")
    sentence4.font.name = 'Arial Narrow'
    sentence4.bold = True
    sentence4.italic = True
    sentence4.font.size = docx.shared.Pt(12)
    
 ###########################################   
 
    #ENtre le titre et le tableau
 
    style3=styles.add_style('Avant_tableau', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style3.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
#    paragraph_format.left_indent = Inches(10)
    fontdebut = style3.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.italic = True
    fontdebut.size = docx.shared.Pt(10) 
    
    paragraph = document.add_paragraph("Préalablement au dépôt du dossier le promoteur obtient un numéro d’enregistrement sur le site internet de l’ANSM. Ce numéro identifie chaque recherche réalisée en France.", style='Avant_tableau')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
  
 ##############################

    #Tableau central
    table=document.add_table(rows=15, cols=1, style='Table Grid')
    

    table.cell(0,0).text=("DOSSIER ADMINISTRATIF")
    table.cell(1,0).text=("Courrier de demande d’avis daté et signé")
    table.cell(2,0).text=("Formulaire de demande d’avis (site internet de la base de données EudraCT)")
    table.cell(3,0).text=("Document additionnel (annexe 1) + supports pour recrutement des personnes")
    table.cell(4,0).text=("Pour les recherches mentionnées au 1o de l’article L. 1121-1, le cas échéant, la copie de la ou des autorisations de lieux de recherches mentionnées à l’article L. 1121-13 du CSP")
    table.cell(5,0).text=("DOSSIER SUR LA RECHERCHE")
    table.cell(6,0).text=("Protocole de recherche (daté + numéro de version)")
    table.cell(7,0).text=("Résumé du protocole (daté + numéro de version)")
    table.cell(8,0).text=("La brochure pour l’investigateur mentionnée à l’article R. 1123-20 du CSP, "
                          "ou la notice d’instruction en vigueur pour tout DM pourvu du marquage CE, "
                          "ou la notice d’utilisation en vigueur pour tout DMDIV pourvu du marquage CE, accompagnée, pour les recherches mentionnées au 1o de l’article L.1121-1 dans le cadre desquelles le dispositif est utilisé dans une destination autre que celle figurant sur cette notice, de la synthèse des données justifiant son utilisation et sa sécurité d’emploi dans le cadre de la recherche "
                          "\nSi la brochure pour l’investigateur appartient à un tiers, l’autorisation du tiers délivrée au promoteur pour l’utiliser.")
    table.cell(9,0).text=("Document d’information sauf situation article L.1122-1-4"
                          "Ce document décrit notamment les contre-indications, les effets indésirables graves et les précautions d’utilisation du dispositif sur lequel porte la recherche. "
                          "Si le DM ou le DMDIV n’est pas pourvu du marquage CE, ce document précise également les risques éventuels liés au DM, DMDIV et, le cas échéant, ceux liés à la procédure de mise en oeuvre du DM. "
                          "Si le DM ou le DMDIV est pourvu du marquage CE, le dossier sur la recherche comprend également une comparaison et la justification des divergences pertinentes en termes de sécurité des personnes, entre le document d’information destiné aux personnes qui se prêtent à la recherche et la notice d’instruction pour les dispositifs médicaux, ou la notice d’utilisation pour les DM et DMDIV")
    table.cell(10,0).text=("Formulaire de consentement sauf lorsqu’il est fait application de l’article L. 1122-1-4")
    table.cell(11,0).text=("Attestation d’assurance (Décret n°2016-1537 du 16 novembre 2016 - art. 3)")
    table.cell(12,0).text=("Une justification de l’adéquation des moyens humains, matériels et techniques au projet de recherche et de leur compatibilité avec les impératifs de sécurité des personnes qui s’y prêtent, sauf si le lieu bénéficie de l’autorisation mentionnée à l’article L. 1121-13 du CSP")
    table.cell(13,0).text=("Curriculum vitae signé du ou des investigateurs datant d’un an maximum")
    table.cell(14,0).text=("La nature de la décision finale de l’ANSM, si disponible.")
    n=1
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    if n==1 or n==6:
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        #run.bold = True
                        font = run.font
                        font.size= docx.shared.Pt(11)
                        font.name = 'Arial'
                    else:
                        font = run.font
                        font.size= docx.shared.Pt(10)
                        font.name = 'Arial Narrow'
                        #paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
                    n=n+1
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="AFAFAF"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="AFAFAF"/>'.format(nsdecls('w')))
    table.rows[5].cells[0]._tc.get_or_add_tcPr().append(shading_elm_2)
    
    style5=styles.add_style('fin_tableau', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style5.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style5.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.italic = True
    fontdebut.size = docx.shared.Pt(10) 
    document.add_paragraph('Forme : 4 dossiers complets + 1 version électronique\n\n', style='fin_tableau')
    
    
def page2_cpp_dm(document):

    
    '''Marge de la page'''
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        
    styles= document.styles
    style1 = styles.add_style('Debut_page2', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style1.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style1.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(12) 
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run('Annexe 1\n')
    sentence.font.name = 'Arial Narrow'
    sentence.font.size = docx.shared.Pt(10.5)
    sentence2=paragraph.add_run('DOCUMENT ADDITIONNEL\n')
    sentence2.bold = True
    sentence2.font.name = 'Arial Narrow'
    sentence2.font.size = docx.shared.Pt(12)
    sentence3=paragraph.add_run('À LA DEMANDE D’AVIS AU COMITÉ DE PROTECTION DES PERSONNES SUR UN PROJET DE \nRECHERCHE MENTIONNÉE AU 1° OU AU 2° DE L’ARTICLE L. 1121-1 PORTANT SUR UN\n')
    sentence3.font.name = 'Arial Narrow'
    sentence3.font.size = docx.shared.Pt(12)
    sentence4=paragraph.add_run('DISPOSITIF MÉDICAL OU UN DISPOSITIF MÉDICAL DE DIAGNOSTIC')
    sentence4.bold = True
    sentence4.font.name = 'Arial Narrow'
    sentence4.font.size = docx.shared.Pt(12)
    sentence5=paragraph.add_run(' IN VITRO')
    sentence5.bold = True
    sentence5.italic = True
    sentence5.font.name = 'Arial Narrow'
    sentence5.font.size = docx.shared.Pt(12)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER   
    
    style3 = styles.add_style('page2_normal', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style3.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style3.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(11)
    
    paragraph=document.add_paragraph('\nCe document doit être complété de façon claire, compréhensible et en français.\n', style='page2_normal')
  
    
    paragraph=document.add_paragraph('1. Numéro d’enregistrement de la recherche :', style='page2_normal')
    paragraph=document.add_paragraph('2. Titre complet de la recherche :', style='page2_normal')
    paragraph=document.add_paragraph('3. Justification de la recherche :', style='page2_normal')
    paragraph=document.add_paragraph('4. Hypothèse principale de la recherche et objectifs :', style='page2_normal')
    paragraph=document.add_paragraph('5. Evaluation des bénéfices et des risques que présente la recherche, notamment les bénéfices escomptés pour les personnes qui se prêtent à la recherche et les risques prévisibles liés à l’utilisation des produits et aux procédures d’investigation de la recherche (incluant notamment la douleur, l’inconfort, l’atteinte à l’intégrité physique des personnes se prêtant à la recherche, les mesures visant à éviter et/ou prendre en charge les événements) :', style='page2_normal')
    paragraph=document.add_paragraph('6. Justifications de l’inclusion de personnes visées aux articles L. 1121-5 à L. 1121-8 et L. 1122-1-2 du code de la santé publique (notamment mineurs, majeurs protégés, recherches mises en oeuvre dans des situations d’urgence) et procédure mise en oeuvre afin d’informer et recueillir le consentement de ces personnes ou de leurs représentants légaux :', style='page2_normal')
    paragraph=document.add_paragraph('7. Description des modalités de recrutement des personnes (joindre notamment tous les supports publicitaires utilisés pour la recherche en vue du recrutement des personnes) :', style='page2_normal')
    paragraph=document.add_paragraph('8. Procédures d’investigation menées et différences par rapport aux conditions habituelles d’utilisation du dispositif médical ou dispositif médical de diagnostic in vitro, le cas échéant :', style='page2_normal')
    paragraph=document.add_paragraph('9. Justification de l’existence ou non : i) d’une interdiction de participer simultanément à une autre recherche ; ii) d’une période d’exclusion pendant laquelle la participation à une autre recherche est interdite.', style='page2_normal')   
    paragraph=document.add_paragraph('10. Modalités et montant de l’indemnisation des personnes se prêtant à la recherche, le cas échéant :', style='page2_normal')
    paragraph=document.add_paragraph('11. Motifs de constitution ou non d’un comité de surveillance indépendant :', style='page2_normal')
    paragraph=document.add_paragraph('12. Nombre prévu de personnes à inclure dans la recherche :\n', style='page2_normal')
    paragraph=document.add_paragraph('Par la présente, j’atteste/j’atteste au nom du promoteur (rayer la mention inutile) que les informations fournies ci-dessus à l’appui de la demande d’avis sont exactes.\n', style='page2_normal')
    paragraph=document.add_paragraph('Nom :\nPrénom :\nAdresse :\nFonction :\nDate :\nSignature :', style='page2_normal')
    