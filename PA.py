# -*- coding: utf-8 -*-

#INSTALLER PACKAGE : dans commande windows (cmd) ecrire : 
#C:\Users\Asuspc\Anaconda3>python -m pip install python-docx
"""
Created on Thu Nov  8 15:49:21 2018

@author: Asuspc
"""
import docx
from docx import Document

def getText():
    'Ouvre le fichier V1 pré-existant, copie le texte et remplace les balises \
     par les bonnes valeurs'
    
    f1 = open('V1.docx', 'rb') #ouvre le premier fichier
    doc = Document(f1)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)  #copie tout le texte dans une liste
    f1.close()                         #ferme le fichier
    print(fullText) 
    for i in range(len(fullText)):
        if fullText[i]=="{Prenom}":
            fullText[i]="Julie"       #modification souhaitée de la liste
        if fullText[i]=="{Nom}":
            fullText[i]="Poilvet"
    doc2=Document()                     #création du nouveau document
    doc2.add_paragraph(fullText)        #écris le nouveau le nouveau texte
    doc2.save("V2.docx")                   #sauvegarde
    return(fullText)

