# -*- coding: utf-8 -*-
"""
Created on Fri Feb  1 18:17:55 2019

@author: Marion
"""

import pandas as pd
import docx
from docx.api import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
#from docx.shared import RGBColor

#docmuents du cpp pour les dispositifs médicaux

def main_cpp_RNI(extract):
    document = docx.Document()
    cpp_rni(document)
    page2_cpp_rni(document, extract)
    document.save("soumission-cpp-RNI.docx")

def cpp_rni(document):
    

    
    '''Marge de la page'''
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    '''Titre CPP'''
  #  paragraph = document.add_paragraph()
    styles= document.styles
    style1 = styles.add_style('Debut', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style1.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style1.font
    fontdebut.name = 'Book Antiqua'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(20) 
    
    
    
    paragraph1 = document.add_paragraph('Comité de Protection des Personnes', style='Debut')
    paragraph1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    #Sous titre
     
    
    paragraph = document.add_paragraph('OUEST III', style='Debut')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    #ajouter le trait et les ombres
    
    #Infos promoteur
    
    style2=styles.add_style('Promoteur', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style2.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style2.font
    fontdebut.name = 'Book Antiqua'
    fontdebut.italic = True
    fontdebut.size = docx.shared.Pt(10) 
      
    paragraph = document.add_paragraph("Agréé par arrêté ministériel en date du 31 mai 2012, \nConstitué selon l'arrêté du Directeur Général de l'ARS Poitou Charentes en date du 25 juin 2012.\n\n"
                                       "C.H.U La Milétrie\nPavillon Administratif - Porte 213\n "
                                       "2 rue de le milétrie - CS 90 577 - 86021 POITIERS CEDEX\n"
                                       "Tel : 05.49.45.21.57\nFax : 05.49.46.12.62 \nE-mail : cpp-ouest3@chu-poitiers.fr \n", style='Promoteur')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    
    #titre du milieu qui dit pour quel proto
    
    paragraph = document.add_paragraph()
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    sentence = paragraph.add_run("Demande d'avis au CPP")
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence.font.name = 'Arial Narrow'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(12)
    sentence2 = paragraph.add_run(" (arrêté du 2 décembre 2016)\n")
    sentence2.font.name = 'Arial Narrow'
    sentence2.font.size = docx.shared.Pt(12)
    sentence3 = paragraph.add_run("sur un projet de recherche mentionnée mentionné au 3° de l’article L. 1121-1 du CSP\n")
    sentence3.font.name = 'Arial Narrow'
    sentence3.bold = True
    sentence3.font.size = docx.shared.Pt(12)
    
    
    
 ###########################################   


    #Tableau central
    table=document.add_table(rows=13, cols=1, style='Table Grid')
    

    table.cell(0,0).text=("DOSSIER ADMINISTRATIF")
    table.cell(1,0).text=("Courrier de demande d’avis daté et signé")
    table.cell(2,0).text=("Formulaire de demande d’avis (site internet de la base de données EudraCT)")
    table.cell(3,0).text=("DOSSIER SUR LA RECHERCHE")
    table.cell(4,0).text=("Protocole de recherche (daté + numéro de version)")
    table.cell(5,0).text=("Résumé du protocole (daté + numéro de version)")
    table.cell(6,0).text=("Le document d’information destiné aux personnes qui se prêtent à la recherche prévue à l’article L. 1122-1 du code de la santé publique, rédigé en français.")
    table.cell(7,0).text=("Le cahier de recueil des données de l’étude et/ou questionnaires")
    table.cell(8,0).text=("Le cas échéant un document attestant que l’étude a été demandée par l’Agence nationale de sécurité du médicament et des produits de santé, la Haute Autorité de santé, le ministère chargé de la santé ou l’Agence européenne des médicaments ;")
    table.cell(9,0).text=("La liste des investigateurs avec Curriculum vitae signé du ou des investigateurs datant d’un an maximum (investigateur principal de chaque site)")
    table.cell(10,0).text=("La description de l’utilisation (exclusive ou non) le cas échéant, de données extraites de systèmes d’information existants ou de bases d’étude déjà réalisées ;")
    table.cell(11,0).text=("Origine et nature des données nominatives recueillies, le cas échéant ; la justification du recours à celles-ci; le mode de circulation des données, les destinataires des données personnelles traitées; la durée de conservation des données; le cas échéant le transfert de données en dehors de l’UE;")
    table.cell(12,0).text=("Le cas échéant, la déclaration de conformité à une méthodologie homologuée de référence par la commission nationale de l’informatique et des libertés.")
    n=1
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    if n==1 or n==4:
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        #run.bold = True
                        font = run.font
                        font.size= docx.shared.Pt(11)
                        font.name = 'Arial'
                    else:
                        font = run.font
                        font.size= docx.shared.Pt(10)
                        font.name = 'Arial Narrow'
                        #paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
                    n=n+1
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="AFAFAF"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="AFAFAF"/>'.format(nsdecls('w')))
    table.rows[3].cells[0]._tc.get_or_add_tcPr().append(shading_elm_2)
    
    style5=styles.add_style('fin_tableau', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style5.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style5.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.italic = True
    fontdebut.size = docx.shared.Pt(10) 
    document.add_paragraph('\nForme : 4 dossiers complets + 1 version électronique\n\n', style='fin_tableau')
    
    document.add_page_break()
    
def page2_cpp_rni(document, extract):

    
    '''Marge de la page'''
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        
    styles= document.styles
    style1 = styles.add_style('Debut_page2', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style1.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style1.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(12) 
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run('Annexe 1\n')
    sentence.font.name = 'Arial Narrow'
    sentence.font.size = docx.shared.Pt(10.5)
    sentence2=paragraph.add_run('FORMULAIRE DE DEMANDE D’AVIS AU COMITÉ DE PROTECTION DES\nPERSONNES\nPOUR UNE RECHERCHE ')
    sentence2.font.name = 'Arial Narrow'
    sentence2.font.size = docx.shared.Pt(12)
    sentence3=paragraph.add_run('MENTIONNÉE AU 3°\nDE L’ARTICLE L. 1121-1 ')
    sentence3.font.name = 'Arial Narrow'
    sentence3.bold = True
    sentence3.font.size = docx.shared.Pt(12)
    sentence4=paragraph.add_run('DU CODE DE LA SANTÉ PUBLIQUE\n')
    sentence4.font.name = 'Arial Narrow'
    sentence4.font.size = docx.shared.Pt(12)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER   
    
    style3 = styles.add_style('page2_normal', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style3.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style3.font
    fontdebut.name = 'Arial Narrow'
    fontdebut.size = docx.shared.Pt(11)
    
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run('Promoteur\n')
    sentence.font.name = 'Arial Narrow'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(11)
    sentence=paragraph.add_run('Nom, raison sociale, sigle :' + extract['promoteur_nom_organisme']+'\nNom du responsable :'+extract['promoteur_nom_personne_contact']+'\nAdresse complète :'+extract['promoteur_adresse']+'\nTéléphone :'+extract['promoteur_num_telephone']+'\nMél :' + extract['promoteur_courriel'] + '\n')
    sentence.font.name = 'Arial Narrow'
    sentence.font.size = docx.shared.Pt(11)
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run('Investigateur coordonnateur :\n')
    sentence.font.name = 'Arial Narrow'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(11)
    sentence=paragraph.add_run('Nom, Prénom : '+extract['investigateur_coordinateur_nom']+' '+extract['investigateur_coordinateur_prenom']+'\nQualité : '+extract['investigateur_coordinateur_qualification']+'\nAdresse complète : '+extract['investigateur_coordinateur_adresse']+'\nTéléphone : '+extract['investigateur_coordinateur_telephone']+'\nMél : '+extract['investigateur_coordinateur_courriel']+'\n')
    sentence.font.name = 'Arial Narrow'
    sentence.font.size = docx.shared.Pt(11)
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run('Le cas échéant, personne responsable du traitement automatisé des données :\n')
    sentence.font.name = 'Arial Narrow'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(11)
    sentence=paragraph.add_run('Nom, Prénom :\nQualité :\nAdresse complète :\nTéléphone :\nMél :\n')
    sentence.font.name = 'Arial Narrow'
    sentence.font.size = docx.shared.Pt(11)
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run('Recherche :\n')
    sentence.font.name = 'Arial Narrow'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(11)
    sentence=paragraph.add_run('Intitulé de la recherche : ' + extract['titre_complet']+'\nNuméro d’enregistrement : '+extract['num_idrcb']+'\nNombre de personnes susceptibles d’être incluses dans la recherche : '+extract['taille_etude_longue']+'\nCaractéristiques de la recherche : \n')
    sentence.font.name = 'Arial Narrow'
    sentence.font.size = docx.shared.Pt(11)
    sentence=paragraph.add_run('              – transversale □\n              – longitudinale □\n              – comparatif ou non □\n              – cohorte □\n              – étude cas – témoin □\n              – exploratoire/démonstratif □\n              – autre □')
    sentence.font.name = 'Arial Narrow'
    sentence.font.size = docx.shared.Pt(11)
    
    paragraph=document.add_paragraph('\nPréciser : Personnes ayant accès au traitement automatisé des données :\n', style='page2_normal')
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run('Date')
    sentence.font.name = 'Arial Narrow'
    sentence.font.size = docx.shared.Pt(11)
    sentence=paragraph.add_run(' Signature du promoteur')
    sentence.font.name = 'Arial Narrow'
    sentence.italic = True
    sentence.font.size = docx.shared.Pt(11)