# -*- coding: utf-8 -*-
"""
Created on Thu Dec  6 16:36:00 2018

@author: Marion
"""
import pandas as pd
import docx
from docx.api import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

#permet d'extraire les tableaux
def table_extract():
    #ouvrir le premier document où est le tableau
    document = Document('V1.docx')
    doc2=Document() 
    #permet de dire au programme qu'il y a un tabaleau dedans
    table = document.tables[0]
    
    #créé une liste des valuers des cases
    data = []
    
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    print (data)
    #parcours la liste
    #on écrit dans le document
    #appelle la fonction qui cherche la valeur avec en paramètre le dictionnaire de valeur et le mot clé de colonne de gauche cherché
    #rassemble les valeurs dans un tableau
    df = pd.DataFrame(data)
    df.iloc[0,:].tolist()
    print(df)
    reco_tableau_bon(doc2, data)
    doc2.save("V2.docx")  


#fonction qui permet de renvoyer une valeur de cellule présente dans ce qui est cherché.
#si ce n'est pas trouvé, il renvoie "rien trouve"
def selec_good_data(le_dico, le_mot_cle):
    data=[]
    data=le_dico
    for i in data:
        #parcours les dico, valeursde la liste
        for keys, values in i.items():
            #selectionne la valeir cherchée
            if keys == le_mot_cle:
                #on retorune les valeurs
                print(keys)
                print(values)
                print('suivant')
                return (values)
    return ("")

#fonction qui compte le nombre de clé du dictionnaire=le nombre de lignes du futur tableau        
def compte_dico(le_dico):
    data=[]
    data=le_dico
    i=0
    for n in data:
        for keys in n.items():
            i=i+1
    return i


#méthode en pseudo-objet qui ne demande pas de statique mais suppose que le tableau donnée est le bon dès le départ et dans le bon ordre
def reco_tableau_bon(doc, le_dico):
    '''Marge de la page'''
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    '''Logos de l'en-tete'''
    header = section.header
    p = header.paragraphs[0]
    r = p.add_run() 
    r.add_picture('imageGauche.png')
    r.add_text('                                                                                                                                     ')
    r.add_text('ACRONYME')
        
    '''Titre'''
    paragraph2 = doc.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run('RESUME DU PROTOCOLE VERSION XX')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(16) 
    data=[]
    data=le_dico
    i=compte_dico(data)
    table=doc.add_table(rows=i, cols=2, style='Table Grid')
    n=0
    while n<i:
        for x in data:
            w=0
            for keys, values in x.items():
                if w == n:
                    table.cell(n,0).text=keys
                    table.cell(n,1).text=values
                    break
                w=w+1
        n=n+1
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= docx.shared.Pt(11)
                    font.name = 'Times New Roman'
  
    '''Pied de page'''
    footer = section.footer
    p = footer.paragraphs[0]
    r = p.add_run('Version n°X du XX/XX/201X	                               CONFIDENTIEL                                                Page 3 sur 14') 
    r.font.name = 'Times New Roman'
    r.font.size = docx.shared.Pt(11)
    
    #checher à mettre un tableau python dans un word
    #problème: on écrit un paragraphe à chaque valeur de la liste
    #récupérer la mise en page dudit tableau
    #le résumé change-t-il? Y a-t-il des images