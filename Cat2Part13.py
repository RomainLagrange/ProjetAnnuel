# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 16:56:19 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie13(document):
    'Creation de la partie 13 du protcole de catégorie 2'
#    document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

#    Style(document)

   
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('13	CONSIDERATIONS ETHIQUES ET REGLEMENTAIRES',document)
    
    TexteGris('prendre contact avec la promotion interne \n pour aide a la redaction de ces chapitres', document)

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le promoteur et l’(es) investigateur(s) s’engagent à ce que cette recherche soit réalisée en conformité avec la loi n°2004-806 du 9 août 2004, ainsi qu’en accord avec les Bonnes Pratiques Cliniques (I.C.H. version 4 du 1er mai 1996 et décision du 24 novembre 2006) et la déclaration d’Helsinki (qui peut être retrouvée dans sa version intégrale sur le site ')
    run1.style='Paragraphe'
    run2=p.add_run('http://www.wma.net).')
    run2.style='Paragraphe'
    run2.font.underline=True

    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('La recherche est conduite conformément au présent protocole. Hormis dans les situations d’urgence nécessitant la mise en place d’actes thérapeutiques précis, l’(es) investigateur(s) s’engage(nt) à respecter le protocole en tous points en particulier en ce qui concerne le recueil du consentement et la notification et le suivi des événements indésirables graves.')
    run1.style='Paragraphe'
    
   # Ecriture du 13.1  
    Titre2('13.1	Approbation de la recherche',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('CPP\n')
    run1.style='Paragraphe'
    run1.font.underline=True
    run2=p.add_run('Cette recherche a reçu l’avis favorable du Comité de Protection des Personnes ')
    run2.style='Paragraphe'
    run3=p.add_run('xxx (CPP xxxxx).')
    run3.style='Paragraphe'
    run3.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Assurance')
    run1.style='Paragraphe'
    run1.font.underline=True
    run2=p.add_run('Le CHU de Poitiers, promoteur de cette recherche, a souscrit un contrat d’assurance en responsabilité civile auprès de SHAM (18 rue Edouard Rochet-69372 LYON Cedex 08), contrat d’assurance n° 148163 conformément aux dispositions de l’article L1121-10 du code de la santé publique.')
    run2.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('CNIL/RGPD')
    run1.style='Paragraphe'
    run1.font.underline=True
    run2=p.add_run('Les données enregistrées à l’occasion de cette recherche font l’objet d’un traitement informatisé par le promoteur dans le respect de la loi n°78-17 du 6 janvier 1978 relative à l’informatique, aux fichiers et aux libertés  modifiée par la loi 2004-801 du 6 août 2004 et la loi n°2018-493 du 20 juin 2018 relative à la protection des données personnelles;. ')
    run2.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’étude respectera le Règlement Général sur la Protection des Données (RGPD) n°2016-679 du 27 avril 2016.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Cette recherche entre dans le cadre de la « Méthodologie de référence » (MR-001) homologué par délibération n°2018-153 du 3 mai 2018, entrée en vigueur le 13 juillet 2018. Le CHU de Poitiers a signé un engagement de conformité à cette « Méthodologie de référence ».')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Clinical Trial')
    run1.style='Paragraphe'
    run1.font.underline=True
    run2=p.add_run('Cette recherche sera enregistrée sur le site http://clinicaltrials.gov/ ')
    run2.style='Paragraphe'
    
  
    # Ecriture du 13.2  
    Titre2('13.2	Modifications au protocole',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toute modification substantielle, c’est à dire toute modification de nature à avoir un impact significatif sur la protection des personnes, sur les conditions de validité et sur les résultats de la recherche, sur la qualité et la sécurité des produits expérimentés, sur l’interprétation des documents scientifiques qui viennent appuyer le déroulement de la recherche ou sur les modalités de conduite de celle-ci, fait l’objet d’un amendement écrit qui est soumis au promoteur ; celui-ci doit obtenir, préalablement à sa mise en œuvre, un avis favorable du CPP.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les modifications non substantielles, c\'est à dire celles n’ayant pas d’impact significatif sur quelque aspect de la recherche que ce soit, sont communiquées au CPP à titre d’information.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toutes les modifications sont validées par le promoteur, et par tous les intervenants de la recherche concernés par la modification, avant soumission au CPP. Cette validation peut nécessiter la réunion de tout comité constitué pour la recherche.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toutes les modifications au protocole doivent être portées à la connaissance de tous les investigateurs qui participent à la recherche. Les investigateurs s’engagent à en respecter le contenu.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toute modification qui modifie la prise en charge des participants ou les bénéfices, risques et contraintes de la recherche fait l’objet d’une nouvelle note d’information et d’un nouveau formulaire de consentement dont le recueil suit la même procédure que celle précitée.')
    run1.style='Paragraphe'
    
     # Ecriture du 13.3 
    Titre2('13.3	Information du patient et formulaire de consentement éclairé écrit',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les patients seront informés de façon complète et loyale, en des termes compréhensibles, des objectifs et des contraintes de l\'étude, des risques éventuels encourus, des mesures de surveillance et de sécurité nécessaires, de leurs droits de refuser de participer à l\'étude ou de la possibilité de se rétracter à tout moment.\nToutes ces informations figurent sur un formulaire d’information et de consentement remis au patient. \nLe consentement libre, éclairé et écrit du patient sera recueilli par l’investigateur, ou un médecin qui le représente avant l’inclusion définitive dans l’étude. \nLe formulaire est signé par les deux parties :')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Un exemplaire original est conservé par l’investigateur,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Un exemplaire (une copie ou un deuxième original) sera remis au patient.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’investigateur devra s’assurer que la personne qui se prête à la recherche aura eu le temps de prendre sa décision librement et aura pu lire et comprendre la notice d’information et le formulaire de consentement.')
    run1.style='Paragraphe'

    
     # Ecriture du 13.4
    Titre2('13.4	Inscription au fichier national des personnes se prêtant à une recherche',document)
    
      #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    
  #  document.save("Cat2Partie13.docx")