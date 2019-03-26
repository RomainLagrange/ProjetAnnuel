# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 15:20:07 2019

@author: Asuspc
"""


import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document) --> écrire en minuscule !!!
#    TexteGrisJustif(texte,document)

def Partie15(document):
    'Creation de la partie 15 du protcole de catégorie 1'
  #  document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

  #  Style(document)
    
#---------------------------------------------------------------ECRITURE
    
    
     #ecriture du premier titre 
    Titre1('15	CONSERVATION DES DOCUMENTS ET DES DONNEES RELATIFS A LA RECHERCHE',document)
    
     #Texte gris centré
    TexteGris('prendre contact avec la promotion interne \n pour aide a la redaction de ces chapitres',document)

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les documents suivants relatifs à cette recherche sont archivés par l’investigateur conformément aux Bonnes Pratiques Cliniques :')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('pour une durée de 10 ans suivant la fin de la recherche ')
    run1.style='Paragraphe'
    run1.font.bold=True
    run1.font.italic=True
    run2=p.add_run('(recherches portant sur des produits cosmétiques),')
    run2.style='Paragraphe'
    run2.font.italic=True

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('pour une durée de 15 ans suivant la fin de la recherche ')
    run1.style='Paragraphe'
    run1.font.bold=True
    run1.font.italic=True
    run2=p.add_run('(recherches portant sur des médicaments, des dispositifs médicaux ou des dispositifs médicaux de diagnostic in vitro ou recherches ne portant pas sur un produit mentionné à l’article L.5311-1 du code de la santé publique),')
    run2.style='Paragraphe'
    run2.font.italic=True

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('pour une durée de 30 ans suivant la fin de la recherche ')
    run1.style='Paragraphe'
    run1.font.bold=True
    run1.font.italic=True
    run2=p.add_run('(recherches portant sur des produits sanguins labiles, des organes, des tissus d’origine humaine ou animale ou des préparations de thérapie cellulaire),')
    run2.style='Paragraphe'
    run2.font.italic=True

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('pour une durée de 40 ans suivant la fin de la recherche ')
    run1.style='Paragraphe'
    run1.font.bold=True
    run1.font.italic=True
    run2=p.add_run('(recherches portant sur des médicaments dérivés du sang ou des dispositifs médicaux incorporant une substance qui est susceptible d’être considérée comme un médicament dérivé du sang),')
    run2.style='Paragraphe'
    run2.font.italic=True
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('Le protocole et les modifications éventuelles au protocole')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('Les cahiers d’observation (copies)')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('Les dossiers source des participants ayant signé un consentement')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('Tous les autres documents et courriers relatifs à la recherche')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('pour une durée de 30 ans suivant la fin de la recherche')
    run1.style='Paragraphe'
    run1.font.bold=True
    run1.font.italic=True
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('L’exemplaire original des consentements éclairés signés des participants')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Tous ces documents sont sous la responsabilité de l’investigateur pendant la durée réglementaire d’archivage.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Aucun déplacement ou destruction ne pourra être effectué sans l’accord du promoteur. Au terme de la durée réglementaire d’archivage, le promoteur sera consulté pour destruction. Toutes les données, tous les documents et rapports pourront faire l’objet d’audit ou d’inspection.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('A l’issue de la période d’utilité pratique, l’ensemble des documents à archiver, tels que définis dans la procédure DRC-DOC-004 « classement et archivage des documents liés aux recherches» du CHU de Poitiers sera transféré sur le site d’archivage (Service Central des Archives – CHU Poitiers) et sera placé sous la responsabilité du Promoteur pendant 15 ans après la fin de l’étude conformément aux pratiques institutionnelles.')
    run1.style='Paragraphe'

#    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    #document.save("Partie7.docx") 
    
  #  document.save("Partie15.docx")
    