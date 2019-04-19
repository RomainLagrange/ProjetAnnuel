# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:22:29 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:13:09 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Mon Feb 18 11:51:47 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie10(document):
    'Creation de la partie 10 du protcole de catégorie 3'
   # document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

   # Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('10	CONTROLE ET ASSURANCE DE LA QUALITE',document)
    
    
   # Ecriture du 10.1  
    Titre2('10.1	Consignes pour le recueil des données',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toutes les informations requises par le protocole doivent être consignées sur les cahiers d’observation et une explication doit être apportée pour chaque donnée manquante. Les données doivent être recueillies au fur et à mesure qu\'elles sont obtenues, et transcrites dans ces cahiers de façon nette et lisible.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les données enregistrées dans l’e-CRF ')
    run1.style='Paragraphe'
    run2=p.add_run('[CRF] ')
    run2.style='Paragraphe'
    run2.font.italic=True
    run3=p.add_run('et provenant des documents sources doivent être cohérentes entre elles ; dans le cas contraire, les différences doivent être justifiées et documentées.')
    run3.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L\'investigateur est responsable de l\'exactitude, de la qualité et de la pertinence de toutes les données saisies.')
    run1.style='Paragraphe'
    
    # Ecriture du 10.2  
    Titre2('10.2	Suivi de la recherche',document)
    
    TexteGris('a completer uniquement si applicable', document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le suivi de la recherche sera assuré par un technicien de recherche clinique. Il sera chargé, auprès de la personne qui dirige et surveille la recherche, de :')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('la logistique et la surveillance de la recherche,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('l’établissement des rapports concernant son état d’avancement,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('la vérification de la mise à jour du cahier d’observation (demande d’informations complémentaires, corrections,…),')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('l’envoi des prélèvements.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Il travaillera conformément aux procédures opératoires standardisées, en collaboration avec l’attaché de recherche clinique délégué par le promoteur.')
    run1.style='Paragraphe'
    
    
    # Ecriture du 10.3  
    Titre2('10.3	Contrôle de Qualité',document)
    
    TexteGris('a completer uniquement si applicable', document)

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Un attaché de recherche clinique mandaté par le promoteur visite de façon régulière chaque centre, lors de la mise en place de la recherche, une ou plusieurs fois en cours de recherche selon le rythme des inclusions et en fin de recherche. Lors de ces visites, les éléments suivants seront revus :')
    run1.style='Paragraphe'
#
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('respect du protocole de la recherche,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('le recueil des EvI si nécessaire pour la recherche,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('qualité des données recueillies dans le cahier d\'observation : exactitude, données manquantes, cohérence des données avec les documents sources (dossiers médicaux, carnets de rendez-vous, originaux des résultats de laboratoire, etc,…).')
    run1.style='Paragraphe'


    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L\'investigateur et les membres de son équipe acceptent de se rendre disponibles lors des visites de Contrôle de Qualité (monitoring) effectuées à intervalles réguliers par l’Attaché de Recherche Clinique.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toute visite fera l’objet d’un rapport de monitorage par compte-rendu écrit.')
    run1.style='Paragraphe'
    
    #Ecriture du titre 10.4
    Titre2('10.4	Gestion des données',document)

#----------------gestion des données ----------
 #style    
    styles = document.styles
    styleBackgroundGrey = styles.add_style('CRF', WD_STYLE_TYPE.CHARACTER)
    styleBackgroundGrey.base_style = styles['No Spacing']
    fontBackgroundGrey = styleBackgroundGrey.font
    fontBackgroundGrey.name = 'Times New Roman'
    fontBackgroundGrey.size = docx.shared.Pt(11)
    fontBackgroundGrey.bold = True
    
    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table = document.add_table(rows = 1, cols = 1)
    row = table.rows[0].cells
    para_text ='Gestion des données pour une étude e-CRF'
    cell = row[0]
    pt = cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    cell._tc.get_or_add_tcPr().append(shading_elm)
    p.style='CRF'
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('L’investigateur devra dater et signer les pages du CRF complétées à la fin du recueil des données ; elles seront considérées comme documents source.\nCe document fera partie intégrante du dossier médical du patient et y sera conservé en permanence. ')
    run1.style='Paragraphe'
    
    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table = document.add_table(rows = 1, cols = 1)
    row = table.rows[0].cells
    para_text ='Gestion des données pour une étude CRF papier'
    cell = row[0]
    pt = cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run(para_text)
    cell._tc.get_or_add_tcPr().append(shading_elm)
    p.style='CRF'
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les données erronées relevées sur les cahiers d\'observation seront clairement barrées et les nouvelles données seront notées, à côté de l\'information barrée, accompagnées des initiales, de la date et éventuellement d’une justification par l’investigateur ou la personne autorisée qui aura fait la correction.')
    run1.style='Paragraphe'
#-----------------------------------------------------------

    #Ecriture du titre 10.5
    Titre2('10.5	Audit et inspection',document)

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Un audit peut être réalisé à tout moment par des personnes mandatées par le promoteur et indépendantes des personnes menant la recherche. Il a pour objectif de vérifier la sécurité des participants et le respect de leurs droits, le respect de la réglementation applicable et la fiabilité des données.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Une inspection peut également être diligentée par une autorité compétente (ANSM pour la France ou EMA dans le cadre d’un essai européen par exemple). ')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’audit, aussi bien que l’inspection, pourront s’appliquer à tous les stades de la recherche, du développement du protocole à la publication des résultats et au classement des données utilisées ou produites dans le cadre de la recherche.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les investigateurs acceptent de se conformer aux exigences du promoteur en ce qui concerne un audit et à l’autorité compétente pour une inspection de la recherche.')
    run1.style='Paragraphe'

    
    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
  
  #  document.save("Cat3Part10.docx")   