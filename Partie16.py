# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 15:21:18 2019

@author: Asuspc
"""


import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document) --> écrire en minuscule !!!
#    TexteGrisJustif(texte,document)

def Partie16(document):
    'Creation de la partie 16 du protcole de catégorie 1'
  #  document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

  #  Style(document)
    
#---------------------------------------------------------------ECRITURE
    
    
     #ecriture du premier titre 
    Titre1('16	RAPPORT FINAL',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Dans un délai d\'un an suivant la fin de la recherche ou son interruption, un rapport final sera établi et signé par le promoteur et l\'investigateur. Ce rapport sera tenu à la disposition de l\'autorité compétente. Le promoteur transmettra au CPP et, le cas échéant, à l\'ANSM les résultats de la recherche sous forme d\'un résumé du rapport final dans un délai d\'un an après la fin de la recherche.')
    run1.style='Paragraphe'

    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    #document.save("Partie7.docx") 
  
  #  document.save("Partie16.docx")