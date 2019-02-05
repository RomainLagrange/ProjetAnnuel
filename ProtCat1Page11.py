# -*- coding: utf-8 -*-
"""
Created on Thu Jan 31 13:32:03 2019

@author: Julie
"""

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt
from docx.shared import RGBColor

#COUleur Titre1 + indentation + soulignae titre2 + organisation texte + saut à la ligne


def Page11():
    'Creation de la page 11 du protcole de catégorie 1'
    document = docx.Document()
    
#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
 
    
    #essai = document.add_paragraph('test',style='Title') TROUVER COMMENT FAIRE LA MEME LIGNE 

    styles = document.styles
#   definition du style Titre1
    styleTitre1 = styles.add_style('Titre1', WD_STYLE_TYPE.PARAGRAPH, WD_ALIGN_PARAGRAPH.CENTER)
    styleTitre1.base_style = styles['Heading1']
    fontTitre1 = styleTitre1.font
    fontTitre1.name = 'Times New Roman' #police
    fontTitre1.size = docx.shared.Pt(12) #taille
    fontTitre1.all_caps = True #toujours en majuscule
    fontTitre1.bold= True #en gras
    fontTitre1.color.rgb = RGBColor(0x0,0x70,0xC0) #couleur bleu, en base 16
    #ecriture du premier titre (1) + rajouter la bordure en dessous !!!
    paragraph=document.add_paragraph('1	JUSTICATION SCIENTIFIQUE ET DESCRIPTION GENERALE\n', style='Titre1') #titre
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER #centrer
                

    
    #definition du style pour le texte indicatif   ESPACEMENT LIGNES
    styleIndic = styles.add_style('TexteItalic', WD_STYLE_TYPE.PARAGRAPH)
    styleIndic.base_style = styles['Normal']
    fontIndic = styleIndic.font
    fontIndic.name = 'Times New Roman'   #INDENTATION
    fontIndic.size = docx.shared.Pt(11)
    fontIndic.italic= True
    
    #Texte indicatif en italique TEST
    paragraph1 = document.add_paragraph ('Présentation du problème et justification étayée par les connaissances actuelles avec leurs références à la littérature scientifique et aux données pertinentes.\
    Indiquer en quoi l’objectif est nouveau et utile, pour le progrès des connaissances médicales et/ou de la prise en charge des malades. Les retombées attendues et perspectives peuvent également être développées dans ce chapitre.\n\
    C’est dans ce paragraphe que vous devez justifier la pertinence de votre étude.', style ='TexteItalic') 
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    #Definition du Titre2    INDENTATION + NON SOULIGNEMENT CHIFFRE
    styleTitre2 = styles.add_style('Titre2', WD_STYLE_TYPE.PARAGRAPH, WD_ALIGN_PARAGRAPH.JUSTIFY)
    styleTitre2.base_style = styles['Heading2']
    fontTitre2 = styleTitre2.font
    fontTitre2.name = 'Times New Roman'
    fontTitre2.size = docx.shared.Pt(14)
    fontTitre2.bold= True
    fontTitre2.color.rgb = RGBColor(0x0,0x0,0x0)
    #ecriture du titre1.1
    document.add_paragraph('1.1	Etat actuel des connaissances', style='Titre2') 

    #Definition du Titre3
    styleTitre3 = styles.add_style('Titre3', WD_STYLE_TYPE.PARAGRAPH)
    styleTitre3.base_style = styles['Heading3']
    fontTitre3 = styleTitre3.font
    fontTitre3.name = 'Times New Roman'
    fontTitre3.size = docx.shared.Pt(12)
    fontTitre3.bold= True
    fontTitre3.underline= True
    fontTitre3.color.rgb = RGBColor(0x0,0x0,0x0)
    #ecriture du titre1.1.1
    document.add_paragraph('1.1.1	Sur la pathologie', style='Titre3')  
            #ajuster style
    
    #Texte indicatif en italique TEST    #UTILISER, style ='TexteItalic'
    paragraph2 = document.add_paragraph() 
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sentence = paragraph2.add_run('Epidémiologie de la pathologie traitée comportant fréquence et gravité, facteurs de risque, conséquences individuelles et médico-économiques')
    sentence.font.italic=True
    sentence.font.size = docx.shared.Pt(11)
    
    #ecriture du titre1.1.2
    document.add_paragraph('1.1.2	Sur les traitements, stratégies et procédures de référence et à l’étude', 
                           style='Titre3')  
            #ajuster style
    

    #Texte indicatif en italique
    paragraph3 = document.add_paragraph ('Décrire les traitements/stratégies/ \
    procédures déjà existants et leurs limites. \
    Décrire les traitements/stratégies/procédures à l’étude\
    pharmacologique et/ou physiopathologique,…)\n Justifier notamment le \
    choix du groupe de comparaison : utilisation d’un traitement/stratégie\
    /procédure de référence, d’un placebo en se fondant sur les connaissances \
    scientifiques actuelles, en particulier les recommandations pour le \
    traitement de la maladie ou de l’état de santé à l’étude.\n Si la recherche\
    porte sur un médicament,préciser et justifier le choix de mener la \
    recherche sur des volontaires sains ou sur des patients, le choix de la \
    forme pharmaceutique, la posologie, le schéma d’administration,la durée du\
    traitement et de la voie d’administration.\n Attention les modalités \
    précises d’administration, le schéma d’administration ne sont pas à \
    détailler ici ; ils seront détaillés dans le paragraphe 7.\n', 
    style='TexteItalic')
    paragraph4 = document.add_paragraph() 
    paragraph4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sentence = paragraph4.add_run('Exemple : Les voies d’administration, \
    posologie, schéma d’administration et durée de traitement sont ceux de \
    l’AMM de [produit à l’étude] et sont détaillés dans les Résumés des \
    Caractéristiques Produits de cette molécule en annexe de ce protocole.')
    sentence.font.size = docx.shared.Pt(11)
                                        
                                        
    #ecriture du titre1.2
    document.add_paragraph('1.2	Hypothèse de la recherche et résultats attendus', style='Titre2')
    #Texte indicatif en italique
    paragraph5 = document.add_paragraph ('Définir précisément l’hypothèse, \
    physiopathologique ou autre, qui justifie la mise en place de la recherche\
    , en mentionnant le traitement/ la stratégie/la procédure à l’étude, \
    la population cible (justifier le choix de mener la recherche sur des \
    volontaires sains/patients) et le critère sur lequel il sera jugé.', 
    style='TexteItalic')

     
     
     #ecriture du titre1.3
    document.add_paragraph('1.3	Justification des choix méthodologiques', style='Titre2')
    #definition style texte surligné en gris   CENTRER + COULEUR 
    styles = document.styles
    styleBackgroundGrey = styles.add_style('BackgroundGrey', WD_STYLE_TYPE.PARAGRAPH, WD_ALIGN_PARAGRAPH.CENTER)
    styleBackgroundGrey.base_style = styles['Normal']
    fontBackgroundGrey = styleBackgroundGrey.font
    fontBackgroundGrey.name = 'Times New Roman'
    fontBackgroundGrey.size = docx.shared.Pt(11)
    fontBackgroundGrey.bold = True
    fontBackgroundGrey.small_caps = True
    fontBackgroundGrey.highlight_color = WD_COLOR_INDEX.GRAY_25
    #Texte sur fond gris
    paragraph6 = document.add_paragraph ('prendre contact avec la plateforme de\ methodologie \n pour aide a la redaction du paragraphe 2.3', style = 'BackgroundGrey')
    #Texte indicatif en italique
    paragraph7 = document.add_paragraph() 
    paragraph7.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sentence = paragraph7.add_run('Si nécessaire, justifier les choix réalisés\
    pour le schéma de la recherche, le critère de jugement principal, le type \
    de comparaison et la conduite de la recherche.')
    sentence.font.italic=True
    sentence.font.size = docx.shared.Pt(12)
    
    
    
     #ecriture du titre1.4
    document.add_paragraph('1.4	Rapport bénéfices / risques prévisibles'
                           , style='Titre2')
    #ecriture du titre1.4.1
    document.add_paragraph('1.4.1	Bénéfices', style='Titre3')  
     #Texte indicatif en italique
    paragraph8 = document.add_paragraph('Expliquer quel(s) est (sont) le(s) \
    bénéfice(s) individuel(s) et collectif(s), le(s) risque(s) prévisible(s) et\
    les contraintes liées à la recherche. Juger le rapport qui permet de \
    proposer ce protocole à l’étude. Indiquer les bénéfices et les risques que \
    présente la recherche, notamment les bénéfices escomptés pour les personnes\
    qui se prêtent à la recherche.', style='TexteItalic')
    
    #ecriture du titre1.4.2
    document.add_paragraph('1.4.2	Risques', style='Titre3')  
     #Texte indicatif en italique
    paragraph9 = document.add_paragraph('Décrire les risques prévisibles liés \
    au traitement et aux procédures d’investigation de la recherche (incluant \
    notamment la douleur, l’inconfort, l’atteinte à l’intégrité physique des\
    personnes se prêtant à la recherche, les mesures visant à éviter et/ou \
    prendre en charge les événements inattendus).', style='TexteItalic')
    
    
     #ecriture du titre1.5
    document.add_paragraph('1.5	Retombées attendues', style='Titre2')
     #Texte indicatif en italique
    paragraph10 = document.add_paragraph('Description détaillée des retombées \
    attendues par cette recherche (en terme d’amélioration des connaissances \
    sur une pathologie, d’augmentation de l’arsenal thérapeutique,…)..',
    style='TexteItalic')
    


    document.save("page11.docx")   