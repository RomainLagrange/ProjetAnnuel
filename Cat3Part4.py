# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:10:31 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:09:16 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:07:33 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:05:49 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 16:19:49 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style, Titre1,Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK

from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)
    

def Partie4(document):
    'Creation de la partie 4 du protcole de catégorie 3'
   # document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)
        
 #   Style(document)

 
    Titre1('4	CONCEPTION DE LA RECHERCHE',document)
    #ecriture du premier titre 
#    paragraph=document.add_paragraph('1	JUSTICATION SCIENTIFIQUE ET DESCRIPTION GENERALE\n', style='Titre1') #titre
#    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER #centrer

    TexteGris('prendre contact avec la plateforme methodologie \n pour aide a la redaction de ce chapitre',document)
  
    
    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
     

   # document.save("Cat3Partie4.docx")   