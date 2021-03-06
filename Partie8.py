# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 14:40:27 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie8(document):
    'Creation de la partie 8 du protcole de catégorie 1'
   # document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

 #   Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('8	TRAITEMENTS ET PROCEDURES ASSOCIE(E)S ',document)

   # Ecriture du 8.1  
    Titre2('8.1	Traitements / procédures associé(e)s autorisés',document)
    
    #Ecriture du titre 8.1.1
    Titre3('8.1.1','Médicaments auxiliaires',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Médicament auxiliaire: médicament utilisé pour les besoins d\'un essai clinique conformément au protocole, mais non comme médicament expérimental (article 2 du règlement européen).')
    run1.style='Paragraphe'
    
    #Ecriture du titre 8.1.2
    Titre3('8.1.2','Autres traitements / procédures',document)
    
    #Titre 8.2
    Titre2('8.2	Traitements / Procédures associé(e)s interdit(e)s',document)
    
    #Titre 8.3
    Titre2('8.3	Interactions médicamenteuses',document)
    
        #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    #document.save("Partie7.docx") 
    
  #  document.save("Partie8.docx")
    
    