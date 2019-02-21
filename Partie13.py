# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 15:13:37 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document) --> écrire en minuscule !!!
#    TexteGrisJustif(texte,document)

def Partie13(document):
    'Creation de la partie 13 du protcole de catégorie 1'
 #   document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

  #  Style(document)
    
#---------------------------------------------------------------ECRITURE
    
    
     #ecriture du premier titre 
    Titre1('13	CONTROLE ET ASSURANCE DE LA QUALITE',document)
    
     #Texte gris centré
    TexteGris('prendre contact avec la promotion interne \n pour aide a la redaction de ce chapitre',document)
    
    
    #Ecriture du 13.1  
    Titre2('13.1	Consignes pour le recueil des données',document)
    
    #Ecriture du 13.2  
    Titre2('13.2	Contrôle de la qualité',document)
    
    #Ecriture du 13.3  
    Titre2('13.3	Gestion des données',document)
    
    TexteGrisJustif('Gestion des données pour une étude e-CRF',document)
    
    TexteGrisJustif('Gestion des données pour une étude CRF papier',document)

     #Ecriture du 13.4 
    Titre2('13.4	Audits et inspections',document)
    
        #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
  #  document.save("Partie13.docx")
    
    
    
    
    
    