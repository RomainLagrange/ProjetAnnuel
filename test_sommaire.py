# -*- coding: utf-8 -*-
"""
Created on Tue Apr 16 11:47:30 2019

@author: Asuspc
"""
from docx.oxml.ns import qn 
from docx.oxml import OxmlElement 

#class som:
#    def __init__(self):
#        paragraph = self.document.add_paragraph() 


def sommaire(document):
    
    paragraph = document.add_paragraph() 
    run = paragraph.add_run() 
    fldChar = OxmlElement('w:fldChar') # creates a new element 
    fldChar.set(qn('w:fldCharType'), 'begin') # sets attribute on element 
    instrText = OxmlElement('w:instrText') 
    instrText.set(qn('xml:space'), 'preserve') # sets attribute on element 
    instrText.text = 'TOC \\o "1-20" \\h \\z \\u' # change 1-3 depending on heading levels you need 
    
    fldChar2 = OxmlElement('w:fldChar') 
    fldChar2.set(qn('w:fldCharType'), 'separate') 
    fldChar3 = OxmlElement('w:t') 
    fldChar3.text = "Right-click to update field." 
    fldChar2.append(fldChar3) 
    
    fldChar4 = OxmlElement('w:fldChar') 
    fldChar4.set(qn('w:fldCharType'), 'end') 
    
    r_element = run._r 
    r_element.append(fldChar) 
    r_element.append(instrText) 
    r_element.append(fldChar2) 
    r_element.append(fldChar4) 
    p_element = paragraph._p 