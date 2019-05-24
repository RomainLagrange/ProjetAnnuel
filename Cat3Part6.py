# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:13:09 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Mon Feb 18 11:51:47 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ROW_HEIGHT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie6(document,extract):
    'Creation de la partie 6 du protcole de catégorie 3'
   # document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

   # Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('6	DEROULEMENT DE LA RECHERCHE',document)
    
    
   # Ecriture du 6.1  
    Titre2('6.1	Calendrier de la recherche',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Durée de la période d’inclusion :' + extract['duree_inclusion'])
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Durée de suivi par participant : '+extract['duree_participation'])
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Durée totale de la recherche ')
    run1.style='Paragraphe'
    run2=p.add_run('(durée de la période d’inclusion + durée de participation) :' + extract['duree_totale_etude'])
    run2.style='Paragraphe'
    run3=p.add_run(' ')
    run3.style='Paragraphe'
    
    
    
    # Ecriture du 6.2  
    Titre2('6.2	Tableau récapitulatif du suivi participant',document)
    
       #TABLEAU
    table = document.add_table(rows = 10, cols = 5)
    table.style = 'Table Grid' #Normal
    table.alignment = WD_TABLE_ALIGNMENT.CENTER 
    row = table.rows[0] #ligne 1
    p=row.cells[1].add_paragraph('Définir les différents temps de recueil')
    p.style.font.name='Times New Roman'
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    row = table.rows[1]
    p=row.cells[0].add_paragraph('Lister les paramètres recueillis')
    p.style.font.name='Times New Roman'
    p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    
    # Ecriture du 6.3  
    Titre2('6.3	Information des personnes concernées',document)
    
   
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le médecin propose au patient de participer à cette recherche et l’informe de l\'objectif:' + extract['objectif_principal'])
    run1.style='Paragraphe'
    
#    p=document.add_paragraph()
#    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
#    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
#    p.style='List Bullet 2'
#    run1=p.add_run('de l’objectif,')
#    run1.style='Paragraphe'
#    run1.font.color.rgb = RGBColor(0x92,0xD0,0x50) 
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('du traitement informatisé des données le concernant qui seront recueillies au cours de cette recherche et lui précise également ses droits d’accès, d’opposition et de rectification à ces données.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le médecin vérifie également les ')
    run1.style='Paragraphe'
    run2=p.add_run('critères d’éligibilité :\n')
    run2.style='Paragraphe'
    run2.font.color.rgb = RGBColor(0x92,0xD0,0x50) 

    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['criteres_inclusion'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['criteres_non_inclusion'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Si la personne est d\'accord pour participer, elle donne oralement son accord et sa non-opposition est documentée dans son dossier médical. Le participant pourra, à tout moment, s’opposer à l’utilisation de ses données, dans le cadre de la recherche.')
    run1.style='Paragraphe'
    
    #Ecriture du titre 6.4
    Titre2('6.4	Visites de suivi',document)


    #Ecriture du titre 6.5
    Titre2('6.5	Visite de fin de la recherche',document)

    
    #Ecriture du titre 6.6
    Titre2('6.6	Collection d’échantillons biologiques',document)

    TexteGris('prendre contact avec la promotion interne \n pour aide a la redaction de ce chapitre', document)

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('SI APPLICABLE')
    run1.font.italic=True
	
    #Ecriture du titre6.6.1
    Titre3('6.6.1','Objectifs',document)


    #Ecriture du titre6.6.2
    Titre3('6.6.2','Description de(s) la collection(s)',document)

    
    #Ecriture du titre6.6.3
    Titre3('6.6.3','Conservation',document)

    
    #Ecriture du titre6.6.4
    Titre3('6.6.4','Devenir de la collection',document)


    
    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
  
  #  document.save("Cat3Part6.docx")   