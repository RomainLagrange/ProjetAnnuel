# -*- coding: utf-8 -*-
"""
Created on Fri Mar  1 22:35:39 2019

@author: Marion
"""

import pandas as pd
import docx
from docx.api import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
import time
from time import gmtime, strftime

def main_ansm_pb(extract):
     document = docx.Document()
     partie_A_B(document, extract)
     partie_C(document, extract)
     partie_D(document, extract)
     partie_E(document, extract)
     partie_F_G(document, extract)
     partie_H_I(document, extract)
     date = (strftime('%d-%m-%Y',time.localtime()))
     document.save("soumission_ansm_pb_"+extract['titre_abrege']+date+".docx")
     
def partie_A_B(document, extract):
    
    '''Marge de la page'''
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.95)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("\n Demande d’autorisation auprès de l’ANSM et demande d’avis a un comité de protection des personnes d'une recherche biomédicale portant sur un produit sanguin labile, un organe, un tissu d’origine humaine ou animale ou une préparation de thérapie cellulaire\n")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.bold = True
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(11)
                    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nCe formulaire est commun pour la demande d’autorisation auprès de l’ANSM et pour la demande d’avis au CPP. Certains items peuvent ne pas être applicables à tous les produits, dans ce cas ne pas en tenir compte.")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10)
    
    '''Partie ANSM/CPP'''
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nPARTIE A COMPLETER PAR L’ANSM / LE CPP\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.underline=True
    fontdebut.color.rgb = RGBColor(0x80,0x0,0x0)
                    
    table = document.add_table(rows=9, cols=3, style='Table Grid')
    table.cell(0,0).text=("RECEVABILITE")
    table.cell(0,1).text=("ÉVALUATION")
    table.cell(0,2).text=("DÉCISION/AVIS")
    table.cell(1,0).text=("Date de réception de la demande :")
    table.cell(1,1).text=("Date de passage devant le groupe d’experts : ")
    table.cell(1,2).text=("Refus d’autorisation / avis défavorable 	□")
    for i in range(2,7):
        for j in range(0,3):
            if i==2 or i==4 or i==6:
                if j==2:
                    table.cell(i,j).text=("Date :   /  /  ")
                else:
                    table.cell(i,j).text=("  /  /  ")
    table.cell(3,0).text=("Date de demande de documents manquants : ")
    table.cell(3,1).text=("Date de demande d’informations complémentaires / objections motivées : ")
    table.cell(3,2).text=("Autorisation / avis favorable 	                □")
    table.cell(5,0).text=("Date d’enregistrement du dossier complet : ")
    table.cell(5,1).text=("Date de réception des informations complémentaires / amendées : ")
    table.cell(5,2).text=("Retrait de la demande 	                □")
    table.cell(7,0).text=("Date du début d'évaluation (J0) :   /  /  ")
    table.cell(8,0).text=("Référence attribuée par l'ANSM : 	     \nRéférence attribuée par le CPP : 	     ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n<3 or n==7 or n==6 or n==13 or n==12 or n==18 or n==19:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    fontdebut.color.rgb = RGBColor(0x80,0x0,0x0)
                    if n<3:
                        fontdebut.bold = True
                    n=n+1
    for i in range(0,3):
        a=table.cell(1,i)
        b=table.cell(6,i)
        a.merge(b)
    a=table.cell(7,0)
    b=table.cell(7,2)
    a.merge(b)
    a=table.cell(8,0)
    b=table.cell(8,2)
    a.merge(b)    
    
    '''Partie A'''
    
    styles= document.styles
    style=styles.add_style('debut_page', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(10) 
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nPARTIE A COMPLETER PAR LE DEMANDEUR\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(10)
    fontdebut.underline=True
    sentence=paragraph.add_run("\nDEMANDE D’AUTORISATION À L’ANSM : 	                                                                         □\n"
                               "DEMANDE D’AVIS AU CPP	                                                                                                   □\n"
                               "\nA. IDENTIFICATION DE LA RECHERCHE BIOMÉDICALE\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(10)
    
    table = document.add_table(rows=3, cols=1, style='Table Grid')
    table.cell(0,0).text=("A.1	    Etat membre dans lequel la demande est soumise : FRANCE")
    table.cell(1,0).text=("A.2	    Numéro d’enregistrement de la recherche en France (ID RCB)  :\n" + extract['num_idrcb'] +
                          "A.3	    Titre complet de la recherche :" + extract['titre_complet'])
    table.cell(2,0).text=("A.4	    Numéro de code du protocole attribué par le promoteur, version et date  : \n" +extract['code_protocole'] +
                          "A.5	    Nom ou titre abrégé de la recherche, le cas échéant : \n" + extract['titre_abrege'] +
                          "A.6	    Numérotation ISRCTN , le cas échéant :      \n"
                          "A.7	    S'agit-il d'une resoumission de la demande ?	□ oui 	□ non\n"
                          "A.7.1	    Si oui, indiquer la lettre de resoumission  :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n<2:
                        fontdebut.bold = True
                        if n==0:
                            fontdebut.color.rgb = RGBColor(0x80,0x80,0x80)
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(2,0)
    a.merge(b)
                    
    '''Partie B'''
    paragraph=document.add_paragraph("\nB. IDENTIFICATION DU PROMOTEUR RESPONSABLE DE LA DEMANDE", style='debut_page')
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("B.1	PROMOTEUR ")
    table.cell(1,0).text=("B.1.1	Organisme : " + extract['promoteur_nom_organisme'] +
                          "\nB.1.2	Nom de la personne à contacter : " + extract['promoteur_nom_personne_contact']+
                          "\nB.1.3	Adresse :  " + extract['promoteur_adresse']+
                          "\nB.1.4	Numéro de téléphone : " +extract['promoteur_num_telephone']+
                          "\nB.1.5	Numéro de télécopie : " + extract['promoteur_num_telecopie']+  
                          "\nB.1.6	Mail :      "+extract['promoteur_courriel'])
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("B.2	REPRÉSENTANT LÉGAL DU PROMOTEUR  DANS LA COMMUNAUTÉ EUROPÉENNE POUR LA RECHERCHE CONCERNÉE")
    table.cell(1,0).text=("B.2.1	Organisme : " + extract['promoteur_UE_nom_organisme']+
                          "\nB.2.2	Nom de la personne à contacter : "  +extract['promoteur_UE_nom_personne_contact']+
                          "\nB.2.3	Adresse :  " + extract['promoteur_UE_adresse']+
                          "\nB.2.4	Numéro de téléphone : "   +extract['promoteur_UE_num_telephone'] +
                          "\nB.2.5	Numéro de télécopie : "   +extract['promoteur_UE_num_telecopie']+
                          "\nB.2.6	Mail :   "+extract['promoteur_UE_courriel'])
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("B.3	STATUT DU PROMOTEUR")
    table.cell(1,0).text=("B.3.1    Privé (commercial)                                                                      	□\n"
                          "B.3.2    Institutionnel  (non commercial)                                              	□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
                    
def partie_C(document, extract):
    
    '''Partie C'''
    paragraph=document.add_paragraph("\nC. IDENTIFICATION DU DEMANDEUR (cocher les cases appropriées)\n", style='debut_page')
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("C.1	DEMANDE AUPRÈS DE L’ANSM	                                                                      □")
    table.cell(1,0).text=("C.1.1       Promoteur                                                                                                                   □\n"
                          "C.1.2       Représentant légal du promoteur                                                                               □\n"
                          "C.1.3	Personne ou organisme délégué par le promoteur pour soumettre la demande	     □\n"
                          "C.1.4 	Préciser ci-après les informations relatives au demandeur, même si elles figurent ailleurs dans le formulaire : Si promoteur, partie B1, si représentant légal du promoteur, partie B2\n"
                          "C.1.4.1 	Organisme :  "+extract['demandeur_nom_organisme']+"    \n"
                          "C.1.4.2 	Nom de la personne à contacter :  "+extract['demandeur_nom_personne_contact']+"    \n"
                          "C.1.4.3 	Adresse :  "+extract['demandeur_UE_adresse']+"    \n"
                          "C.1.4.4 	Numéro de téléphone : "+extract['demandeur_UE_num_telephone']+"     \n"
                          "C.1.4.5 	Numéro de télécopie : "+extract['demandeur_UE_num_telecopie']+"     \n"
                          "C.1.4.6	Mail : "+extract['demandeur_UE_courriel'])
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("C.2	DEMANDE AUPRÈS DU CPP                                                                                      □")
    table.cell(1,0).text=("C.2.1       Promoteur                                                                                                                   □\n"
                          "C.2.2       Représentant légal du promoteur                                                                               □\n"
                          "C.2.3	Personne ou organisme délégué par le promoteur pour soumettre la demande	     □\n"
                          "C.2.4 	Préciser ci-après les informations relatives au demandeur, même si elles figurent ailleurs dans le formulaire : Si promoteur, partie B1, si représentant légal du promoteur, partie B2\n"
                          "C.2.4.1	Organisme :  "+extract['demandeur_nom_organisme']+"   \n"
                          "C.2.4.2	Nom de la personne à contacter :  "+extract['demandeur_nom_personne_contact']+"    \n"
                          "C.2.4.3	Adresse :  "+extract['demandeur_UE_adresse']+"     \n"
                          "C.2.4.4	Numéro de téléphone :   "+extract['demandeur_UE_num_telephone']+"    \n"
                          "C.2.4.5	Numéro de télécopie : "+extract['demandeur_UE_num_telecopie']+"\n"     
                          "C.2.4.6	Mail :  "+extract['demandeur_UE_courriel'])
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
    document.add_page_break()
    
def partie_D(document, extract):
    
    '''Partie D1'''
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nD. DONNEES RELATIVES A CHAQUE PRODUIT SUR LEQUEL PORTE LA RECHERCHE")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(10) 
    sentence=paragraph.add_run("\nLes informations concernant chaque produit doivent être indiquées dans cette section :\n"
                               "       -	pour chaque produit sur lequel porte la recherche\n"
                               "       -	pour chaque produit utilisé comme comparateur \n"
                               "       -	et pour chaque placebo, le cas échéant.\n"
                               "Si la recherche biomédicale porte sur plusieurs produits, répéter cette section, en attribuant à chaque produit un numéro d’ordre à l'item D.1.1. \n"
                               "Si le produit est une association, les informations doivent être données pour chaque substance active ou produit concerné.\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10) 
    
    table = document.add_table(rows=4, cols=1, style='Table Grid')
    table.cell(0,0).text=("D.1	 IDENTIFICATION DU PRODUIT SUR LEQUEL PORTE LA RECHERCHE ")
    table.cell(1,0).text=("Indiquer ci-dessous quel produit est décrit dans cette section D. Le cas échéant, répéter cette section autant de fois qu'il y a de produits utilisés dans la recherche (numéroter chaque produit de 1 à n)")
    table.cell(2,0).text=("D.1.1         Cette section concerne le produit numéro :	     \n"
                          "D.1.2         Produit étudié                                                                       	□\n"
                          "D.1.3         Produit utilisé comme comparateur                                   	□")
    table.cell(3,0).text=("Pour le placebo, aller directement en section D.7")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==2:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,0)
    a.merge(b)
    a=table.cell(2,0)
    b=table.cell(3,0)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    '''Partie D2'''
    
    table = document.add_table(rows=10, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.2	   STATUT DU PRODUIT SUR LEQUEL PORTE LA RECHERCHE ")
    table.cell(1,0).text=("D.2.1	   Le produit utilisé dans la recherche dispose-t-il d'une autorisation en France ou est-il enregistré en France  ?")
    table.cell(2,0).text=("D.2.1.1	   Si oui en D.2.1, préciser pour le produit utilisé dans la recherche :")
    table.cell(3,0).text=("D.2.1.1.1    Nom du produit autorisé ou enregistré ou nom commercial, le cas échéant :      ")
    table.cell(4,0).text=("D.2.1.1.2    Nom du titulaire de l’autorisation :      ")
    table.cell(5,0).text=("D.2.1.1.3    Numéro d’autorisation ou d’enregistrement :      ")
    table.cell(6,0).text=("D.2.1.1.4    Le produit sur lequel porte la recherche est-il modifié par rapport à son autorisation ? ")
    table.cell(7,0).text=("D.2.1.1.4.1 Si oui, veuillez préciser :      ")
    table.cell(8,0).text=("D.2.1.2       Le produit dispose-t-il d’une autorisation ou d’un enregistrement dans un autre pays ?")
    table.cell(9,0).text=("D.2.1.2.1    Si oui, veuillez préciser le pays et le nom de l’autorité qui a autorisé le produit :      ")
    for i in range(1,10):
        if i==1 or i==8:
            table.cell(i,5).text=("□ oui  □ non")
        elif i==6:
            table.cell(i,5).text=("□ oui  □ non\n")
        else:
            table.cell(i,5).text=(" ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2 or n==12 or n==16:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(9,4)
    a.merge(b)
    a=table.cell(1,5)
    b=table.cell(9,5)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.2.2	 Dossier du produit sur lequel porte la recherche :")
    table.cell(1,0).text=("D.2.2.1     Dossier complet\n"
                          "D.2.2.2	 Dossier simplifié ")
    table.cell(0,5).text=(" ")
    table.cell(1,5).text=("□ oui  □ non\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==3:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,4)
    a.merge(b)
    a=table.cell(0,5)
    b=table.cell(1,5)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.2.3	L’utilisation du produit a-t-elle déjà été autorisée dans le cadre d'une recherche biomédicale précédente conduite par le promoteur dans la Communauté européenne ? ")
    table.cell(1,0).text=("D.2.3.1	Si oui, préciser dans quel(s) État(s) membre(s) et par quelle autorité :      ")
    table.cell(0,5).text=("\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==1:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,4)
    a.merge(b)
    a=table.cell(0,5)
    b=table.cell(1,5)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.2.4	Le produit est-il désigné, dans l’indication étudiée dans la recherche, comme un médicament orphelin dans la Communauté européenne ?")
    table.cell(1,0).text=("D.2.4.1	Si oui, indiquer le numéro de désignation du médicament orphelin  :      ")
    table.cell(0,5).text=("□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==1:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,4)
    a.merge(b)
    a=table.cell(0,5)
    b=table.cell(1,5)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.2.5	Un avis scientifique a-t-il été rendu sur le produit dans le cadre de cette recherche ?")
    table.cell(1,0).text=("D.2.5.1	Si oui en D.2.5, veuillez préciser qui a rendu l'avis et en joindre une copie à votre dossier :\n"
                          "D.2.5.1.1 Avis du CHMP  ? \n"
                          "D.2.5.1.2 Avis d'une autorité compétente d'un Etat membre ?")
    table.cell(0,5).text=("□ oui  □ non")
    table.cell(1,5).text=("\n\n\n□ oui  □ non\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==1 or n==3:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,4)
    a.merge(b)
    a=table.cell(0,5)
    b=table.cell(1,5)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=4, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.3	 DESCRIPTION DU PRODUIT SUR LEQUEL PORTE LA RECHERCHE")
    table.cell(1,0).text=("D.3.1	  Nom du produit, le cas échéant  :      \n"
                          "D.3.2	  Nom de code, le cas échéant  :    \n"  
                          "D.3.3	  Code ATC , si enregistré officiellement :     ")
    table.cell(2,0).text=("D.3.4	  Type de produit :")
    table.cell(3,0).text=("Le produit est-il :\n"
                          "D.3.4.1	   Une préparation de thérapie cellulaire\n"
                          "D.3.4.2	   Un tissu\n"
                          "D.3.4.3	   Un organe ou un tissu composite\n"
                          "D.3.4.4	   Un produit sanguin labile\n"
                          "D.3.4.5	   Autre\n"
                          "D.3.4.5.1	 Si autre, préciser :      ")
    table.cell(2,5).text=("\n\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==3:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1 or n==2:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(1,5)
    a.merge(b)
    a=table.cell(2,0)
    b=table.cell(3,4)
    a.merge(b)
    a=table.cell(2,5)
    b=table.cell(3,5)
    a.merge(b)
    
    '''Partie D4'''
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=8, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.4	   PRÉPARATION DE THÉRAPIE CELLULAIRE")
    table.cell(1,0).text=("D.4.1       Origine des cellules")
    table.cell(2,0).text=("D.4.1.1    Autologue\n"
                          "D.4.1.2    Allogénique\n"
                          "D.4.1.3	   Xénogénique\n"
                    	  "D.4.1.3.1 Si oui, préciser les espèces d’origine :      ")
    table.cell(2,5).text=("\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non")
    table.cell(3,0).text=("D.4.2       Type de cellules")
    table.cell(4,0).text=("D.4.2.1    Cellules souches\n"
                          "D.4.2.2	   Cellules différenciées	\n"
                          "D.4.2.2.1  Préciser le type de cellules (exemple : kératinocytes, fibroblastes, chondrocytes…) :  \n"    
                          "D.4.2.3	    Les cellules sont-elles associées à une matrice ou un support	\n"
                          "D.4.2.3.1  Si oui, préciser :     \n"
                          "D.4.2.4	    Autre	\n"
                          "D.4.2.4.1  Si oui, préciser :      ")
    table.cell(4,5).text=("\n□ oui  □ non\n□ oui  □ non\n\n\n□ oui  □ non\n\n□ oui  □ non\n")
    table.cell(5,0).text=("D.4.3	   Forme pharmaceutique :      \n"
                          "D.4.4	   Durée maximale du traitement pour une personne prévue par le protocole :      \n"
                          "D.4.5	   Dose maximale permise (préciser : dose journalière ou dose cumulée ; unités et voie d'administration) :      \n"
                          "D.4.6	   Voie d’administration :   \n"   
                          "D.4.7	   Nom de chaque substance active :      ")
    table.cell(6,0).text=("D.4.8	   Dosage (préciser tous les dosages utilisés)")
    table.cell(7,0).text=("D.4.8.1	   Unité de concentration :   \n"   
                          "D.4.8.2	   Type de concentration (« nombre exact », « intervalle », « plus que » ou « jusqu’à ») :      \n"
                          "D.4.8.3	   Concentration (nombre) :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==3 or n==6:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1 or n==4 or n==7 or n==8:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(2,4)
    a.merge(b)
    a=table.cell(1,5)
    b=table.cell(2,5)
    a.merge(b)
    a=table.cell(3,0)
    b=table.cell(4,4)
    a.merge(b)
    a=table.cell(3,5)
    b=table.cell(4,5)
    a.merge(b)
    a=table.cell(5,0)
    b=table.cell(5,5)
    a.merge(b)
    a=table.cell(6,0)
    b=table.cell(7,5)
    a.merge(b)
    
    document.add_page_break()
    
    '''Partie D5'''
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=5, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.5	   TISSU OU ORGANE")
    table.cell(1,0).text=("D.5.1        Origine du tissu, du tissu composite ou de l’organe")
    table.cell(2,0).text=("D.5.1.1     Autologue\n"
                          "D.5.1.2     Allogénique\n"
                          "D.5.1.3	    Xénogénique	\n"
                          "D.5.1.3.1 Préciser les espèces d’origine :      ")
    table.cell(3,0).text=("D.5.2	  Type de tissu ou d’organe")
    table.cell(4,0).text=("D.5.2.1	  Tissu	\n"
                          "D.5.2.1.1 Préciser le type de tissu (cornée, peau, os, …) :     \n" 
                          "D.5.2.2	  Tissu composite	\n"
                          "D.5.2.2.1 Préciser :      \n"
                          "D.5.2.3	   Organe	\n"
                          "D.5.2.3.1 Préciser :      ")
    table.cell(1,5).text=("\n□ oui  □ non\n□ oui  □ non\n□ oui  □ non\n")
    table.cell(3,5).text=("\n□ oui  □ non\n\n□ oui  □ non\n\n□ oui  □ non\n")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2 or n==5:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1 or n==4:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(2,4)
    a.merge(b)
    a=table.cell(1,5)
    b=table.cell(2,5)
    a.merge(b)
    a=table.cell(3,0)
    b=table.cell(4,4)
    a.merge(b)
    a=table.cell(3,5)
    b=table.cell(4,5)
    a.merge(b)
    
    '''Partie D6'''
    
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=11, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.6	PRODUIT SANGUIN LABILE ")
    table.cell(1,0).text=("D.6.1      Origine du produit sanguin labile")
    table.cell(2,0).text=("D.6.1.1   Autologue\n"
                          "D.6.1.2	  Homologue")
    table.cell(2,5).text=("\n□ oui  □ non\n□ oui  □ non")
    table.cell(3,0).text=("D.6.2	   Type de produit sanguin labile")
    table.cell(4,0).text=("D.6.2.1	   Concentrés de Plaquettes	\n"
                          "D.6.2.1.1 Préciser issus de sang total ou issus d’aphérèse :      \n"
                          "D.6.2.2	   Plasma	\n"
                          "D.6.2.2.1 Préciser issu de sang total ou issu d’aphérèse :      \n"
                          "D.6.2.3	   Concentrés de Globules rouges	\n"
                          "D.6.2.3.1  Préciser issus de sang total ou issus d’aphérèse :    \n"  
                          "D.6.2.4	  Sang total	\n"
                          "D.6.2.5	   Autre	\n"
                          "D.6.2.5.1 Si autre, préciser :      ")
    table.cell(4,5).text=("\n□ oui  □ non\n\n□ oui  □ non\n\n□ oui  □ non\n\n□ oui  □ non\n□ oui  □ non")
    table.cell(5,0).text=("D.6.3	   Le produit sanguin labile est-il soumis à un procédé d’inactivation")
    table.cell(5,5).text=("□ oui  □ non")
    table.cell(6,0).text=("D.6.3.1	   Si oui, préciser :      ")
    table.cell(7,0).text=("D.6.4	   Durée maximale du traitement pour une personne prévue par le protocole :      \n"
                          "D.6.5	   Dose maximale permise (préciser : dose journalière ou dose cumulée ; unités) :    \n"  
                          "D.6.6	   Dosage (préciser tous les dosages utilisés) :")
    table.cell(8,0).text=("D.6.6.1	   Unité de concentration :      \n"
                          "D.6.6.2	   Type de concentration (« nombre exact », « intervalle », « plus que » ou « jusqu’à ») :      \n"
                          "D.6.6.3	   Concentration (nombre) :      \n")
    table.cell(9,0).text=("D.6.7	   Nom du(des) dispositif(s) médical(aux) associé(s) au produit sanguin labile :      ")
    
    table.cell(10,0).text=("D.6.7.1	   Ce(s) dispositif(s) médical(aux) dispose(nt) d’un marquage CE\n"
                           "D.6.7.2	   Si oui, est-il (sont-ils) utilisé(s) dans la (les) même(s) indication(s) que celle(s) du (de leurs) marquage(s) CE ?")
    table.cell(10,5).text=("\n□ oui  □ non\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==3 or n==6 or n==8 or n==14:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1 or n==4 or n==7 or n==10 or n==12:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(2,4)
    a.merge(b)
    a=table.cell(1,5)
    b=table.cell(2,5)
    a.merge(b)
    a=table.cell(3,0)
    b=table.cell(4,4)
    a.merge(b)
    a=table.cell(3,5)
    b=table.cell(4,5)
    a.merge(b)
    a=table.cell(5,0)
    b=table.cell(6,4)
    a.merge(b)
    a=table.cell(5,5)
    b=table.cell(6,5)
    a.merge(b)
    a=table.cell(7,0)
    b=table.cell(8,4)
    a.merge(b)
    a=table.cell(7,5)
    b=table.cell(8,5)
    a.merge(b)
    a=table.cell(9,0)
    b=table.cell(10,4)
    a.merge(b)
    a=table.cell(9,5)
    b=table.cell(10,5)
    a.merge(b)
    
    '''Partie D7'''
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=3, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.7  DONNEES RELATIVES AU PLACEBO (répéter la section autant de fois que nécessaire, le cas échéant)")
    table.cell(1,0).text=("D.7.1	  Un placebo est-il utilisé ?\n"
                          "D.7.2	  Cette section concerne le placebo numéro : (     )\n"
                          "D.7.3	  Forme pharmaceutique :      \n"
                          "D.7.4	  Voie d’administration :      \n"
                          "D.7.5	  De quel produit est-ce le placebo ? Préciser le numéro du produit, tel qu'indiqué en D.1 : (     )")
    table.cell(2,0).text=("D.7.5.1	  Composition, hormis la ou les substances actives :      \n"
                          "D.7.5.2	  Est-elle identique à celle du produit étudié ?	\n"
                          "D.7.5.2.1  Si non, préciser les principaux composants :      \n")
    table.cell(1,5).text=("□ oui  □ non\n")
    table.cell(2,5).text=("\n\n\n\n\n□ oui  □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2 or n==4:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(2,4)
    a.merge(b)
    a=table.cell(1,5)
    b=table.cell(2,5)
    a.merge(b)
    
    '''Partie D8'''
    
    paragraph=document.add_paragraph()
    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=3, cols=1, style='Table Grid')
    table.cell(0,0).text=("D.8	DONNEES SUR LES ETABLISSEMENTS DE PRELEVEMENT, PREPARATION, CONSERVATION, LIBERATION, ADMINISTRATION ")
    table.cell(1,0).text=("D.8.1	  PRODUIT SANGUIN LABILE :\n"
                          "D.8.1.1	  Établissement où le produit est libéré :")
    table.cell(2,0).text=("D.8.1.1.1 Nom de l’établissement, code de l’établissement, le cas échéant :      \n"
                          "D.8.1.1.2 Adresse :     \n" 
                          "D.8.1.1.3 Indiquer le numéro d’autorisation, ou la date d’agrément : \n"     
                          "D.8.1.1.4 Si pas d’autorisation, préciser les motifs :       ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(1,0)
    b=table.cell(2,0)
    a.merge(b)

    paragraph=document.add_paragraph()
    
    table = document.add_table(rows=11, cols=6, style='Table Grid')
    table.cell(0,0).text=("D.8.2       PRÉPARATIONS DE THÉRAPIE CELLULAIRE, TISSUS, ORGANES :")
    table.cell(1,0).text=("D.8.2.1	  Sites de prélèvement")
    table.cell(2,0).text=("D.8.2.1.1 Nom de l’établissement, code de l’établissement, le cas échéant :      \n"
                          "D.8.2.1.2 Adresse :      ")
    table.cell(3,0).text=("D.8.2.2	  Sites de préparation, conservation, libération")
    table.cell(4,0).text=("D.8.2.2.1 Nom de l’établissement, code de l’établissement, le cas échéant :  \n"    
                          "D.8.2.2.2 Adresse :    \n"  
                          "D.8.2.2.3 Indiquer le numéro d’autorisation, ou la date d’agrément  :      \n"
                          "D.8.2.2.4 Si pas d’autorisation, préciser les motifs :       \n"
                          "D.8.2.2.5 Pour les préparations de thérapie cellulaire : le local où est réalisée la préparation du produit a-t-il déjà fait l’objet d’une inspection par l’Afssaps lors d’une autre recherche biomédicale ?")
    table.cell(4,5).text=("\n\n\n\n\n\n\n\n\n□ oui  □ non")
    table.cell(5,0).text=("D.8.2.3	  Liste des sous-traitants")
    table.cell(6,0).text=("D.8.2.3.1 Étape réalisée :   \n"  
                          "D.8.2.3.2 Nom :     \n"
                          "D.8.2.3.3 Adresse :      ")
    table.cell(7,0).text=("D.8.2.4	  Sites d’administration")
    table.cell(8,0).text=("D.8.2.4.1 Nom de l’établissement, code de l’établissement, le cas échéant :  \n"    
                          "D.8.2.4.2 Adresse :      ")
    table.cell(9,0).text=("D.8.2.5	  Établissement importateur, le cas échéant ")
    table.cell(10,0).text=("D.8.2.5.1 Nom de l’établissement :      \n"
                           "D.8.2.5.2 Adresse :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==5:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==1 or n==3 or n==6 or n==8 or n==10:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(10,4)
    a.merge(b)
    a=table.cell(0,5)
    b=table.cell(10,5)
    a.merge(b)
    
    
    
def partie_E(document, extract):

    document.add_page_break()
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nE. INFORMATIONS GENERALES RELATIVES A LA RECHERCHE BIOMEDICALE")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(10) 
    sentence=paragraph.add_run("\nCette section est destinée à fournir des informations concernant les objectifs, domaine et méthodologie de la recherche. Si le protocole prévoit la réalisation d'une sous-étude en France, indiquer les informations relatives à cette sous-étude en section E.2.3. \n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10) 
    
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("E.1	CONDITION MEDICALE OU PATHOLOGIE ETUDIEE")
    table.cell(1,0).text=("E.1.1	Préciser la ou les conditions médicales / pathologies étudiées  (texte libre) :      \n"
                          "E.1.2	Version MedDRA, niveau, terme et classification  (répéter autant de fois que nécessaire) :    \n"  
                          "E.1.3	L'une des conditions médicales étudiées est-elle une maladie rare  ?	            □ oui   □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie E2'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("E.2	OBJECTIF(S) DE LA RECHERCHE")
    table.cell(1,0).text=("E.2.1	Objectif principal : "+extract['objectif_principal']+"\n"
                          "E.2.2	Objectifs secondaires : "+extract['objectif_secondaire']+"\n"
                          "E.2.3	Une sous-étude est-elle prévue ?	                                                        □ oui   □ non\n"
                          "E.2.3.1	Si oui, préciser le titre complet, la date et la version de chaque sous-étude et leurs objectifs :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie E3'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("E.3	PRINCIPAUX CRITERES D’INCLUSION (énumérer les plus importants)")
    table.cell(1,0).text=(extract['critere_inclusion_courte'])
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
    '''Partie E4'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("E.4	PRINCIPAUX CRITERES DE NON INCLUSION (énumérer les plus importants)")
    table.cell(1,0).text=(extract['critere_non_inclusion_courte'])
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
                    
    '''Partie E5'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=3, cols=1, style='Table Grid')
    table.cell(0,0).text=("E.5	CRITERE(S) D’EVALUATION PRINCIPAL(AUX)")
    table.cell(1,0).text=("Critère de jugement principal :"+extract['critere_jugement_principal_longue'])
    table.cell(1,0).text=("Critères de jugement secondaires :"+extract['critere_jugement_secondaire_courte'])
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie E6'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=3, style='Table Grid')
    table.cell(0,0).text=("E.6	 DOMAINE(S) DE LA RECHERCHE – Cocher la ou les cases appropriées")
    table.cell(1,0).text=("E.6.1     Diagnostic\n"
                          "E.6.2     Prophylaxie\n"
                          "E.6.3     Thérapeutique\n"
                          "E.6.4     Sécurité\n"
                          "E.6.5     Efficacité\n"
                          "E.6.6     Pharmacocinétique\n"
                          "E.6.7     Pharmacodynamie\n"
                          "E.6.8     Bioéquivalence\n"
                          "E.6.9     Dose-effet\n"
                          "E.6.10    Pharmaco-économie\n"
                          "E.6.11    Autre\n"
                          "E.6.11.1	Si autre, préciser :      ")
    table.cell(1,2).text=("□ \n□ \n□ \n□ \n□ \n□ \n□ \n□ \n□ \n□ \n□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,2)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(1,1)
    a.merge(b)
    
    document.add_page_break()
    
    '''Partie E7'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=3, style='Table Grid')
    table.cell(0,0).text=("E.7	  TYPE DE RECHERCHE  ET PHASE")
    table.cell(1,0).text=("E.7.1	  Pharmacologie humaine (Phase I)\n"
                          "Il s'agit de :\n"
                          "E.7.1.1    La première administration ou première greffe à l’homme\n"
                          "E.7.1.2    Une étude de bioéquivalence\n"
                          "E.7.1.3	  Autre\n"
                          "E.7.1.3.1 Si autre, préciser :      \n"
                          "E.7.2      Essai thérapeutique exploratoire (Phase II)\n"
                          "E.7.3      Essai thérapeutique de confirmation (Phase III)\n"
                          "E.7.4	  Essai thérapeutique conformément à l’autorisation (Phase IV)")
    table.cell(1,2).text=("□ \n\n□ \n□ \n□ \n\n□ \n□ \n□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,2)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(1,1)
    a.merge(b)
                    
    '''Partie E8'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=4, cols=6, style='Table Grid')
    table.cell(0,0).text=("E.8	METHODOLOGIE DE LA RECHERCHE")
    table.cell(1,0).text=("E.8.1	  Comparatif	\n"
                          "Si oui, préciser :\n"
                          "E.8.1.1    Tirage au sort\n"
                          "E.8.1.2    Ouvert\n"
                          "E.8.1.3    Simple insu\n"
                          "E.8.1.4    Double insu\n"
                          "E.8.1.5    A groupes parallèles\n"
                          "E.8.1.6    Plan croisé\n"
                          "E.8.1.7	 Autre	\n"
                          "E.8.1.7.1 Si autre, préciser :      \n"
                          "E.8.2	  Si comparatif, préciser le comparateur utilisé\n"
                          "E.8.2.1    Autre(s) produit(s)\n"
                          "E.8.2.2    Médicament\n"
                          "E.8.2.3    Placebo, le cas échéant\n"
                          "E.8.2.4	  Autre\n"
                          "E.8.2.4.1 Si autre, préciser :      ")
    table.cell(1,5).text=("□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n\n\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non")
    table.cell(2,0).text=("E.8.3      La recherche est-elle monocentrique (voir aussi section G) ?\n"
                          "E.8.4	  La recherche est-elle multicentrique (voir aussi section G) ?\n"
                          "E.8.4.1	  Nombre prévu de lieux de recherche en France :    \n"  
                          "E.8.5	  Est-il prévu de mener la recherche dans plusieurs états membres ?\n"
                          "E.8.5.1	  Nombre prévu de lieux de recherche dans la Communauté européenne :  \n"    
                          "E.8.6      Est-il prévu de mener la recherche dans des pays tiers ?\n"
                          "E.8.7	  Un comité de surveillance indépendant a-t-il été constitué ?")
    table.cell(2,5).text=("□ oui   □ non\n□ oui   □ non\n\n□ oui   □ non\n\n□ oui   □ non\n□ oui   □ non")
    table.cell(3,0).text=("E.8.8	  Définition de la fin de la recherche, et justification si celle-ci ne correspond pas à la date de la dernière visite de la dernière personne participant à la recherche   :    \n"  
                          "E.8.9      Estimation initiale de la durée de la recherche  (en années, mois et jours) : \n"+extract['duree_totale_etude']+"\n"
                          "E.8.9.1   en France : 	      années       mois       jours\n"
                          "E.8.9.2	 dans tous les pays concernés par l’essai : 	      années       mois       jours")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2 or n==4:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(1,4)
    a.merge(b)
    a=table.cell(2,0)
    b=table.cell(2,4)
    a.merge(b)
    a=table.cell(3,0)
    b=table.cell(3,4)
    a.merge(b)
    
    '''Partie E9'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=1, cols=6, style='Table Grid')
    table.cell(0,0).text=("E.9	La constitution d’une (ou plusieurs) collections d’échantillons biologiques est-elle prévue ?")
    table.cell(0,5).text=("□ oui   □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==1:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,4)
    a.merge(b)
    
    document.add_page_break()
    
def partie_F_G(document, extract):
    
    '''Partie F'''
    paragraph=document.add_paragraph("\nF. PERSONNES PARTICIPANT A LA RECHERCHE\n", style='debut_page')
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("F.1	TRANCHE D'ÂGE ÉTUDIÉE")
    table.cell(1,0).text=("F.1.1	Moins de 18 ans\n"
                          "Si oui, préciser :\n"
                          "F.1.1.1  In Utero\n"
                          "F.1.1.2  Nouveaux-nés prématurés (jusqu’à l’âge gestationnel ≤ 37 semaines)\n"
                          "F.1.1.3  Nouveau-nés (0-27 jours)\n"
                          "F.1.1.4  Nourrissons (28 jours - 23 mois)\n"
                          "F.1.1.5  Enfants (2-11 ans)\n"
                          "F.1.1.6  Adolescents (12-17 ans)\n"
                          "F.1.2     De 18 à 65 ans\n"
                          "F.1.3	Plus de 65 ans")
    table.cell(1,5).text=("□ oui   □ non\n\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(1,4)
    a.merge(b)
    
    '''Partie F2'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("F.2	SEXE")
    table.cell(1,0).text=("F.2.1     Femmes    □\n"
                          "F.2.2	Hommes    □")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1

    '''Partie F3'''
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("F.3	GROUPE DE PERSONNES PARTICIPANT A LA RECHERCHE")
    table.cell(1,0).text=("F.3.1         Volontaires sains\n"
                          "F.3.2         Volontaires malades\n"
                          "F.3.3         Populations particulières\n"
                          "F.3.3.1      Femmes en âge de procréer\n"
                          "F.3.3.2      Femmes en âge de procréer utilisant un moyen de contraception\n"
                          "F.3.3.3      Femmes enceintes\n"
                          "F.3.3.4      Femmes allaitantes\n"
                          "F.3.3.5      Personnes en situation d’urgence\n"
                          "F.3.3.6	    Personnes incapables de donner personnellement leur consentement	\n"
                          "F.3.3.6.1   Si oui, préciser :   \n"   
                          "F.3.3.7	    Autre\n"
                          "F.3.3.7.1   Si oui, préciser :      ")
    table.cell(1,5).text=("□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n\n□ oui   □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(1,4)
    a.merge(b)
    
    '''Partie F4'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("F.4	NOMBRE PREVU DE PERSONNES A INCLURE")
    table.cell(1,0).text=("F.4.1	En France\n"
                          "F.4.2	En cas de recherche  menée dans plusieurs pays :\n"
                          "F.4.2.1	Dans la Communauté européenne\n"
                          "F.4.2.2	Pour l’ensemble de la recherche")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie F5'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("F.5	TRAITEMENT(S) OU SOIN(S) PREVU(S) POUR LES PERSONNES A LA FIN DE LEUR PARTICIPATION A LA RECHERCHE  (texte libre) :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    document.add_page_break()
    
    '''Partie G'''
    paragraph=document.add_paragraph("\nG. LIEUX DE RECHERCHES ENVISAGES / INVESTIGATEURS EN FRANCE\n", style='debut_page')
    
    '''Partie G1'''
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("G.1	INVESTIGATEUR COORDONNATEUR (si recherche multicentrique) et investigateur principal (si recherche monocentrique) ")
    table.cell(1,0).text=("G.1.1	Nom : "+extract['investigateur_coordinateur_nom']+"     \n"
                          "G.1.3	Prénom : "+extract['investigateur_coordinateur_prenom']+"     \n"
                          "G.1.4	Qualification, spécialité : "+extract['investigateur_coordinateur_qualification']+"     \n"
                          "G.1.5	Adresse professionnelle : "+extract['investigateur_coordinateur_adresse_professionnelle'])
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie G2'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("G.2	INVESTIGATEURS PRINCIPAUX (si recherche multicentrique ; répéter cette section autant de fois que nécessaire) ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    for i in range(len(extract['autre_investigateur_nom'])):
        table = document.add_table(rows=1, cols=1, style='Table Grid')
        table.cell(0,0).text=("G.2.1	Nom :  "+extract['autre_investigateur_nom'][i]+"    \n"
                              "G.2.3	Prénom :  "+extract['autre_investigateur_prenom'][i]+"   \n" 
                              "G.2.4	Qualification, spécialité :  "+extract['autre_investigateur_qualification'][i]+"    \n"
                              "G.2.5	Adresse professionnelle :  "+extract['autre_investigateur_adresse_professionnelle'][i])
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        fontdebut = run.font
                        fontdebut.name = 'Arial'
                        fontdebut.size = docx.shared.Pt(10)
    
    '''Partie G3'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("G.3	PLATEAU TECHNIQUE UTILISE AU COURS DE LA RECHERCHE\nLaboratoire ou autre plateau technique où sont effectuées de façon centralisée les mesures ou évaluations des paramètres ou critères principaux étudiés dans la recherche (à compléter pour chaque organisme, répéter la section si nécessaire)")
    table.cell(1,0).text=("G.3.1	Organisme :  "+extract['plateau_technique_organisme']+"    \n"
                          "G.3.2	Nom de la personne à contacter : "+extract['plateau_technique_personne_contact']+"     \n"
                          "G.3.3	Adresse :   "+extract['plateau_technique_adresse']+"   \n"
                          "G.3.4	Numéro de téléphone :  "+extract['plateau_technique_num_telephone']+"    \n"
                          "G.3.5	Tâches confiées :  "+extract['plateau_technique_taches_confiees'])
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    '''Partie G4'''
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    table.cell(0,0).text=("G.4	PRESTATAIRE A QUI LE PROMOTEUR A CONFIE CERTAINES OBLIGATIONS ET FONCTIONS AFFERENTES A LA RECHERCHE (à compléter pour chaque organisme, répéter la section si nécessaire)")
    table.cell(1,0).text=("G.4.1	    Le promoteur a-t-il confié en partie ou en totalité des obligations et des fonctions majeures lui incombant au titre de la recherche à un autre organisme ou à un tiers ?\n"
                          "Préciser pour chaque organisme :\n"
                          "G.4.1.1	    Organisme :      \n"
                          "G.4.1.2	    Nom de la personne à contacter :    \n"  
                          "G.4.1.3	    Adresse :      \n"
                          "G.4.1.4	    Numéro de téléphone :      \n"
                          "Obligations / fonctions confiées :\n"
                          "G.4.1.5     Ensemble des tâches du promoteur\n"
                          "G.4.1.6     Monitoring\n"
                          "G.4.1.7     Réglementaire (ex : préparation des dossiers soumis à l'Afssaps et au CPP)\n"
                          "G.4.1.8     Recrutement des investigateurs\n"
                          "G.4.1.9     IVRS  - tirage au sort du traitement\n"
                          "G.4.1.10   Gestion/collecte des données\n"
                          "G.4.1.11   Saisie électronique des données\n"
                          "G.4.1.12   Déclaration des effets indésirables graves et/ou incidents graves \n"
                          "G.4.1.13   Audit de l'assurance qualité\n"
                          "G.4.1.14   Analyses statistiques\n"
                          "G.4.1.15   Rédaction médicale\n"
                          "G.4.1.16    Autres tâches confiées\n"
                          "G.4.1.16.1  Si oui, veuillez préciser :      ")
    table.cell(1,5).text=("□ oui   □ non\n\n\n\n\n\n\n\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non\n□ oui   □ non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==2:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,5)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(1,4)
    a.merge(b)
    
    document.add_page_break()
    
def partie_H_I(document, extract):
    
    '''Partie H'''
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nH. AFSSAPS / CPP CONCERNE PAR LA DEMANDE")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(10) 
    sentence=paragraph.add_run("\nSi cette demande est adressée à l’Afssaps remplir les informations ci-dessous :")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10) 
    
    '''Partie H'''
    paragraph=document.add_paragraph()
    table = document.add_table(rows=4, cols=3, style='Table Grid')
    table.cell(0,0).text=("H.1	   CPP CONCERNÉ OU PRESSENTI")
    table.cell(1,0).text=("H.1.1	   Nom et adresse : "+extract['CPP'])
    table.cell(2,0).text=("H.2	   AVIS DU CPP :")
    table.cell(3,0).text=("H.2.1       A demander\n"
                          "H.2.2	   En cours\n"
                          "H.2.2.1	   Si en cours, préciser la date de soumission :   /  /    \n"
                          "H.2.3	   Obtenu(e)	\n"
                          "	   Si obtenu(e), préciser :	\n"
                          "H.2.3.1     Date de l’avis :   /  /    \n"
                          "H.2.3.2     Avis favorable\n"
                          "H.2.3.3	   Avis défavorable	\n"
                          "	   Si avis défavorable, préciser :\n"
                          "H.2.3.3.1 Les motifs :      \n"
                          "H.2.3.3.2 La date éventuelle envisagée de resoumission de la demande :      ")
    table.cell(3,2).text=("\n□\n□\n\n□\n\n\n□\n□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==2:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,2)
    a.merge(b)
    a=table.cell(2,0)
    b=table.cell(3,1)
    a.merge(b)
    a=table.cell(2,2)
    b=table.cell(3,2)
    a.merge(b)
    
    '''Partie H3'''
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("\nSi cette demande est adressée au CPP remplir les informations ci-dessous :")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(10) 
    
    table = document.add_table(rows=2, cols=3, style='Table Grid')
    table.cell(0,0).text=("H.3	AUTORISATION AUPRES DE L’ANSM :\n")
    table.cell(1,0).text=("	Agence française de sécurité sanitaire des produits de santé\n"
                          "	Direction des thérapies innovantes, des produits issus du corps humain et des vaccins\n"
                          "	143 / 147 BD ANATOLE FRANCE\n"
                          "	93285 SAINT-DENIS CEDEX\n"
                          "H.3.1    A demander\n"
                          "H.3.2	En cours\n"
                          "H.3.2.1	Si en cours, préciser la date de soumission :   /  /    \n"
                          "H.3.3	Obtenu(e)	\n"
                          "	Si obtenu(e), préciser :	\n"
                          "H.3.3.1  Date de la décision :   /  /  \n"
                          "H.3.3.2  Autorisation\n"
                          "H.3.3.3	  Refus d'autorisation	\n"
                          "	Si refus d'autorisation, préciser :\n"
                          "H.3.3.3.1	Les motifs :      \n"
                          "H.3.3.3.2	La date éventuelle envisagée de resoumission de le demande :      ")
    table.cell(1,2).text=("\n\n\n\n\n□\n□\n\n□\n\n\n□\n□")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0 or n==2:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(0,2)
    a.merge(b)
    a=table.cell(1,0)
    b=table.cell(1,1)
    a.merge(b)
    
    document.add_page_break()
    
    '''Partie I'''
    
    paragraph=document.add_paragraph("\nI. SIGNATURE DU DEMANDEUR EN FRANCE", style='debut_page')
    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=1, cols=8, style='Table Grid')
    table.cell(0,0).text=("I.1")
    table.cell(0,1).text=("Par la présente, j’atteste / j’atteste au nom du promoteur (rayer la mention inutile) ce qui suit :\n"
                          "-	les informations fournies ci-dessus à l’appui de la demande sont exactes ;\n"
                          "-	la recherche sera réalisée conformément au protocole, à la réglementation nationale et aux principes de bonnes pratiques ;\n"
                          "-	il est raisonnable de mettre en œuvre la recherche proposée ; \n"
                          "-	je m'engage à déclarer les effets indésirables graves et/ou incidents  et à soumettre les rapports de sécurité, conformément à la réglementation applicable ;\n"
                          "-	je m'engage à soumettre un résumé du rapport final de la recherche à l’ANSM au plus tard 1 an après la fin de la recherche dans tous les pays.")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    a=table.cell(0,1)
    b=table.cell(0,7)
    a.merge(b)
    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("I.2	DEMANDEUR AUPRES DE L'ANSM (tel qu'indiqué en C.1)")
    table.cell(0,1).text=("I.2.1     Date :   /  /    \n"
                          "I.2.2     Signature  :   \n"
                          "I.2.3	Nom :      ")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    paragraph=document.add_paragraph()
    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.cell(0,0).text=("I.3	DEMANDEUR AUPRES DU CPP (tel qu'indiqué en C.2)")
    table.cell(0,1).text=("I.3.1     Date :   /  /    \n"
                          "I.3.2     Signature  :   \n"
                          "I.3.3	Nom :      ")   
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    if n==0:
                        fontdebut.bold = True
                    n=n+1
    
    
    
    
    