# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 16:45:53 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie8(document):
    'Creation de la partie 8 du protcole de catégorie 2'
 #   document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

 #   Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('8	GESTION DES ÉVÉNEMENTS INDÉSIRABLES / EFFETS INDESIRABLES / INCIDENTS',document)
#    
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Les évènements indésirables / effets indésirables / incidents seront à déclarer aux différents circuits de vigilances sanitaires applicables à chaque produit ou pratique concernée (vigilance du soin, pharmacovigilance, matériovigilance, hémovigilance, cosmétovigilance…) en conformité avec la règlementation en vigueur.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Il est recommandé aux déclarants de spécifier que le patient est inclus dans un essai clinique et d’identifier précisément l’essai clinique concerné.')
    run1.style='Paragraphe'

#        #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

    
  #  document.save("Cat2Partie8.docx")