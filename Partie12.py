# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 15:11:16 2019

@author: Asuspc
"""


import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document) --> écrire en minuscule !!!
#    TexteGrisJustif(texte,document)

def Partie12(document):
    'Creation de la partie 12 du protcole de catégorie 1'
  #  document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

  #  Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
    #Ecriture du premier titre 
    Titre1('12	DROIT D’ACCES AUX DONNEES ET DOCUMENTS SOURCE ',document)
    
    
    #Ecriture du 12.1  
    Titre2('12.1	Accès aux données',document)
    
    #Ecriture du 12.2  
    Titre2('12.2	Données source',document)
    
    #Ecriture du 12.3  
    Titre2('12.3	Confidentialité des données',document)
    
        #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    #document.save("Partie7.docx") 
    
 #   document.save("Partie12.docx")
    
    
    
    