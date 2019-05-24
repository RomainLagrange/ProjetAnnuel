# -*- coding: utf-8 -*-
"""
Created on Sat Feb 16 14:05:20 2019

@author: Julie
"""

import docx
import StyleProt1
from StyleProt1 import Style, TexteGris, Titre1, Titre2,Titre3, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie4(document,extract):
    'Creation de la partie 4 du protcole de catégorie 1'
  #  document = docx.Document()

    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)
 
 #   Style(document)

    Titre1('4	CONCEPTION DE LA RECHERCHE',document)
    
    #Texte sur fond gris  
    TexteGris('prendre contact avec la plateforme de methodologie \n pour aide a la redaction de ce chapitre',document)

   # Ecriture du 4.1  
    Titre2('4.1	Schéma de la recherche',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['traitement_strategie_courte'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(11)
    # Ecriture du 4.2  
    Titre2('4.2	Méthode pour la randomisation',document)
    
    
                #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
   # document.save("Partie4.docx")   


