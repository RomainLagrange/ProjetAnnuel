# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:32:49 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:17:09 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:11:16 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:09:16 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:07:33 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:05:49 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 16:19:49 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style, Titre1,Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK

from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)
    

def Partie12(document):
    'Creation de la partie 12 du protcole de catégorie 3'
   # document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)
        
   # Style(document)

 
    Titre1('12	CONSERVATION DES DOCUMENTS ET DES DONNEES RELATIVES A LA RECHERCHE',document)
    #ecriture du premier titre 

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les documents suivants seront archivés par le nom de l’étude (')
    run1.style='Paragraphe'
    run2=p.add_run('xxxxx')
    run2.style='Paragraphe'
    run2.font.color.rgb = RGBColor(0x92,0xD0,0x50) 
    run3=p.add_run(') dans les locaux du service ')
    run3.style='Paragraphe'
    run4=p.add_run('xxxxx ')
    run4.style='Paragraphe'
    run4.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run5=p.add_run('du CHU de Poitiers jusqu’à la fin de la période d’utilité pratique.')
    run5.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Ces documents sont : ')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Protocole et annexes,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('CRF papiers : données individuelles (copies authentifiées de données brutes)')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Documents de suivi')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Analyses statistiques')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Rapport final de l’étude')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('A l’issue de la période d’utilité pratique, l’ensemble des documents à archiver, tels que définis dans la procédure de « classement et archivage des documents liés aux recherches biomédicales » du CHU de Poitiers sera transféré sur le site d’archivage (Service Central des Archives – Hôpital de Poitiers) et sera placé sous la responsabilité du Promoteur pendant 15 ans après la fin de l’étude conformément aux pratiques institutionnelles.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Aucun déplacement ou destruction ne pourront être effectués sans l’accord du Promoteur. Au terme des 15 ans, le promoteur sera consulté pour destruction.')
    run1.style='Paragraphe'
    

    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
     

   # document.save("Cat3Partie12.docx")   