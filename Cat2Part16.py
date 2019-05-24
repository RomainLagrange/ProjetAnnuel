# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 17:01:28 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie16(document,extract):
    'Creation de la partie 16 du protcole de catégorie 2'
 #   document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

  #  Style(document)

   
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('16	FAISABILITE DE L’ETUDE',document)

    
   # Ecriture du 16.1  
    Titre2('16.1	Expertise scientifique',document)
    
    
    # Ecriture du 16.2  
    Titre2('16.2	Collaborations ',document)
    
     # Ecriture du 16.3 
    Titre2('16.3	Financement du projet (si ce point ne fait pas partie d’un document distinct)',document)
    
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['titre_abrege'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    
     #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)  
 
#    document.save("Cat2Partie16.docx")