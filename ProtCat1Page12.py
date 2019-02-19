# -*- coding: utf-8 -*-
"""
Created on Sat Feb 16 13:34:49 2019

@author: Julie
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre3,TexteGris
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#revoir titre1 et texte encardé gris

def Page12():
    'Creation de la page 12 du protcole de catégorie 1'
    document = docx.Document()



#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
 
    StyleProt1.Style(document)


    
    #Texte sur fond gris   
    TexteGris('prendre contact avec la plateforme de methodologie \n pour aide a la redaction de ce chapitre', document)


    

    
   # Ecriture du 2.1  
    document.add_paragraph('2.1	Objectif principal\n', style='Titre2') 
    
    # Ecriture du 2.2  
    document.add_paragraph('2.1	Objectifs secondaires\n', style='Titre2') 
    
    document.save("page12.docx")   


