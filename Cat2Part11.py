# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 16:51:14 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie11(document):
    'Creation de la partie 11 du protcole de catégorie 2'
  #  document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

   # Style(document)

   
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('11	DROIT D’ACCES AUX DONNEES ET DOCUMENTS SOURCE ',document)
    
    
   # Ecriture du 11.1  
    Titre2('11.1	Accès aux données',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’acceptation de la participation au protocole implique que les investigateurs mettront à disposition les documents et données individuelles strictement nécessaires au suivi, au contrôle de qualité et à l’audit de la recherche, à la disposition des personnes ayant un accès à ces documents conformément aux dispositions législatives et réglementaires en vigueur (articles L.1121-3 et R.5121-13 du code de la santé publique).')
    run1.style='Paragraphe'
    
    # Ecriture du 11.2  
    Titre2('11.2	Données source',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Ensemble des informations figurant dans des documents originaux, ou dans des copies authentifiées de ces documents, relatif aux examens cliniques, aux observations ou à d’autres activités menées dans le cadre d’une recherche et nécessaires à la reconstitution et à l’évaluation de la recherche. Les documents dans lesquels les données sources sont enregistrées sont appelés les documents sources.')
    run1.style='Paragraphe'
    
    # Ecriture du 11.3
    Titre2('11.3	Confidentialité des données',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Conformément aux dispositions législatives en vigueur, les personnes ayant un accès direct aux données source prendront toutes les précautions nécessaires en vue d\'assurer la confidentialité des informations relatives aux médicaments expérimentaux, aux recherches, aux personnes qui s\'y prêtent et notamment en ce qui concerne leur identité ainsi qu’aux résultats obtenus.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Ces personnes, au même titre que les investigateurs eux-mêmes, sont soumises au secret professionnel.')
    run1.style='Paragraphe'


    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Pendant la recherche ou à son issue, les données recueillies sur les personnes qui s’y prêtent et transmises au promoteur par les investigateurs (ou tous autres intervenants spécialisés) seront rendues anonymes. Elles ne doivent en aucun cas faire apparaître en clair les noms des personnes concernées ni leur adresse. ')
    run1.style='Paragraphe'
#
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Seules les initiales du nom et du prénom du patient seront enregistrées, accompagnées d’un numéro codé propre à l’étude indiquant l’ordre d’inclusion des sujets.\n')
    run1.style='Paragraphe'
    run2=p.add_run('Ex : n° de centre – n° d’inclusion du patient dans l’ordre chronologique, cf paragraphe 7.3.1.')
    run2.style='Paragraphe'
    run2.font.italic=True
    
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le promoteur s’assurera que chaque personne qui se prête à la recherche a donné son accord par écrit pour l’accès aux données individuelles la concernant et strictement nécessaires au contrôle de qualité de la recherche.')
    run1.style='Paragraphe'
    
     #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

    
  #  document.save("Cat2Partie11.docx")