# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:30:44 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 10:13:09 2019

@author: Asuspc
"""

# -*- coding: utf-8 -*-
"""
Created on Mon Feb 18 11:51:47 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie11(document):
    'Creation de la partie 11 du protcole de catégorie 3'
   # document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

   # Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('11	CONSIDERATIONS ETHIQUES ET REGLEMENTAIRES',document)
    
    TexteGris('prendre contact avec la promotion interne \n pour aide a la redaction de ce chapitre', document)

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Cette étude sera conduite conformément aux Recommandations et Bonnes pratiques en épidémiologie. L’étude est observationnelle, les problèmes éthiques qu’elle soulève sont donc peu nombreux puisqu’elle n’intervient pas dans le traitement des patients. ')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Cette étude observationnelle ne change rien à la prise en charge habituelle des patients. ')
    run1.style='Paragraphe'

   # Ecriture du 11.1  
    Titre2('11.1	Information des patients',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les patients seront informés de l’utilisation de données médicales les concernant dans le cadre de cette étude et de l’inclusion de ces données personnelles, rendues anonymes, dans une base informatique. \nUne note d’information et de non-opposition sera remise au patient et le fait que le patient a été informé et qu’il ne s’oppose pas à participer à cette étude sera noté dans le dossier médical.')
    run1.style='Paragraphe'
    
    # Ecriture du 11.2  
    Titre2('11.2	Anonymat des patients et conformité aux textes de référence',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les informations recueillies sont traitées confidentiellement conformément à la loi relative à l’informatique, aux fichiers et aux libertés.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Cette étude entre dans le cadre de la « Méthodologie de Référence » (MR-003) en application des dispositions de l’article 54 alinéa 5 de la loi n°78-17 du 6 janvier 1978 modifiée relative à l’informatique, aux fichiers et aux libertés. Ce changement a été homologué par délibération n°2016-263 du 21 juillet 2016. Le CHU de Poitiers, promoteur de l’étude, a signé un engagement de conformité à cette « Méthodologie de Référence ». ')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les déclarations devant être effectuées dans le cadre de la MR-003 seront traitées par le référent CIL (Comité Informatique et Liberté) du CHU de Poitiers.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Si la recherche ne rentre pas dans le champ d’application de la MR-003 : Le nom de la structure responsable du traitement des données ')
    run1.style='Paragraphe'
    run1.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run2=p.add_run('a déclaré la recherche à la Commission Nationale de l’Informatique et des Libertés (CNIL).')
    run2.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('L’investigateur s’assure que l’anonymat de chaque sujet participant à l’étude est respecté : seule la première lettre du nom et la première lettre du prénom ainsi qu’un numéro de patient figurent dans le cahier d’observation et autre document de l’étude. \nCe numéro de patient sera constitué du numéro de centre suivi du numéro d’entrée dans l’étude.\nUne liste de correspondance sera conservée par l’investigateur dans le classeur de l’étude.\nCette liste de correspondance fait partie des documents de l’étude et sera conservée pour une durée de 15 ans après la fin de la recherche.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Aucune information permettant l’identification des personnes n’est communiquée à des tiers autres que ceux, représentants du promoteur et des Autorités compétentes, réglementairement habilitées à détenir cette information et qui sont tous tenus au secret professionnel.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Le promoteur et la(les) personne(s) qui dirige(nt) et surveille(nt) la recherche s’engagent à ce que cette recherche soit réalisée en conformité avec la loi n°2012-300 du 5 mars 2012 relative aux recherches impliquant la personne humaine et la déclaration d’Helsinki (qui peut être retrouvée dans sa version intégrale sur le site ')
    run1.style='Paragraphe'
    run2=p.add_run('http://www.wma.net/en/30publications/10policies/b3/).')
    run2.style='Paragraphe'
    run2.font.underline = True
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Cette recherche a reçu l’avis favorable du Comité de Protection des Personnes (CPP) de ')
    run1.style='Paragraphe'
    run2=p.add_run('nom du CPP.')
    run2.style='Paragraphe'
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run2.font.italic = True
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Si applicable :\n')
    run1.style='Paragraphe'
    run1.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run1.font.italic = True
    run2=p.add_run('Cette recherche est enregistrée sur le site http://clinicaltrials.gov/ ')
    run2.style='Paragraphe'
    
    
    # Ecriture du 11.3  
    Titre2('11.3	Modifications au protocole ',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toute modification substantielle, c’est à dire toute modification de nature à avoir un impact significatif sur la protection des personnes, sur les conditions de validité et sur les résultats de la recherche, sur la qualité et la sécurité des produits expérimentés, sur l’interprétation des documents scientifiques qui viennent appuyer le déroulement de la recherche ou sur les modalités de conduite de celle-ci, fait l’objet d’un amendement écrit qui est soumis au promoteur ; celui-ci doit obtenir, préalablement à sa mise en œuvre, un avis favorable du CPP.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les modifications non substantielles, c\'est à dire celles n’ayant pas d’impact significatif sur quelque aspect de la recherche que ce soit, sont communiquées au CPP à titre d’information.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toutes les modifications sont validées par le promoteur, et par tous les intervenants de la recherche concernés par la modification, avant soumission au CPP. Cette validation peut nécessiter la réunion de tout comité constitué pour la recherche.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toutes les modifications au protocole doivent être portées à la connaissance de tous les investigateurs qui participent à la recherche. Les investigateurs s’engagent à en respecter le contenu.')
    run1.style='Paragraphe'
  
    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
  
  #  document.save("Cat3Part11.docx")   