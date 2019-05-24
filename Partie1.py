# -*- coding: utf-8 -*-
"""
Created on Thu Jan 31 13:32:03 2019

@author: Julie
"""

import docx
import StyleProt1
from StyleProt1 import Style, Titre1,Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_BREAK


#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)
    
def Partie1(document,extract):
#def Partie1():
    'Creation de la partie 1 du protcole de catégorie 1'
 #   document = docx.Document()


#   Marge de la page
#    sections = document.sections
#    for section in sections:
#        section.top_margin = Cm(2)
#        section.bottom_margin = Cm(2)
#        section.left_margin = Cm(2)
#        section.right_margin = Cm(2)
        
  #  Style(document)

     #ecriture du premier titre 
    Titre1('1	JUSTICATION SCIENTIFIQUE ET DESCRIPTION GENERALE',document)

    
    # Ecriture du 1.1  
    Titre2('1.1	Etat actuel des connaissances',document)
  

    #Ecriture du titre1.1.1
    Titre3('1.1.1','Sur la pathologie',document)

       
    
    #ecriture du titre1.1.2
    Titre3('1.1.2','Sur les traitements, stratégies et procédures de référence et à l’étude',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['traitement_strategie_longue'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(11) 

  
                                        
    #ecriture du titre1.2
    Titre2('1.2	Hypothèse de la recherche et résultats attendus',document)

    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['critere_jugement_principal_courte'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(11)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['traitement_strategie_courte'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(11) 
    

    #Ecriture du titre1.3
    Titre2('1.3 Justification des choix méthodologiques',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['justification_etude_longue'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(11)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['critere_jugement_principal_courte'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(11) 
    
    #Texte sur fond gris  
    TexteGris('prendre contact avec la plateforme de methodologie \n pour aide a la redaction du paragraphe 2.3', document)


 
   
     #Ecriture du titre1.4
    Titre2('1.4 Rapport bénéfices / risques prévisibles',document)
    

    
    #Ecriture du titre1.4.1
    Titre3('1.4.1','Bénéfices',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['benefices'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(11)


   
    #Ecriture du titre1.4.2
    Titre3('1.4.2','Risques',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['risques'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(11)

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’investigateur doit constamment surveiller, évaluer et documenter les risques et doit s’assurer qu’ils pourront être gérés de manière satisfaisante.')
    run1.style='Paragraphe'
    
     #Ecriture du titre1.5
    Titre2('1.5 Retombées attendues',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['retombee_attenduees_longue'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(11)

    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
     


  #  document.save("Partie1.docx")   
    