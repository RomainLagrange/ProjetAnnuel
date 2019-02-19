# -*- coding: utf-8 -*-
"""
Created on Sat Feb 16 13:42:50 2019

@author: Julie
"""

import docx
import StyleProt1
from StyleProt1 import Style, Titre3, TexteGris
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#revoir titre1 et texte encardé gris

def Page13():
    'Creation de la page 13 du protcole de catégorie 1'
    document = docx.Document()



#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
 

    
    StyleProt1.Style(document)

    
   # Ecriture du 3.1 
    document.add_paragraph('3.1	Critère d’évaluation principal\n', style='Titre2') 
    
    # Ecriture du 3.2  
    document.add_paragraph('3.2	Critères d’évaluation secondaires\n', style='Titre2') 
    

    
    
    document.save("page13.docx")   


