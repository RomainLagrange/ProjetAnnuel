# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 14:54:04 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ROW_HEIGHT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document) --> écrire en minuscule !!!
#    TexteGrisJustif(texte,document)

#def Partie9(document):
def Partie9():
    'Creation de la partie 9 du protcole de catégorie 1'
    document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

    Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
     #ecriture du premier titre 
    Titre1('9	EVALUATION DE LA SECURITE',document)
    
     #Texte gris centré
    TexteGris('prendre contact avec l\'unite de vigilance des essais cliniques \n pour aide a la redaction de ce chapitre',document)
    
    
    #Ecriture du 9.1  
    Titre2('9.1	Définitions',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Evénement indésirable ')
    run1.style='Paragraphe'
    run1.font.bold= True
    run2=p.add_run('(article R1123-46 du code de la santé publique)\nToute manifestation nocive survenant chez une personne qui se prête à une recherche impliquant la personne humaine, que cette manifestation soit liée ou non à la recherche ou au produit sur lequel porte cette recherche.')
    run2.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Effet indésirable ')
    run1.style='Paragraphe'
    run1.font.bold= True
    run2=p.add_run('(article R1123-46 du code de la santé publique)\nEvénement indésirable survenant chez une personne qui se prête à une recherche impliquant la personne humaine, lorsque cet événement est lié à la recherche ou au produit sur lequel porte cette recherche.')
    run2.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Evénement ou effet indésirable grave ')
    run1.style='Paragraphe'
    run1.font.bold= True
    run2=p.add_run('(article R1123-46 du code de la santé publique et guide ICH E2B)\nTout événement ou effet indésirable qui :')
    run2.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('entraîne la mort,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('met en danger la vie de la personne qui se prête à la recherche,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('nécessite une hospitalisation ou la prolongation de l\'hospitalisation,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('provoque une incapacité ou un handicap important ou durable,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('se traduit par une anomalie ou une malformation congénitale,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('ou tout événement considéré comme médicalement grave,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('et s\'agissant du médicament, quelle que soit la dose administrée.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’expression « mise en jeu du pronostic vital » est réservée à une menace vitale immédiate, au moment de l’événement indésirable.')
    run1.style='Paragraphe'
    
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Effet indésirable inattendu ')
    run1.style='Paragraphe'
    run1.font.bold= True
    run2=p.add_run('(article R1123-46 du code de la santé publique)')
    run2.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Pour les recherches portant sur un médicament, effet indésirable inattendu : tout effet indésirable du produit dont la nature, la sévérité, la fréquence ou l\'évolution ne concorde pas avec les informations de référence sur la sécurité mentionnées dans le résumé des caractéristiques du produit ou dans la brochure pour l’investigateur lorsque le produit n’est pas autorisé.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Pour les recherches portant sur un dispositif médical ou sur un dispositif médical de diagnostic in vitro, effet indésirable inattendu : tout effet du dispositif dont la nature, la sévérité ou l’évolution ne concordent pas avec les informations de référence figurant respectivement dans la notice d’instruction ou dans la notice d’utilisation du dispositif lorsque celui-ci fait l’objet d’un marquage CE, et dans le protocole ou la brochure pour l’investigateur lorsqu’il ne fait pas l’objet d’un tel marquage.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Pour les autres recherches impliquant la personne humaine, effet indésirable inattendu : tout effet indésirable dont la nature, la sévérité ou l’évolution ne concorde pas avec les informations relatives aux produits, actes pratiqués et méthodes utilisées au cours de la recherche')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Fait nouveau')
    run1.style='Paragraphe'
    run1.font.bold= True
    run2=p.add_run('(article R1123-46 du code de la santé publique)')
    run2.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Toute nouvelle donnée pouvant conduire à une réévaluation du rapport des bénéfices et des risques de la recherche ou du produit objet de la recherche, à des modifications dans l’utilisation de ce produit, dans la conduite de la recherche, ou des documents relatifs à la recherche, ou à suspendre ou interrompre ou modifier le protocole de la recherche ou des recherches similaires. ')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Pour les essais portant sur la première administration ou utilisation d’un produit de santé chez des personnes qui ne présentent aucune affection: tout effet indésirable grave.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Imputabilité : ')
    run1.style='Paragraphe'
    run1.font.bold= True
    run2=p.add_run('relation entre l’EvI et la recherche. L’EvI lié à la recherche deviendra un EI. Les facteurs à prendre en compte pour la détermination de l’imputabilité sont : ')
    run2.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('la chronologie des évènements,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('la disparition de l’EvI lors de l’arrêt du (des) médicament(s) et/ou la réapparition en cas de ré-administration,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('la pharmacodynamie et la pharmacocinétique des médicaments,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('la notion d’antécédent d\'évènement similaire lors de l’administration du médicament ou d’un médicament de la même classe,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('l’existence d’une autre étiologie.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Intensité : ')
    run1.style='Paragraphe'
    run1.font.bold= True
    run2=p.add_run('l’intensité des EvI est évaluée par l’investigateur, soit en s’aidant d’une échelle de gradation des évènements indésirables annexée au protocole (ex : classification NCI-CTC pour les essais en cancérologie), soit selon la classification suivante :')
    run2.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('léger de grade 1 : EvI généralement transitoire et sans retentissement sur les activités normales,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('modéré de grade 2 : EvI suffisamment gênant pour retentir sur les activités normales,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 3'
    run1=p.add_run('sévère de grade 3 : EvI modifiant considérablement le cours normal des activités du patient, ou invalidant, ou constituant une menace pour la vie du patient.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Remarque : le critère d’intensité ne doit pas être confondu avec le critère de gravité qui sert de guide pour définir les obligations de déclaration.')
    run1.style='Paragraphe'
    
#    p=document.add_paragraph()
#    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
#    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
#    p.style='List Bullet 2'
#    run1=p.add_run('')
#    run1.style='Paragraphe'
 
    #Ecriture du 9.2  
    Titre2('9.2	Description des événements indésirables graves attendus',document)

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Tout évènement indésirable grave ne concordant pas avec ceux décrits  dans la liste des événements attendus ou dans la dernière version du document de référence est qualifié d’inattendu.')
    run1.style='Paragraphe'

    #Ecriture du 9.3  
    Titre2('9.3	Conduite à tenir par l’investigateur en cas d’événement indésirable, de fait nouveau ou de grossesse',document)
    
    #Ecriture du titre 9.3.1
    Titre3('9.3.1','Recueil des événements indésirables (EvI)',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Dès la signature du consentement, l’investigateur est responsable du recueil de tous les évènements indésirables. Il rapporte tous les événements indésirables graves et non graves (EvI biologiques et cliniques) qui surviennent entre la signature du consentement et la fin de participation du patient ou la fin de recueil des événements indésirables, dans le cahier d’observation.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Ces évènements indésirables seront évalués à chaque visite au cours de l’étude par un interrogatoire et lors de l’examen clinique du patient.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Exception au recueil : ')
    run1.style='Paragraphe'
    run1.font.bold = True
    
    #Ecriture du titre 9.3.2
    Titre3('9.3.2','Déclaration des événements indésirables graves (EvIG), des événements indésirables d’intérêt et des faits nouveaux ',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’investigateur évalue chaque événement indésirable au regard de sa gravité. ')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Délais de déclaration')
    run1.style='Paragraphe'
    run1.font.bold = True
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’investigateur doit déclarer au promoteur, ')
    run1.style='Paragraphe'
    run2=p.add_run('sans délai ')
    run2.style='Paragraphe'
    run2.font.bold = True
    run2.font.underline = True
    run3=p.add_run('à partir du jour où il en a connaissance, tout événement indésirable grave (EvIG), tout événement indésirable d’intérêt ou tout fait nouveau de sécurité s’il survient :')
    run3.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('à partir de la date de signature du consentement,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('pendant toute la durée de suivi du patient prévue par la recherche,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('jusqu’à ')
    run1.style='Paragraphe'
    run2=p.add_run('X ')
    run2.style='Paragraphe'
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('jours ')
    run3.style='Paragraphe'
    run3.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run4=p.add_run('(à définir entre l’unité de vigilance et l’investigateur) ')
    run4.style='Paragraphe'
    run4.font.bold = True
    run5=p.add_run('après la fin du suivi du participant prévue par la recherche')
    run5.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Par ailleurs, quel que soit le délai de survenue après la fin de l’étude, tout EvIG susceptible d’être dû à la recherche doit être déclaré sans délai au promoteur dès lors qu’aucune autre cause que la recherche ne peut raisonnablement lui être attribuée (par exemple des effets graves pouvant apparaître à grande distance de l’exposition au médicament, tels des cancers ou des anomalies congénitales).')
    run1.style='Paragraphe'
    
    #IMAGE
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Exception à la déclaration sans délai')
    run1.style='Paragraphe'
    run1.font.bold = True

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Par exemple, les circonstances suivantes ne seront pas à déclarer immédiatement au promoteur mais seront saisies dans le CRF:')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('passage en hôpital de jour,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('hospitalisation pour traitement de routine ou surveillance de la pathologie étudiée non associé à une détérioration de l’état du patient,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('etc… ')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Modalités de déclaration au promoteur')
    run1.style='Paragraphe'
    run1.font.bold = True
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Tout EvIG, quelle que soit sa relation de causalité avec le traitement de l’essai ou la recherche, doit être déclaré par fax au  ')
    run1.style='Paragraphe'
    run2=p.add_run('05 49 44 30 58.')
    run2.style='Paragraphe'
    run2.font.bold = True
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Coordonnées de l’unité de vigilance du promoteur :')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Unité de Vigilance des Essais Cliniques du CHU Poitiers (Direction de la Recherche)\nDr. Sophie DURANTON\ntéléphone au 05 49 44 30 50\nfax au 05 49 44 30 58\nmail au sophie.duranton@chu-poitiers.fr')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Informations à transmettre au promoteur')
    run1.style='Paragraphe'
    run1.font.bold = True
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’investigateur doit documenter au mieux l’événement, en donner si possible, le diagnostic médical. ')
    run1.style='Paragraphe'
    run2=p.add_run('L’investigateur doit suivre le patient ayant présenté un EvIG jusqu’à sa résolution, ')
    run2.style='Paragraphe'
    run2.font.bold = True
    run3=p.add_run('sa stabilisation à un niveau jugé par lui comme médicalement acceptable ou le retour à l’état antérieur, ')
    run3.style='Paragraphe'
    run4=p.add_run('même si le patient a arrêté la procédure de la recherche. ')
    run4.style='Paragraphe'
    run4.font.bold = True
    run5=p.add_run('Un complément d’information concernant le suivi et l’évolution de l’événement, si elle n’est pas mentionnée dans le premier rapport, sera envoyé au promoteur par l’investigateur dès que possible.')
    run5.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Chaque EvIG sera décrit sur le formulaire prévu à cet effet (« Déclaration d’Evènement Indésirable Grave ») en essayant d’être le plus exhaustif possible. Les informations à transmettre sont les suivantes : ')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('identification du patient (numéro, code, date de naissance, date d’inclusion, sexe, poids, taille),')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('gravité de l’EvI,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('date de début et de fin de l’EvI,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('description claire et détaillée de l’EvI (diagnostic, symptômes, intensité, chronologie, actions entreprises et résultats),')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('évolution de l’EvI,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('maladies en cours ou antécédents pertinents du patient,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('traitements reçus par le patient au moment de la survenue de l’EvIG,')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('lien de causalité de l’EvI avec le(s) médicament(s) expérimental (aux), le(s) comparateur(s), les éventuels traitements associés, la recherche ou d’autres critères (défini par l’investigateur).')
    run1.style='Paragraphe'

    
    #Ecriture du titre 9.3.3
    Titre3('9.3.3','Déclaration des grossesses',document)
    
    p=document.add_paragraph()
    run1=p.add_run('Toute grossesse devra être déclarée immédiatement au promoteur.\nL’investigateur informe l’unité de Vigilance des Essais Cliniques du promoteur et envoie par fax le formulaire d’EvIG sur lequel doit être noté la date prévisible d’accouchement, les coordonnées de l’obstétricien et de la maternité prévue pour l’accouchement si la grossesse se poursuit. \nL’investigateur doit suivre la patiente jusqu’au terme de la grossesse ou de son interruption (interruption volontaire de grossesse (IVG), interruption médicale de grossesse (IMG), fausse couche, etc...) et en notifier l’issue au promoteur. S’il s’agit d’une exposition paternelle, l’investigateur doit obtenir l’accord de la parturiente pour recueillir les informations sur la grossesse.\nSi l’issue de la grossesse entre dans le cadre de la définition des événements indésirables graves (avortement spontané avec hospitalisation, mort fœtale, anomalie congénitale, ...) l’investigateur doit suivre la procédure de déclaration des EvIG.')
    run1.style='Paragraphe'
    
    #Ecriture du titre 9.3.4
    Titre3('9.3.4','Tableau récapitulatif du circuit de déclaration par type d’événement',document)
    
    table = document.add_table(5, 3)
    table.style = 'Table Grid'
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_2)

    row = table.rows[0].cells
    
    cell = row[0]
    pt = cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Type d\'evenement')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
    p.bold = True
    p.font.small_caps = True

    cell = row[1]
    pt = cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('modalites de notification')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
    p.bold = True
    p.font.small_caps = True

    cell = row[2]
    pt = cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('delai de notification au promoteur')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
    p.bold = True
    p.font.small_caps = True

    row = table.rows[1].cells
    
    cell = row[0]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Evénement indésirable ')
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
    p = pt.add_run('non grave')
    p.font.name = 'Times New Roman'
    p.font.bold=True
    
    cell = row[1]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Dans le cahier d’observation')
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'

    cell = row[2]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Pas de notification immédiate')
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'  
    
    row = table.rows[2].cells
    
    cell = row[0]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Evénement indésirable ')
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
    p = pt.add_run('grave')
    p.font.name = 'Times New Roman'
    p.font.bold=True
    p = pt.add_run('ou')
    p.font.name = 'Times New Roman'
    p = pt.add_run('événement d’intérêt ')
    p.font.name = 'Times New Roman'
    p.font.bold=True
    p = pt.add_run('(si applicable)')
    p.font.name = 'Times New Roman'
    
    
    cell = row[1]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Formulaire de déclaration initiale d’EvIG + follow up si nécessaire + recueil dans le cahier d’observation')
    p.font.name = 'Times New Roman'
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER

    cell = row[2]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Déclaration sans délai au promoteur')
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
    p.font.bold= True
    
    row = table.rows[3].cells
    
    cell = row[0]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Fait nouveau')
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
     
    cell = row[1]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Rapport écrit')
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'

    cell = row[2]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Déclaration sans délai au promoteur')
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
    p.font.bold= True
    
    row = table.rows[4].cells
    
    cell = row[0]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Grossesse')
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
     
    cell = row[1]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Formulaire de déclaration d’EvIG')
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'

    cell = row[2]
    pt =  cell.paragraphs[0]
    t = pt.text = ''
    p = pt.add_run('Dès confirmation de la grossesse')
    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'  

    
    #Ecriture du 9.4
    Titre2('9.4 Déclaration par le promoteur des effets indésirables graves inattendus, des faits nouveaux et autres évènements',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Le promoteur évalue, indépendamment de l’investigateur, le lien de causalité entre l’événement indésirable grave, les traitements expérimentaux ')
    run1.style='Paragraphe'
    run2=p.add_run('(A adapter en fonction de l’étude : médicament, dispositif médical/procédure de mise en place du DM…), ')
    run2.style='Paragraphe'
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0) 
    run3=p.add_run('les traitements associés et la recherche.\nTous les événements indésirables graves pour lesquels l’investigateur ou le promoteur estime qu’une relation de causalité peut être raisonnablement envisagée sont considérés comme des suspicions d’effets indésirables graves.')
    run3.style='Paragraphe'
    

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Le promoteur évalue si l’effet indésirable grave est attendu ou inattendu en se basant sur la liste des évènements indésirables graves attendus décrits dans le paragraphe 9.2 du protocole et sur le document de référence tel que défini dans le protocole.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Le promoteur déclare selon les délais en vigueur les informations de sécurité aux autorités compétentes et au CPP selon les exigences réglementaires spécifique à chaque type d’essai.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('S’agissant d’une étude portant sur un médicament, le promoteur enregistre dans la base de données EudraVigilance tous les effets indésirables graves inattendus.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1=p.add_run('Dans le cas d’une recherche en insu, le promoteur déclare les EIGI à l’ANSM après avoir levé l’insu.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Tableau récapitulatif  des déclarations par type d’étude')
    run1.style='Paragraphe'
    run1.font.bold = True
    run1.font.underline = True    
    
    table = document.add_table(9, 4)
    table.style = 'Table Grid'
    for cell in table.columns[0].cells: 
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for cell in table.columns[1].cells: 
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for cell in table.columns[2].cells: 
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for cell in table.columns[3].cells: 
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    row = table.rows[0].cells
    
    cell = row[0]
    pt = cell.paragraphs[0]
    p = pt.add_run('Type d’étude et type d’EvI')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
    p.bold = True

    cell = row[1]
    pt = cell.paragraphs[0]
 #   t = pt.text = ''
    p = pt.add_run('Déclaration aux autorités compétentes concernées')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
    p.bold = True

    cell = row[2]
    pt = cell.paragraphs[0]
    p = pt.add_run('Délai de déclaration initiale')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
    p.bold = True
    
    cell = row[3]
    pt = cell.paragraphs[0]
    p = pt.add_run('Délai du follow up')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'
    p.bold = True
    
    row = table.rows[1].cells
    
    cell = row[0]
    pt = cell.paragraphs[0]
    p = pt.add_run('Médicament : SUSAR')
    p.font.name = 'Times New Roman'

    cell = row[1]
    pt = cell.paragraphs[0]
 #   t = pt.text = ''
    p = pt.add_run('EMA, ANSM')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'

    cell = row[2]
    pt = cell.paragraphs[0]
    p = pt.add_run('- décès ou mise en jeu du pronostic vital : sans délai\n- autre critère : max 15 j')
    p.font.name = 'Times New Roman'
    
    cell = row[3]
    pt = cell.paragraphs[0]
    p = pt.add_run('Max 8j')
    p.font.name = 'Times New Roman'
    
    row = table.rows[2].cells
    
    cell = row[0]
    pt = cell.paragraphs[0]
    p = pt.add_run('DM : SUSAR et EvIG lié au geste de mise en œuvre ')
    p.font.name = 'Times New Roman'

    cell = row[1]
    pt = cell.paragraphs[0]
 #   t = pt.text = ''
    p = pt.add_run('ANSM')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'

    cell = row[2]
    pt = cell.paragraphs[0]
    p = pt.add_run('- décès ou mise en jeu du pronostic vital : sans délai\n\n- autre critère : max 15 j')
    p.font.name = 'Times New Roman'
    
    cell = row[3]
    pt = cell.paragraphs[0]
    p = pt.add_run('- Max 8 j\n\n-Max 15 j')
    p.font.name = 'Times New Roman'
    
    row = table.rows[3].cells
    
    cell = row[0]
    pt = cell.paragraphs[0]
    p = pt.add_run('HPS : SUSAR')
    p.font.name = 'Times New Roman'

    cell = row[1]
    pt = cell.paragraphs[0]
 #   t = pt.text = ''
    p = pt.add_run('ANSM')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'

    cell = row[2]
    pt = cell.paragraphs[0]
    p = pt.add_run('- décès ou mise en jeu du pronostic vital : sans délai\n\n- autre critère : max 15 j')
    p.font.name = 'Times New Roman'
    
    cell = row[3]
    pt = cell.paragraphs[0]
    p = pt.add_run('Max 8j')
    p.font.name = 'Times New Roman'

    row = table.rows[4].cells
    
    cell = row[0]
    pt = cell.paragraphs[0]
    p = pt.add_run('Cosméto/Tatouage : effet indésirable grave')
    p.font.name = 'Times New Roman'

    cell = row[1]
    pt = cell.paragraphs[0]
 #   t = pt.text = ''
    p = pt.add_run('ANSM')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'

    cell = row[2]
    pt = cell.paragraphs[0]
    p = pt.add_run('Max 7 j')
    p.font.name = 'Times New Roman'
    
    cell = row[3]
    pt = cell.paragraphs[0]
    p = pt.add_run('Sans délai')
    p.font.name = 'Times New Roman'

    row = table.rows[5].cells
    
    cell = row[0]
    pt = cell.paragraphs[0]
    p = pt.add_run('Cosméto/Tatouage : effet indésirable ayant nécessité un traitement médical et effet\nindésirable paraissant revêtir un caractère de gravité justifiant une déclaration')
    p.font.name = 'Times New Roman'

    cell = row[1]
    pt = cell.paragraphs[0]
 #   t = pt.text = ''
    p = pt.add_run('ANSM')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'

    cell = row[2]
    pt = cell.paragraphs[0]
    p = pt.add_run('Max 15j')
    p.font.name = 'Times New Roman'
    
    cell = row[3]
    pt = cell.paragraphs[0]
    p = pt.add_run('Sans délai')
    p.font.name = 'Times New Roman'
    
    row = table.rows[6].cells
    
    cell = row[0]
    pt = cell.paragraphs[0]
    p = pt.add_run('Cosméto/Tatouage : autres effets indésirables')
    p.font.name = 'Times New Roman'

    cell = row[1]
    pt = cell.paragraphs[0]
 #   t = pt.text = ''
    p = pt.add_run('ANSM')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'

    cell = row[2]
    pt = cell.paragraphs[0]
    p = pt.add_run('A la fin de l’essai si durée < 3 mois ou trimestriellement')
    p.font.name = 'Times New Roman'
    
    row = table.rows[7].cells
    
    cell = row[0]
    pt = cell.paragraphs[0]
    p = pt.add_run('Organes/Tissus : effet indésirable grave et incident grave')
    p.font.name = 'Times New Roman'

    cell = row[1]
    pt = cell.paragraphs[0]
 #   t = pt.text = ''
    p = pt.add_run('ANSM')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'

    cell = row[2]
    pt = cell.paragraphs[0]
    p = pt.add_run('immédiatement')
    p.font.name = 'Times New Roman'
    
    cell = row[3]
    pt = cell.paragraphs[0]
    p = pt.add_run('imméditement')
    p.font.name = 'Times New Roman'

    row = table.rows[8].cells
    
    cell = row[0]
    pt = cell.paragraphs[0]
    p = pt.add_run('PSL : effet indésirable et incident grave')
    p.font.name = 'Times New Roman'

    cell = row[1]
    pt = cell.paragraphs[0]
 #   t = pt.text = ''
    p = pt.add_run('ANSM')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.font.name = 'Times New Roman'

    cell = row[2]
    pt = cell.paragraphs[0]
    p = pt.add_run('immédiatement')
    p.font.name = 'Times New Roman'
    
    cell = row[3]
    pt = cell.paragraphs[0]
    p = pt.add_run('imméditement')
    p.font.name = 'Times New Roman'
    
    document.add_paragraph(' ')
           
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Le promoteur déclare sans délai les faits nouveaux survenus au cours de la recherche :')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('à l’ANSM,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('au Comité de Protection des Personnes,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('au directeur général de l’Agence Régionale de Santé (si applicable : voir 9.5).')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Le promoteur et l’investigateur prennent les mesures urgentes appropriées. Le promoteur en informe l’autorité compétente et le comité de protection des personnes.')
    run1.style='Paragraphe'
    
    
    #Ecriture du 9.5
    Titre2('9.5	Essai chez un volontaire sain',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Articles R1123-46 et R1123-62 du code de la santé publique, Circulaire DGS/PP1/2016/61 du 1er mars 2016.')
    run1.style='Paragraphe'
    run1.font.italic= True
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Pour les essais portant sur les volontaires sains, en cas de première administration ou utilisation d’un produit de santé expérimental, tout effet indésirable grave sera considéré comme un fait nouveau. Il devra être déclaré sans délai à l’ANSM, au CPP et au Directeur général de l’ARS. ')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Le promoteur :')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('suspend l’administration ou l’utilisation du produit expérimental chez les personnes participant à la recherche dans l’attente de l’adoption de mesures définitives et jusqu’à la démonstration de l’absence de danger ; les personnes participant à la recherche devront systématiquement être informées et leur consentement obtenu avant toute nouvelle administration du produit à l’étude.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('prend les mesures de sécurité appropriées,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('-	informe sans délai l’autorité compétente et le comité de protection des personnes ')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Pour les recherches sur le médicament ou sur les préparations de thérapie cellulaire, et portant sur des volontaires sains, le promoteur déclare sans délai à l\'ANSM tous les événements ou effets indésirables graves (article R1123-54 du code de la santé publique).')
    run1.style='Paragraphe'
    
    #Ecriture du 9.6
    Titre2('9.6 Rapport annuel de sécurité',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Article R1123-61 du code de la santé publique.')
    run1.style='Paragraphe'
    run1.font.italic= True
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run ('A la date anniversaire de ')
    run1.style='Paragraphe'
    run2=p.add_run('l’autorisation de la recherche (pour les études portant sur un médicament)/la première inclusion (pour tous les autres types de recherche), ')
    run2.style='Paragraphe'
    run2.font.italic= True
    run3 = p.add_run ('le promoteur rédige un rapport de sécurité comprenant :')
    run3.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('la liste des effets indésirables graves susceptibles d’être liés au(x) traitement(s) expérimental(aux) de la recherche incluant les effets graves attendus et inattendus, survenus dans l’essai concerné pendant la période couverte par le rapport,')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('une analyse concise et critique de la sécurité des participants se prêtant à la recherche.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('les tableaux de synthèse de tous les effets indésirables graves survenus dans l’essai concerné depuis le début de la recherche')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run ('Ce rapport est envoyé à l’ANSM et au CPP dans les 60 jours suivant la date anniversaire de ')
    run1.style='Paragraphe'
    run2=p.add_run('l’autorisation de la recherche/la première inclusion.')
    run2.style='Paragraphe'
    run2.font.italic= True

    
        #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
  
    document.save("Partie9.docx")
    
    
    