# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 16:26:36 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style, TexteGris, Titre1, Titre2,Titre3, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie4(document):
    'Creation de la partie 4 du protocole de catégorie 2'
  #  document = docx.Document()

    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
 
 #   Style(document)

    Titre1('4	CRITERES D’ELIGIBILITE',document)
    
   # Ecriture du 4.1  
    Titre2('4.1	Critères d’inclusion',document)
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Tous les patients inclus dans cette recherche devront vérifier tous les critères d’inclusion listés ci-dessous :')
    run1.style='Paragraphe'
    

    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Justifier l’inclusion de personnes visées aux articles L.1121-5 à L.1121-8 et L.1122-1-2 du code de la santé publique (ex : mineurs, majeurs protégés, femmes enceintes, allaitantes, femmes en âge de procréer, personnes en situation d’urgence, personnes incapables de donner personnellement leur consentement, etc) et procédure mise en œuvre afin d’informer et de recueillir le consentement de ces personnes ou de leurs représentants légaux.')
    run1.style='Paragraphe'
    run1.font.italic = True
    
    
    
    # Ecriture du 4.2  
    Titre2('4.2	Critères de non inclusion',document)

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Tous les patients inclus dans cette recherche ne devront avoir aucun des critères de non inclusion listés ci-dessous :')
    run1.style='Paragraphe'
    
      # Ecriture du 4.3  
    Titre2('4.3	Faisabilité et modalités de recrutement',document)
    
#        #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

    
#    document.save("Cat2Partie4.docx")   