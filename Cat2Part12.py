# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 16:52:34 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie12(document):
    'Creation de la partie 12 du protcole de catégorie 2'
  #  document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

#    Style(document)

   
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('12	CONTROLE ET ASSURANCE DE LA QUALITE',document)
    
    TexteGris('prendre contact avec la promotion interne \n pour aide a la redaction de ce chapitre', document)

    
   # Ecriture du 12.1  
    Titre2('12.1	Consignes pour le recueil des données',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toutes les informations requises par le protocole doivent être consignées sur les cahiers d’observation et une explication doit être apportée pour chaque donnée manquante. Les données doivent être recueillies au fur et à mesure qu\'elles sont obtenues, et transcrites dans ces cahiers de façon nette et lisible.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les données enregistrées dans l’e-CRF ')
    run1.style='Paragraphe'
    run2=p.add_run('[CRF] ')
    run2.font.color.rgb = RGBColor(0x0,0xB0,0xF0)
    run2.style='Paragraphe'
    run2.font.italic=True
    run3=p.add_run('et provenant des documents sources doivent être cohérentes entre elles ; dans le cas contraire, les différences doivent être justifiées et documentées.')
    run3.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L\'investigateur est responsable de l\'exactitude, de la qualité et de la pertinence de toutes les données saisies.')
    run1.style='Paragraphe'
    
    # Ecriture du 12.2  
    Titre2('12.2	Contrôle de la qualité',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Un attaché de recherche clinique mandaté par le promoteur visite de façon régulière chaque centre investigateur, lors de la mise en place de la recherche, une ou plusieurs fois en cours de recherche selon le rythme des inclusions et en fin de recherche. Lors de ces visites, et conformément au plan de monitorage basé sur le risque (participant, logistique, impact, ressources), les éléments suivants seront revus :')
    run1.style='Paragraphe'
#
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Les consentements éclairés')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('Le respect du protocole de l\'étude et des procédures qui y sont définies')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('La déclaration des EvIG (si applicable)')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('La qualité des données recueillies dans le cahier d’observation : exactitude, données manquantes, cohérence des données avec les documents "source" (dossiers médicaux, carnets de rendez-vous, originaux des résultats de laboratoire, etc…)')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.style='List Bullet 2'
    run1=p.add_run('La gestion des produits éventuels.')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L\'investigateur et les membres de son équipe acceptent de se rendre disponibles lors des visites de Contrôle de Qualité (monitoring) effectuées à intervalles réguliers par l’Attaché de Recherche Clinique.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Toute visite fera l’objet d’un rapport de monitorage par compte-rendu écrit.')
    run1.style='Paragraphe'
    
     # Ecriture du 12.3 
    Titre2('12.3	Gestion des données',document)
    
    #Texte gris justifié
    TexteGrisJustif('Gestion des données pour une étude e-CRF',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’investigateur devra dater et signer les pages du CRF complétées à la fin du recueil des données ; elles seront considérées comme documents source.\nCe document fera partie intégrante du dossier médical du patient et y sera conservé en permanence. ')
    run1.style='Paragraphe'
    
    #Texte gris justifié
    TexteGrisJustif('Gestion des données pour une étude CRF papier',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les données erronées relevées sur les cahiers d\'observation seront clairement barrées et les nouvelles données seront notées, à côté de l\'information barrée, accompagnées des initiales, de la date et éventuellement d’une justification par l’investigateur ou la personne autorisée qui aura fait la correction.')
    run1.style='Paragraphe'
    
     # Ecriture du 12.4
    Titre2('12.4	Audits et inspections',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Un audit peut être réalisé à tout moment par des personnes mandatées par le promoteur et indépendantes des personnes menant la recherche. Il a pour objectif de vérifier la sécurité des participants et le respect de leurs droits, le respect de la réglementation applicable et la fiabilité des données')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Une inspection peut également être diligentée par une autorité compétente (ANSM pour la France ou EMA dans le cadre d’un essai européen par exemple). ')
    run1.style='Paragraphe'

    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('L’audit, aussi bien que l’inspection, pourront s’appliquer à tous les stades de la recherche, du développement du protocole à la publication des résultats et au classement des données utilisées ou produites dans le cadre de la recherche.')
    run1.style='Paragraphe'
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Les investigateurs acceptent de se conformer aux exigences du promoteur en ce qui concerne un audit et à l’autorité compétente pour une inspection de la recherche.')
    run1.style='Paragraphe'
    
         #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

 
  #  document.save("Cat2Partie12.docx")