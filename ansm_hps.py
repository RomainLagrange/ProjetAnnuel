# -*- coding: utf-8 -*-
"""
Created on Tue Feb 26 12:13:42 2019

@author: Utilisateur
"""

import pandas as pd
import docx
from docx.api import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
import time
from time import gmtime, strftime

def main_ansm_hps(extract):
    document = docx.Document()
    parties_ABC(document, extract)
    partie_D_a_G(document, extract)
    partie_H_fin(document, extract)
    date = (strftime('%d-%m-%Y',time.localtime()))
    document.save("soumission_ansm_hps_"+extract['titre_abrege']+"_"+date+".docx")
    
def parties_ABC(document, extract):
    
    '''Marge de la page'''
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.59)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.63)
        
    
    '''Titre du document en bleu'''
    
    
    
    table = document.add_table(rows=1, cols=6, style='Table Grid')
    cell=table.cell(0,0)
    paragraph = cell.paragraphs[0]
    ca = paragraph.add_run()
    ca.add_picture('ansm.jpg')
    table.cell(0,1).text=("FORMULAIRE DE DEMANDE D'AUTORISATION AUPRES DE L’AGENCE NATIONALE DE SECURITE DU MEDICAMENT ET DES PRODUITS DE SANTE "
                          "OU DE DEMANDE D’AVIS AU COMITÉ DE PROTECTION DES PERSONNES POUR UNE RECHERCHE MENTIONNEE AU 1° OU AU 2° DE L’ARTICLE L. 1121-1 DU CODE DE LA "
                          "SANTE PUBLIQUE NE PORTANT PAS SUR UN PRODUIT MENTIONNE A L’ARTICLE L. 5311-1 DU CODE DE LA SANTE PUBLIQUE - ")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.bold = True
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(10)
                    fontdebut.color.rgb = RGBColor(0x0,0x70,0xC0)
    a=table.cell(0,1)
    b=table.cell(0,5)
    a.merge(b)
                    
    '''Titre du document'''
    styles= document.styles
    style=styles.add_style('debut_page', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before
    paragraph_format.space_after
    fontdebut = style.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(11) 
    
    paragraph=document.add_paragraph("\nPartie réservée à l’ANSM / au Comité de protection des personnes (CPP) ", style='debut_page')
    table = document.add_table(rows=4, cols=3, style='Table Grid')
    table.cell(0,0).text=("Date de réception de la demande : ")
    table.cell(0,1).text=("Date de demande d’informations complémentaires : ")
    table.cell(0,2).text=("Autorisation de l'ANSM:")
    table.cell(1,2).text=("    Oui         Non\n")
    table.cell(2,2).text=("Date :")
    table.cell(3,0).text=("Date d’enregistrement de la demande considérée complète : \n\nDate du début de la procédure :")
    table.cell(3,1).text=("Date de réception des informations complémentaires / amendées :")
    table.cell(3,2).text=("Avis du CPP :")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial Narrow'
                    fontdebut.size = docx.shared.Pt(9)
                    if n==3:
                        fontdebut.bold = True
                    n=n+1
    for i in range(0,3):
        a=table.cell(0,i)
        b=table.cell(2,i)
        a.merge(b)
    
    paragraph=document.add_paragraph()
    sentence=paragraph.add_run("Partie à compléter par le demandeur :\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(11)
    sentence=paragraph.add_run("Recherche interventionnelle mentionnée au 1° de l’article L. 1121-1 du code de la santé publique :")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(9)
    sentence=paragraph.add_run(":    OUI      NON\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(9)
    sentence=paragraph.add_run("Recherche interventionnelle ne comportant que des risques et contraintes minimes mentionnée au 2° de l’article L. 1121-2 du code de la santé publique :")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(9)
    sentence=paragraph.add_run(":    OUI      NON\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(9)
    sentence=paragraph.add_run("\nDEMANDE D'AUTORISATION A L’ANSM :")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(9)
    sentence=paragraph.add_run(":   ouiI      non\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(9)
    sentence=paragraph.add_run("DEMANDE D’AVIS AU CPP :")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.size = docx.shared.Pt(9)
    sentence=paragraph.add_run(":   ouiI      non\n")
    fontdebut = sentence.font
    fontdebut.name = 'Arial'
    fontdebut.bold = True
    fontdebut.size = docx.shared.Pt(9)
    

    '''Partie A'''
    paragraph=document.add_paragraph("\nA. IDENTIFICATION DE LA RECHERCHE ", style='debut_page')
    table = document.add_table(rows=6, cols=6, style='Table Grid')
    #merge des cellules
    for i in range(0,6):
        if i==4:
            a=table.cell(i,0)
            b=table.cell(i,2)
            a.merge(b)
            a=table.cell(i,3)
            b=table.cell(i,5)
            a.merge(b)
        elif i==2 or i==3:
            a=table.cell(i,0)
            b=table.cell(i,1)
            a.merge(b)
            a=table.cell(i,2)
            b=table.cell(i,3)
            a.merge(b)
            a=table.cell(i,4)
            b=table.cell(i,5)
            a.merge(b)
        else:
            a=table.cell(i,0)
            b=table.cell(i,5)
            a.merge(b)
    table.cell(0,0).text=("Titre complet de la recherche : "+extract['titre_complet'])
    table.cell(1,0).text=("Numéro IDRCB d'enregistrement de la recherche :"+extract['num_idrcb'])
    table.cell(2,0).text=("Numéro de code du protocole de la recherche donné par le promoteur")
    table.cell(3,0).text=(extract['code_protocole'])
    table.cell(2,2).text=("Version")
    table.cell(2,4).text=("Date :")
    table.cell(4,0).text=("Nom ou titre abrégé de la recherche, le cas échéant : ")
    table.cell(4,3).text=(extract['titre_abrege'])
    table.cell(5,0).text=("Inscription au fichier VRB           oui            non")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
    
    '''Partie B'''
    paragraph=document.add_paragraph("\nB. IDENTIFICATION DU PROMOTEUR RESPONSABLE DE LA DEMANDE\n      B1. Promoteur  ", style='debut_page')
    table = document.add_table(rows=3, cols=2, style='Table Grid')
    for i in range(0,2):
        a=table.cell(i,0)
        b=table.cell(i,1)
        a.merge(b)
    table.cell(0,0).text=("Nom de l'organisme : "+extract['promoteur_nom_organisme'])
    table.cell(1,0).text=("Nom de la personne à contacter : "+extract['promoteur_nom_personne_contact'])
    table.cell(2,0).text=("Adresse : "+extract['promoteur_adresse'])
    table.cell(2,1).text=("Numéro de téléphone : "+extract['promoteur_num_telephone']+"\n\nNuméro de télécopie : "+extract['promoteur_num_telecopie']+"\n\nCourriel : "+extract['promoteur_courriel'])
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)

    paragraph=document.add_paragraph("\n       B2. Représentant légal  du promoteur dans l’Union européenne pour la recherche (si différent du promoteur) ",style='debut_page')
    table = document.add_table(rows=4, cols=2, style='Table Grid')
    for i in range(0,4):
        if not(i==2):
            a=table.cell(i,0)
            b=table.cell(i,1)
            a.merge(b)
    table.cell(0,0).text=("Nom de l'organisme : "+extract['promoteur_UE_nom_organisme'])
    table.cell(1,0).text=("Nom de la personne à contacter : "+extract['promoteur_UE_nom_personne_contact'])
    table.cell(2,0).text=("Adresse : "+extract['promoteur_UE_adresse'])
    table.cell(2,1).text=("Numéro de téléphone : "+extract['promoteur_UE_num_telephone']+"\n\nNuméro de télécopie : "+extract['promoteur_UE_num_telecopie']+"\n\nCourriel : "+extract['promoteur_UE_courriel'])
    table.cell(3,0).text=("Statut du promoteur :        commercial          non commercial")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)

    '''Partie C'''
    paragraph=document.add_paragraph("\n C. IDENTIFICATION DU DEMANDEUR \nC1. Demande pour l’ANSM", style='debut_page')
    table = document.add_table(rows=4, cols=2, style='Table Grid')
    for i in range(0,3):
        a=table.cell(i,0)
        b=table.cell(i,1)
        a.merge(b)
    table.cell(0,0).text=("Si promoteur, partie B1, si représentant légal du promoteur, partie B2 ")
    table.cell(1,0).text=("Nom de l'organisme : "+extract['demandeur_nom_organisme'])
    table.cell(2,0).text=("Nom de la personne à contacter : "+extract['demandeur_nom_personne_contact'])
    table.cell(3,0).text=("Adresse : "+extract['demandeur_UE_adresse'])
    table.cell(3,1).text=("Numéro de téléphone : "+extract['demandeur_UE_num_telephone']+"\n\nNuméro de télécopie : "+extract['demandeur_UE_num_telecopie']+"\n\nCourriel : "+extract['demandeur_UE_courriel'])
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
    
    paragraph=document.add_paragraph("\n  C2. Demande pour le Comité de protection des personnes", style='debut_page')
    table = document.add_table(rows=4, cols=2, style='Table Grid')
    for i in range(0,3):
        a=table.cell(i,0)
        b=table.cell(i,1)
        a.merge(b)
    table.cell(0,0).text=("Si promoteur, partie B1, si représentant légal du promoteur, partie B2 ")
    table.cell(1,0).text=("Nom de l'organisme : "+extract['demandeur_nom_organisme'])
    table.cell(2,0).text=("Nom de la personne à contacter : "+extract['demandeur_nom_personne_contact'])
    table.cell(3,0).text=("Adresse : "+extract['demandeur_UE_adresse'])
    table.cell(3,1).text=("Numéro de téléphone : "+extract['demandeur_UE_num_telephone']+"\n\nNuméro de télécopie : "+extract['demandeur_UE_num_telecopie']+"\n\nCourriel : "+extract['demandeur_UE_courriel'])
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
    
def partie_D_a_G(document, extract):
    
    '''Partie D'''
    paragraph=document.add_paragraph("\n D. DONNÉES SUR LE(S) PRODUITS(S) EXPÉRIMENTAL (AUX) UTILISÉ(S) DANS LA RECHERCHE : PRODUITS(S) ÉTUDIÉ(S) OU UTILISÉ(S) COMME COMPARATEUR(S) ", style='debut_page')
    table = document.add_table(rows=4, cols=1, style='Table Grid')
    table.cell(0,0).text=("Indiquer ici quel PE est concerné par cette section D ; si nécessaire, utiliser d’autres fiches pour chaque PE utilisé dans l’essai (à numéroter de 1 à n) :")
    table.cell(1,0).text=("Cette section concerne le PE numéro :  ")
    table.cell(2,0).text=("PE étudié               oui              non")
    table.cell(3,0).text=("PE utilisé comme comparateur            oui           non")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
    
    '''Partie D1'''
    paragraph=document.add_paragraph("\n  D.1. DESCRIPTION DU PRODUIT EXPÉRIMENTAL", style='debut_page')
    table = document.add_table(rows=5, cols=1, style='Table Grid')
    table.cell(0,0).text=("Nom du produit, le cas échéant : "+extract['produit_nom'])
    table.cell(1,0).text=("Nom de code, le cas échéant : "+extract['produit_nom_code'])
    table.cell(2,0).text=("Voie d’administration (utiliser les termes standard) : "+extract['produit_voie_administration'])
    table.cell(3,0).text=("Dosage (préciser tous les dosages utilisés) : \n- Concentration (nombre) : "+extract['produit_dosage_concentration']+"\n- Unité de concentration : "+extract['produit_dosage_unite_concentration'])
    table.cell(4,0).text=("Le produit expérimental contient-il une substance active :\n\n- d’origine chimique ?            oui          non	\n- d’origine biologique ?          oui          non\n\n"
                          "Est-ce :\n- un produit à base de plantes ?         oui         non\n- un produit contenant des organismes génétiquement modifiés ?		 oui  	 non\n"
                          "- un autre type de produit ?   oui        non	\n\n• Si oui, préciser :")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
    
    
    paragraph=document.add_paragraph("\n Type de produit ", style='debut_page')
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
    
    paragraph=document.add_paragraph("\n Fabricant du produit utilisé", style='debut_page')
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("Fabricant\n- Nom de l’établissement : "+extract['fabriquant_dispositif_nom']+"\n- Adresse : "+extract['fabriquant_dispositif_adresse'])
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
    
    '''Partie E'''
    paragraph=document.add_paragraph("\n E. INFORMATIONS SUR LE PLACEBO (le cas échéant) (répéter la section si nécessaire) ", style='debut_page')
    table = document.add_table(rows=6, cols=1, style='Table Grid')
    table.cell(0,0).text=("Cette section se rapporte au placebo n° :  "+extract['placebo_numero'])
    table.cell(1,0).text=("Un placebo est-il utilisé ? 		   oui           non")
    table.cell(2,0).text=("De quel produit expérimental est-ce un placebo ? "+extract['placebo_numero_ME'])
    table.cell(3,0).text=("Préciser le(s) numéro(s) de PE selon la section D.")
    table.cell(4,0).text=("Voie d’administration : "+extract['placebo_voie_administration'])
    table.cell(5,0).text=("Composition, hormis la (les) substance(s) active(s) : \n- est-elle identique à celle du produit expérimental étudié?       oui           non\n\n•  Si non, préciser les principaux composants :   ")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
    
    paragraph=document.add_paragraph("\n FABRICANT DU PLACEBO", style='debut_page')
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("Fabricant\n- Nom de l’établissement : "+extract['fabriquant_placebo_nom']+"\n- Adresse :"+extract['fabriquant_placebo_adresse'])
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
    
    '''Partie G'''
    paragraph=document.add_paragraph("\n G. INFORMATIONS GÉNÉRALES SUR L’ESSAI", style='debut_page')
    table = document.add_table(rows=7, cols=1, style='Table Grid')
    table.cell(0,0).text=("Condition médicale ou pathologie étudiée")
    table.cell(1,0).text=("Préciser la condition médicale : "+extract['pathologie_etudiee']+"\nClassification CIM  : 	"+extract['classification_cim']+"\n\nEst-ce une maladie rare ?       oui           non\n\nObjectif(s) de l’essai\nObjectif principal : "+extract['objectif_principal']+"\nObjectifs secondaires : "+extract['objectif_secondaire'])
    table.cell(2,0).text=("Principaux critères d’inclusion (énumérer les plus importants) " +extract['critere_inclusion_longue'])
    table.cell(3,0).text=("Principaux critères de non inclusion (énumérer les plus importants "+extract['critere_non_inclusion_longue'])
    table.cell(4,0).text=("Critère(s) d’évaluation principal (aux) " +extract['critere_jugement_principal_longue']+" "+extract['critere_jugement_secondaire_longue'])
    table.cell(5,0).text=("Domaine(s) d’étude :\n")
    table.cell(6,0).text=("- Physiologie\n- Physiopathologie\n- Epidémiologie\n- Génétique\n- Science du comportement\n- Produits à visée nutritionnelle\n- Stratégies diagnostiques\n- Stratégies thérapeutiques et préventives\n\n                • Si autres préciser :")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
                    if n==0 or n==5:
                        fontdebut.bold=True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,0)
    a.merge(b)
    a=table.cell(5,0)
    b=table.cell(6,0)
    a.merge(b)
    
    paragraph=document.add_paragraph("\n Méthodologie de l’essai", style='debut_page')
    table = document.add_table(rows=7, cols=4, style='Table Grid')
    table.cell(0,0).text=("Tirage au sort :\nLa recherche comporte-t-elle une comparaison de  groupes?")
    table.cell(0,3).text=("   oui           non\n   oui           non")
    for i in range (0,8):
        if i==0 or i==4:
            a=table.cell(i,0)
            b=table.cell(i,2)
            a.merge(b)
        else:
            a=table.cell(i,0)
            b=table.cell(i,3)
            a.merge(b)
    table.cell(1,0).text=("Autre méthodologie")
    table.cell(2,0).text=("Préciser le(s) comparateur(s) utilisé(s) :\n")
    table.cell(3,0).text=("- (d’) autre(s) produits(s)\n- placebo \n- autre \n\n            • Si oui, préciser :")
    table.cell(4,0).text=("La recherche  est-elle multicentrique ?\nLa recherche  est-elle prévue pour être menée dans plusieurs Etat membres ?\nCette recherche implique-t-elle des pays tiers ?")
    table.cell(4,3).text=("   oui           non\n   oui           non\n   oui           non")
    table.cell(5,0).text=("Durée maximale de participation pour un sujet selon le protocole : "+extract['duree_participation']+"\n")
    table.cell(6,0).text=("Définition de la fin de la recherche et justification, si celle-ci ne correspond pas à la date de la dernière visite de la dernière personne participant à la recherche :\n\nEstimation initiale de la durée de la recherche   :\n             • en France :       						 "+extract['duree_totale_etude']+" ans  mois\n             • dans tous les pays concernés par la recherche :  	 	 "+extract['duree_totale_etude']+" ans  mois")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==3 or n==19:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
                    if n==3 or n==11 or n==19 or n==23:
                        fontdebut.bold=True
                    n=n+1
    a=table.cell(2,0)
    b=table.cell(3,0)
    a.merge(b)
    a=table.cell(5,0)
    b=table.cell(6,0)
    a.merge(b)
    
def partie_H_fin(document, extract):
    
    '''Partie H'''
    paragraph=document.add_paragraph("\nH. PERSONNES PARTICIPANT A LA RECHERCHE BIOMEDICALE", style='debut_page')
    table = document.add_table(rows=3, cols=8, style='Table Grid')
    for i in range (0,2):
        for j in range (0,8):
            if j%2==0:
                a=table.cell(i,j)
                b=table.cell(i,j+1)
                a.merge(b)
    a=table.cell(2,1)
    b=table.cell(2,7)
    a.merge(b)
    table.cell(0,0).text=("Tranche d’âge étudiée")
    table.cell(0,2).text=("< 18 ans")
    table.cell(0,4).text=("18-65 ans")
    table.cell(0,6).text=("> 65 ans")
    table.cell(1,1).text=("Nouveaux-nés prématurés\n(jusqu’à l’âge gestationnel \n≤ 37 semaines)\nNouveau-nés (0-27 jours) \nNourrissons (28 jours - 23 mois)\nEnfants (2-11 ans)\nAdolescents (12-17 ans)")
    table.cell(2,0).text=("Sexe")
    table.cell(2,1).text=("              Femmes                Hommes")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    if n==0 or n==2 or n==4 or n==6:
                        fontdebut.bold=True
                        fontdebut.size = docx.shared.Pt(9)
                    elif n==8:
                        fontdebut.size = docx.shared.Pt(8)
                    n=n+1
    
    paragraph=document.add_paragraph("\nPersonnes participant à la recherche", style='debut_page')
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    table.cell(0,0).text=("Volontaires sains\nMalades\nFemmes enceintes\nFemmes allaitantes\nPersonnes en situation  d’urgence\nPersonnes incapables de donner personnellement  leur consentement\n dont majeurs sous tutelle")
    table.cell(0,1).text=("oui           non\noui           non\noui           non\noui           non\noui           non\noui           non\noui           non")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if n==1:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
                    if n==1:
                        fontdebut.bold=True
                    n=n+1
    
    paragraph=document.add_paragraph("\nNombre prévu de personnes à inclure :", style='debut_page')
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("• en France :\n"
                          "En cas d’essai mené dans plusieurs pays :\n"
                          "         • dans l’Union européenne : \n"
                          "         • pour l’ensemble des pays participant à la recherche: ")
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
        
    '''Partie I'''
    paragraph=document.add_paragraph("\nI. INVESTIGATEURS ET LIEUX DE RECHERCHE\n"
                                     "          I.1. Investigateur coordonnateur ", style='debut_page')
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    a=table.cell(0,1)
    b=table.cell(1,1)
    a.merge(b)
    table.cell(0,0).text=("Nom :	"+extract['investigateur_coordinateur_nom']+"	\nPrénoms : "+extract['investigateur_coordinateur_prenom']+"\nQualification, spécialité :  "+extract['investigateur_coordinateur_qualification']+"\nCourriel : "+extract['investigateur_coordinateur_courriel'])
    table.cell(0,1).text=("Adresse : "+extract['investigateur_coordinateur_adresse'])
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
    
    paragraph=document.add_paragraph("\n        I.2. Autres investigateurs ", style='debut_page')
    for i in range(len(extract['autre_investigateur_nom'])):
        table = document.add_table(rows=1, cols=2, style='Table Grid')
        a=table.cell(0,1)
        b=table.cell(1,1)
        a.merge(b)
        table.cell(0,0).text=("Nom :	"+extract['autre_investigateur_nom'][i]+"	\nPrénoms : "+extract['autre_investigateur_prenom'][i]+"\nQualification, spécialité :  "+extract['autre_investigateur_qualification'][i]+"\nCourriel : "+extract['autre_investigateur_courriel'][i])
        table.cell(0,1).text=("Adresse : "+extract['autre_investigateur_adresse'][i])
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        fontdebut = run.font
                        fontdebut.name = 'Arial'
                        fontdebut.size = docx.shared.Pt(9)
   
    paragraph=document.add_paragraph("\n        I.3 Lieu de recherche ( le cas échéant, si la recherche doit se dérouler dans un lieu nécessitant une 	     autorisation mentionnée à l’article L. 1121-13 du code de la santé publique) :", style='debut_page')
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.cell(0,0).text=("Intitulé du lieu: "+extract['lieu_recherche_intitule']+"\nN° d’autorisation: "+extract['lieu_recherche_num_autorisation']+"\ndélivré le: "+extract['lieu_recherche_delivre_le']+"\ndate de limite de validité: "+extract['lieu_recherche_date_limite_validite'])
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
                    
    '''Partie J'''
    paragraph=document.add_paragraph("\nJ. INFORMATION SUR LE COMITE DE PROTECTION DES PERSONNES (CPP)", style='debut_page')
    table = document.add_table(rows=6, cols=4, style='Table Grid')
    a=table.cell(0,0)
    b=table.cell(0,3)
    a.merge(b)
    table.cell(0,0).text=("Nom et adresse : "+extract['CPP'])
    table.cell(1,0).text=("Avis :")
    table.cell(1,1).text=("à demander")
    table.cell(1,2).text=("en cours\nDate de soumission :")
    table.cell(1,3).text=("donné")
    a=table.cell(2,1)
    b=table.cell(2,3)
    a.merge(b)
    table.cell(2,0).text=("Si donné, préciser :")
    a=table.cell(3,2)
    b=table.cell(3,3)
    a.merge(b)
    a=table.cell(4,2)
    b=table.cell(4,3)
    a.merge(b)
    a=table.cell(3,0)
    b=table.cell(4,0)
    a.merge(b)
    a=table.cell(3,1)
    b=table.cell(4,1)
    a.merge(b)
    table.cell(3,0).text=("Date de l’avis:")
    table.cell(3,1).text=("Avis favorable\n         Oui        non")
    table.cell(3,2).text=("Avis défavorable ")
    table.cell(4,2).text=("un second examen a-t-il été demandé à un autre CPP ? :")
    table.cell(5,1).text=("Si oui lequel ?")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
                    if not(n==0 or n==13):
                        fontdebut.bold=True
                    n=n+1
    a=table.cell(3,2)
    b=table.cell(4,2)
    a.merge(b)
    
    
    '''Partie K'''
    paragraph=document.add_paragraph("\nK. SIGNATURE DU DEMANDEUR EN FRANCE", style='debut_page')
    table = document.add_table(rows=3, cols=2, style='Table Grid')
    table.cell(0,0).text=("Par la présente, j’atteste / j’atteste au nom du promoteur (rayer la mention inutile) ce qui suit :\n")
    table.cell(1,0).text=(" - les informations fournies ci-dessus à l’appui de la demande sont exactes ;\n"
                          "- la recherche sera réalisée conformément au protocole, à la réglementation nationale et aux principes de bonnes pratiques cliniques ;\n"
                          "- il est raisonnable d’entreprendre la recherche proposée ; \n"
                          "- je soumettrai un résumé du rapport final de la recherche à l’ANSM et au Comité de protection des personnes concerné au plus tard 1 an après la fin de l’essai dans tous les pays ;\n"
                          "- je déclarerai la date effective du commencement de l’essai à l’ANSM et au Comité de protection des personnes concerné dès qu’elle sera connue.")
    table.cell(2,0).text=("DEMANDEUR auprès de l’ANSM\n(comme indiqué à la section C1) :\n\nDate : \n\nSignature : \n\nNom :")
    table.cell(2,1).text=("DEMANDEUR auprès du CPP\n(comme indiqué à la section C2) :\n\nDate : \n\nSignature : \n\nNom :")
    n=0
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    fontdebut = run.font
                    fontdebut.name = 'Arial'
                    fontdebut.size = docx.shared.Pt(9)
                    if n==0:
                        fontdebut.bold=True
                    n=n+1
    a=table.cell(0,0)
    b=table.cell(1,1)
    a.merge(b)