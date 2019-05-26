# -*- coding: utf-8 -*-
"""
Created on Sat Feb 16 14:09:42 2019

@author: Romain
"""


import re
from docx import Document


def extract1(dico):
    
    #dabord on extrait tout qu'on ajoute dans une liste
    f1 = open(dico['le_chemin'], 'rb') #ouvre le premier fichier
    #f1 = open("trame-simplifiée-cat-1.docx", 'rb')
    doc = Document(f1)
    fullText=[]
    for para in doc.paragraphs:
        para.text+="#SAUT5089"
        if para.text=="#SAUT5089":
            para.text=""
        fullText.append(para.text)   
    f1.close()
    
    #puis on met tous les éléments de la liste bout a bout dans un immense string
    texte1=""
    for i in fullText:
        texte1+=i
    texte2=texte1.replace("\xa0","")
    texte=texte2.replace("\n","#SAUT5089")



    #creation du dico de donnes
    infos={}
    #on ajoute les éléments 1 par 1
    x=re.search(r"(?<=Titre complet de la recherche:).*(?=Nom ou titre)",texte).group()
    x=x.replace("#SAUT5089", "")
    infos["titre_complet"]=x
    x=re.search(r"(?<=Nom ou titre abrégé:).*(?=N° de code du protocole)",texte).group()
    x=x.replace("#SAUT5089", "")
    infos["titre_abrege"]=x
    x=re.search(r"(?<=Protocole ... version n°… du .../.../…\):).*(?=N°EudraCT)",texte).group()
    x=x.replace("#SAUT5089", "")
    infos["code_protocole"]=x
    x=re.search(r"(?<=N°EudraCT:).*(?=N°IDRCB:)",texte).group()
    x=x.replace("#SAUT5089", "")
    infos["num_eudract"]=x
    x=re.search(r"(?<=N°IDRCB:).*(?=Classification CIM: )",texte).group()
    x=x.replace("#SAUT5089", "")
    infos["num_idrcb"]=x
    x=re.search(r"(?<=Classification CIM:).*(?=Préciser la condition médicale)",texte).group()
    x=x.replace("#SAUT5089", "")
    infos["classification_cim"]=x
    x=re.search(r"(?<=condition médicale ou pathologie étudiée:).*(?=Identification du promoteur responsable)",texte).group()
    x=x.replace("#SAUT5089", "")
    infos["pathologie_etudiee"]=x
    
    #en premier on récupère tout le bloc promoteur
    x=re.search(r"(?<=Identification du promoteur responsable de la demande).*(?=Représentant légal du promoteur dans l’UE)",texte).group()
    #puis tous les éléments du promoteur 1 par 1
    y=re.search(r"(?<=Nom de l’organisme:).*(?=Nom de la personne à contacter:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['promoteur_nom_organisme']=y
    y=re.search(r"(?<=Nom de la personne à contacter:).*(?=Adresse:)",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['promoteur_nom_personne_contact']=y
    y=re.search(r"(?<=Adresse:).*(?=N° téléphone:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['promoteur_adresse']=y
    y=re.search(r"(?<=N° téléphone:).*(?=N° télécopie:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['promoteur_num_telephone']=y
    y=re.search(r"(?<=N° télécopie:).*(?=Courriel:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['promoteur_num_telecopie']=y
    y=re.search(r"(?<=Courriel:).*",x).group()
    y=y.replace("#SAUT5089", "")
    infos['promoteur_courriel']=y
    
    x=re.search(r"(?<=Représentant légal du promoteur dans l’UE ).*(?=Identification des investigateurs)",texte).group()
    #puis tous les éléments du promoteur 1 par 1
    y=re.search(r"(?<=Nom de l’organisme:).*(?=Nom de la personne à contacter:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['promoteur_UE_nom_organisme']=y
    y=re.search(r"(?<=Nom de la personne à contacter:).*(?=Adresse:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['promoteur_UE_nom_personne_contact']=y
    y=re.search(r"(?<=Adresse:).*(?=N° téléphone:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['promoteur_UE_adresse']=y
    y=re.search(r"(?<=N° téléphone:).*(?=N° télécopie:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['promoteur_UE_num_telephone']=y
    y=re.search(r"(?<=N° télécopie:).*(?=Courriel:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['promoteur_UE_num_telecopie']=y
    y=re.search(r"(?<=Courriel:).*",x).group()
    y=y.replace("#SAUT5089", "")
    infos['promoteur_UE_courriel']=y
    
    #idem pour investigateur coordinateur
    x=re.search(r"(?<=Investigateur coordinateur:).*(?=Autres investigateurs:)",texte).group()
    #puis tous les éléments de l'investigateur coordinateur
    y=re.search(r"(?<=Nom:).*(?=Prénom:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['investigateur_coordinateur_nom']=y
    y=re.search(r"(?<=Prénom:).*(?=Qualification, spécialité: )",x).group()
    y=y.replace("#SAUT5089", "")
    infos['investigateur_coordinateur_prenom']=y
    y=re.search(r"(?<=Qualification, spécialité:).*(?=Adresse professionnelle:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['investigateur_coordinateur_qualification']=y
    y=re.search(r"(?<=Adresse professionnelle:).*(?=Nom de l’établissement:)",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['investigateur_coordinateur_adresse_professionnelle']=y
    y=re.search(r"(?<=Nom de l’établissement:).*(?=Service:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_nom_etablissement']=y
    y=re.search(r"(?<=Service:).*(?=Adresse:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_service']=y
    y=re.search(r"(?<=Adresse:).*(?=N° téléphone:)",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['investigateur_coordinateur_adresse']=y
    y=re.search(r"(?<=N° téléphone:).*(?=N° télécopie:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['investigateur_coordinateur_telephone']=y
    y=re.search(r"(?<=N° télécopie:).*(?=Courriel:)",x).group()
    y=y.replace("#SAUT5089", "")
    infos['investigateur_coordinateur_telecopie']=y
    y=re.search(r"(?<=Courriel:).*",x).group()
    y=y.replace("#SAUT5089", "")
    infos['investigateur_coordinateur_courriel']=y
    
    #pour les autres investigateurs, plus subtile
    x=re.search(r"(?<=Autres investigateurs:).*(?=Identification du demandeur)",texte).group()
    #je m'en sers pour les regex
    x+="Nom: "
    x=x.replace("#SAUT5089", " ")
    #initialise avec des listes vides
    infos['autre_investigateur_nom']=[]
    infos['autre_investigateur_prenom']=[]
    infos['autre_investigateur_qualification']=[]
    infos['autre_investigateur_adresse_professionnelle']=[]
    infos['autre_investigateur_nom_etablissement']=[]
    infos['autre_investigateur_service']=[]
    infos['autre_investigateur_adresse']=[]
    infos['autre_investigateur_telephone']=[]
    infos['autre_investigateur_telecopie']=[]
    infos['autre_investigateur_courriel']=[]
    #compte le nombre d'autres investigateurs
    n=x.count('Nom: ')
    #on boucle pour chaque investigateur
    for i in range(n-1): 
        infos['autre_investigateur_nom'].append(re.search(r"(?<=Nom:).*?(?=Prénom:)",x).group())
        infos['autre_investigateur_prenom'].append(re.search(r"(?<=Prénom:).*?(?=Qualification, spécialité: )",x).group())
        infos['autre_investigateur_qualification'].append(re.search(r"(?<=Qualification, spécialité:).*?(?=Adresse professionnelle:)",x).group())
        infos['autre_investigateur_adresse_professionnelle'].append(re.search(r"(?<=Adresse professionnelle:).*?(?=Nom de l’établissement:)",x).group())
        infos['autre_investigateur_nom_etablissement'].append(re.search(r"(?<=Nom de l’établissement:).*?(?=Service: )",x).group())
        infos['autre_investigateur_service'].append(re.search(r"(?<=Service:).*?(?=Adresse:)",x).group())
        infos['autre_investigateur_adresse'].append(re.search(r"(?<=Adresse:).*?(?=N° téléphone:)",x).group())
        infos['autre_investigateur_telephone'].append(re.search(r"(?<=N° téléphone:).*?(?=N° télécopie:)",x).group())
        infos['autre_investigateur_telecopie'].append(re.search(r"(?<=N° télécopie:).*?(?=Courriel:)",x).group())
        infos['autre_investigateur_courriel'].append(re.search(r"(?<=Courriel:).*?(?=Nom:)",x).group())
        #on recup la taille du premier bloc investigateur
        z=len(re.search(r"(?<=Nom:).*?(?=Nom: )",x).group())
        #on enleve ce bloc a x qui contient tous les investigateurs, ainsi a la prochaine boucle la regex se fera sur l'investigateur suivant
        x=x[(z-1):]
        
    #idem pour demande
    x=re.search(r"(?<=Identification du demandeur:).*(?=Justification de l’étude)",texte).group()
    #puis tous les éléments du demandeur
    y=re.search(r"(?<=Nom de l’organisme:).*(?=Nom de la personne à contacter:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['demandeur_nom_organisme']=y
    y=re.search(r"(?<=Nom de la personne à contacter:).*(?=Adresse:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['demandeur_nom_personne_contact']=y
    y=re.search(r"(?<=Adresse:).*(?=N° téléphone:)",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['demandeur_UE_adresse']=y
    y=re.search(r"(?<=N° téléphone:).*(?=N° télécopie:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['demandeur_UE_num_telephone']=y
    y=re.search(r"(?<=N° télécopie:).*(?=Courriel:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['demandeur_UE_num_telecopie']=y
    y=re.search(r"(?<=Courriel:).*",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['demandeur_UE_courriel']=y  
    #justification de létude
    table = doc.tables[0]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage dans courte
    courte=re.search(r"(?<=Bref rappel \(données de la littérature scientifique, pathologie, domaine d’étude\)\#SAUT5089).*",courte).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")
    infos['justification_etude_courte']=courte
    infos['justification_etude_longue']=longue
  
    #benefices de l'étude
    benefice=re.search(r"(?<=notamment les bénéfices escomptés pour les personnes qui se prêtent à la recherche\.).*(?=Risques:)",texte).group()
    benefice=benefice.replace("#SAUT5089", "\n")
    infos['benefices']=benefice
    
#    #benefices de l'étude
#    benefice=re.search(r"(?<=notamment les bénéfices escomptés pour les personnes qui se prêtent à la recherche\.).*(?=Risques:)",texte1).group()
#    infos['benefices']=benefice
#    
#    #risques de l'étude
#    risque=re.search(r"(?<=visant à éviter et/ou prendre en charge les événements inattendus\)\.).*(?=Retombées attendues)",texte1).group()
#    infos['risques']=risque
    
     #risques de l'étude
    risque=re.search(r"(?<=visant à éviter et/ou prendre en charge les événements inattendus\)\.).*(?=Retombées attendues)",texte).group()
    risque=risque.replace("#SAUT5089", "\n")
    infos['risques']=risque
    
    #retombées attendues
    table = doc.tables[1]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    longue=longue.replace("\n","#SAUT5089")
    courte=courte.replace("\n","#SAUT5089")
    #on ajoute '\n' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"

    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Description des retombées attendues par cette recherche\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=d’augmentation de l’arsenal thérapeutique,…\)\.\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")    
    infos['retombee_attenduees_courte']=courte
    infos['retombee_attenduees_longue']=longue
    #objectif principal
#    principal=re.search(r"(?<=Objectif Principal:).*(?=Objectif secondaires:)",texte1).group()
#    infos['objectif_principal']=principal
    
    principal=re.search(r"(?<=Objectif Principal:).*(?=Objectif secondaires:)",texte).group()
    principal=principal.replace("#SAUT5089", "\n")
    infos['objectif_principal']=principal
    
#    #objectif secondaire
#    secondaire=re.search(r"(?<=Objectif secondaires:).*?(?=Critères de Jugement)",texte1).group()
#    infos['objectif_secondaire']=secondaire
    
    #objectif secondaire
    secondaire=re.search(r"(?<=Objectif secondaires:).*?(?=Critères de Jugement)",texte).group()
    secondaire=secondaire.replace("#SAUT5089", "\n")
    infos['objectif_secondaire']=secondaire
    
    #critères de jugement principal
    table = doc.tables[2]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    longue=longue.replace("\n","#SAUT5089")
    courte=courte.replace("\n","#SAUT5089")
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Un seul critère correspondant à l’objectif principal \#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=Il permettra également le calcul de l’effectif de l’étude\. \#SAUT5089).*",longue).group()
    longue=longue.replace("#SAUT5089","\n")
    courte=courte.replace("#SAUT5089","\n")    
    infos['critere_jugement_principal_courte']=courte
    infos['critere_jugement_principal_longue']=longue
    
    #critères de jugement secondaire
    table = doc.tables[3]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    longue=longue.replace("\n","#SAUT5089")
    courte=courte.replace("\n","#SAUT5089") 
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Liste de tous les critères de jugement secondaires\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=répondant aux objectifs secondaires\.\#SAUT5089).*",longue).group()
    longue=longue.replace("#SAUT5089","\n")
    courte=courte.replace("#SAUT5089","\n")    
    infos['critere_jugement_secondaire_courte']=courte
    infos['critere_jugement_secondaire_longue']=longue
    
    #critères d'inclusion
    table = doc.tables[4]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    longue=longue.replace("\n","#SAUT5089")
    courte=courte.replace("\n","#SAUT5089")     
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=à la partie correspondante dans le corps du protocole § 6\.1\)\#SAUT5089).*",courte).group()
    longue=longue.replace("#SAUT5089","\n")
    courte=courte.replace("#SAUT5089","\n")     
    infos['critere_inclusion_courte']=courte
    infos['critere_inclusion_longue']=longue
    
    #critères de non inclusion
    table = doc.tables[5]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    longue=longue.replace("\n","#SAUT5089")
    courte=courte.replace("\n","#SAUT5089")       
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=à la partie correspondante dans le corps du protocole § 6\.2\)\#SAUT5089).*",courte).group()
    longue=longue.replace("#SAUT5089","\n")
    courte=courte.replace("#SAUT5089","\n")       
    infos['critere_non_inclusion_courte']=courte
    infos['critere_non_inclusion_longue']=longue
    
#    #justification inclusion
#    justif=re.search(r"(?<=Justifications de l’inclusion de personnes visées:).*(?=Modalités de recrutements)",texte1).group()
#    infos['justification_inclusion']=justif
    
    #justification inclusion
    justif=re.search(r"(?<=Justifications de l’inclusion de personnes visées:).*(?=Modalités de recrutements)",texte).group()
    justif=justif.replace("#SAUT5089", "\n")
    infos['justification_inclusion']=justif
    
#    #modalités_recrutement
#    recru=re.search(r"(?<=Modalités de recrutements:).*(?=Traitement et stratégie)",texte1).group()
#    infos['modalite_recrutement']=recru
    
    #modalités_recrutement
    recru=re.search(r"(?<=Modalités de recrutements:).*(?=Traitement et stratégie)",texte).group()
    recru=recru.replace("#SAUT5089", "\n")
    infos['modalite_recrutement']=recru
    
    #traitement et stratégie
    table = doc.tables[6]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    longue=longue.replace("\n","#SAUT5089")
    courte=courte.replace("\n","#SAUT5089")  
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=traitements/stratégies/procédures\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=la durée du traitement et de la voie d’administration\.\#SAUT5089).*",longue).group()
    longue=longue.replace("#SAUT5089","\n")
    courte=courte.replace("#SAUT5089","\n")  
    infos['traitement_strategie_courte']=courte
    infos['traitement_strategie_longue']=longue
    
    #fabriquant du dispositif
    x=re.search(r"(?<=Fabriquant du dispositif étudié:).*(?=Fabriquant du placebo)",texte).group()
    #puis tous les éléments du fabriquant
    y=re.search(r"(?<=Nom:).*(?=Adresse)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['fabriquant_dispositif_nom']=y
    y=re.search(r"(?<=Adresse:).*(?=N° téléphone:)",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['fabriquant_dispositif_adresse']=y
    y=re.search(r"(?<=N° téléphone:).*(?=N° télécopie:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['fabriquant_dispositif_num_telephone']=y
    y=re.search(r"(?<=N° télécopie:).*(?=Courriel:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['fabriquant_dispositif_num_telecopie']=y
    y=re.search(r"(?<=Courriel:).*",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['fabriquant_dispositif_courriel']=y  
    
    #fabriquant du placebo
    x=re.search(r"(?<=Fabriquant du placebo:).*(?=Description du produit/médicament)",texte).group()
    #puis tous les éléments du fabriquant
    y=re.search(r"(?<=Nom:).*(?=Adresse)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['fabriquant_placebo_nom']=y
    y=re.search(r"(?<=Adresse:).*(?=N° téléphone:)",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['fabriquant_placebo_adresse']=y
    y=re.search(r"(?<=N° téléphone:).*(?=N° télécopie:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['fabriquant_placebo_num_telephone']=y
    y=re.search(r"(?<=N° télécopie:).*(?=Courriel:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['fabriquant_placebo_num_telecopie']=y
    y=re.search(r"(?<=Courriel:).*",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['fabriquant_placebo_courriel']=y  
    
    #description produit
    x=re.search(r"(?<=Description du produit/médicament expérimental:).*(?=Informations sur le placebo)",texte).group()
    #puis tous les éléments du produit
    y=re.search(r"(?<=Nom du produit:).*(?=Nom de code)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['produit_nom']=y
    y=re.search(r"(?<=Nom de code:).*(?=Voie d’administration)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['produit_nom_code']=y
    y=re.search(r"(?<=Voie d’administration:).*(?=Dosage concentration)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['produit_voie_administration']=y
    y=re.search(r"(?<=Dosage concentration :).*(?=Dosage unité de concentration)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['produit_dosage_concentration']=y
    y=re.search(r"(?<=Dosage unité de concentration:).*",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['produit_dosage_unite_concentration']=y 
    
    #description placebo
    x=re.search(r"(?<=Informations sur le placebo).*(?=Etude)",texte).group()
    #puis tous les éléments du placebo
    y=re.search(r"(?<=Numéro du placebo:).*(?=De quel produit expérimental)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['placebo_numero']=y
    y=re.search(r"(?<=préciser le numéro du ME:).*(?=Voie d’administration)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['placebo_numero_ME']=y
    y=re.search(r"(?<=Voie d’administration:).*",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['placebo_voie_administration']=y
   
    #taille de l'étude
    table = doc.tables[7]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    longue=longue.replace("\n","#SAUT5089")
    courte=courte.replace("\n","#SAUT5089")      
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Nombre de personnes à inclure:).*",courte).group()
    longue=longue.replace("#SAUT5089","\n")
    courte=courte.replace("#SAUT5089","\n")      
    infos['taille_etude_courte']=courte
    infos['taille_etude_longue']=longue
    
#    #modalités de l'indemnisation
#    indem=re.search(r"(?<=Modalités et montant de l’indemnisation des personnes se prêtant à la recherche:).*(?=Justification de l’existence)",texte1).group()
#    infos['indemnisation']=indem
    
        #modalités de l'indemnisation
    indem=re.search(r"(?<=Modalités et montant de l’indemnisation des personnes se prêtant à la recherche:).*(?=Justification de l’existence)",texte).group()
    indem=indem.replace("#SAUT5089", " ")
    infos['indemnisation']=indem
    
#    #justification existence
#    justi=re.search(r"(?<=Justification de l’existence:).*?(?=Durée)",texte1).group()
#    infos['justification_existence']=justi
    
     #justification existence
    justi=re.search(r"(?<=Justification de l’existence:).*?(?=Durée)",texte).group()
    justi=justi.replace("#SAUT5089", "\n")
    infos['justification_existence']=justi
    
#    #durée des inclusions
#    x=re.search(r"(?<=Durée prévue des inclusions:).*(?=Durée de participation pour une personne se prêtant à la recherche)",texte1).group()
#    infos['duree_inclusion']=x
    
     #durée des inclusions
    x=re.search(r"(?<=Durée prévue des inclusions:).*(?=Durée de participation pour une personne se prêtant à la recherche)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos['duree_inclusion']=x
    
#    #durée de participation
#    x=re.search(r"(?<=c'est-à-dire la dernière visite du dernier patient inclus.).*(?=Durée totale de l’étude)",texte1).group()
#    infos['duree_participation']=x
    
     #durée de participation
    x=re.search(r"(?<=c'est-à-dire la dernière visite du dernier patient inclus.).*(?=Durée totale de l’étude)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos['duree_participation']=x
    
#    #durée totale
#    x=re.search(r"(?<=Durée totale de l’étude:).*(?=Analyse statistiques des données)",texte1).group()
#    infos['duree_totale_etude']=x
    
    #durée totale
    x=re.search(r"(?<=Durée totale de l’étude:).*(?=Analyse statistiques des données)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos['duree_totale_etude']=x
    
    #analyse statistique
    table = doc.tables[8]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Bref rappel des méthodes statistiques\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=données manquantes, inutilisées ou non valides\.\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089", "\n")
    longue=longue.replace("#SAUT5089", "\n")
    infos['analyse_statistique_courte']=courte
    infos['analyse_statistique_longue']=longue
    
    #lieu de recherche
    x=re.search(r"(?<=dans un lieu nécessitant une autorisation de l’ARS\)).*(?=Plateau technique)",texte).group()
    #puis tous les éléments du placebo
    y=re.search(r"(?<=Intitulé du lieu:).*(?=N° d’autorisation:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['lieu_recherche_intitule']=y
    y=re.search(r"(?<=N° d’autorisation:).*(?=Délivré le:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['lieu_recherche_num_autorisation']=y
    y=re.search(r"(?<=Délivré le:).*(?=Date de limite de validité:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['lieu_recherche_delivre_le']=y
    y=re.search(r"(?<=Date de limite de validité:).*(?=Nom et adresse:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['lieu_recherche_date_limite_validite']=y
    y=re.search(r"(?<=Nom et adresse:).*",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['lieu_recherche_nom_adresse']=y
    
    #plateau technique
    x=re.search(r"(?<=Plateau technique).*(?=Nom du CPP:)",texte).group()
    #puis tous les éléments du placebo
    y=re.search(r"(?<=Organisme:).*(?=Nom de la personne à contacter)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['plateau_technique_organisme']=y
    y=re.search(r"(?<=Nom de la personne à contacter:).*(?=Adresse)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['plateau_technique_personne_contact']=y
    y=re.search(r"(?<=Adresse:).*(?=N° téléphone)",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['plateau_technique_adresse']=y
    y=re.search(r"(?<=N° téléphone:).*(?=N° télécopie)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['plateau_technique_num_telephone']=y
    y=re.search(r"(?<=N° télécopie:).*(?=Courriel)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['plateau_technique_num_telecopie']=y
    y=re.search(r"(?<=Courriel:).*(?=Tâches confiées)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['plateau_technique_courriel']=y
    y=re.search(r"(?<=Tâches confiées:).*(?=CPP)",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['plateau_technique_taches_confiees']=y
    
#    #CPP
#    x=re.search(r"(?<=Nom du CPP:).*(?=Modalités de constitution ou non d’un comité de surveillance indépendant)",texte1).group()
#    infos['CPP']=x
    
    #CPP
    x=re.search(r"(?<=Nom du CPP:).*(?=Modalités de constitution ou non d’un comité de surveillance indépendant)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos['CPP']=x
    
#    #CPP
#    x=re.search(r"(?<=Modalités de constitution ou non d’un comité de surveillance indépendant:).*",texte1).group()
#    infos['comite_surveillance_independant']=x
    
    #CPP
    x=re.search(r"(?<=Modalités de constitution ou non d’un comité de surveillance indépendant:).*",texte).group()
    x=x.replace("#SAUT5089", "\n")
    infos['comite_surveillance_independant']=x
    
    #on formate pour enlever les potentiels blancs avant la valeur
    for k, v in infos.items():
        if type(v) is str:
            infos[k] = v.lstrip(" ")
        else:
            for i in range(len(v)):
                infos[k][i] = v[i].lstrip(" ")
          
    return infos
    
    
def extract2(dico):   
    
    #dabord on extrait tout qu'on ajoute dans une liste
    #f1 = open('Trame-simplifiée-cat-2.docx', 'rb') #ouvre le premier fichier
    f1 = open(dico['le_chemin'], 'rb')
    doc = Document(f1)
    fullText=[]
    for para in doc.paragraphs:
        para.text+="#SAUT5089"
        fullText.append(para.text)   
    f1.close()
    
    #puis on met tous les éléments de la liste bout a bout dans un immense string
    texte1=""
    for i in fullText:
        texte1+=i
    texte2=texte1.replace("\xa0","")
    texte=texte2.replace("\n","#SAUT5089")

    #creation du dico de donnes
    infos={}
    #on ajoute les éléments 1 par 1
    x=re.search(r"(?<=Titre complet de la recherche:).*(?=Nom ou titre)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos["titre_complet"]=x
    x=re.search(r"(?<=Nom ou titre abrégé:).*(?=N° de code du protocole)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos["titre_abrege"]=x
    x=re.search(r"(?<=Protocole ... version n°… du .../.../…\):).*(?=N°IDRCB:)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos["code_protocole"]=x
    x=re.search(r"(?<=N°IDRCB:).*(?=Identification du promoteur responsable)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos["num_idrcb"]=x
    
    #en premier on récupère tout le bloc promoteur
    x=re.search(r"(?<=Identification du promoteur responsable de la demande).*(?=Identification des investigateurs)",texte).group()
    #puis tous les éléments du promoteur 1 par 1
    y=re.search(r"(?<=Nom de l’organisme:).*(?=Nom de la personne à contacter:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['promoteur_nom_organisme']=y
    y=re.search(r"(?<=Nom de la personne à contacter:).*(?=Adresse:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['promoteur_nom_personne_contact']=y
    y=re.search(r"(?<=Adresse:).*(?=N° téléphone:)",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['promoteur_adresse']=y
    y=re.search(r"(?<=N° téléphone:).*(?=N° télécopie:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['promoteur_num_telephone']=y
    y=re.search(r"(?<=N° télécopie:).*(?=Courriel:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['promoteur_num_telecopie']=y
    y=re.search(r"(?<=Courriel:).*",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['promoteur_courriel']=y

    
    #idem pour investigateur coordinateur
    x=re.search(r"(?<=Investigateur coordinateur:).*(?=Autres investigateurs)",texte).group()
    #puis tous les éléments de l'investigateur coordinateur
    y=re.search(r"(?<=Nom:).*(?=Prénom:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_nom']=y
    y=re.search(r"(?<=Prénom:).*(?=Nom de l’établissement:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_prenom']=y
    y=re.search(r"(?<=Nom de l’établissement:).*(?=Service: )",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_nom_etablissement']=y
    y=re.search(r"(?<=Service:).*(?=Adresse: )",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_service']=y
    y=re.search(r"(?<=Adresse:).*(?=N° téléphone:)",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['investigateur_coordinateur_adresse']=y
    y=re.search(r"(?<=N° téléphone:).*(?=N° télécopie:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_telephone']=y
    y=re.search(r"(?<=N° télécopie:).*(?=Courriel:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_telecopie']=y
    y=re.search(r"(?<=Courriel:).*",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_courriel']=y
    
    #pour les autres investigateurs, plus subtile
    x=re.search(r"(?<=Autres investigateurs:).*(?=Justification de l’étude)",texte).group()
    #je m'en sers pour les regex
    x+="Nom: "
    x=x.replace("#SAUT5089", " ")
    #initialise avec des listes vides
    infos['autre_investigateur_nom']=[]
    infos['autre_investigateur_prenom']=[]
    infos['autre_investigateur_qualification']=[]
    infos['autre_investigateur_adresse_professionnelle']=[]
    infos['autre_investigateur_nom_etablissement']=[]
    infos['autre_investigateur_service']=[]
    infos['autre_investigateur_adresse']=[]
    infos['autre_investigateur_telephone']=[]
    infos['autre_investigateur_telecopie']=[]
    infos['autre_investigateur_courriel']=[]
    #compte le nombre d'autres investigateurs
    n=x.count('Nom: ')
    #on boucle pour chaque investigateur
    for i in range(n-1): 
        infos['autre_investigateur_nom'].append(re.search(r"(?<=Nom:).*?(?=Prénom:)",x).group())
        infos['autre_investigateur_prenom'].append(re.search(r"(?<=Prénom:).*?(?=Nom de l’établissement:)",x).group())
        infos['autre_investigateur_nom_etablissement'].append(re.search(r"(?<=Nom de l’établissement:).*?(?=Service:)",x).group())
        infos['autre_investigateur_service'].append(re.search(r"(?<=Service:).*?(?=Adresse:)",x).group())
        infos['autre_investigateur_adresse'].append(re.search(r"(?<=Adresse:).*?(?=N° téléphone:)",x).group())
        infos['autre_investigateur_telephone'].append(re.search(r"(?<=N° téléphone:).*?(?=N° télécopie:)",x).group())
        infos['autre_investigateur_telecopie'].append(re.search(r"(?<=N° télécopie:).*?(?=Courriel:)",x).group())
        infos['autre_investigateur_courriel'].append(re.search(r"(?<=Courriel:).*?(?=Nom:)",x).group())
        infos['autre_investigateur_qualification'].append("")
        #on recup la taille du premier bloc investigateur
        z=len(re.search(r"(?<=Nom:).*?(?=Nom: )",x).group())
        #on enleve ce bloc a x qui contient tous les investigateurs, ainsi a la prochaine boucle la regex se fera sur l'investigateur suivant
        x=x[(z-1):]
        
    
    #justification de létude
    table = doc.tables[0]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage dans courte
    courte=re.search(r"(?<=Bref rappel \(données de la littérature scientifique, pathologie, domaine d’étude\)\#SAUT5089).*",courte).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")
    infos['justification_etude_courte']=courte
    infos['justification_etude_longue']=longue
  
#    #benefices de l'étude
#    benefice=re.search(r"(?<=notamment les bénéfices escomptés pour les personnes qui se prêtent à la recherche\.).*(?=Risques:)",texte1).group()
#    infos['benefices']=benefice
#    
#    #risques de l'étude
#    risque=re.search(r"(?<=visant à éviter et/ou prendre en charge les événements inattendus\)\.).*(?=Retombées attendues)",texte1).group()
#    infos['risques']=risque
    
    #benefices de l'étude
    benefice=re.search(r"(?<=notamment les bénéfices escomptés pour les personnes qui se prêtent à la recherche\.).*(?=Risques:)",texte).group()
    benefice=benefice.replace("#SAUT5089", "\n")
    infos['benefices']=benefice
    
    #risques de l'étude
    risque=re.search(r"(?<=visant à éviter et/ou prendre en charge les événements inattendus\)\.).*(?=Retombées attendues)",texte).group()
    risque=risque.replace("#SAUT5089", "\n")
    infos['risques']=risque
    
    #retombées attendues
    table = doc.tables[1]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")
    #on ajoute '\n' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"

    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Description des retombées attendues par cette recherche\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=d’augmentation de l’arsenal thérapeutique,…\)\.\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")
    infos['retombee_attenduees_courte']=courte
    infos['retombee_attenduees_longue']=longue
#    #objectif principal
#    principal=re.search(r"(?<=Objectif Principal:).*(?=Objectif secondaires:)",texte1).group()
#    infos['objectif_principal']=principal
#    
#    #objectif secondaire
#    secondaire=re.search(r"(?<=Objectif secondaires:).*?(?=Critères de Jugement)",texte1).group()
#    infos['objectif_secondaire']=secondaire
#    infos['justification_existence']=""
    
    #objectif principal
    principal=re.search(r"(?<=Objectif Principal:).*(?=Objectif secondaires:)",texte).group()
    principal=principal.replace("#SAUT5089", "\n")
    infos['objectif_principal']=principal
    
    #objectif secondaire
    secondaire=re.search(r"(?<=Objectif secondaires:).*?(?=Critères de Jugement)",texte).group()
    secondaire=secondaire.replace("#SAUT5089", "\n")
    infos['objectif_secondaire']=secondaire
    infos['justification_existence']=""
    
    #critères de jugement principal
    table = doc.tables[2]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089") 
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Un seul critère correspondant à l’objectif principal \#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=Il permettra également le calcul de l’effectif de l’étude\.\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")    
    infos['critere_jugement_principal_courte']=courte
    infos['critere_jugement_principal_longue']=longue
    
    #critères de jugement secondaire
    table = doc.tables[3]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")     
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Liste de tous les critères de jugement secondaires\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=répondant aux objectifs secondaires\.\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")      
    infos['critere_jugement_secondaire_courte']=courte
    infos['critere_jugement_secondaire_longue']=longue
    
    #critères d'inclusion
    table = doc.tables[4]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")      
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=à la partie correspondante dans le corps du protocole § 4\.1\)\#SAUT5089).*",courte).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")                          
    infos['critere_inclusion_courte']=courte
    infos['critere_inclusion_longue']=longue
    
    #critères de non inclusion
    table = doc.tables[5]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")     
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=à la partie correspondante dans le corps du protocole § 4\.2\)\#SAUT5089).*",courte).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")             
    infos['critere_non_inclusion_courte']=courte
    infos['critere_non_inclusion_longue']=longue
    
#    #justification inclusion
#    justif=re.search(r"(?<=Justifications de l’inclusion de personnes visées:).*(?=Modalités de recrutements)",texte1).group()
#    infos['justification_inclusion']=justif
    
    #justification inclusion
    justif=re.search(r"(?<=Justifications de l’inclusion de personnes visées:).*(?=Modalités de recrutements)",texte).group()
    justif=justif.replace("#SAUT5089", "\n")
    infos['justification_inclusion']=justif
#    
#    #modalités_recrutement
#    recru=re.search(r"(?<=Modalités de recrutements:).*(?=Traitement/stratégie/procédures :)",texte1).group()
#    infos['modalite_recrutement']=recru
    
    #modalités_recrutement
    recru=re.search(r"(?<=Modalités de recrutements:).*(?=Traitement/stratégie/procédures :)",texte).group()
    recru=recru.replace("#SAUT5089", "\n")
    infos['modalite_recrutement']=recru
    
    #traitement et stratégie
    table = doc.tables[6]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")     
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=traitements/stratégies/procédures\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=la durée du traitement et de la voie d’administration\.\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")                         
    infos['traitement_strategie_courte']=courte
    infos['traitement_strategie_longue']=longue
   
    #taille de l'étude
    table = doc.tables[7]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")     
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Nombre de personnes à inclure: \#SAUT5089).*",courte).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")     
    infos['taille_etude_courte']=courte
    infos['taille_etude_longue']=longue
    
#    #modalités de l'indemnisation
#    indem=re.search(r"(?<=Modalités et montant de l’indemnisation des personnes se prêtant à la recherche:).*?(?=Durée)",texte1).group()
#    infos['indemnisation']=indem
#   
#    #durée des inclusions
#    x=re.search(r"(?<=Durée prévue des inclusions:).*(?=Durée de participation pour une personne se prêtant à la recherche)",texte1).group()
#    infos['duree_inclusion']=x
#    
#    #durée de participation
#    x=re.search(r"(?<=c'est-à-dire la dernière visite du dernier patient inclus.).*(?=Durée totale de l’étude)",texte1).group()
#    infos['duree_participation']=x
#    
#    #durée totale
#    x=re.search(r"(?<=Durée totale de l’étude:).*(?=Analyse statistiques des données)",texte1).group()
#    infos['duree_totale_etude']=x
    
    #modalités de l'indemnisation
    indem=re.search(r"(?<=Modalités et montant de l’indemnisation des personnes se prêtant à la recherche:).*?(?=Durée)",texte).group()
    indem=indem.replace("#SAUT5089", "\n")
    infos['indemnisation']=indem
   
    #durée des inclusions
    x=re.search(r"(?<=Durée prévue des inclusions:).*(?=Durée de participation pour une personne se prêtant à la recherche)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos['duree_inclusion']=x
    
    #durée de participation
    x=re.search(r"(?<=c'est-à-dire la dernière visite du dernier patient inclus.).*(?=Durée totale de l’étude)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos['duree_participation']=x
    
    #durée totale
    x=re.search(r"(?<=Durée totale de l’étude:).*(?=Analyse statistiques des données)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos['duree_totale_etude']=x
    
    #analyse statistique
    table = doc.tables[8]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")     
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Bref rappel des méthodes statistiques\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=données manquantes, inutilisées ou non valides\.\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")                         
    infos['analyse_statistique_courte']=courte
    infos['analyse_statistique_longue']=longue
    
#    #CPP
#    x=re.search(r"(?<=Motifs de constitution ou non d’un comité de surveillance indépendant:).*",texte1).group()
#    infos['comite_surveillance_independant']=x
    
    #CPP
    x=re.search(r"(?<=Motifs de constitution ou non d’un comité de surveillance indépendant:).*",texte).group()
    x=x.replace("#SAUT5089", "\n")
    infos['comite_surveillance_independant']=x
    
    infos['promoteur_UE_nom_organisme']=""
    infos['promoteur_UE_nom_personne_contact']=""
    infos['promoteur_UE_adresse']=""
    infos['promoteur_UE_num_telephone']=""
    infos['promoteur_UE_num_telecopie']=""
    infos['promoteur_UE_courriel']=""
    infos['demandeur_nom_organisme']=""
    infos['demandeur_UE_adresse']=""
    infos['demandeur_UE_num_telephone']=""
    infos['demandeur_UE_num_telecopie']=""
    infos['demandeur_UE_courriel']=""
    infos['demandeur_nom_personne_contact']=""
    infos['produit_nom']=""
    infos['produit_nom_code']=""
    infos['produit_voie_administration']=""
    infos['produit_dosage_concentration']=""
    infos['produit_dosage_unite_concentration']=""
    infos['placebo_numero']=""
    infos['placebo_numero_ME']=""
    infos['placebo_voie_administration']=""
    infos['fabriquant_placebo_nom']=""
    infos['fabriquant_placebo_adresse']=""
    infos['pathologie_etudiee']=""
    infos['classification_cim']=""
    infos['investigateur_coordinateur_adresse_professionnelle']=""
    infos['investigateur_coordinateur_qualification']=""
    infos['lieu_recherche_intitule']=""
    infos['lieu_recherche_num_autorisation']=""
    infos['lieu_recherche_delivre_le']=""
    infos['lieu_recherche_date_limite_validite']=""
    infos['lieu_recherche_nom_adresse']=""
            
    for k, v in infos.items():
        if type(v) is str:
            infos[k] = v.lstrip(" ")
        else:
            for i in range(len(v)):
                infos[k][i] = v[i].lstrip(" ")
            
    
    return infos
    
def extract3(dico):   
    
    #dabord on extrait tout qu'on ajoute dans une liste
    #f1 = open('Trame-simplifiée-cat-3.docx', 'rb') #ouvre le premier fichier
    f1 = open(dico['le_chemin'], 'rb')
    doc = Document(f1)
    fullText=[]
    for para in doc.paragraphs:
        para.text+="#SAUT5089"
        fullText.append(para.text)   
    f1.close()
    
    #puis on met tous les éléments de la liste bout a bout dans un immense string
    texte1=""
    for i in fullText:
        texte1+=i
    texte2=texte1.replace("\xa0","")
    texte=texte2.replace("\n","#SAUT5089")
    #creation du dico de donnes
    infos={}
    
    #on ajoute les éléments 1 par 1
    x=re.search(r"(?<=Titre complet de la recherche:).*(?=Acronyme)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos["titre_complet"]=x
    x=re.search(r"(?<=Acronyme:).*(?=Protocole version n°)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos["titre_abrege"]=x
    x=re.search(r"(?<=Protocole version n°... en date du .../.../….).*(?=Identification du promoteur responsable de la demande)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos["code_protocole"]=x
    
    #en premier on récupère tout le bloc promoteur
    x=re.search(r"(?<=Identification du promoteur responsable de la demande).*(?=Identification investigateur coordonnateur)",texte).group()
    #puis tous les éléments du promoteur 1 par 1
    y=re.search(r"(?<=Nom de l’organisme :).*(?=Adresse complète)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['promoteur_nom_organisme']=y
    y=re.search(r"(?<=Adresse complète:).*(?=N° téléphone:)",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['promoteur_adresse']=y
    y=re.search(r"(?<=N° téléphone:).*(?=N° télécopie:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['promoteur_num_telephone']=y
    y=re.search(r"(?<=N° télécopie:).*(?=Courriel:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['promoteur_num_telecopie']=y
    y=re.search(r"(?<=Courriel:).*",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['promoteur_courriel']=y

    
    #idem pour investigateur coordinateur
    x=re.search(r"(?<=Identification investigateur coordonnateur :).*(?=Justification de la recherche)",texte).group()
    #puis tous les éléments de l'investigateur coordinateur
    y=re.search(r"(?<=Nom investigateur:).*(?=Service:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_nom']=y
    y=re.search(r"(?<=Service:).*(?=Qualité:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_service']=y
    y=re.search(r"(?<=Qualité:).*(?=Adresse complète)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_qualite']=y
    y=re.search(r"(?<=Adresse complète:).*(?=N° téléphone:)",x).group()
    y=y.replace("#SAUT5089", "\n")
    infos['investigateur_coordinateur_adresse']=y
    y=re.search(r"(?<=N° téléphone:).*(?=N° télécopie:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_telephone']=y
    y=re.search(r"(?<=N° télécopie:).*(?=Courriel:)",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_telecopie']=y
    y=re.search(r"(?<=Courriel:).*",x).group()
    y=y.replace("#SAUT5089", " ")
    infos['investigateur_coordinateur_courriel']=y
    
    #justification de létude
    table = doc.tables[0]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")    
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage dans courte
    courte=re.search(r"(?<=Bref rappel \(données de la littérature scientifique, pathologie, domaine d’étude\)\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=justifier la pertinence de votre étude.\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")    
    infos['justification_etude_courte']=courte
    infos['justification_etude_longue']=longue
    
    #retombées attendues
    table = doc.tables[1]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")
    #on ajoute '\n' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"

    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Description des retombées attendues par cette recherche\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=Description détaillée des retombées attendues par cette recherche.\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")      
    infos['retombee_attenduees_courte']=courte
    infos['retombee_attenduees_longue']=longue
  
#    #objectif principal
#    principal=re.search(r"(?<=Objectif principal:).*(?=Objectif secondaire:)",texte1).group()
#    infos['objectif_principal']=principal
#    
#    #objectif secondaire
#    secondaire=re.search(r"(?<=Objectif secondaire:).*?(?=Critères de jugement)",texte1).group()
#    infos['objectif_secondaire']=secondaire
    
    #objectif principal
    principal=re.search(r"(?<=Objectif principal:).*(?=Objectif secondaire:)",texte).group()
    principal=principal.replace("#SAUT5089", "\n")
    infos['objectif_principal']=principal
    
    #objectif secondaire
    secondaire=re.search(r"(?<=Objectif secondaire:).*?(?=Critères de jugement)",texte).group()
    secondaire=secondaire.replace("#SAUT5089", "\n")
    infos['objectif_secondaire']=secondaire
    
    #critères de jugement principal
    table = doc.tables[2]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")    
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Un seul critère correspondant à l’objectif principal\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=la nécessité d’être validé par un comité.\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")                         
    infos['critere_jugement_principal_courte']=courte
    infos['critere_jugement_principal_longue']=longue
    
    #critères de jugement secondaire
    table = doc.tables[3]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")    
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Liste de tous les critères de jugement secondaires\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=la forme du critère,\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")
    infos['critere_jugement_secondaire_courte']=courte
    infos['critere_jugement_secondaire_longue']=longue
    
#    inclu=re.search(r"(?<=Critères d’inclusion:).*(?=Critères de non inclusion)",texte1).group()
#    infos['criteres_inclusion']=inclu
#    
#    noninclu=re.search(r"(?<=Critères de non inclusion:).*(?=Traitements/Stratégies/Procédures)",texte1).group()
#    infos['criteres_non_inclusion']=noninclu
    
    inclu=re.search(r"(?<=Critères d’inclusion:).*(?=Critères de non inclusion)",texte).group()
    inclu=inclu.replace("#SAUT5089", "\n")
    infos['criteres_inclusion']=inclu
    
    noninclu=re.search(r"(?<=Critères de non inclusion:).*(?=Traitements/Stratégies/Procédures)",texte).group()
    noninclu=noninclu.replace("#SAUT5089", "\n")
    infos['criteres_non_inclusion']=noninclu
    
    #traitement et stratégie
    table = doc.tables[4]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")    
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=traitements/stratégies/procédures\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=la procédure à l’étude\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")                        
    infos['traitement_strategie_courte']=courte
    infos['traitement_strategie_longue']=longue
   
    #taille de l'étude
    table = doc.tables[5]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")    
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Nombre de Patients\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=dans chaque lieu de recherches, le cas échéant\.\#SAUT5089).*",longue).group()  
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")    
    infos['taille_etude_courte']=courte
    infos['taille_etude_longue']=longue

#    #durée des inclusions
#    x=re.search(r"(?<=Durée de la période d’inclusion:).*(?=Durée de la participation pour chaque participant)",texte1).group()
#    infos['duree_inclusion']=x
#    
#    #durée de participation
#    x=re.search(r"(?<=Durée de la participation pour chaque participant:).*(?=Durée totale de l’étude)",texte1).group()
#    infos['duree_participation']=x
#    
#    #durée totale
#    x=re.search(r"(?<=Durée totale de l’étude:).*(?=Analyse statistique des données)",texte1).group()
#    infos['duree_totale_etude']=x
    
    #durée des inclusions
    x=re.search(r"(?<=Durée de la période d’inclusion:).*(?=Durée de la participation pour chaque participant)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos['duree_inclusion']=x
    
    #durée de participation
    x=re.search(r"(?<=Durée de la participation pour chaque participant:).*(?=Durée totale de l’étude)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos['duree_participation']=x
    
    #durée totale
    x=re.search(r"(?<=Durée totale de l’étude:).*(?=Analyse statistique des données)",texte).group()
    x=x.replace("#SAUT5089", " ")
    infos['duree_totale_etude']=x
    
    #analyse statistique
    table = doc.tables[6]  
    #créé une liste des valeurs des cases
    data = [] 
    keys = None
    #on récupère toutes les valeur dans une liste dont chaque valeur est un dico
    #la clé est la valeur de la colonne de gauche et la valeur celle de la cellule de droite
    #chaque valeur du dico est de la fçon suivante: key, value où key=cellule de gauche et value= cellule de droite
    for i, column in enumerate(table.columns):
        text = (cell.text for cell in column.cells)
    #créé le dictionnaire
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    tab=data[0]
    for key, value in tab.items():
        courte=key
        longue=value
    courte=courte.replace("\xa0","")
    longue=longue.replace("\xa0","")
    courte=courte.replace("\n","#SAUT5089")
    longue=longue.replace("\n","#SAUT5089")        
    #on ajoute ' ' pour eviter l'erreur avec les regex en cas de non remplissage par l'investigateur
    courte+="\n"
    longue+="\n"
    #on retire l'aide au remplissage 
    courte=re.search(r"(?<=Bref rappel des méthodes statistiques\#SAUT5089).*",courte).group()
    longue=re.search(r"(?<=compris le calendrier des analyses intermédiaires prévues\.\#SAUT5089).*",longue).group()
    courte=courte.replace("#SAUT5089","\n")
    longue=longue.replace("#SAUT5089","\n")                         
    infos['analyse_statistique_courte']=courte
    infos['analyse_statistique_longue']=longue
    
    for k, v in infos.items():
        if type(v) is str:
            infos[k] = v.lstrip(" ")
        else:
            for i in range(len(v)):
                infos[k][i] = v[i].lstrip(" ")
            
    
    return infos
    
    