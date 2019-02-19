# -*- coding: utf-8 -*-
"""
Created on Mon Feb 18 11:51:47 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style, Titre3, TexteGris
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#revoir titre1 et texte encardé gris

def Page16():
    'Creation de la page 16 du protcole de catégorie 1'
    document = docx.Document()
   # styles = document.styles

#    from docx.oxml.ns import nsdecls
#    from docx.oxml import parse_xml

#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 
   # shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))) #CREER LE FOND GRIS

    StyleProt1.Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    #AJOUTER SCHEMA DEROULEMENT
    
    
   # Ecriture du 6.1  
    document.add_paragraph('6.1 Calendrier de la recherche\n', style='Titre2') 
    
    # Ecriture du 6.2  
    document.add_paragraph('6.2	Tableau récapitulatif du suivi d’un participant à la recherche\n', style='Titre2') 
    
    #AJOUTER TABLEAU
    
    # Ecriture du 6.3  
    document.add_paragraph('6.3	Visites de pré-inclusion / inclusion = Visite V0\n', style='Titre2') 
    
    
    
    #Ecriture du titre6.3.1
    StyleProt1.Titre3('6.3.1','Recueil du consentement',document)
   
    
    TexteGris('prendre contact avec la promotion interne \n pour aide a la redaction de ce chapitre', document)
    

    #TEXTE
    
    #Ecriture du titre6.3.2
    StyleProt1.Titre3('6.3.2','Déroulement de la visite',document)

    
    #Ecriture du titre 6.4
    document.add_paragraph('6.4	Visite de randomisation = Visite (Vx, ou Jx, ou Mx…)', style='Titre2') 

    #Ecriture du titre6.4.1
    StyleProt1.Titre3('6.4.1','Description des examens',document)

    
    #Ecriture du titre6.4.2
    StyleProt1.Titre3('6.4.2','Randomisation du patient',document)



    #Ecriture du titre 6.5
    document.add_paragraph('6.5	Visites de suivi = visite (Vx, ou Jx ou Sx ou Mx…)', style='Titre2') 

    #Ecriture du titre6.5.1
    StyleProt1.Titre3('6.5.1','Visite (Vx, ou Sx, ou Jx, ou Mx…)',document)

    
    #Ecriture du titre6.5.2
    StyleProt1.Titre3('6.5.2','Visite (Vx, ou Sx, ou Jx, ou Mx…)',document)

    
    #Ecriture du titre 6.6
    document.add_paragraph('6.6	Visite de fin de la recherche', style='Titre2') 
    
    #Ecriture du titre 6.7
    document.add_paragraph('6.7	Règles d’arrêt de la participation d’une personne à la recherche', style='Titre2') 

    TexteGris('prendre contact avec la promotion interne \n pour aide a la redaction de ce chapitre', document)

#	 #Texte sur fond gris   
#    table = document.add_table(rows = 1, cols = 1)
#    row = table.rows[0].cells
#    para_text = 'prendre contact avec la promotion interne \n pour aide a la redaction de ce chapitre'
#    cell = row[0]
#    pt = cell.paragraphs[0]
#    t = pt.text = ''
#    p = pt.add_run(para_text)
#    cell._tc.get_or_add_tcPr().append(shading_elm)
#    p.style='BackgroundGrey'
#    pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
	
    #Ecriture du titre6.7.1
    StyleProt1.Titre3('6.7.1','Arrêt de participation définitif ou temporaire d’un patient dans l’étude)',document)


    #Ecriture du titre6.7.2
    StyleProt1.Titre3('6.7.2','Modalités de remplacement des patients exclus, le cas échéant',document)

    
    #Ecriture du titre6.7.3
    StyleProt1.Titre3('6.7.3','Modalités et calendrier de recueil pour ces données',document)

    
    #Ecriture du titre6.7.4
    StyleProt1.Titre3('6.7.4','Modalités de suivi de ces personnes',document)

    #Ecriture du titre 6.8
    document.add_paragraph('6.8	Contraintes liées à la recherche et indemnisation éventuelle des participants', style='Titre2') 
    
    #Ecriture du titre 6.9
    document.add_paragraph('6.9	Collection d’échantillons biologiques', style='Titre2') 
    
    
    TexteGris('prendre contact avec la promotion interne \n pour aide a la redaction de ce chapitre', document)

    
    
    StyleProt1.Titre3('6.9.1','Objectifs',document)
    StyleProt1.Titre3('6.9.2','Description de(s) (la) collection(s) ',document)
    StyleProt1.Titre3('6.9.3','Conservation',document)
    StyleProt1.Titre3('6.9.4','Devenir de la collection',document)
    
    #Ecriture du titre 6.10
    document.add_paragraph('6.10	Arrêt d’une partie ou de la totalité de la recherche', style='Titre2') 
    

    
    
    document.save("page16.docx")   