# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 15:06:49 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document) --> écrire en minuscule !!!
#    TexteGrisJustif(texte,document)

def Partie11(document,extract):
    'Creation de la partie 11 du protcole de catégorie 1'
 #   document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

  #  Style(document)


#    
#---------------------------------------------------------------ECRITURE
    
    
     #ecriture du premier titre 
    Titre1('11	ASPECTS STATISTIQUES',document)
    
     #Texte gris centré
    TexteGris('prendre contact avec la plateforme de methodologie \n pour aide a la redaction de ces chapitres',document)
    
    
    #Ecriture du 11.1  
    Titre2('11.1	Calcul de la taille d’étude',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['critere_jugement_principal_courte'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(11)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['taille_etude_longue'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(11)
#    #Ecriture du 11.2  
    Titre2('11.2	Méthodes statistiques employées',document)
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['analyse_statistique_longue'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(11)
    #Ecriture du 11.3  
    Titre2('11.3	Analyse de la sécurité',document)
    
    #Texte gris centré
    TexteGris('prendre contact avec la cellule de vigilance \n pour aide a la redaction de ce chapitre',document)
    
    p=document.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run1=p.add_run('Un plan d’analyse détaillé sera défini et fera l’objet d’une validation par le Conseil Scientifique de l’étude, s’il existe. Les modifications ultérieures devront intervenir avant la levée d’insu sur la base de données et seront systématiquement validées par le Conseil Scientifique.')
    run1.style='Paragraphe'
    
    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    #document.save("Partie7.docx") 
   
  #  document.save("Partie11.docx")