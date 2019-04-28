# -*- coding: utf-8 -*-
"""
Created on Tue Apr 16 11:21:37 2019

@author: Asuspc

"""
import docx
import extraction
from docx.enum.text import WD_ALIGN_PARAGRAPH
#from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_SECTION
import re
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import StyleProt1
from StyleProt1 import Style, Titre1,Titre2, Titre3, TexteGris, TexteGrisJustif


#extract=extraction.extract1(dico)
document = docx.Document()

'''Marge des page'''
sections = document.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

def PageGarde(document,extract):
    
    
 #   document = docx.Document()
    Style(document)
    
    sections = document.sections
    page_garde = sections[0]
        
    
    '''Logos de l'en-tete'''
    header = page_garde.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    r = p.add_run() 
    r.add_picture('imageGauche.png')
    r.add_text('                                                                                                                                     ')
    r.add_picture('imageDroite.png')
    
    
    '''Titre de la recherche'''
    paragraph = document.add_paragraph()
    sentence = paragraph.add_run('  \n\n'+extract['titre_complet'])
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence.font.name = 'Times New Roman'
    sentence.italic = True
    sentence.font.size = docx.shared.Pt(14) 
    
    #ajouter code
    
#    '''Acronyme''' #absent
#    paragraph2 = document.add_paragraph()
#    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
#    sentence = paragraph2.add_run(extract['titre_abrege'])
#    '''Then format the sentence'''
#    sentence.font.name = 'Times New Roman'
#    sentence.bold = True
#    sentence.font.size = docx.shared.Pt(22) 
    
    '''Version protocole'''
    paragraph2 = document.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run(extract['code_protocole'])
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(14) 
    sentence.bold = True
    
    '''Protcole cat 3'''
    paragraph = document.add_paragraph()
    sentence = paragraph.add_run('PROTOCOLE DE RECHERCHE NON INTERVETIONNELLE IMPLIQUANT LA PERSONNE HUMAINE (RNIPH) (Catégorie 3 - recherche sur données prospectives)')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence.font.name = 'Times New Roman'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(14) 
    
    '''N° ID-RCB  : '''
    paragraph2 = document.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run('Numéro ID-RCB : \n')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(12) 
    sentence.bold = True
    
    p=document.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run1=p.add_run('Cette recherche a obtenu le financement')
    run1.style='Paragraphe'
    
    '''Promoteur'''
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = paragraph.add_run('PROMOTEUR :\n')
    run1.font.name = 'Times New Roman'
    run1.font.size = docx.shared.Pt(12) 
    run1.bold = True
    run1.underline = True
    run2 = paragraph.add_run(extract['promoteur_nom_organisme']+'\n'+extract['promoteur_adresse']+'\nTél : '+extract['promoteur_num_telephone']+' / Fax : '+extract['promoteur_num_telecopie']+'\n')
    run2.font.name = 'Times New Roman'
    run2.font.size = docx.shared.Pt(12) 
    
    '''investigateur'''
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = paragraph.add_run('INVESTIGATEUR COORDONNATEUR :\n')
    run1.font.name = 'Times New Roman'
    run1.font.size = docx.shared.Pt(12) 
    run1.bold = True
    run1.underline = True
    run2 = paragraph.add_run(extract['investigateur_coordinateur_nom']+'\nService de : '+extract['investigateur_coordinateur_service']+'\n'+extract['investigateur_coordinateur_adresse']+'\nTél : '+extract['investigateur_coordinateur_telephone']+' / Fax : '+extract['investigateur_coordinateur_telecopie']+'\nE-mail : '+extract['investigateur_coordinateur_courriel'])
    run2.font.name = 'Times New Roman'
    run2.font.size = docx.shared.Pt(12) 
    
    '''GIRCI SOHO'''
    paragraph2 = document.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run('Ce protocole a été conçu et rédigé à partir de la version 3.0 du 01/02/2017\ndu protocole-type du GIRCI SOHO\n')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(12) 
    sentence.bold = True
    
    '''Confidentiel'''
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = paragraph.add_run('CE DOCUMENT CONFIDENTIEL')
    run1.font.name = 'Times New Roman'
    run1.font.size = docx.shared.Pt(10) 
    run1.underline = True
    run2 = paragraph.add_run(' EST LA PROPRIETE DU CHU DE POITIERS.\nAUCUNE INFORMATION NON PUBLIEE FIGURANT DANS CE DOCUMENT NE PEUT ETRE DIVULGUEE SANS AUTORISATION ECRITE PREALABLE DU CHU DE POITIERS')
    run2.font.name = 'Times New Roman'
    run2.font.size = docx.shared.Pt(10)
    
    '''Pied de page'''
    footer = page_garde.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    r = p.add_run(' ')
    

    '''Fin de page'''
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)

    
   
    #document.save("page_garde.docx")                   #sauvegarde
   
def Page_version(document,extract):
    
    '''Logos de l'en-tete'''
    
    page_sign = document.add_section()
    header2 = page_sign.header
    header2.is_linked_to_previous = False
    p = header2.paragraphs[0]
    p.alignment = 2
    r = p.add_run() 
    r.add_text('\t\t'+extract['titre_abrege'])
    p2 = header2.add_paragraph()
    r2 = p2.add_run() 
    r2.add_picture('imageGauche3.png')
    
    '''Titre'''
    paragraph2 = document.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run('HISTORIQUE DES MISES A JOUR DU PROTOCOLE')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(16)
    
    table = document.add_table(rows=5, cols=3, style='Table Grid')
    table.autofit = False
    for cell in table.columns[0].cells:
        cell.width =Cm(4)
    for cell in table.columns[1].cells:
        cell.width =Cm(4)
    for cell in table.columns[2].cells:
        cell.width =Cm(10)
    table.cell(0,0).text = 'Version'
    table.cell(0,1).text = 'Date' 
    table.cell(0,2).text = 'Raison de la Mise à Jour' 
    
    '''Pied de page'''
    footer = page_sign.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    r = p.add_run(extract['code_protocole']+'\tCONFIDENTIEL')
    r.font.name = 'Times New Roman'
    r.font.size = docx.shared.Pt(11)
     
    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def PageSignature(document,extract):
    
    #document = docx.Document()

    
#    p.style = document.styles["Header"]
        
    '''Titre'''
    paragraph2 = document.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run('PAGE DE SIGNATURE DU PROTOCOLE')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(16) 
    
    '''Signature investigateur'''
    paragraph2 = document.add_paragraph()
    sentence = paragraph2.add_run('Signature de l’investigateur')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(11) 
    sentence.bold = True
    sentence.underline = True
    
    '''Premiere case'''
    table = document.add_table(rows=1, cols=1, style='Table Grid')
    text1 = ' \nJ\'ai lu ce protocole d’essai clinique dont le CHU de Poitiers est le promoteur. Je confirme qu\'il contient toutes les informations nécessaires à la conduite de l’essai. Je m\'engage à mener cet essai en respectant ses directives et les termes et conditions qui y sont définis.\n'
    text2 = 'Je m\'engage à réaliser l’essai en respectant :\n\n'
    text3 = '    -  les principes de la “Déclaration d’Helsinki”, \n\
    -  les règles et recommandations de bonnes pratiques cliniques internationales (ICH-E6) et française      (règles de bonnes pratiques cliniques pour les recherches portant sur des médicaments à usage humain - décisions du 24 novembre 2006), \n\
    -  la législation nationale et la réglementation relative aux essais cliniques,\n\
    -  la conformité avec la Directive Essais Cliniques de l’UE [2001/20/EC].\n\n\n'
    text4 = "Je m'engage également à ce que les investigateurs et les autres membres qualifiés de mon équipe aient accès au protocole et aux documents relatifs à la conduite de l’essai pour leur permettre de travailler dans le respect des dispositions figurant dans ces documents.\n"
    text5 = "Investigateur : Dr/ Pr XXXXX\n(Prénom NOM)\n\n\n\n"
    text6 = "Signature : ……………………………………………..                          Date : ___________________\n"       

    table.cell(0,0).text = text1 +text2 + text3 + text4 +text5+text6
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= docx.shared.Pt(11)
                    font.name = 'Times New Roman'
                    
    '''Signature investigateur coordonnateur'''
    paragraph2 = document.add_paragraph()
    sentence = paragraph2.add_run(' \nSignature de l’Investigateur Coordonnateur')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(11) 
    sentence.bold = True
    sentence.underline = True
    
    '''Deuxieme case'''
    table2 = document.add_table(rows=1, cols=1, style='Table Grid')
    text5 = "Investigateur Coordonnateur : Dr/ Pr XXXXX\n(Prénom NOM)\n\n\n"
    text6 = "Signature : ……………………………………………..                          Date : ___________________\n" 
    table2.cell(0,0).text = text5 + text6
    for row in table2.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= docx.shared.Pt(11)
                    font.name = 'Times New Roman'
                    
                    
    '''Signature investigateur coordonnateur'''
    paragraph2 = document.add_paragraph()
    sentence = paragraph2.add_run(' \nSignature de l’Investigateur Coordonnateur')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(11) 
    sentence.bold = True
    sentence.underline = True
    
    '''Troisieme case'''
    table3 = document.add_table(rows=1, cols=1, style='Table Grid')
    text5 = "Promoteur : Jean-Pierre DEWITTE\nPour le Directeur Général et par délégation\nle Directeur de la Recherche,\n\n\n"
    text6 = "Signature : ……………………………………………..                          Date : ___________________\n" 
    table3.cell(0,0).text = text5 + text6
    for row in table3.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= docx.shared.Pt(11)
                    font.name = 'Times New Roman'
  
    
    
    
    '''Fin de page'''
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    
    
    #document.save("page_signature.docx")                   #sauvegarde
    
def PageCorespondant(document,extract):
    
    #document = docx.Document()
    
#    p.style = document.styles["Header"]
        
    '''Titre'''
    paragraph2 = document.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run('PRINCIPAUX CORRESPONDANTS')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(16)
    
    page_corespondant = document.add_section(WD_SECTION.CONTINUOUS)
    sectPr = page_corespondant._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')
    
    paragraph2 = document.add_paragraph()
    sentence1 = paragraph2.add_run('Investigateur coordonnateur/principal\n')
    '''Then format the sentence'''
    sentence1.font.name = 'Times New Roman'
    sentence1.font.size = docx.shared.Pt(12) 
    sentence1.bold = True  
    sentence2 = paragraph2.add_run(extract['investigateur_coordinateur_nom']+'\nService: '+extract['investigateur_coordinateur_service']+'\n'+extract['investigateur_coordinateur_adresse']+'\nTél : '+extract['investigateur_coordinateur_telephone']+'\nFax : '+extract['investigateur_coordinateur_telecopie']+'\nE-mail : ')
    '''Then format the sentence'''
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10) 
    sentence3 = paragraph2.add_run(extract['investigateur_coordinateur_courriel'])
    '''Then format the sentence'''
    sentence3.font.name = 'Times New Roman'
    sentence3.font.size = docx.shared.Pt(10)
    sentence3.underline = True

    paragraph2 = document.add_paragraph()
    sentence = paragraph2.add_run('Autres Spécialités')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.font.size = docx.shared.Pt(12) 
    sentence.bold = True 
    
    paragraph2 = document.add_paragraph()
    sentence1 = paragraph2.add_run('Plateforme  Méthodologie et Biostatistiques\n')
    '''Then format the sentence'''
    sentence1.font.name = 'Times New Roman'
    sentence1.font.size = docx.shared.Pt(12) 
    sentence1.bold = True  
    
    docu = document.add_section(WD_SECTION.NEW_PAGE)
    sectPr = docu._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'1')
    
#    '''Fin de page'''
#    paragraph = document.add_paragraph()
#    run = paragraph.add_run()
#    run.add_break(WD_BREAK.PAGE)
    

def liste_abreviation(document,extract):
   
      
      '''Titre'''
      paragraph2 = document.add_paragraph()
      paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
      sentence = paragraph2.add_run('LISTE DES ABREVIATIONS')
      '''Then format the sentence'''
      sentence.font.name = 'Times New Roman'
      sentence.bold = True
      sentence.font.size = docx.shared.Pt(16) 
      paragraph2 = document.add_paragraph('')
      texte=''
      ansm=amm=arc=bpc=cis=cnil=cpp=crf=e_crf=evi=evig=eig=eigi=ide=mr=rcp=susar=tec=False #pour savoir si on les a deja trouvés une fois
      '''Parcours tout l'extract pour voir si on trouve le mot quelque part dans le texte'''
      for para in extract.values():
          if para !="" and para !=" " and isinstance(para, str):
          #ANSM
              if (re.search(r"ANSM",para) and ansm==False): #si on trouve le mot (re.search renvoi True s'il trouve le mot) pour la première fois
                  texte+='ANSM\t\t\tAgence Nationale de Sécurité du Médicaments et des produits de santé\n' #alors on ajoute la ligne a texte
                  ansm=True #on a trouve ANSM une fois pas besoin de l'ajouter a nouveau les fois suivantes
          #AMM
              if re.search(r"AMM",para) and amm==False:
                 texte+='AMM\t\t\tAutorisation de Mise sur le Marché\n'
                 amm=True
          
          #ARC
              if re.search(r"ARC",para) and arc==False: 
                  texte+='ARC\t\t\tAttaché de Recherche Clinique\n'
                  arc=True
          #BPC
              if re.search(r"BPC",para) and bpc==False:
                  texte+='BPC\t\t\tBonnes Pratiques Cliniques\n'
                  bpc=True
          
          #CNIL
              if re.search(r"CNIL",para) and cnil==False:
                  texte+='CNIL\t\t\tCommission Nationale de l’Informatique et des Libertés\n'
                  cnil=True
          #CPP
              if re.search(r"CPP",para) and cpp==False: 
                  texte+='CPP\t\t\tComité de Protection des Personnes\n' 
                  cpp=True
          #CRF
              if re.search(r"CRF",para) and crf==False: 
                  texte+='CRF\t\t\tCase Report Form (cahier d’observation)\n' 
                  crf=True
          #e-CRF
              if re.search(r"e-CRF",para) and e_crf==False: 
                  texte+='e-CRF\t\t\tCahier d’observation électronique\n'
                  e_crf=True
          
          #EvI
              if re.search(r"EvI",para) and evi==False: 
                  texte+='EvI\t\t\tEvènement Indésirable\n'
                  evi=True
          
          #IDE
              if re.search(r"IDE",para) and ide==False: 
                  texte+='IDE\t\t\tInfirmier (ère) Diplômé(e) d\'Etat\n'
                  ide=True
          
          #MR
              if re.search(r"MR",para) and mr==False: 
                  texte+='MR\t\t\tMéthodologie de Référence\n'
                  mr=True
          
          #RCP
              if re.search(r"RCP",para) and rcp==False: 
                  texte+='RCP\t\t\tRésumé des Caractéristiques d\'un Produit\n'
                  rcp=True
          
          #TEC
              if re.search(r"TEC",para) and tec==False: 
                  texte+='TEC\t\t\tTechnicien d\'Etude Clinique\n'
                  tec=True
                    
      paragraph2 = document.add_paragraph()
      sentence = paragraph2.add_run(texte)
      '''Then format the sentence'''
      sentence.font.name = 'Times New Roman'
      sentence.font.size = docx.shared.Pt(11)
     
     #FIN DU DOC 
      paragraph = document.add_paragraph()
      run = paragraph.add_run()
      run.add_break(WD_BREAK.PAGE)
    

     
def resume_protocole(document,extract):

     '''Titre'''
     paragraph2 = document.add_paragraph()
     paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
     sentence = paragraph2.add_run('RESUME DU PROTOCOLE VERSION XX')
     '''Then format the sentence'''
     sentence.font.name = 'Times New Roman'
     sentence.bold = True
     sentence.font.size = docx.shared.Pt(16)
     
     table = document.add_table(rows=16, cols=2, style='Table Grid')
     table.autofit = False
     for cell in table.columns[0].cells:
         cell.width =Cm(4)
     for cell in table.columns[1].cells:
         cell.width =Cm(14.5)
   
     table.cell(0,0).text = 'Titre'       
     table.cell(0,1).text = extract['titre_complet']
     table.cell(1,0).text = 'Promoteur'
     table.cell(1,1).text = extract['promoteur_nom_organisme']+'\n'+extract['promoteur_adresse']+'\nTél : '+extract['promoteur_num_telephone']+' / Fax : '+extract['promoteur_num_telecopie']
     table.cell(2,0).text = 'Investigateur Coordonnateur'
     table.cell(2,1).text = extract['investigateur_coordinateur_nom']+'\nService: '+extract['investigateur_coordinateur_service']+'\n'+extract['investigateur_coordinateur_adresse']+'\nTél : '+extract['investigateur_coordinateur_telephone']+' / Fax : '+extract['investigateur_coordinateur_telecopie']+'\n'+extract['investigateur_coordinateur_courriel']
     table.cell(3,0).text = 'Justification / contexte'
     table.cell(3,1).text = extract['justification_etude_courte']
     table.cell(4,0).text = 'Objectif Principal'
     table.cell(4,1).text = extract['objectif_principal']
     table.cell(5,0).text = 'Objectifs Secondaires'
     table.cell(5,1).text = extract['objectif_secondaire']
     table.cell(6,0).text = 'Critère de Jugement Principal'
     table.cell(6,1).text = extract['critere_jugement_principal_courte']
     table.cell(7,0).text = 'Critères de Jugement Secondaires'
     table.cell(7,1).text = extract['critere_jugement_secondaire_courte']
     table.cell(8,0).text = 'Schéma de la recherche'
     table.cell(9,0).text = 'Critères d\'Inclusion'
     table.cell(9,1).text = extract['criteres_inclusion']
     table.cell(10,0).text = 'Critères de Non Inclusion des Sujets'
     table.cell(10,1).text = extract['criteres_non_inclusion']
     table.cell(11,0).text = 'Traitements / Stratégies / Procédures'
     table.cell(11,1).text = extract['traitement_strategie_courte']
     table.cell(12,0).text = 'Taille d\'étude'
     table.cell(12,1).text = extract['taille_etude_courte']
     table.cell(13,0).text = 'Durée de la Recherche '
     table.cell(13,1).text = 'Durée de la période d\’inclusion : '+extract['duree_inclusion']+'\nDurée de la participation pour chaque participant : '+extract['duree_participation']+'\nDurée totale de l’étude : '+extract['duree_totale_etude']
     table.cell(14,0).text = 'Analyse statistique des données'
     table.cell(14,1).text = extract['analyse_statistique_courte']
     table.cell(15,0).text = 'Retombées attendues '
     table.cell(15,1).text = extract['retombee_attenduees_courte']
    
    
     for row in table.rows:
         for cell in row.cells:
             paragraphs = cell.paragraphs
             for paragraph in paragraphs:
                 for run in paragraph.runs:
                     font = run.font
                     font.size= docx.shared.Pt(11)
                     font.name = 'Times New Roman'
                     
    #FIN DU DOC 
     paragraph = document.add_paragraph()
     run = paragraph.add_run()
     run.add_break(WD_BREAK.PAGE)
     
        
def Page_abstract(document,extract):
    
    '''Titre'''
    paragraph2 = document.add_paragraph()
    paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sentence = paragraph2.add_run('ABSTRACT')
    '''Then format the sentence'''
    sentence.font.name = 'Times New Roman'
    sentence.bold = True
    sentence.font.size = docx.shared.Pt(16)
     
    document.add_paragraph('This research has been registered in http://www.clinicaltrials.gov/ the date under the n° numéro.')
    document.add_paragraph('Titre complet de la recherche en anglais et acronyme.')
    document.add_paragraph('Titre simplifié de la recherche de 120 caractères maximum en anglais.')
    document.add_paragraph('Nom du promoteur is the sponsor of this research.')
    document.add_paragraph('This research will be conducted with the support of nom de la firme pharmaceutique / source of grants (PHRC,…).')
    document.add_paragraph('Brief summary : courte description de la recherche et de son objectif principal en anglais, en 5 lignes environ. ')
    document.add_paragraph('Detailed description : résumé de la recherche en anglais comportant une partie justification scientifique détaillée de 10 lignes environ, description du traitement/stratégie/procédure en 3 lignes environ et description du suivi en 5 lignes environ. ')
    document.add_paragraph('Primary outcome: critère de jugement principal  et visite au cours de laquelle celui-ci est recueilli en anglais (exemples : at inclusion (D0) ou 6 months after inclusion).')
    document.add_paragraph('Secondary outcomes: liste de tous les critères de jugement secondaires et visites durant lesquels ceux-ci sont recueillis en anglais.')
    document.add_paragraph('•	Study design : description des principales caractéristiques de la recherche selon le type de recherche.')
    document.add_paragraph('•	Eligibility criteria: \no	inclusion criteria: liste des principaux critères d’inclusion en anglais.\no	exclusion criteria: liste des principaux critères de non inclusion en anglais.')
    document.add_paragraph('•	Arm number or label and arm type : brève description des bras du protocole (experimental/active comparator/placebo, comparator/sham comparator/no intervention/other.')
    document.add_paragraph('•	Interventions : description succincte des traitements/stratégies/procédures de la recherche, pour chacun des bras le cas échéant. ')
    document.add_paragraph('•	Number of subjects : taille d’étude.')
    document.add_paragraph('•	Statistical analysis : bref rappel des méthodes statistiques.')
    document.add_paragraph('•	Conditions : pathologie ou objet de la recherche. Utiliser des termes du MeSH (National Library of Medecine’s Medical Subject Headings).')
    document.add_paragraph('•	Key-words : mot-clés décrivant la recherche.')
     
    #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
     
def test():
    PageGarde(document,extract)
    Page_version(document,extract)
    PageCorespondant(document,extract)
    PageSignature(document,extract)
    liste_abreviation(document,extract)
    resume_protocole(document,extract)
    Page_abstract(document,extract)
    
    
                    
    document.save("test_cat1.docx") 