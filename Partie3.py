# -*- coding: utf-8 -*-
"""
Created on Sat Feb 16 13:42:50 2019

@author: Julie
"""

import docx
import StyleProt1
from StyleProt1 import Style, Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie3(document):
    'Creation de la partie 3 du protcole de catégorie 1'
   # document = docx.Document()



#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
 

    
   # Style(document)

    Titre1('3	CRITERES DE JUGEMENT',document)

    
   # Ecriture du 3.1 
    Titre2('3.1	Critère d’évaluation principal',document)
    
    # Ecriture du 3.2  
    Titre2('3.2	Critères d’évaluation secondaires',document)

            #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    
 #   document.save("Partie3.docx")   


