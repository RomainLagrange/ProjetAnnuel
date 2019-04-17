# -*- coding: utf-8 -*-
"""
Created on Thu Feb 21 16:41:19 2019

@author: Asuspc
"""

import docx
import StyleProt1
from StyleProt1 import Style,Titre1, Titre2, Titre3, TexteGris, TexteGrisJustif
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor, Inches

#MEMO POUR ECRIRE LES TITRES :
#    Titre1('num + texte du protocole',document)
#    Titre2('num + texte du protocole',document)
#    Titre3('numero','texte',document)
#    TexteGris(texte,document)
#    TexteGrisJustif(texte,document)

def Partie6(document):
    'Creation de la partie 6 du protcole de catégorie 2'
   # document = docx.Document()


#   Marge de la page
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

#---------------------------DEFINITIONS DES STYLES
 

  #  Style(document)

   
#---------------------------------------------------------------ECRITURE
    
    
    #ecriture du premier titre 
    Titre1('6	STRATEGIE(S) / PROCEDURES DE LA RECHERCHE',document)
    
    paragraph2 = document.add_paragraph()
    sentence2 = paragraph2.add_run(extract['traitement_strategie_longue'])
    sentence2.font.name = 'Times New Roman'
    sentence2.font.size = docx.shared.Pt(10)
    
   # Ecriture du 6.1  
    Titre2('6.1	Stratégie / procédure expérimental(e)',document)
    
    #Texte gris justifié
    TexteGrisJustif('Pour une stratégie/procédure',document)
    
    # Ecriture du 6.2  
    Titre2('6.2	Stratégie / Procédure de comparaison',document)
       
    #Texte gris justifié
    TexteGrisJustif('Pour une stratégie/procédure',document)
    
#        #FIN DU DOC 
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    
 #  document.save("Cat2Partie6")